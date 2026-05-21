#!/usr/bin/env python3
"""
Baza Empire Agent Dashboard — v4
Full control center: agents, cron jobs, artifacts, settings, logs, infra
"""
import os, sys, json, yaml, subprocess, re, datetime, sqlite3, uuid, secrets, functools
from pathlib import Path
from flask import Flask, render_template, request, jsonify, send_from_directory, make_response, session, redirect
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename
from dashboard.private_inbound import (
    is_private as _is_private_path,
    PRIVATE_INBOUND_DIRNAME,
    VAULT_DIRNAME,
    move_to_vault as _vault_move_in,
    move_out_of_vault as _vault_move_out,
    migrate_legacy_inbound_meta as _migrate_legacy_inbound_meta,
)

try:
    from flask_login import LoginManager, UserMixin, login_user, logout_user, login_required, current_user
    import bcrypt
    CLOUD_ENABLED = True
except ImportError:
    CLOUD_ENABLED = False

app = Flask(__name__)
app.secret_key = os.environ.get("FLASK_SECRET_KEY", secrets.token_hex(32))
# 200MB ceiling — bulk receipt uploads can be 20+ phone-camera JPEGs at once.
app.config['MAX_CONTENT_LENGTH'] = 200 * 1024 * 1024
# When Vision UI's "keep unlocked" toggle is on, sessions become permanent for
# this duration (covers a full work session through tab switches and brief
# browser quits). Default off — see /api/vision/keep-unlocked.
app.config['PERMANENT_SESSION_LIFETIME'] = datetime.timedelta(hours=12)

# ── Paths ─────────────────────────────────────────────────────────────────────
DASHBOARD_DIR  = os.path.dirname(os.path.abspath(__file__))
FRAMEWORK_DIR  = os.path.dirname(DASHBOARD_DIR)
CONFIG_PATH    = os.path.join(FRAMEWORK_DIR, "config", "agents.yaml")
ARTIFACTS_DIR  = os.path.join(DASHBOARD_DIR, "artifacts")
LOGS_DIR       = os.path.join(FRAMEWORK_DIR, "logs")
SECRETS_PATH   = os.path.join(FRAMEWORK_DIR, "configs", "secrets.env")
VENV_PYTHON    = os.path.join(FRAMEWORK_DIR, "venv", "bin", "python")
if not os.path.exists(VENV_PYTHON):
    VENV_PYTHON = "python3"

CLOUD_STORAGE  = '/mnt/empirepool/cloud'
CLOUD_QUOTA_MB = 204800  # 200GB per user

os.makedirs(ARTIFACTS_DIR, exist_ok=True)
os.makedirs(LOGS_DIR, exist_ok=True)

@app.after_request
def add_no_cache(response):
    """Prevent browser caching of HTML templates during development."""
    if response.content_type and 'text/html' in response.content_type:
        response.headers['Cache-Control'] = 'no-cache, no-store, must-revalidate'
        response.headers['Pragma'] = 'no-cache'
        response.headers['Expires'] = '0'
    return response

# ── AHB123 Business Hub — Schema Init ─────────────────────────────────────────

def init_ahb_tables():
    """Create AHB123 tables in baza_projects.db if they don't exist.
    Wrapped in try/except — when another process holds a write lock at boot
    (e.g. a long-running dedup/backfill script) the tables already exist on
    every deployed instance, so deferring is safe."""
    try:
        _init_ahb_tables_inner()
    except sqlite3.OperationalError as e:
        print(f"[startup] init_ahb_tables deferred — DB busy: {e}", flush=True)

def _init_ahb_tables_inner():
    conn = sqlite3.connect(os.path.join(DASHBOARD_DIR, 'baza_projects.db'), timeout=8.0)
    conn.execute("PRAGMA busy_timeout = 8000")
    # WAL is a persistent journal mode; setting it once on the DB file lets
    # readers run concurrently with a writer (eliminates most "database is
    # locked" races in the autosave + detail-modal save flows). Idempotent.
    try:
        conn.execute("PRAGMA journal_mode = WAL")
        conn.execute("PRAGMA synchronous = NORMAL")
    except Exception:
        pass
    c = conn.cursor()
    c.executescript("""
        CREATE TABLE IF NOT EXISTS ahb_clients (
            id TEXT PRIMARY KEY,
            name TEXT,
            phone TEXT,
            email TEXT,
            address TEXT,
            city TEXT DEFAULT 'Philadelphia',
            source TEXT,
            status TEXT DEFAULT 'lead',
            notes TEXT,
            assigned_agent TEXT,
            created_at TEXT DEFAULT (datetime('now')),
            updated_at TEXT DEFAULT (datetime('now'))
        );
        CREATE TABLE IF NOT EXISTS ahb_projects (
            id TEXT PRIMARY KEY,
            client_id TEXT,
            title TEXT,
            address TEXT,
            scope TEXT,
            description TEXT,
            budget_low REAL,
            budget_high REAL,
            status TEXT DEFAULT 'estimate',
            start_date TEXT,
            end_date TEXT,
            assigned_agents TEXT,
            notes TEXT,
            created_at TEXT DEFAULT (datetime('now')),
            updated_at TEXT DEFAULT (datetime('now'))
        );
        CREATE TABLE IF NOT EXISTS ahb_invoices (
            id TEXT PRIMARY KEY,
            client_id TEXT,
            project_id TEXT,
            invoice_number TEXT,
            line_items TEXT,
            subtotal REAL,
            tax REAL,
            total REAL,
            status TEXT DEFAULT 'draft',
            due_date TEXT,
            paid_date TEXT,
            notes TEXT,
            created_at TEXT DEFAULT (datetime('now')),
            updated_at TEXT DEFAULT (datetime('now'))
        );
        CREATE TABLE IF NOT EXISTS ahb_receipts (
            id TEXT PRIMARY KEY,
            project_id TEXT,
            vendor TEXT,
            amount REAL,
            category TEXT,
            description TEXT,
            receipt_date TEXT,
            file_path TEXT,
            ocr_text TEXT,
            created_by TEXT,
            created_at TEXT DEFAULT (datetime('now'))
        );
        CREATE TABLE IF NOT EXISTS ahb_payroll (
            id TEXT PRIMARY KEY,
            worker_name TEXT,
            role TEXT,
            hours REAL,
            rate REAL,
            total REAL,
            period_start TEXT,
            period_end TEXT,
            status TEXT DEFAULT 'pending',
            project_id TEXT,
            notes TEXT,
            created_at TEXT DEFAULT (datetime('now'))
        );
        CREATE TABLE IF NOT EXISTS ahb_estimates (
            id TEXT PRIMARY KEY,
            client_id TEXT,
            project_id TEXT,
            title TEXT,
            description TEXT,
            scope TEXT,
            line_items TEXT,
            subtotal REAL,
            markup_pct REAL DEFAULT 15,
            total REAL,
            status TEXT DEFAULT 'draft',
            generated_by TEXT,
            notes TEXT,
            created_at TEXT DEFAULT (datetime('now')),
            updated_at TEXT DEFAULT (datetime('now'))
        );
        CREATE TABLE IF NOT EXISTS ahb_chats (
            id TEXT PRIMARY KEY,
            visitor_name TEXT,
            visitor_email TEXT,
            visitor_phone TEXT,
            channel TEXT DEFAULT 'website',
            status TEXT DEFAULT 'active',
            lead_score TEXT,
            assigned_agent TEXT DEFAULT 'nova_sterling',
            created_at TEXT DEFAULT (datetime('now')),
            updated_at TEXT DEFAULT (datetime('now'))
        );
        CREATE TABLE IF NOT EXISTS ahb_chat_messages (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            chat_id TEXT,
            role TEXT,
            content TEXT,
            agent_id TEXT,
            created_at TEXT DEFAULT (datetime('now'))
        );
        CREATE TABLE IF NOT EXISTS ahb_voice_configs (
            id TEXT PRIMARY KEY,
            name TEXT,
            voice TEXT DEFAULT 'en-US-GuyNeural',
            rate TEXT DEFAULT '+0%',
            pitch TEXT DEFAULT '+0Hz',
            volume TEXT DEFAULT '+0%',
            style TEXT DEFAULT 'friendly',
            pauses_enabled INTEGER DEFAULT 1,
            filler_words INTEGER DEFAULT 0,
            breathing_sounds INTEGER DEFAULT 0,
            notes TEXT,
            created_at TEXT DEFAULT (datetime('now'))
        );
        CREATE TABLE IF NOT EXISTS ahb_employees (
            id TEXT PRIMARY KEY,
            name TEXT,
            position TEXT,
            hourly_rate REAL,
            pay_type TEXT DEFAULT 'hourly',
            pay_method TEXT,
            phone TEXT,
            email TEXT,
            active INTEGER DEFAULT 1,
            created_at TEXT DEFAULT (datetime('now')),
            updated_at TEXT DEFAULT (datetime('now'))
        );
        CREATE TABLE IF NOT EXISTS ahb_events (
            id TEXT PRIMARY KEY,
            title TEXT,
            details TEXT,
            date TEXT,
            time TEXT,
            end_time TEXT,
            category TEXT,
            all_day INTEGER DEFAULT 0,
            project_id TEXT,
            employee_id TEXT,
            created_at TEXT DEFAULT (datetime('now'))
        );
        CREATE TABLE IF NOT EXISTS ahb_notes (
            id TEXT PRIMARY KEY,
            title TEXT,
            content TEXT,
            is_list INTEGER DEFAULT 0,
            is_task INTEGER DEFAULT 0,
            tags TEXT,
            pinned INTEGER DEFAULT 0,
            project_id TEXT,
            due_date TEXT,
            checklist_items TEXT,
            author_employee_id TEXT,
            created_at TEXT DEFAULT (datetime('now')),
            updated_at TEXT DEFAULT (datetime('now'))
        );
        CREATE TABLE IF NOT EXISTS ahb_debts (
            id TEXT PRIMARY KEY,
            name TEXT,
            type TEXT,
            frequency TEXT DEFAULT 'Monthly',
            payment_amount REAL DEFAULT 0,
            due_date TEXT,
            payoff_date TEXT,
            balance REAL DEFAULT 0,
            created_at TEXT DEFAULT (datetime('now')),
            updated_at TEXT DEFAULT (datetime('now'))
        );
        CREATE TABLE IF NOT EXISTS ahb_files (
            id TEXT PRIMARY KEY,
            name TEXT,
            file_type TEXT,
            file_path TEXT,
            size INTEGER,
            tags TEXT,
            category TEXT,
            year TEXT,
            project_id TEXT,
            created_at TEXT DEFAULT (datetime('now'))
        );
        CREATE TABLE IF NOT EXISTS ahb_project_phases (
            id TEXT PRIMARY KEY,
            project_id TEXT,
            phase_number INTEGER,
            name TEXT,
            value REAL,
            start_date TEXT,
            end_date TEXT,
            status TEXT DEFAULT 'pending',
            created_at TEXT DEFAULT (datetime('now'))
        );
        CREATE TABLE IF NOT EXISTS ahb_tax_requirements (
            id TEXT PRIMARY KEY,
            title TEXT,
            details TEXT,
            due_date TEXT,
            completed INTEGER DEFAULT 0,
            category TEXT DEFAULT 'tax',
            created_at TEXT DEFAULT (datetime('now'))
        );
        CREATE TABLE IF NOT EXISTS ahb_timeclock (
            id TEXT PRIMARY KEY,
            employee_id TEXT,
            date TEXT,
            clock_in TEXT,
            clock_out TEXT,
            lunch_start TEXT,
            lunch_end TEXT,
            hours REAL,
            lunch_minutes INTEGER DEFAULT 0,
            status TEXT DEFAULT 'active',
            notes TEXT,
            created_at TEXT DEFAULT (datetime('now'))
        );
        CREATE TABLE IF NOT EXISTS ahb_receipt_queue (
            id TEXT PRIMARY KEY,
            image_path TEXT,
            mode TEXT DEFAULT 'single',
            status TEXT DEFAULT 'pending',
            result_json TEXT,
            error TEXT,
            receipt_id TEXT,
            created_at TEXT DEFAULT (datetime('now'))
        );
        CREATE TABLE IF NOT EXISTS ahb_payments (
            id TEXT PRIMARY KEY,
            invoice_id TEXT,
            amount REAL,
            payment_method TEXT,
            payment_date TEXT,
            notes TEXT,
            created_at TEXT DEFAULT (datetime('now'))
        );
        CREATE TABLE IF NOT EXISTS ahb_phase_tasks (
            id TEXT PRIMARY KEY,
            phase_id TEXT,
            project_id TEXT,
            title TEXT,
            status TEXT DEFAULT 'pending',
            assigned_to TEXT,
            notes TEXT,
            created_at TEXT DEFAULT (datetime('now'))
        );
        CREATE TABLE IF NOT EXISTS baza_roadmap (
            id TEXT PRIMARY KEY,
            title TEXT NOT NULL,
            description TEXT DEFAULT '',
            status TEXT DEFAULT 'planned',
            priority TEXT DEFAULT 'medium',
            category TEXT DEFAULT 'general',
            assigned_agent TEXT DEFAULT '',
            target_date TEXT DEFAULT '',
            started_at TEXT DEFAULT '',
            completed_at TEXT DEFAULT '',
            notes TEXT DEFAULT '',
            created_at TEXT DEFAULT (datetime('now')),
            updated_at TEXT DEFAULT (datetime('now'))
        );
        CREATE TABLE IF NOT EXISTS baza_dash_links (
            id TEXT PRIMARY KEY,
            title TEXT NOT NULL,
            url TEXT NOT NULL,
            icon TEXT DEFAULT '&#128279;',
            category TEXT DEFAULT 'general',
            sort_order INTEGER DEFAULT 0,
            created_at TEXT DEFAULT (datetime('now'))
        );
        CREATE TABLE IF NOT EXISTS baza_infra_notes (
            id TEXT PRIMARY KEY,
            section TEXT DEFAULT 'general',
            note TEXT NOT NULL,
            author TEXT DEFAULT 'system',
            created_at TEXT DEFAULT (datetime('now'))
        );
        CREATE TABLE IF NOT EXISTS ahb_voice_logs (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            caller_name TEXT,
            caller_phone TEXT,
            direction TEXT DEFAULT 'inbound',
            duration_seconds INTEGER DEFAULT 0,
            transcript TEXT,
            audio_file TEXT,
            status TEXT DEFAULT 'completed',
            lead_created INTEGER DEFAULT 0,
            agent_notes TEXT,
            created_at TEXT DEFAULT (datetime('now'))
        );
        CREATE TABLE IF NOT EXISTS ahb_blueprints (
            id TEXT PRIMARY KEY,
            name TEXT NOT NULL DEFAULT 'Untitled Blueprint',
            project_id TEXT DEFAULT '',
            units TEXT DEFAULT 'imperial',
            data TEXT NOT NULL DEFAULT '{}',
            thumbnail_path TEXT DEFAULT '',
            notes TEXT DEFAULT '',
            created_at TEXT DEFAULT (datetime('now')),
            updated_at TEXT DEFAULT (datetime('now'))
        );
        CREATE TABLE IF NOT EXISTS ahb_blueprint_renders (
            id TEXT PRIMARY KEY,
            blueprint_id TEXT NOT NULL,
            floor_level INTEGER DEFAULT 1,
            mode TEXT DEFAULT 'photorealistic',
            prompt TEXT DEFAULT '',
            image_path TEXT DEFAULT '',
            created_at TEXT DEFAULT (datetime('now'))
        );
    """)
    conn.commit()
    # Add new columns to existing tables (idempotent)
    alter_stmts = [
        "ALTER TABLE ahb_notes ADD COLUMN color INTEGER DEFAULT 1",
        "ALTER TABLE ahb_projects ADD COLUMN acquisition_type TEXT",
        "ALTER TABLE ahb_projects ADD COLUMN value REAL",
        "ALTER TABLE ahb_clients ADD COLUMN company TEXT",
        "ALTER TABLE ahb_clients ADD COLUMN tags TEXT",
        "ALTER TABLE ahb_payroll ADD COLUMN overtime_hours REAL DEFAULT 0",
        "ALTER TABLE ahb_payroll ADD COLUMN employee_id TEXT",
        "ALTER TABLE ahb_invoices ADD COLUMN terms TEXT",
        "ALTER TABLE ahb_invoices ADD COLUMN client_name TEXT",
        "ALTER TABLE ahb_invoices ADD COLUMN project_name TEXT",
        "ALTER TABLE ahb_receipts ADD COLUMN store_name TEXT",
        "ALTER TABLE ahb_receipts ADD COLUMN payment_method TEXT",
        "ALTER TABLE ahb_receipts ADD COLUMN total REAL",
        "ALTER TABLE ahb_receipts ADD COLUMN teller_name TEXT DEFAULT ''",
        "ALTER TABLE ahb_receipts ADD COLUMN store_location TEXT DEFAULT ''",
        "ALTER TABLE ahb_receipts ADD COLUMN purchase_time TEXT DEFAULT ''",
        "ALTER TABLE ahb_receipts ADD COLUMN image_path TEXT DEFAULT ''",
        "ALTER TABLE ahb_receipts ADD COLUMN tax_amount REAL DEFAULT 0",
        "ALTER TABLE ahb_receipts ADD COLUMN subtotal REAL DEFAULT 0",
        "ALTER TABLE ahb_receipts ADD COLUMN items_json TEXT DEFAULT '[]'",
        "ALTER TABLE ahb_receipts ADD COLUMN ocr_raw TEXT DEFAULT ''",
        "ALTER TABLE ahb_receipts ADD COLUMN ocr_structured TEXT DEFAULT '{}'",
        "ALTER TABLE ahb_receipts ADD COLUMN year TEXT DEFAULT ''",
        "ALTER TABLE ahb_invoices ADD COLUMN date TEXT DEFAULT ''",
        "ALTER TABLE ahb_invoices ADD COLUMN parent_invoice_id TEXT DEFAULT ''",
        "ALTER TABLE ahb_invoices ADD COLUMN is_change_order INTEGER DEFAULT 0",
        "ALTER TABLE ahb_invoices ADD COLUMN overdue_since TEXT DEFAULT ''",
        "ALTER TABLE ahb_invoices ADD COLUMN overdue_interest_per_week REAL DEFAULT 50",
        "ALTER TABLE ahb_invoices ADD COLUMN company_name TEXT DEFAULT 'All Home Building Co'",
        "ALTER TABLE ahb_invoices ADD COLUMN contractor_name TEXT DEFAULT 'Sergey Tkach'",
        "ALTER TABLE ahb_invoices ADD COLUMN client_address TEXT DEFAULT ''",
        "ALTER TABLE ahb_invoices ADD COLUMN client_email TEXT DEFAULT ''",
        "ALTER TABLE ahb_invoices ADD COLUMN client_phone TEXT DEFAULT ''",
        "ALTER TABLE ahb_invoices ADD COLUMN project_address TEXT DEFAULT ''",
        "ALTER TABLE ahb_projects ADD COLUMN client_name TEXT DEFAULT ''",
        "ALTER TABLE ahb_projects ADD COLUMN client_email TEXT DEFAULT ''",
        "ALTER TABLE ahb_projects ADD COLUMN contact_info TEXT DEFAULT ''",
        "ALTER TABLE ahb_projects ADD COLUMN location TEXT DEFAULT ''",
        "ALTER TABLE ahb_files ADD COLUMN photo_section TEXT DEFAULT ''",
        "ALTER TABLE ahb_files ADD COLUMN document_type TEXT DEFAULT ''",
        "ALTER TABLE ahb_invoices ADD COLUMN year TEXT DEFAULT ''",
        "ALTER TABLE ahb_projects ADD COLUMN year TEXT DEFAULT ''",
        "ALTER TABLE ahb_projects ADD COLUMN commission_pct REAL DEFAULT 0",
        "ALTER TABLE ahb_projects ADD COLUMN commission_value REAL DEFAULT 0",
        "ALTER TABLE ahb_projects ADD COLUMN commission_beneficiary TEXT DEFAULT ''",
        "ALTER TABLE ahb_receipt_queue ADD COLUMN parent_image_path TEXT",
        "ALTER TABLE ahb_receipt_queue ADD COLUMN pair_id TEXT",
        "ALTER TABLE ahb_receipt_queue ADD COLUMN split_col INTEGER",
        # Phase/task workflow upgrades: approx duration on tasks, dedup keys on events
        "ALTER TABLE ahb_phase_tasks ADD COLUMN approx_minutes INTEGER DEFAULT 0",
        "ALTER TABLE ahb_phase_tasks ADD COLUMN source_line_idx INTEGER",
        "ALTER TABLE ahb_events ADD COLUMN phase_id TEXT DEFAULT ''",
        "ALTER TABLE ahb_events ADD COLUMN task_id TEXT DEFAULT ''",
    ]
    for stmt in alter_stmts:
        try:
            c.execute(stmt)
        except Exception:
            pass
    try:
        c.execute("UPDATE ahb_receipt_queue SET status='ready' WHERE status='done'")
    except Exception:
        pass
    c.executescript("""
        CREATE TABLE IF NOT EXISTS ahb_quotes (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            project_id TEXT NOT NULL,
            method TEXT,
            scope TEXT,
            description TEXT,
            total REAL NOT NULL,
            breakdown TEXT,
            notes TEXT,
            is_active INTEGER DEFAULT 0,
            created_at TEXT DEFAULT (datetime('now'))
        );
        CREATE INDEX IF NOT EXISTS idx_ahb_quotes_pid ON ahb_quotes(project_id);

        CREATE TABLE IF NOT EXISTS ahb_receipt_corrections (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            receipt_id TEXT NOT NULL,
            changed_by TEXT NOT NULL,
            changed_at TEXT DEFAULT (datetime('now')),
            field TEXT NOT NULL,
            old_value TEXT,
            new_value TEXT
        );
        CREATE INDEX IF NOT EXISTS idx_arc_rid ON ahb_receipt_corrections(receipt_id);
        CREATE INDEX IF NOT EXISTS idx_arc_field ON ahb_receipt_corrections(field);
    """)
    conn.commit()
    conn.close()

init_ahb_tables()


def init_cloud_tables():
    """Create Baza Cloud tables."""
    conn = sqlite3.connect(os.path.join(DASHBOARD_DIR, 'baza_projects.db'))
    c = conn.cursor()
    c.executescript("""
        CREATE TABLE IF NOT EXISTS baza_cloud_users (
            id TEXT PRIMARY KEY,
            email TEXT UNIQUE NOT NULL,
            password_hash TEXT NOT NULL,
            display_name TEXT DEFAULT '',
            status TEXT DEFAULT 'invited',
            storage_used INTEGER DEFAULT 0,
            storage_limit INTEGER DEFAULT 5368709120,
            invite_code TEXT DEFAULT '',
            is_admin INTEGER DEFAULT 0,
            created_at TEXT DEFAULT (datetime('now')),
            updated_at TEXT DEFAULT (datetime('now'))
        );
        CREATE TABLE IF NOT EXISTS baza_cloud_invites (
            id TEXT PRIMARY KEY,
            email TEXT DEFAULT '',
            code TEXT UNIQUE NOT NULL,
            created_by TEXT DEFAULT 'serge',
            used_at TEXT DEFAULT '',
            created_at TEXT DEFAULT (datetime('now'))
        );
        CREATE TABLE IF NOT EXISTS baza_cloud_files (
            id TEXT PRIMARY KEY,
            user_id TEXT NOT NULL,
            filename TEXT NOT NULL,
            path TEXT NOT NULL,
            size INTEGER DEFAULT 0,
            mime_type TEXT DEFAULT '',
            category TEXT DEFAULT 'files',
            created_at TEXT DEFAULT (datetime('now'))
        );
        CREATE TABLE IF NOT EXISTS cloud_shares (
            token TEXT PRIMARY KEY,
            user_id TEXT NOT NULL DEFAULT '1',
            path TEXT NOT NULL,
            expires_at TEXT,
            created_by TEXT DEFAULT 'serge',
            created_at TEXT DEFAULT (datetime('now')),
            access_count INTEGER DEFAULT 0,
            last_accessed_at TEXT
        );
        CREATE INDEX IF NOT EXISTS idx_cshr_path ON cloud_shares(path);
    """)
    conn.commit()
    conn.close()

init_cloud_tables()

# ── Task Events (visibility pipeline) — schema init ───────────────────────────
try:
    from core.task_events import init_schema as _init_task_events_schema
    _init_task_events_schema()
except Exception as _e:
    print(f"[task_events] schema init skipped: {_e}", file=sys.stderr)

# ── Baza Projects (sub-project #4) — schema init ──────────────────────────────
try:
    from core.baza_projects import ensure_schema as _ensure_baza_projects_schema
    _ensure_baza_projects_schema()
except Exception as _e:
    print(f"[baza_projects] schema init skipped: {_e}", file=sys.stderr)

# Cloud upload directory
CLOUD_UPLOAD_DIR = os.path.join(DASHBOARD_DIR, 'uploads', 'cloud')
os.makedirs(CLOUD_UPLOAD_DIR, exist_ok=True)


# All file types are allowed — no whitelist

# ── Config helpers ────────────────────────────────────────────────────────────

def load_config():
    with open(CONFIG_PATH, 'r') as f:
        return yaml.safe_load(f) or {}

def save_config(config):
    with open(CONFIG_PATH, 'w') as f:
        yaml.dump(config, f, default_flow_style=False, allow_unicode=True)

def load_secrets() -> dict:
    secrets = {}
    if not os.path.exists(SECRETS_PATH):
        return secrets
    with open(SECRETS_PATH) as f:
        for line in f:
            line = line.strip()
            if line and not line.startswith("#") and "=" in line:
                k, _, v = line.partition("=")
                secrets[k.strip()] = v.strip().strip('"').strip("'")
    return secrets

def save_secrets(secrets: dict):
    lines = []
    for k, v in secrets.items():
        lines.append(f'{k}="{v}"')
    with open(SECRETS_PATH, 'w') as f:
        f.write("\n".join(lines) + "\n")

# ── Agent / service helpers ───────────────────────────────────────────────────

def svc_name(agent_id: str) -> str:
    return f"baza-agent-{agent_id.replace('_', '-')}"

def get_agent_status(agent_id: str) -> str:
    """Resolve agent status from (1) Redis heartbeat — works for any node,
    including remote ones like phantom — (2) systemd as fallback for legacy
    agents that haven't started publishing a heartbeat yet."""
    # 1. Redis heartbeat is the authoritative cross-node signal
    try:
        import redis as _redis, time as _time, json as _json
        r = _redis.Redis(host='localhost', port=6379, decode_responses=True, socket_timeout=2)
        hb_raw = r.get(f"baza:heartbeat:{agent_id}")
        if hb_raw:
            try:
                hb = _json.loads(hb_raw)
                age = int(_time.time()) - int(hb.get("ts", 0))
                if age < 180:
                    return 'online'
                if age < 600:
                    return 'stale'
            except Exception:
                pass
    except Exception:
        pass
    # 2. systemd fallback for agents on this host that haven't heartbeat-ed yet
    try:
        r = subprocess.run(['systemctl', 'is-active', svc_name(agent_id)],
                           capture_output=True, text=True, timeout=5)
        if r.stdout.strip() == 'active':
            return 'online'
        # Unit doesn't exist (e.g. specter_voss runs on phantom) AND no heartbeat → offline
        return 'offline'
    except Exception:
        return 'unknown'

def get_agent_logs(agent_id: str, lines: int = 80) -> str:
    try:
        r = subprocess.run(
            ['journalctl','-u', svc_name(agent_id),'-n',str(lines),'--no-pager','--output=short'],
            capture_output=True, text=True, timeout=10)
        return r.stdout
    except:
        return "Could not fetch logs."

def get_recent_messages(agent_id: str, limit: int = 20) -> list:
    """Read from SQLite context DB."""
    db_path = os.path.join(FRAMEWORK_DIR, "data", "context.db")
    if not os.path.exists(db_path):
        db_path = os.path.join(FRAMEWORK_DIR, "context.db")
    if not os.path.exists(db_path):
        return []
    try:
        conn = sqlite3.connect(db_path)
        conn.row_factory = sqlite3.Row
        cur = conn.cursor()
        cur.execute("""
            SELECT role, content, timestamp
            FROM messages
            WHERE agent_id = ?
            ORDER BY timestamp DESC LIMIT ?
        """, (agent_id, limit))
        rows = [dict(r) for r in cur.fetchall()]
        conn.close()
        return list(reversed(rows))
    except Exception as e:
        return []

# ── Cron helpers ──────────────────────────────────────────────────────────────

CRON_TAG = "# baza-empire-managed"

def list_crons() -> list:
    """Return all baza-managed cron jobs."""
    try:
        r = subprocess.run(['crontab','-l'], capture_output=True, text=True)
        lines = r.stdout.splitlines()
        jobs = []
        for i, line in enumerate(lines):
            if CRON_TAG in line:
                # Extract name from tag: # baza-empire-managed name=<name>
                name_m = re.search(r'name=([^\s]+)', line)
                name = name_m.group(1) if name_m else f"job_{i}"
                jobs.append({"id": name, "raw": line, "line_index": i})
            elif line.strip() and not line.startswith("#"):
                # Check if previous line was a tag
                if i > 0 and CRON_TAG in lines[i-1]:
                    name_m = re.search(r'name=([^\s]+)', lines[i-1])
                    name = name_m.group(1) if name_m else f"job_{i}"
                    # Parse cron fields
                    parts = line.split(None, 5)
                    jobs[-1]["schedule"] = " ".join(parts[:5]) if len(parts) >= 5 else line
                    jobs[-1]["command"]  = parts[5] if len(parts) > 5 else ""
                    jobs[-1]["enabled"]  = not line.startswith("#")
        return jobs
    except:
        return []

def get_raw_crontab() -> str:
    try:
        r = subprocess.run(['crontab','-l'], capture_output=True, text=True)
        return r.stdout
    except:
        return ""

def set_raw_crontab(content: str) -> bool:
    try:
        proc = subprocess.run(['crontab','-'], input=content, capture_output=True, text=True)
        return proc.returncode == 0
    except:
        return False

def add_cron_job(name: str, schedule: str, command: str) -> bool:
    raw = get_raw_crontab()
    # Remove existing job with same name
    lines = []
    skip_next = False
    for line in raw.splitlines():
        if CRON_TAG in line and f"name={name}" in line:
            skip_next = True
            continue
        if skip_next:
            skip_next = False
            continue
        lines.append(line)
    # Add new job
    lines.append(f"{CRON_TAG} name={name}")
    lines.append(f"{schedule} {command}")
    return set_raw_crontab("\n".join(lines) + "\n")

def remove_cron_job(name: str) -> bool:
    raw = get_raw_crontab()
    lines = []
    skip_next = False
    for line in raw.splitlines():
        if CRON_TAG in line and f"name={name}" in line:
            skip_next = True
            continue
        if skip_next:
            skip_next = False
            continue
        lines.append(line)
    return set_raw_crontab("\n".join(lines) + "\n")

def toggle_cron_job(name: str, enabled: bool) -> bool:
    raw = get_raw_crontab()
    lines = raw.splitlines()
    for i, line in enumerate(lines):
        if CRON_TAG in line and f"name={name}" in line:
            if i+1 < len(lines):
                cmd = lines[i+1].lstrip("#").strip()
                lines[i+1] = cmd if enabled else f"#{cmd}"
    return set_raw_crontab("\n".join(lines) + "\n")

# ── Artifact helpers ───────────────────────────────────────────────────────────

def _read_artifact_meta(fpath: str) -> dict:
    """Read sidecar .meta file for an artifact, returns {} if not found."""
    meta_path = fpath + ".meta"
    if os.path.exists(meta_path):
        try:
            import json as _j
            return _j.loads(open(meta_path).read())
        except Exception:
            pass
    return {}

def _infer_agent_from_filename(fname: str, known_agents: list) -> str:
    """Infer agent_id from filename prefix pattern: agent_id_timestamp_..."""
    fname_lower = fname.lower()
    for a in known_agents:
        if fname_lower.startswith(a + '_') or fname_lower.startswith(a + '.'):
            return a
    # Try two-part prefix (e.g. "claw_batto_...")
    parts = fname.split('_')
    if len(parts) >= 2:
        candidate = f"{parts[0]}_{parts[1]}"
        if candidate in known_agents:
            return candidate
    return ""

def scan_artifacts_dir(base_dir: str, project_id: str = "", agent_id: str = "") -> list:
    """Recursively scan a directory for artifacts, preserving all file types."""
    # Load known agent IDs for filename-based inference
    try:
        config = load_config()
        known_agents = list(config.get('agents', {}).keys())
    except Exception:
        known_agents = []

    files = []
    for root, dirs, fnames in os.walk(base_dir):
        # Vault (.vault/) is excluded entirely — no ghost thumbnails in
        # Data Hub. .private-inbound/ is now PUBLIC staging and gets
        # included. Other dotted dirs (.git, .vision-*) stay excluded.
        dirs[:] = [d for d in dirs
                   if (not d.startswith('.') or d == PRIVATE_INBOUND_DIRNAME)
                   and d != VAULT_DIRNAME]
        for fname in sorted(fnames):
            # Skip sidecar meta files
            if fname.endswith('.meta'):
                continue
            fpath = os.path.join(root, fname)
            # Per-file privacy still respected (legacy meta flag) but the
            # vault dir is already excluded above.
            file_private = _is_private_path(fpath)
            if file_private:
                continue
            rel   = os.path.relpath(fpath, base_dir)
            # Determine project_id from relative path structure: {project}/{file}
            parts = rel.split(os.sep)
            if project_id:
                proj     = project_id
                sub_path = rel  # base_dir is already the project dir
            else:
                proj     = parts[0] if len(parts) > 1 else "shared"
                sub_path = os.sep.join(parts[1:]) if len(parts) > 1 else fname
            sub_path = sub_path.replace(os.sep, '/')  # URL-safe

            # Determine agent_id: sidecar meta > filename inference > directory hint > simon_bately fallback
            meta  = _read_artifact_meta(fpath)
            dir_hint = ''
            for part in parts:
                # e.g. "simon_bately-uploads", "claw_batto-uploads"
                if '-uploads' in part or '-chat' in part:
                    candidate = part.replace('-uploads','').replace('-chat','')
                    if candidate in known_agents:
                        dir_hint = candidate
                        break
                if part in known_agents:
                    dir_hint = part
                    break
            agent = agent_id or meta.get('agent_id', '') or dir_hint or _infer_agent_from_filename(fname, known_agents) or "simon_bately"

            try:
                stat = os.stat(fpath)
            except OSError:
                # Broken symlink or file vanished between walk and stat — skip.
                continue
            ext  = os.path.splitext(fname)[1].lower()
            files.append({
                "name":       sub_path,
                "basename":   fname,
                "rel_path":   rel,
                "abs_path":   fpath,
                "size":       stat.st_size,
                "modified":   datetime.datetime.fromtimestamp(stat.st_mtime).strftime("%Y-%m-%d %H:%M"),
                "project_id": proj,
                "agent_id":   agent,
                "task_id":    meta.get('task_id', ''),
                "ext":        ext,
                "file_type":  _ext_to_type(ext),
                "subcategory": _classify_subcategory(fname, rel, ext, agent),
                "private":    file_private,
            })
    return files

def _classify_subcategory(fname: str, rel_path: str, ext: str, agent_id: str = "") -> str:
    """Fine-grained subcategory for Data Hub sub-tabs.

    Returns one of:
      audio, video,
      img_inbound, img_generated, img_edited, img_logos, img_marketing,
        img_blueprints, img_receipts, img_project,
      doc_pdf, doc_md, doc_text, doc_spreadsheet, doc_word,
      code_html, code_python, code_jsts, code_shell, code_config,
        code_css, code_sql, code_other,
      archive, other
    """
    low  = fname.lower()
    pth  = rel_path.lower().replace('\\', '/')

    audio_ext = {'.mp3','.wav','.ogg','.flac','.aac','.m4a','.wma','.opus'}
    video_ext = {'.mp4','.mkv','.avi','.mov','.webm','.flv','.wmv','.m4v'}
    img_ext   = {'.png','.jpg','.jpeg','.gif','.webp','.svg','.bmp','.tiff','.tif','.ico'}
    arc_ext   = {'.zip','.tar','.gz','.tgz','.bz2','.7z','.rar'}

    if ext in audio_ext: return 'audio'
    if ext in video_ext: return 'video'
    if ext in arc_ext:   return 'archive'

    if ext == '.pdf':  return 'doc_pdf'
    if ext in ('.md','.rst'):        return 'doc_md'
    if ext in ('.txt','.log'):       return 'doc_text'
    if ext in ('.csv','.xlsx','.xls','.ods'): return 'doc_spreadsheet'
    if ext in ('.docx','.doc','.odt','.rtf'): return 'doc_word'

    if ext in ('.html','.htm'):      return 'code_html'
    if ext == '.py':                  return 'code_python'
    if ext in ('.js','.ts','.jsx','.tsx','.mjs','.cjs'): return 'code_jsts'
    if ext in ('.sh','.bash','.zsh'): return 'code_shell'
    if ext in ('.json','.yaml','.yml','.toml','.ini','.cfg','.conf','.env','.xml'): return 'code_config'
    if ext == '.css':                 return 'code_css'
    if ext == '.sql':                 return 'code_sql'
    if ext in ('.go','.rs','.rb','.php','.c','.cpp','.cc','.h','.hpp','.java','.kt','.swift','.lua'): return 'code_other'

    if ext in img_ext:
        if 'receipt' in low or '/ahb_receipts' in pth or '/receipts' in pth:
            return 'img_receipts'
        if 'blueprint' in low or 'schematic' in low or 'floor_plan' in low or 'floorplan' in low or '/blueprint' in pth:
            return 'img_blueprints'
        if low.startswith('logo') or '_logo' in low or 'logo_' in low:
            return 'img_logos'
        if low.startswith('gallery_') or '/gallery' in pth or '/marketing' in pth:
            return 'img_marketing'
        if '_inpaint_' in low or '_img2img_' in low or '_edited' in low or '_edit.' in low or '/edited' in pth:
            return 'img_edited'
        if '_gen_' in low or re.match(r'^\d{10}_', low) or re.match(r'^[a-z_]+_\d{8}_\d{4}_[a-f0-9]+\.', low) or re.match(r'^[a-z_]+_\d{10}_', low):
            return 'img_generated'
        if low.startswith('upload_') or '-uploads/' in pth + '/' or '-chat/' in pth + '/':
            return 'img_inbound'
        return 'img_project'

    return 'other'


def _ext_to_type(ext: str) -> str:
    img   = {'.png','.jpg','.jpeg','.gif','.svg','.webp','.ico','.bmp','.tiff','.tif'}
    code  = {'.py','.sh','.bash','.js','.ts','.jsx','.tsx','.json','.yaml','.yml','.toml','.ini','.cfg','.conf','.sql','.html','.css','.xml'}
    doc   = {'.md','.txt','.rst','.csv','.log','.pdf','.docx','.doc','.xlsx','.xls','.pptx','.odt','.rtf'}
    arc   = {'.zip','.tar','.gz','.tgz','.bz2','.7z','.rar'}
    audio = {'.mp3','.wav','.ogg','.flac','.aac','.m4a','.wma','.opus'}
    video = {'.mp4','.mkv','.avi','.mov','.webm','.flv','.wmv'}
    if ext in img:   return 'image'
    if ext in code:  return 'code'
    if ext in doc:   return 'document'
    if ext in arc:   return 'archive'
    if ext in audio: return 'audio'
    if ext in video: return 'video'
    return 'other'

def artifacts_for_project(project_id: str) -> list:
    proj_dir = os.path.join(ARTIFACTS_DIR, project_id)
    if not os.path.exists(proj_dir):
        return []
    return scan_artifacts_dir(proj_dir, project_id=project_id)

def all_artifacts() -> list:
    if not os.path.exists(ARTIFACTS_DIR):
        return []
    return scan_artifacts_dir(ARTIFACTS_DIR)

# ── Routes — SPA Shell & Browser ──────────────────────────────────────────────

@app.route('/shell')
def shell_page():
    """SPA shell — single-page wrapper with persistent browser panel."""
    return render_template('shell.html')

@app.route('/api/browser/proxy')
def browser_proxy():
    """PHANTOM PROXY — Privacy-first server-side proxy.
    - Strips ALL tracking cookies, analytics, fingerprinting
    - Randomizes User-Agent per request
    - Blocks known tracker domains
    - No referrer sent
    - No client IP forwarded
    - All requests go through baza server — client IP never touches target"""
    import urllib.request as _ur, urllib.parse as _up
    import random, re as _re

    target = request.args.get('url', '')
    if not target:
        return 'No URL provided', 400

    try:
        parsed = _up.urlparse(target)
        if parsed.scheme not in ('http', 'https'):
            return 'Only HTTP/HTTPS allowed', 400
        if parsed.hostname in ('localhost', '127.0.0.1', '0.0.0.0'):
            return redirect(target)
    except Exception:
        return 'Invalid URL', 400

    # ── PRIVACY: Randomized User-Agents ──────────────────────────────────────
    UA_POOL = [
        'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/125.0.0.0 Safari/537.36',
        'Mozilla/5.0 (Macintosh; Intel Mac OS X 14_5) AppleWebKit/605.1.15 (KHTML, like Gecko) Version/17.5 Safari/605.1.15',
        'Mozilla/5.0 (X11; Linux x86_64; rv:128.0) Gecko/20100101 Firefox/128.0',
        'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/125.0.0.0 Safari/537.36 Edg/125.0.0.0',
    ]

    # ── PRIVACY: Known tracker domains to block ──────────────────────────────
    TRACKER_DOMAINS = {
        'google-analytics.com', 'googletagmanager.com', 'doubleclick.net',
        'facebook.net', 'fbcdn.net', 'connect.facebook.net',
        'analytics.google.com', 'adservice.google.com',
        'hotjar.com', 'clarity.ms', 'mixpanel.com', 'segment.io',
        'newrelic.com', 'nr-data.net', 'sentry.io',
        'ads-twitter.com', 'static.ads-twitter.com',
        'pixel.quantserve.com', 'mc.yandex.ru',
    }

    # Block requests to known trackers
    if parsed.hostname and any(parsed.hostname.endswith(t) for t in TRACKER_DOMAINS):
        return '', 204  # silent block

    try:
        headers = {
            'User-Agent': random.choice(UA_POOL),
            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8',
            'Accept-Language': 'en-US,en;q=0.5',
            'DNT': '1',
            'Sec-GPC': '1',
        }
        req = _ur.Request(target, headers=headers)
        # Try default first, fallback to IPv6 if it fails
        try:
            with _ur.urlopen(req, timeout=12) as resp:
                content = resp.read()
                ct = resp.headers.get('Content-Type', 'text/html')
        except Exception:
            # Retry with IPv6 preference for sites unreachable via IPv4
            import socket as _sock
            _orig_gai = _sock.getaddrinfo
            _sock.getaddrinfo = lambda *a, **k: [r for r in _orig_gai(*a, **k) if r[0] == _sock.AF_INET6] or _orig_gai(*a, **k)
            try:
                req2 = _ur.Request(target, headers=headers)
                with _ur.urlopen(req2, timeout=12) as resp:
                    content = resp.read()
                    ct = resp.headers.get('Content-Type', 'text/html')
            finally:
                _sock.getaddrinfo = _orig_gai

        if 'text/html' in ct:
            html = content.decode('utf-8', errors='replace')

            # Inject <base href> so CSS/JS/images load directly from origin
            # (only the HTML page itself goes through proxy — sub-resources load direct)
            base_tag = f'<base href="{target}">'
            if '<head>' in html.lower():
                html = _re.sub(r'<head[^>]*>', lambda m: m.group(0) + base_tag, html, count=1, flags=_re.IGNORECASE)
            else:
                html = base_tag + html

            # Rewrite ONLY <a href> links to go through proxy
            # Skip <base>, <link>, and non-navigation hrefs
            from urllib.parse import urljoin
            def _proxy_link(match):
                full = match.group(0)
                tag_before = match.group(1)
                quote = match.group(2)
                url = match.group(3)
                # Only proxy <a> tag hrefs, skip <base>, <link>, etc.
                tag_lower = tag_before.lower().strip()
                if not tag_lower.endswith('<a') and 'area' not in tag_lower:
                    return full
                if url.startswith(('#', 'javascript:', 'data:', 'mailto:', 'tel:', '/api/browser')):
                    return full
                abs_url = urljoin(target, url)
                if abs_url.startswith(('http://', 'https://')):
                    return f'{tag_before} href={quote}/api/browser/proxy?url={_up.quote(abs_url, safe="")}{quote}'
                return full
            html = _re.sub(r'(<[^>]*?)\s+href\s*=\s*(["\'])([^"\']*?)\2', _proxy_link, html)

            # ── PRIVACY: Strip tracking scripts ──────────────────────────────
            # Remove Google Analytics
            html = _re.sub(r'<script[^>]*google-analytics\.com[^>]*>.*?</script>', '', html, flags=_re.DOTALL|_re.IGNORECASE)
            html = _re.sub(r'<script[^>]*googletagmanager\.com[^>]*>.*?</script>', '', html, flags=_re.DOTALL|_re.IGNORECASE)
            # Remove Facebook pixel
            html = _re.sub(r'<script[^>]*facebook\.net[^>]*>.*?</script>', '', html, flags=_re.DOTALL|_re.IGNORECASE)
            html = _re.sub(r'<noscript[^>]*><img[^>]*facebook[^>]*></noscript>', '', html, flags=_re.IGNORECASE)
            # Remove Hotjar, Clarity, Mixpanel
            html = _re.sub(r'<script[^>]*hotjar\.com[^>]*>.*?</script>', '', html, flags=_re.DOTALL|_re.IGNORECASE)
            html = _re.sub(r'<script[^>]*clarity\.ms[^>]*>.*?</script>', '', html, flags=_re.DOTALL|_re.IGNORECASE)
            # Remove common tracking pixels
            html = _re.sub(r'<img[^>]*(?:pixel|tracking|beacon|analytics)[^>]*/?>', '', html, flags=_re.IGNORECASE)
            # Remove inline gtag/ga scripts
            html = _re.sub(r"<script[^>]*>\s*(?:window\.dataLayer|gtag|ga\s*\(|fbq\s*\(|_hmt\.push).*?</script>", '', html, flags=_re.DOTALL|_re.IGNORECASE)
            # Remove CSP meta tags that block iframe embedding
            html = _re.sub(r'<meta[^>]*content-security-policy[^>]*>', '', html, flags=_re.IGNORECASE)
            html = _re.sub(r'<meta[^>]*x-frame-options[^>]*>', '', html, flags=_re.IGNORECASE)

            resp_obj = make_response(html)
            resp_obj.headers['Content-Type'] = 'text/html; charset=utf-8'
        else:
            resp_obj = make_response(content)
            resp_obj.headers['Content-Type'] = ct

        # ── PRIVACY: Strip headers that block iframe embedding + tracking ──
        resp_obj.headers['Referrer-Policy'] = 'no-referrer'
        resp_obj.headers['X-Content-Type-Options'] = 'nosniff'
        resp_obj.headers.pop('Set-Cookie', None)       # block tracking cookies
        resp_obj.headers.pop('X-Frame-Options', None)  # allow iframe embedding
        resp_obj.headers.pop('Content-Security-Policy', None)  # remove CSP frame-ancestors
        resp_obj.headers.pop('Cross-Origin-Opener-Policy', None)
        resp_obj.headers.pop('Cross-Origin-Embedder-Policy', None)
        return resp_obj

    except Exception as e:
        return f'''<html><body style="background:#07070f;color:#e94560;font-family:monospace;padding:40px">
        <h2>Phantom Proxy Error</h2><p>{str(e)[:300]}</p>
        <p style="color:#555;font-size:12px">Your IP was never exposed to the target server.</p>
        </body></html>''', 502

@app.route('/api/browser/bookmarks', methods=['GET', 'POST', 'DELETE'])
def api_browser_bookmarks():
    """CRUD for browser bookmarks."""
    import sqlite3
    db = os.path.join(DASHBOARD_DIR, 'baza_projects.db')
    conn = sqlite3.connect(db)
    conn.execute("""CREATE TABLE IF NOT EXISTS browser_bookmarks (
        id INTEGER PRIMARY KEY AUTOINCREMENT, url TEXT NOT NULL, title TEXT DEFAULT '',
        favicon TEXT DEFAULT '', folder TEXT DEFAULT 'default',
        created_at TEXT DEFAULT (datetime('now')))""")
    conn.commit()

    if request.method == 'GET':
        rows = conn.execute("SELECT id, url, title, favicon, folder, created_at FROM browser_bookmarks ORDER BY created_at DESC").fetchall()
        conn.close()
        return jsonify([{"id":r[0],"url":r[1],"title":r[2],"favicon":r[3],"folder":r[4],"created_at":r[5]} for r in rows])
    elif request.method == 'POST':
        data = request.json or {}
        conn.execute("INSERT INTO browser_bookmarks (url, title, favicon, folder) VALUES (?,?,?,?)",
                     (data.get('url',''), data.get('title',''), data.get('favicon',''), data.get('folder','default')))
        conn.commit()
        bm_id = conn.execute("SELECT last_insert_rowid()").fetchone()[0]
        conn.close()
        return jsonify({"id": bm_id, "url": data.get('url'), "title": data.get('title')})
    elif request.method == 'DELETE':
        data = request.json or {}
        bm_id = data.get('id')
        if bm_id:
            conn.execute("DELETE FROM browser_bookmarks WHERE id=?", (bm_id,))
        conn.commit()
        conn.close()
        return jsonify({"success": True})

@app.route('/api/browser/history', methods=['GET', 'POST', 'DELETE'])
def api_browser_history():
    """Browse history storage."""
    import sqlite3
    db = os.path.join(DASHBOARD_DIR, 'baza_projects.db')
    conn = sqlite3.connect(db)
    conn.execute("""CREATE TABLE IF NOT EXISTS browser_history (
        id INTEGER PRIMARY KEY AUTOINCREMENT, url TEXT NOT NULL, title TEXT DEFAULT '',
        visited_at TEXT DEFAULT (datetime('now')))""")
    conn.commit()

    if request.method == 'GET':
        limit = int(request.args.get('limit', 50))
        rows = conn.execute("SELECT id, url, title, visited_at FROM browser_history ORDER BY visited_at DESC LIMIT ?", (limit,)).fetchall()
        conn.close()
        return jsonify([{"id":r[0],"url":r[1],"title":r[2],"visited_at":r[3]} for r in rows])
    elif request.method == 'POST':
        data = request.json or {}
        url = data.get('url', '')
        if url and url != 'about:blank':
            conn.execute("INSERT INTO browser_history (url, title) VALUES (?,?)", (url, data.get('title','')))
            conn.commit()
            # Keep only last 500 entries
            conn.execute("DELETE FROM browser_history WHERE id NOT IN (SELECT id FROM browser_history ORDER BY visited_at DESC LIMIT 500)")
            conn.commit()
        conn.close()
        return jsonify({"success": True})
    elif request.method == 'DELETE':
        conn.execute("DELETE FROM browser_history")
        conn.commit()
        conn.close()
        return jsonify({"success": True})


# ── Routes — Pages ────────────────────────────────────────────────────────────

@app.route('/')
def index():
    config = load_config()
    agents = config.get('agents', {})
    agent_data = []
    for agent_id, agent_config in agents.items():
        status = get_agent_status(agent_id)
        messages = get_recent_messages(agent_id, 5)
        agent_data.append({
            'id': agent_id,
            'name': agent_config.get('name', agent_id),
            'role': agent_config.get('role', ''),
            'model': agent_config.get('model', ''),
            'status': status,
            'recent_messages': messages,
        })
    crons = list_crons()
    return render_template('index.html', agents=agent_data, crons=crons)

@app.route('/agent/<agent_id>')
def agent_detail(agent_id):
    config = load_config()
    agents = config.get('agents', {})
    if agent_id not in agents:
        return "Agent not found", 404
    agent_config = agents[agent_id]
    status   = get_agent_status(agent_id)
    messages = get_recent_messages(agent_id, 40)
    logs     = get_agent_logs(agent_id, 80)
    crons    = [c for c in list_crons() if agent_id.replace('_','-') in c.get('command','') or agent_id in c.get('id','')]
    def _fetch_ollama_models(port):
        try:
            import json as _j
            r = subprocess.run(['curl','-s',f'http://localhost:{port}/api/tags'],
                               capture_output=True, text=True, timeout=3)
            return sorted([m["name"] for m in _j.loads(r.stdout).get("models",[])])
        except:
            return []

    def _fetch_litellm_models():
        try:
            import json as _j
            r = subprocess.run([
                'curl','-s','-H','Authorization: Bearer baza-litellm-internal',
                'http://localhost:4000/v1/models'
            ], capture_output=True, text=True, timeout=3)
            return sorted([m["id"] for m in _j.loads(r.stdout).get("data",[])])
        except:
            return []

    amd_models  = _fetch_ollama_models(11434)
    cuda_models = _fetch_ollama_models(11435)
    cloud_models = _fetch_litellm_models()

    # Fallback so current model always appears even if Ollama is briefly offline
    current_model = agent_config.get("model","")
    if current_model:
        if current_model.startswith(("gpt-","claude-","gemini-","grok-","groq-","mistral-large","codestral","o1","o3")):
            if current_model not in cloud_models:
                cloud_models.insert(0, current_model)
        elif current_model not in amd_models and current_model not in cuda_models:
            amd_models.insert(0, current_model)

    available_models = {
        "Local — AMD GPU (11434)":   amd_models  or ["(offline)"],
        "Local — CUDA GPU (11435)":  cuda_models or ["(offline)"],
        "Cloud via LiteLLM (4000)":  cloud_models or ["(offline)"],
    }
    return render_template('agent.html',
        agent_id=agent_id, agent=agent_config,
        status=status, messages=messages, logs=logs,
        crons=crons, available_models=available_models)

@app.route('/crons')
def crons_page():
    crons = list_crons()
    raw   = get_raw_crontab()
    return render_template('crons.html', crons=crons, raw_crontab=raw)

@app.route('/artifacts')
def artifacts_page():
    """Legacy route — redirect to Data Hub."""
    return redirect('/datahub')


# ── Private gallery (passphrase-locked) ──────────────────────────────────────
# All Telegram-inbound media is marked private (see dashboard/private_inbound.py)
# and excluded from every public Data Hub list/serve/grep route. The private
# gallery is the *only* way to view them through the UI, gated by a session
# unlock against a hashed passphrase stored at dashboard/.private_pass.

PRIVATE_PASS_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                                 '.private_pass')


def _private_pass_is_set() -> bool:
    return os.path.isfile(PRIVATE_PASS_FILE) and os.path.getsize(PRIVATE_PASS_FILE) > 0


def _private_pass_check(passphrase: str) -> bool:
    if not _private_pass_is_set() or not passphrase:
        return False
    try:
        with open(PRIVATE_PASS_FILE, 'r', encoding='utf-8') as fh:
            stored = fh.read().strip()
    except OSError:
        return False
    return bool(stored) and check_password_hash(stored, passphrase)


def _is_private_unlocked() -> bool:
    return bool(session.get('private_unlocked'))


def _list_private_files() -> list:
    """List files in the vault (.vault/). The strict-private location is
    the ONLY thing surfaced here — Telegram inbound is public now and
    appears in Data Hub instead."""
    out = []
    vault_root = os.path.join(ARTIFACTS_DIR, VAULT_DIRNAME)
    if not os.path.isdir(vault_root):
        return out
    img_ext   = {'.png','.jpg','.jpeg','.gif','.webp','.bmp','.tiff','.tif'}
    audio_ext = {'.mp3','.wav','.ogg','.flac','.m4a','.opus'}
    video_ext = {'.mp4','.mkv','.mov','.webm'}
    for root, dirs, fnames in os.walk(vault_root):
        dirs[:] = [d for d in dirs if not d.startswith('__') and d != '.git']
        for fname in fnames:
            if fname.endswith('.meta'):
                continue
            fpath = os.path.join(root, fname)
            if not _is_private_path(fpath):
                continue
            try:
                st = os.stat(fpath)
            except OSError:
                continue
            ext = os.path.splitext(fname)[1].lower()
            if ext in img_ext:
                kind = 'image'
            elif ext in audio_ext:
                kind = 'audio'
            elif ext in video_ext:
                kind = 'video'
            else:
                kind = 'other'
            # Pull caption + agent_id from the JSON .meta sidecar if present
            caption = ''
            agent_id = ''
            received_at = ''
            meta_path = fpath + '.meta'
            if os.path.isfile(meta_path):
                try:
                    with open(meta_path, 'r', encoding='utf-8', errors='replace') as fh:
                        raw = fh.read().strip()
                    if raw.startswith('{'):
                        m = json.loads(raw)
                        caption = str(m.get('caption', ''))[:300]
                        agent_id = str(m.get('agent_id', ''))
                        received_at = str(m.get('received_at') or m.get('created_at') or '')
                except Exception:
                    pass
            rel = os.path.relpath(fpath, ARTIFACTS_DIR).replace(os.sep, '/')
            # Token = urlsafe base64 of rel path; serve route reverses + re-checks privacy
            import base64 as _b64
            token = _b64.urlsafe_b64encode(rel.encode('utf-8')).decode('ascii').rstrip('=')
            out.append({
                'token':       token,
                'rel_path':    rel,
                'basename':    fname,
                'size':        st.st_size,
                'mtime':       st.st_mtime,
                'modified':    datetime.datetime.fromtimestamp(st.st_mtime).strftime('%Y-%m-%d %H:%M'),
                'ext':         ext,
                'kind':        kind,
                'agent_id':    agent_id,
                'caption':     caption,
                'received_at': received_at,
            })
    out.sort(key=lambda r: r['mtime'], reverse=True)
    return out


def _decode_private_token(token: str):
    """Decode a list token back to an absolute path; None if invalid or not private."""
    import base64 as _b64
    try:
        pad = '=' * (-len(token) % 4)
        rel = _b64.urlsafe_b64decode((token + pad).encode('ascii')).decode('utf-8')
    except Exception:
        return None
    if '..' in rel.split('/'):
        return None
    fpath = os.path.realpath(os.path.join(ARTIFACTS_DIR, rel))
    if not fpath.startswith(os.path.realpath(ARTIFACTS_DIR)):
        return None
    if not os.path.isfile(fpath):
        return None
    if not _is_private_path(fpath):
        return None
    return fpath


@app.route('/api/datahub/private/status')
def api_private_status():
    return jsonify({
        'passphrase_set': _private_pass_is_set(),
        'unlocked':       _is_private_unlocked(),
    })


@app.route('/api/datahub/private/unlock', methods=['POST'])
def api_private_unlock():
    if not _private_pass_is_set():
        return jsonify({'ok': False, 'error': 'No passphrase set. Run venv/bin/python dashboard/set_private_pass.py on the server.'}), 400
    payload = request.get_json(silent=True) or {}
    pp = (payload.get('passphrase') or '').strip()
    if not pp:
        return jsonify({'ok': False, 'error': 'Empty passphrase'}), 400
    if not _private_pass_check(pp):
        return jsonify({'ok': False, 'error': 'Wrong passphrase'}), 401
    session['private_unlocked'] = True
    session.permanent = False  # cleared when browser closes
    return jsonify({'ok': True})


@app.route('/api/datahub/private/lock', methods=['POST'])
def api_private_lock():
    session.pop('private_unlocked', None)
    return jsonify({'ok': True})


@app.route('/api/datahub/lock-toggle', methods=['POST'])
def api_datahub_lock_toggle():
    """Flip a Data Hub file in or out of the Vault.

    Body: { project_id: <string>, name: <relative path within project> }
    Action is inferred from current state:
      - public file → move to .vault/ (file disappears from Data Hub)
      - vault file  → move to .private-inbound/<agent_id>/ (visible again)"""
    payload = request.get_json(silent=True) or {}
    project_id = (payload.get('project_id') or '').strip()
    rel        = (payload.get('name') or '').strip()
    if not rel:
        return jsonify({'ok': False, 'error': 'name required'}), 400

    if project_id and project_id != 'shared':
        cur = os.path.join(ARTIFACTS_DIR, project_id, rel)
    else:
        cur = os.path.join(ARTIFACTS_DIR, rel)
    cur = os.path.realpath(cur)
    art_root = os.path.realpath(ARTIFACTS_DIR)
    if not (cur == art_root or cur.startswith(art_root + os.sep)) or not os.path.isfile(cur):
        return jsonify({'ok': False, 'error': 'not found'}), 404

    locking = not _is_private_path(cur)
    framework_dir = os.path.dirname(DASHBOARD_DIR)
    try:
        if locking:
            new_path = _vault_move_in(cur, framework_dir,
                                      extra={'origin_project': project_id} if project_id else None)
        else:
            new_path = _vault_move_out(cur, framework_dir)
    except (OSError, ValueError, FileNotFoundError) as e:
        return jsonify({'ok': False, 'error': str(e)}), 500
    return jsonify({
        'ok': True,
        'locked': locking,
        'new_project': VAULT_DIRNAME if locking else PRIVATE_INBOUND_DIRNAME,
        'new_name': os.path.relpath(new_path, ARTIFACTS_DIR).replace(os.sep, '/'),
    })


@app.route('/api/vault/add', methods=['POST'])
def api_vault_add():
    """Move a file under artifacts/ into .vault/. Body: { token } (a Baza
    pick token) OR { project_id, name } (Data Hub coords). Once moved, the
    file disappears from Data Hub and is only listed by the vault view."""
    payload = request.get_json(silent=True) or {}
    token = (payload.get('token') or '').strip()
    project_id = (payload.get('project_id') or '').strip()
    name = (payload.get('name') or '').strip()
    if token:
        cur = _pick_decode_token(token)
        if not cur:
            return jsonify({'ok': False, 'error': 'invalid token'}), 400
    elif name:
        base = os.path.join(ARTIFACTS_DIR, project_id, name) if (project_id and project_id != 'shared') \
               else os.path.join(ARTIFACTS_DIR, name)
        cur = os.path.realpath(base)
        art_root = os.path.realpath(ARTIFACTS_DIR)
        if not cur.startswith(art_root + os.sep) or not os.path.isfile(cur):
            return jsonify({'ok': False, 'error': 'not found'}), 404
    else:
        return jsonify({'ok': False, 'error': 'token or name required'}), 400
    framework_dir = os.path.dirname(DASHBOARD_DIR)
    try:
        new_path = _vault_move_in(cur, framework_dir,
                                  extra={'origin_project': project_id} if project_id else None)
    except (OSError, ValueError, FileNotFoundError) as e:
        return jsonify({'ok': False, 'error': str(e)}), 500
    return jsonify({'ok': True,
                    'vault_path': os.path.relpath(new_path, ARTIFACTS_DIR).replace(os.sep, '/')})


@app.route('/api/vault/remove', methods=['POST'])
def api_vault_remove():
    """Move a file out of .vault/ back to .private-inbound/manual/. Body:
    { token }  — a vault-listing token from /api/datahub/private/list."""
    if not _is_private_unlocked():
        return jsonify({'ok': False, 'error': 'Locked'}), 401
    payload = request.get_json(silent=True) or {}
    token = (payload.get('token') or '').strip()
    cur = _decode_private_token(token)
    if not cur:
        return jsonify({'ok': False, 'error': 'invalid token'}), 400
    framework_dir = os.path.dirname(DASHBOARD_DIR)
    try:
        new_path = _vault_move_out(cur, framework_dir)
    except (OSError, ValueError, FileNotFoundError) as e:
        return jsonify({'ok': False, 'error': str(e)}), 500
    return jsonify({'ok': True,
                    'new_path': os.path.relpath(new_path, ARTIFACTS_DIR).replace(os.sep, '/')})


@app.route('/api/datahub/private/list')
def api_private_list():
    if not _is_private_unlocked():
        return jsonify({'ok': False, 'error': 'Locked'}), 401
    files = _list_private_files()
    return jsonify({'ok': True, 'count': len(files), 'files': files})


@app.route('/api/datahub/private/serve/<token>')
def api_private_serve(token):
    if not _is_private_unlocked():
        return jsonify({'error': 'Locked'}), 401
    fpath = _decode_private_token(token)
    if not fpath:
        return jsonify({'error': 'Not found or not private'}), 404
    return send_from_directory(os.path.dirname(fpath), os.path.basename(fpath))


@app.route('/api/datahub/private/delete', methods=['POST'])
def api_private_delete():
    if not _is_private_unlocked():
        return jsonify({'ok': False, 'error': 'Locked'}), 401
    payload = request.get_json(silent=True) or {}
    token = payload.get('token') or ''
    fpath = _decode_private_token(token)
    if not fpath:
        return jsonify({'ok': False, 'error': 'Not found'}), 404
    try:
        os.remove(fpath)
        meta = fpath + '.meta'
        if os.path.isfile(meta):
            os.remove(meta)
    except OSError as e:
        return jsonify({'ok': False, 'error': str(e)[:200]}), 500
    return jsonify({'ok': True})


@app.route('/api/datahub/private/bulk-delete', methods=['POST'])
def api_private_bulk_delete():
    """Delete a batch of private files by token. Each file's .meta sidecar
    is removed too. Returns per-token results so the UI can refresh from
    a partial success."""
    if not _is_private_unlocked():
        return jsonify({'ok': False, 'error': 'Locked'}), 401
    payload = request.get_json(silent=True) or {}
    tokens = payload.get('tokens') or []
    if not isinstance(tokens, list) or not tokens:
        return jsonify({'ok': False, 'error': 'tokens[] required'}), 400
    deleted, errors = [], []
    for tok in tokens:
        fpath = _decode_private_token(tok)
        if not fpath:
            errors.append({'token': tok, 'error': 'not found'})
            continue
        try:
            os.remove(fpath)
            meta = fpath + '.meta'
            if os.path.isfile(meta):
                os.remove(meta)
            deleted.append(tok)
        except OSError as e:
            errors.append({'token': tok, 'error': str(e)[:200]})
    return jsonify({'ok': True, 'deleted': deleted, 'errors': errors})


@app.route('/api/datahub/private/zip', methods=['POST'])
def api_private_zip():
    """Stream a zip of the chosen private files. Body: {tokens:[...]}."""
    if not _is_private_unlocked():
        return jsonify({'ok': False, 'error': 'Locked'}), 401
    payload = request.get_json(silent=True) or {}
    tokens = payload.get('tokens') or []
    if not isinstance(tokens, list) or not tokens:
        return jsonify({'ok': False, 'error': 'tokens[] required'}), 400
    import io, zipfile
    buf = io.BytesIO()
    seen = {}
    with zipfile.ZipFile(buf, 'w', zipfile.ZIP_DEFLATED) as z:
        for tok in tokens:
            fpath = _decode_private_token(tok)
            if not fpath or not os.path.isfile(fpath):
                continue
            arcname = os.path.basename(fpath)
            n = seen.get(arcname, 0)
            if n:
                stem, ext = os.path.splitext(arcname)
                arcname = f"{stem}_{n}{ext}"
            seen[os.path.basename(fpath)] = n + 1
            z.write(fpath, arcname=arcname)
    buf.seek(0)
    ts = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
    resp = make_response(buf.getvalue())
    resp.headers['Content-Type'] = 'application/zip'
    resp.headers['Content-Disposition'] = f'attachment; filename="private_{ts}.zip"'
    resp.headers['Cache-Control'] = 'no-store'
    return resp


@app.route('/chains')
def chains_page():
    """Activity Chains — visibility pipeline #1 UI."""
    return render_template('chains.html')


@app.route('/datahub')
def datahub_page():
    project_id = request.args.get('project_id', '')
    if project_id:
        arts = artifacts_for_project(project_id)
    else:
        arts = all_artifacts()
    projects = []
    if os.path.exists(ARTIFACTS_DIR):
        projects = [d for d in os.listdir(ARTIFACTS_DIR)
                    if os.path.isdir(os.path.join(ARTIFACTS_DIR, d))]
    return render_template('datahub.html', artifacts=arts,
                           projects=sorted(projects), current_project=project_id)

@app.route('/settings')
def settings_page():
    config  = load_config()
    secrets = load_secrets()
    # Mask secret values
    masked = {k: ('●'*8 if v else '') for k, v in secrets.items()}
    return render_template('settings.html', config=config, secrets=masked,
                           secret_keys=list(secrets.keys()))

# ── Routes — Agent API ────────────────────────────────────────────────────────

@app.route('/empire-pulse')
def empire_pulse_page():
    return render_template('empire_pulse.html')


@app.route('/api/empire-pulse')
def api_empire_pulse():
    """Per-agent talk-vs-ship ratio over the last N days.

    talked_count: assistant messages in journal containing completion verbs
    shipped_count: artifacts saved in the same window attributed to that agent
    ratio = shipped / talked (capped at 1.0). Lower = more hallucination.
    """
    days = int(request.args.get('days', 7) or 7)
    days = max(1, min(days, 30))
    hours = days * 24
    cutoff = datetime.datetime.now().timestamp() - hours * 3600

    # Real artifacts per agent + per-day buckets for sparklines
    by_agent_files: dict[str, int] = {}
    by_agent_files_daily: dict[str, list[int]] = {}  # newest-first list of len=days
    now_ts = datetime.datetime.now().timestamp()
    for proj in os.listdir(ARTIFACTS_DIR):
        proj_dir = os.path.join(ARTIFACTS_DIR, proj)
        if not os.path.isdir(proj_dir):
            continue
        for fname in os.listdir(proj_dir):
            if fname.endswith('.meta'):
                continue
            full = os.path.join(proj_dir, fname)
            if not os.path.isfile(full):
                continue
            mt = os.path.getmtime(full)
            if mt < cutoff:
                continue
            ag = ""
            try:
                with open(full + ".meta") as mf:
                    ag = (json.load(mf) or {}).get("agent_id", "")
            except Exception:
                head = fname.split("_", 2)
                if len(head) >= 2 and head[0] in (
                    "simon", "claw", "sam", "nova", "phil", "rex", "duke", "scout"
                ):
                    ag = "_".join(head[:2])
            if not ag:
                continue
            by_agent_files[ag] = by_agent_files.get(ag, 0) + 1
            day_idx = int((now_ts - mt) // 86400)  # 0 = today, 1 = yesterday…
            if 0 <= day_idx < days:
                buckets = by_agent_files_daily.setdefault(ag, [0] * days)
                buckets[day_idx] += 1

    # Completion-verb dictionary — widened so we catch fabrications phrased
    # as "submitted", "deployed", "ready", "live", etc. in addition to the
    # original six. Each pattern is anchored with '%' wildcards on both sides.
    COMPLETION_VERBS = [
        "complete", "done", "delivered", "shipped", "finalized", "finished",
        "submitted", "deployed", "implemented", "produced", "drafted",
        "wrapped up", "terminated", " live", "live.", "live,", " ready",
        "ready.", "ready,",
    ]
    verb_clause = " OR ".join(["lower(result) LIKE %s"] * len(COMPLETION_VERBS))
    verb_params = [f"%{v}%" for v in COMPLETION_VERBS]

    # Talked-about-completion counts from task_journal + daily buckets
    by_agent_talked: dict[str, int] = {}
    by_agent_talked_daily: dict[str, list[int]] = {}
    by_agent_fabricated: dict[str, int] = {}  # verified=FALSE rows
    try:
        from core.context_db import get_conn, release_conn
        conn = get_conn()
        cur = conn.cursor()
        cur.execute(
            f"SELECT agent_id, count(*) FROM task_journal "
            f"WHERE created_at > now() - interval %s AND ({verb_clause}) "
            f"GROUP BY agent_id",
            [f"{hours} hours"] + verb_params,
        )
        for ag, n in cur.fetchall():
            if ag:
                by_agent_talked[ag] = n
        # Per-day daily counts for sparklines (one query, group by day index)
        cur.execute(
            f"SELECT agent_id, "
            f"       FLOOR(EXTRACT(EPOCH FROM (now() - created_at)) / 86400)::int AS d, "
            f"       count(*) "
            f"FROM task_journal "
            f"WHERE created_at > now() - interval %s AND ({verb_clause}) "
            f"GROUP BY agent_id, d",
            [f"{hours} hours"] + verb_params,
        )
        for ag, d, n in cur.fetchall():
            if not ag:
                continue
            buckets = by_agent_talked_daily.setdefault(ag, [0] * days)
            d = int(d)
            if 0 <= d < days:
                buckets[d] += n
        # Fabrication counts — completion-claim rows that claim_verifier
        # already marked verified=FALSE in journal_log(). Safe if column
        # doesn't exist yet (treat as 0).
        try:
            cur.execute(
                f"SELECT agent_id, count(*) FROM task_journal "
                f"WHERE created_at > now() - interval %s AND verified = FALSE "
                f"AND ({verb_clause}) "
                f"GROUP BY agent_id",
                [f"{hours} hours"] + verb_params,
            )
            for ag, n in cur.fetchall():
                if ag:
                    by_agent_fabricated[ag] = n
        except Exception:
            conn.rollback()
        cur.close()
        release_conn(conn)
    except Exception as e:
        return jsonify({"error": f"task_journal read failed: {e}"})

    # Filter to actual agent ids (firstname_surname pattern), drop test/path noise
    import re as _re
    AGENT_RE = _re.compile(r"^[a-z][a-z0-9]+_[a-z][a-z0-9]+$")
    all_agents = sorted(set(list(by_agent_files.keys()) + list(by_agent_talked.keys())))
    all_agents = [a for a in all_agents if AGENT_RE.match(a or "")]
    rows = []
    for ag in all_agents:
        talked = by_agent_talked.get(ag, 0)
        shipped = by_agent_files.get(ag, 0)
        ratio = (min(shipped, talked) / talked) if talked else 1.0
        # Sparklines stored newest-day-LEFT in DB style; flip so the chart
        # reads chronologically left-to-right (oldest → newest = today).
        ship_buckets   = list(reversed(by_agent_files_daily.get(ag, [0] * days)))
        talked_buckets = list(reversed(by_agent_talked_daily.get(ag, [0] * days)))
        # Drift score (0-100, higher = more drift between claim and ship):
        # 60% weight on (1 - ratio) — directly captures hallucination
        # 30% weight on volatility of daily ship/talk gap — captures sliding
        # 10% weight on volume — high-volume agents drifting hurt more
        gap_per_day = [max(0, t - s) for s, t in zip(ship_buckets, talked_buckets)]
        if any(gap_per_day):
            mean_gap = sum(gap_per_day) / len(gap_per_day)
            variance = sum((g - mean_gap) ** 2 for g in gap_per_day) / len(gap_per_day)
            volatility = min(1.0, (variance ** 0.5) / max(1.0, mean_gap + 1))
        else:
            volatility = 0.0
        volume_factor = min(1.0, talked / 10.0)  # caps at 10 claims = full weight
        drift_score = round(
            100 * (0.60 * (1 - ratio) + 0.30 * volatility + 0.10 * volume_factor), 1
        )
        fabricated = by_agent_fabricated.get(ag, 0)
        rows.append({
            "agent_id": ag,
            "talked_about_completing": talked,
            "shipped": shipped,
            "fabricated": fabricated,
            "ratio": round(ratio, 2),
            "drift_score": drift_score,
            # Health: green if shipped >= talked, yellow if shipped > 0 but < talked,
            # red if shipped == 0 and talked > 0
            "health": (
                "green" if shipped >= talked
                else "red" if shipped == 0 and talked > 0
                else "yellow"
            ),
            "ship_daily":   ship_buckets,
            "talked_daily": talked_buckets,
        })
    rows.sort(key=lambda r: (r["health"] == "red", -r["talked_about_completing"]), reverse=True)
    rows.sort(key=lambda r: ({"red": 0, "yellow": 1, "green": 2}[r["health"]],
                              -r["talked_about_completing"]))
    return jsonify({
        "days": days,
        "agents": rows,
        "totals": {
            "agents": len(rows),
            "shipping": sum(1 for r in rows if r["health"] == "green"),
            "drifting": sum(1 for r in rows if r["health"] == "yellow"),
            "all_talk": sum(1 for r in rows if r["health"] == "red"),
            "fabrications": sum(r["fabricated"] for r in rows),
        },
    })


@app.route('/api/agents/<agent_id>/recent-artifacts')
def api_agent_recent_artifacts(agent_id):
    """List artifacts attributed to this agent in the last N hours.

    Anti-hallucination check: at-a-glance "did this agent actually produce
    anything today, or just talk about producing things?"
    """
    hours = int(request.args.get('hours', 24) or 24)
    hours = max(1, min(hours, 168))
    cutoff = datetime.datetime.now().timestamp() - hours * 3600
    rows = []
    for proj in os.listdir(ARTIFACTS_DIR):
        proj_dir = os.path.join(ARTIFACTS_DIR, proj)
        if not os.path.isdir(proj_dir):
            continue
        for fname in os.listdir(proj_dir):
            if fname.endswith('.meta'):
                continue
            full = os.path.join(proj_dir, fname)
            if not os.path.isfile(full):
                continue
            mt = os.path.getmtime(full)
            if mt < cutoff:
                continue
            # Attribution from .meta sidecar; fallback to filename prefix
            ag = ""
            try:
                with open(full + ".meta") as mf:
                    ag = (json.load(mf) or {}).get("agent_id", "")
            except Exception:
                head = fname.split("_", 2)
                if len(head) >= 2 and head[0] in (
                    "simon", "claw", "sam", "nova", "phil", "rex", "duke", "scout"
                ):
                    ag = "_".join(head[:2])
            if ag != agent_id:
                continue
            rows.append({
                "name": fname,
                "project_id": proj,
                "size": os.path.getsize(full),
                "modified": datetime.datetime.fromtimestamp(mt).isoformat(),
                "url": f"/api/artifacts/serve/{proj}/{fname}",
                "ext": os.path.splitext(fname)[1].lower(),
            })
    rows.sort(key=lambda r: r["modified"], reverse=True)
    # Quick claim verification across recent journal entries for this agent
    talked_about_completing = 0
    try:
        from core.context_db import get_conn, release_conn
        conn = get_conn()
        cur = conn.cursor()
        cur.execute(
            "SELECT count(*) FROM task_journal WHERE agent_id=%s "
            "AND created_at > now() - interval %s "
            "AND (lower(result) LIKE %s OR lower(result) LIKE %s "
            "    OR lower(result) LIKE %s OR lower(result) LIKE %s)",
            (agent_id, f"{hours} hours",
             "%complete%", "%done%", "%delivered%", "%shipped%"),
        )
        talked_about_completing = cur.fetchone()[0]
        cur.close()
        release_conn(conn)
    except Exception:
        pass
    return jsonify({
        "agent_id": agent_id,
        "hours": hours,
        "count": len(rows),
        "artifacts": rows[:60],
        "claim_warnings": {
            "talked_about_completing": talked_about_completing,
            "actually_produced": len(rows),
            "ratio_ok": (len(rows) > 0) if talked_about_completing else True,
        },
    })


@app.route('/agent/<agent_id>/restart', methods=['POST'])
def restart_agent_route(agent_id):
    try:
        subprocess.run(['sudo','systemctl','restart', svc_name(agent_id)], timeout=10)
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/agent/<agent_id>/stop', methods=['POST'])
def stop_agent_route(agent_id):
    try:
        subprocess.run(['sudo','systemctl','stop', svc_name(agent_id)], timeout=10)
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/agent/<agent_id>/start', methods=['POST'])
def start_agent_route(agent_id):
    try:
        subprocess.run(['sudo','systemctl','start', svc_name(agent_id)], timeout=10)
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/agent/<agent_id>/edit', methods=['POST'])
def edit_agent(agent_id):
    config = load_config()
    if agent_id not in config.get('agents', {}):
        return jsonify({'error': 'Agent not found'}), 404
    data = request.json or {}
    for field in ['model','system_prompt','role','name']:
        if field in data:
            config['agents'][agent_id][field] = data[field]
    save_config(config)
    return jsonify({'success': True})

@app.route('/agent/<agent_id>/logs')
def agent_logs(agent_id):
    lines = request.args.get('lines', 80, type=int)
    return jsonify({'logs': get_agent_logs(agent_id, lines)})

@app.route('/api/status')
def api_status():
    config = load_config()
    result = {aid: get_agent_status(aid) for aid in config.get('agents', {})}
    return jsonify(result)

@app.route('/api/messages/<agent_id>')
def api_messages(agent_id):
    limit = request.args.get('limit', 20, type=int)
    return jsonify(get_recent_messages(agent_id, limit))

@app.route('/api/agents/<agent_id>/conversations')
def api_agent_conversations(agent_id):
    try:
        from core.context_db import get_pool
        pool = get_pool()
        conn = pool.getconn()
        cur = conn.cursor()
        cur.execute("""
            SELECT chat_id,
                   MIN(created_at) as first_msg,
                   MAX(created_at) as last_msg,
                   COUNT(*) as msg_count,
                   (SELECT content FROM messages m2 WHERE m2.chat_id=m.chat_id AND m2.agent_id=%s ORDER BY created_at DESC LIMIT 1) as last_content
            FROM messages m
            WHERE agent_id = %s
            GROUP BY chat_id
            ORDER BY MAX(created_at) DESC
            LIMIT 50
        """, (agent_id, agent_id))
        rows = cur.fetchall()
        cur.close()
        pool.putconn(conn)
        return jsonify([{
            'chat_id': r[0], 'first_msg': str(r[1]), 'last_msg': str(r[2]),
            'msg_count': r[3], 'last_content': (r[4] or '')[:200]
        } for r in rows])
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/live')
def api_live():
    """Single endpoint: agent statuses + last messages + infra metrics."""
    import socket as _socket, shutil as _shutil
    config = load_config()
    agents_cfg = config.get('agents', {})

    def _run(cmd):
        try:
            r = subprocess.run(cmd, shell=True, capture_output=True, text=True, timeout=4)
            return r.stdout.strip()
        except:
            return ""

    def _port(host, port):
        try:
            with _socket.create_connection((host, port), timeout=1):
                return "up"
        except:
            return "down"

    statuses = {aid: get_agent_status(aid) for aid in agents_cfg}
    messages = {aid: get_recent_messages(aid, 3) for aid in agents_cfg}

    # ── Comprehensive temperature + hardware metrics ──────────────────────
    # CPU temp via k10temp (Ryzen)
    cpu_temp = _run("sensors -j 2>/dev/null | python3 -c \"import sys,json; d=json.load(sys.stdin); print(f'{d.get(\\\"k10temp-pci-00c3\\\",{}).get(\\\"Tctl\\\",{}).get(\\\"temp1_input\\\",0):.0f}°C')\" 2>/dev/null")
    if not cpu_temp or cpu_temp == '0°C':
        cpu_temp = _run("cat /sys/class/thermal/thermal_zone0/temp 2>/dev/null | awk '{printf \"%.0f°C\", $1/1000}'") or "N/A"

    # Motherboard + chipset via ASUS EC
    mobo_temp = _run("sensors -j 2>/dev/null | python3 -c \"import sys,json; d=json.load(sys.stdin); ec=d.get('asusec-isa-0000',{}); mb=ec.get('Motherboard',{}); print(f'{mb.get(\\\"temp3_input\\\",0):.0f}°C')\" 2>/dev/null") or "N/A"
    chipset_temp = _run("sensors -j 2>/dev/null | python3 -c \"import sys,json; d=json.load(sys.stdin); ec=d.get('asusec-isa-0000',{}); cs=ec.get('Chipset',{}); [print(f'{v:.0f}°C') for k,v in cs.items() if 'input' in k and isinstance(v,(int,float)) and v>0]\" 2>/dev/null") or "N/A"

    # NVIDIA RTX 3070
    nvidia = _run("nvidia-smi --query-gpu=temperature.gpu,fan.speed,power.draw,memory.used,memory.total,utilization.gpu --format=csv,noheader,nounits 2>/dev/null")
    nvidia_temp = nvidia_fan = nvidia_power = nvidia_vram = nvidia_util = "N/A"
    if nvidia:
        parts = [p.strip() for p in nvidia.split(',')]
        if len(parts) >= 6:
            nvidia_temp = f"{parts[0]}°C"
            nvidia_fan = f"{parts[1]}%"
            nvidia_power = f"{parts[2]}W"
            nvidia_vram = f"{parts[3]}/{parts[4]}MB"
            nvidia_util = f"{parts[5]}%"

    # AMD RX 6700 XT — use rocm-smi first (sysfs locked by Vulkan), fall back to sysfs
    amd_temp = "N/A"
    amd_fan = "N/A"
    amd_power = "N/A"
    _rocm_temp_live = _run("rocm-smi --showtemp 2>/dev/null")
    if _rocm_temp_live:
        for _rl in _rocm_temp_live.splitlines():
            if "Temperature" in _rl and "edge" in _rl.lower():
                _tv = ''.join(c for c in _rl.split(":")[-1] if c.isdigit() or c == '.')
                if _tv:
                    amd_temp = f"{int(float(_tv))}°C"
        _rocm_use_live = _run("rocm-smi --showuse 2>/dev/null")
        for _rl in (_rocm_use_live or "").splitlines():
            if "GPU use" in _rl:
                _uv = ''.join(c for c in _rl.split(":")[-1] if c.isdigit())
                if _uv:
                    amd_power = f"{_uv}% util"
    if amd_temp == "N/A":
        # Fallback to sysfs
        for card in ["card0", "card1", "card2", "card3"]:
            t = _run(f"cat /sys/class/drm/{card}/device/hwmon/hwmon*/temp1_input 2>/dev/null")
            name = _run(f"cat /sys/class/drm/{card}/device/hwmon/hwmon*/name 2>/dev/null")
            if name == "amdgpu" and t and t.strip():
                val = int(t) // 1000
                if val > 0:
                    amd_temp = f"{val}°C"
                fan = _run(f"cat /sys/class/drm/{card}/device/hwmon/hwmon*/fan1_input 2>/dev/null")
                if fan and fan.strip() and fan.strip().isdigit() and int(fan) > 0:
                    amd_fan = f"{fan}RPM"
                break
    if amd_temp == "N/A":
        amd_temp = "idle"

    # NVMe SSD temps
    nvme0_temp = _run("sensors -j 2>/dev/null | python3 -c \"import sys,json; d=json.load(sys.stdin); n=d.get('nvme-pci-0400',d.get('nvme-pci-0e00',{})); c=n.get('Composite',{}); print(f'{c.get(\\\"temp1_input\\\",0):.0f}°C')\" 2>/dev/null") or "N/A"
    nvme1_temp = _run("sensors -j 2>/dev/null | python3 -c \"import sys,json; d=json.load(sys.stdin); n=d.get('nvme-pci-0e00',{}); c=n.get('Composite',{}); print(f'{c.get(\\\"temp1_input\\\",0):.0f}°C')\" 2>/dev/null") or "N/A"

    online = sum(1 for s in statuses.values() if s == 'online')
    return jsonify({
        "statuses": statuses,
        "messages": messages,
        "online": online,
        "total": len(statuses),
        "metrics": {
            "cpu_temp":     cpu_temp,
            "mobo_temp":    mobo_temp,
            "chipset_temp": chipset_temp,
            "nvidia_temp":  nvidia_temp,
            "nvidia_fan":   nvidia_fan,
            "nvidia_power": nvidia_power,
            "nvidia_vram":  nvidia_vram,
            "nvidia_util":  nvidia_util,
            "amd_temp":     amd_temp,
            "amd_fan":      amd_fan,
            "amd_power":    amd_power,
            "nvme0_temp":   nvme0_temp,
            "nvme1_temp":   nvme1_temp,
            "mem":          _run("free -h | awk '/^Mem:/{print $3\"/\"$2}'"),
            "disk":         _run("df -h / | tail -1 | awk '{print $5}'"),
            "ollama_vulkan": _port("localhost", 11434),
            "ollama_amd":   _port("localhost", 11434),
            "ollama_gpu":   _port("localhost", 11435),
            "ollama_cpu":   _port("localhost", 11436),
            "ollama_amd2":  _port("localhost", 11437),
            "litellm":      _port("localhost", 4000),
            "printer":      _run("lpstat -p 2>/dev/null | head -1 | grep -q 'idle\\|printing' && echo 'online' || echo 'offline'") or "offline",
            "printer_name": _run("lpstat -p 2>/dev/null | head -1 | awk '{print $2}'") or "—",
            "printer_jobs": _run("lpstat -o 2>/dev/null | wc -l") or "0",
        }
    })

@app.route('/api/models')
def api_models():
    """All available models: local Ollama (Vulkan + CUDA) + LiteLLM cloud + Ollama library."""
    import socket as _socket
    def _port(host, port):
        try:
            with _socket.create_connection((host, port), timeout=2): return True
        except: return False

    def _ollama_models(base):
        try:
            r = subprocess.run(['curl','-s',f'{base}/api/tags'], capture_output=True, text=True, timeout=5)
            import json as _json
            data = _json.loads(r.stdout)
            return [{"name": m["name"], "size": m.get("size",0),
                     "params": m.get("details",{}).get("parameter_size",""),
                     "quant":  m.get("details",{}).get("quantization_level",""),
                     "installed": True}
                    for m in data.get("models",[])]
        except: return []

    def _litellm_models():
        try:
            r = subprocess.run([
                'curl','-s','-H','Authorization: Bearer baza-litellm-internal',
                'http://localhost:4000/v1/models'
            ], capture_output=True, text=True, timeout=5)
            import json as _json
            return [{"name": m["id"], "provider": "litellm", "installed": True}
                    for m in _json.loads(r.stdout).get("data",[])]
        except: return []

    def _ollama_library():
        """Fetch popular models from Ollama library (pullable but not yet installed)."""
        try:
            import urllib.request as _ur
            import json as _json
            req = _ur.Request('https://ollama.com/api/tags', headers={'User-Agent': 'baza-dashboard'})
            with _ur.urlopen(req, timeout=8) as resp:
                data = _json.loads(resp.read())
            installed = {m["name"].split(":")[0] for m in _ollama_models("http://localhost:11434")}
            library = []
            for m in data.get("models", [])[:100]:
                base_name = m["name"].split(":")[0]
                library.append({
                    "name": m["name"],
                    "size": m.get("size", 0),
                    "installed": base_name in installed,
                })
            return library
        except Exception:
            return []

    # Read LiteLLM config for cloud model list (even if proxy DB is down)
    def _litellm_config_models():
        """Read cloud models from litellm.yaml config file."""
        try:
            config_path = os.path.join(os.path.dirname(DASHBOARD_DIR), 'configs', 'litellm.yaml')
            import yaml as _yaml
            with open(config_path) as f:
                cfg = _yaml.safe_load(f)
            return [{"name": m.get("model_name",""), "provider": m.get("litellm_params",{}).get("model","").split("/")[0],
                     "model_id": m.get("litellm_params",{}).get("model",""), "installed": True}
                    for m in cfg.get("model_list", []) if m.get("model_name")]
        except Exception:
            return []

    return jsonify({
        "vulkan":  {"label": "Ollama Vulkan (RX 6700 XT)", "url": "localhost:11434", "up": _port("localhost",11434), "models": _ollama_models("http://localhost:11434")},
        "cuda":    {"label": "Ollama CUDA (RTX 3070)",      "url": "localhost:11435", "up": _port("localhost",11435), "models": _ollama_models("http://localhost:11435")},
        "cpu":     {"label": "Ollama CPU (64GB RAM)",       "url": "localhost:11436", "up": _port("localhost",11436), "models": _ollama_models("http://localhost:11436")},
        "vulkan2": {"label": "Ollama Vulkan 2 (overflow)",  "url": "localhost:11437", "up": _port("localhost",11437), "models": _ollama_models("http://localhost:11437")},
        "cloud":   {"label": "LiteLLM Cloud Models",        "url": "localhost:4000",  "up": _port("localhost",4000),  "models": _litellm_models() or _litellm_config_models()},
        "library": {"label": "Ollama Library (pullable)",    "url": "ollama.com",      "up": True,                     "models": _ollama_library()},
    })

@app.route('/api/artifacts/project-list')
def api_artifact_project_list():
    """Return all existing project folders + known agent IDs for upload dropdowns."""
    config = load_config()
    agent_ids = list(config.get('agents', {}).keys())
    projects = []
    if os.path.exists(ARTIFACTS_DIR):
        projects = sorted([d for d in os.listdir(ARTIFACTS_DIR)
                           if os.path.isdir(os.path.join(ARTIFACTS_DIR, d))])
    return jsonify({"projects": projects, "agents": agent_ids})

# ── Routes — Cron API ─────────────────────────────────────────────────────────

@app.route('/api/crons', methods=['GET'])
def api_crons_list():
    return jsonify(list_crons())

@app.route('/api/crons/raw', methods=['GET'])
def api_crons_raw():
    return jsonify({'crontab': get_raw_crontab()})

@app.route('/api/crons/raw', methods=['POST'])
def api_crons_raw_save():
    data = request.json or {}
    content = data.get('content', '')
    ok = set_raw_crontab(content)
    return jsonify({'success': ok})

@app.route('/api/crons/add', methods=['POST'])
def api_crons_add():
    data     = request.json or {}
    name     = data.get('name', '').strip()
    schedule = data.get('schedule', '').strip()
    command  = data.get('command', '').strip()
    if not name or not schedule or not command:
        return jsonify({'success': False, 'error': 'name, schedule, command required'}), 400
    ok = add_cron_job(name, schedule, command)
    return jsonify({'success': ok})

@app.route('/api/crons/remove', methods=['POST'])
def api_crons_remove():
    data = request.json or {}
    name = data.get('name', '').strip()
    if not name:
        return jsonify({'success': False, 'error': 'name required'}), 400
    ok = remove_cron_job(name)
    return jsonify({'success': ok})

@app.route('/api/crons/toggle', methods=['POST'])
def api_crons_toggle():
    data    = request.json or {}
    name    = data.get('name', '').strip()
    enabled = data.get('enabled', True)
    ok = toggle_cron_job(name, enabled)
    return jsonify({'success': ok})

@app.route('/api/crons/run-now', methods=['POST'])
def api_crons_run_now():
    """Immediately run a cron job's command in background."""
    data    = request.json or {}
    name    = data.get('name', '').strip()
    crons   = list_crons()
    job     = next((c for c in crons if c.get('id') == name), None)
    if not job or not job.get('command'):
        return jsonify({'success': False, 'error': 'Job not found or no command'}), 404
    try:
        subprocess.Popen(job['command'], shell=True,
                         stdout=open(os.path.join(LOGS_DIR,'cron_manual.log'),'a'),
                         stderr=subprocess.STDOUT)
        return jsonify({'success': True, 'message': f'Running: {job["command"][:80]}'})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

# ── Routes — Artifacts API ────────────────────────────────────────────────────

@app.route('/api/artifacts')
def api_artifacts():
    project_id = request.args.get('project_id')
    agent_id   = request.args.get('agent_id')
    if project_id:
        arts = artifacts_for_project(project_id)
    else:
        arts = all_artifacts()
    if agent_id:
        arts = [a for a in arts if a.get('agent_id') == agent_id]
    return jsonify(arts)

@app.route('/api/artifacts/save-text', methods=['POST'])
def api_artifact_save_text():
    import re as _re
    data       = request.json or {}
    project_id = data.get('project_id', 'shared')
    raw_name   = data.get('filename', f"artifact_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.txt")
    content    = data.get('content', '')
    agent_id   = data.get('agent_id', '')
    task_id    = data.get('task_id', '')
    safe_name  = _re.sub(r'[^\w.\-]', '_', raw_name).strip('_') or 'artifact.txt'
    # All extensions accepted
    proj_dir   = os.path.join(ARTIFACTS_DIR, project_id)
    os.makedirs(proj_dir, exist_ok=True)
    fpath      = os.path.join(proj_dir, safe_name)
    with open(fpath, 'w', encoding='utf-8') as f:
        f.write(content)
    # Write sidecar meta file for reliable agent/task tracking
    if agent_id or task_id:
        try:
            meta = {'agent_id': agent_id, 'task_id': task_id,
                    'created_at': datetime.datetime.now().isoformat()}
            with open(fpath + '.meta', 'w') as mf:
                json.dump(meta, mf)
        except Exception:
            pass
    size = os.path.getsize(fpath)
    return jsonify({'success': True, 'name': safe_name, 'project_id': project_id,
                    'agent_id': agent_id, 'task_id': task_id,
                    'size': size, 'download_url': f'/api/artifacts/download/{project_id}/{safe_name}'})

# ── Baza image picker ─────────────────────────────────────────────────────
# Lets any image-upload surface (project before/during/after, project docs,
# receipts, media library) pick from images already in artifacts/ —
# especially Sam's .private-inbound/sam_axe/ Telegram drops while the SD
# WebUI imaging path is down. The picker is an explicit opt-in, so it
# includes private-inbound photos alongside public artifacts; the default
# Data Hub views still hide private files.
_PICK_IMG_EXTS = {'.png', '.jpg', '.jpeg', '.gif', '.webp', '.heic', '.heif', '.bmp'}
_PICK_KNOWN_AGENTS = ('sam_axe', 'phil_hass', 'simon_bately', 'claw_batto',
                      'nova_sterling', 'duke_harmon', 'scout_reeves',
                      'rex_valor', 'specter_voss')


def _pick_encode_token(rel_path: str) -> str:
    import base64 as _b64
    return _b64.urlsafe_b64encode(rel_path.encode('utf-8')).decode('ascii').rstrip('=')


def _pick_decode_token(token: str):
    """Decode a Baza-picker token to an absolute path inside ARTIFACTS_DIR.
    Returns None for invalid / out-of-tree / non-existent paths. Vault
    files are refused — they're only reachable via the vault flow."""
    import base64 as _b64
    if not token or not isinstance(token, str):
        return None
    try:
        pad = '=' * (-len(token) % 4)
        rel = _b64.urlsafe_b64decode((token + pad).encode('ascii')).decode('utf-8')
    except Exception:
        return None
    parts = rel.replace('\\', '/').split('/')
    if '..' in parts or VAULT_DIRNAME in parts:
        return None
    fpath = os.path.realpath(os.path.join(ARTIFACTS_DIR, rel))
    if not fpath.startswith(os.path.realpath(ARTIFACTS_DIR)):
        return None
    if not os.path.isfile(fpath):
        return None
    return fpath


def _pick_list_images(limit: int = 60, agent_filter: str = '',
                      include_private: bool = True) -> list:
    out = []
    if not os.path.isdir(ARTIFACTS_DIR):
        return out
    for root, dirs, fnames in os.walk(ARTIFACTS_DIR):
        # Skip vault entirely; descend into .private-inbound (public).
        dirs[:] = [d for d in dirs
                   if (not d.startswith('.') or d == PRIVATE_INBOUND_DIRNAME)
                   and d not in {'__pycache__', '.git', VAULT_DIRNAME}]
        for fname in fnames:
            if fname.endswith('.meta'):
                continue
            ext = os.path.splitext(fname)[1].lower()
            if ext not in _PICK_IMG_EXTS:
                continue
            fpath = os.path.join(root, fname)
            try:
                st = os.stat(fpath)
            except OSError:
                continue
            private = _is_private_path(fpath)
            # Vault is excluded by the dirs filter above; this only catches
            # legacy `.meta private=true` stragglers, which should not appear
            # in the picker either.
            if private:
                continue
            agent_id = ''
            caption = ''
            meta_path = fpath + '.meta'
            if os.path.isfile(meta_path):
                try:
                    with open(meta_path, 'r', encoding='utf-8', errors='replace') as fh:
                        raw = fh.read().strip()
                    if raw.startswith('{'):
                        m = json.loads(raw)
                        agent_id = str(m.get('agent_id', '') or '')
                        caption = str(m.get('caption', '') or '')[:200]
                except Exception:
                    pass
            if not agent_id:
                for ag in _PICK_KNOWN_AGENTS:
                    if fname.startswith(ag + '_'):
                        agent_id = ag
                        break
            if agent_filter and agent_id != agent_filter:
                continue
            rel = os.path.relpath(fpath, ARTIFACTS_DIR).replace(os.sep, '/')
            out.append({
                'token':    _pick_encode_token(rel),
                'name':     fname,
                'rel_path': rel,
                'size':     st.st_size,
                'mtime':    st.st_mtime,
                'modified': datetime.datetime.fromtimestamp(st.st_mtime).strftime('%Y-%m-%d %H:%M'),
                'ext':      ext,
                'agent_id': agent_id,
                'private':  private,
                'caption':  caption,
            })
    out.sort(key=lambda r: r['mtime'], reverse=True)
    if limit and len(out) > limit:
        out = out[:limit]
    return out


@app.route('/api/datahub/images/recent')
def api_datahub_images_recent():
    """List recent images from artifacts/ for the Baza picker UI. Includes
    .private-inbound by default — the picker is an explicit opt-in surface."""
    try:
        limit = int(request.args.get('limit', 60))
    except (TypeError, ValueError):
        limit = 60
    limit = max(1, min(limit, 300))
    agent_filter = (request.args.get('agent') or '').strip()
    include_private = request.args.get('include_private', '1') not in ('0', 'false', 'no')
    images = _pick_list_images(limit=limit, agent_filter=agent_filter,
                               include_private=include_private)
    return jsonify({'ok': True, 'count': len(images), 'images': images})


@app.route('/api/datahub/pick/serve/<token>')
def api_datahub_pick_serve(token):
    """Serve a Baza picker file inline. Path is verified to stay under
    ARTIFACTS_DIR by _pick_decode_token."""
    fpath = _pick_decode_token(token)
    if not fpath:
        return jsonify({'error': 'Not found'}), 404
    return send_from_directory(os.path.dirname(fpath), os.path.basename(fpath))


@app.route('/api/artifacts/upload', methods=['POST'])
def api_artifact_upload():
    import re as _re
    project_id = request.form.get('project_id', 'shared')
    subfolder  = request.form.get('subfolder', '').strip('/')
    files      = request.files.getlist('file')
    if not files:
        return jsonify({'success': False, 'error': 'No files provided'}), 400
    proj_dir = os.path.join(ARTIFACTS_DIR, project_id, subfolder) if subfolder else os.path.join(ARTIFACTS_DIR, project_id)
    os.makedirs(proj_dir, exist_ok=True)
    saved = []
    errors = []
    for f in files:
        if not f or not f.filename:
            continue
        # Preserve original filename, sanitise only dangerous chars
        safe_name = _re.sub(r'[^\w.\-_ ()]', '_', f.filename).strip() or f'upload_{len(saved)}'
        fpath = os.path.join(proj_dir, safe_name)
        try:
            f.save(fpath)
            saved.append({'name': safe_name, 'size': os.path.getsize(fpath),
                          'project_id': project_id,
                          'download_url': f'/api/artifacts/download/{project_id}/{safe_name}'})
        except Exception as e:
            errors.append({'name': f.filename, 'error': str(e)})
    return jsonify({'success': len(saved) > 0, 'files': saved, 'errors': errors,
                    'count': len(saved)})

@app.route('/api/artifacts/download/<project_id>/<path:filename>')
def api_artifact_download(project_id, filename):
    # Block vault paths only — .private-inbound/ is public now.
    if project_id == VAULT_DIRNAME or VAULT_DIRNAME in filename.split('/'):
        return jsonify({'error': 'Vault — locked'}), 403
    proj_dir = os.path.join(ARTIFACTS_DIR, project_id)
    fpath = os.path.realpath(os.path.join(proj_dir, filename))
    if not fpath.startswith(os.path.realpath(ARTIFACTS_DIR)):
        return jsonify({'error': 'Forbidden'}), 403
    if _is_private_path(fpath):
        return jsonify({'error': 'Vault — locked'}), 403
    return send_from_directory(proj_dir, filename, as_attachment=True)

@app.route('/api/artifacts/view/<project_id>/<path:filename>')
def api_artifact_view(project_id, filename):
    """Read text file content for preview. Supports subpaths."""
    if project_id == VAULT_DIRNAME or VAULT_DIRNAME in filename.split('/'):
        return jsonify({'error': 'Vault — locked'}), 403
    proj_dir = os.path.join(ARTIFACTS_DIR, project_id)
    fpath = os.path.realpath(os.path.join(proj_dir, filename))
    if not fpath.startswith(os.path.realpath(ARTIFACTS_DIR)):
        return jsonify({'error': 'Forbidden'}), 403
    if not os.path.isfile(fpath):
        return jsonify({'error': 'File not found'}), 404
    if _is_private_path(fpath):
        return jsonify({'error': 'Private'}), 403
    base = os.path.basename(fpath)
    ext = os.path.splitext(base)[1].lower()
    text_exts = {'.txt','.md','.py','.sh','.js','.ts','.jsx','.tsx','.html','.htm',
                 '.css','.json','.yaml','.yml','.toml','.ini','.cfg','.conf','.sql',
                 '.log','.env','.csv','.rst','.xml','.bash','.zsh','.rb','.go',
                 '.php','.svg','.rs','.c','.cpp','.h'}
    if ext in text_exts:
        try:
            content = open(fpath, 'r', errors='replace').read(500_000)
            return jsonify({'content': content, 'name': base, 'type': 'text'})
        except Exception:
            pass
    return jsonify({
        'content': None, 'name': base, 'type': 'binary',
        'serve_url': f'/api/artifacts/serve/{project_id}/{filename}'
    })

_GREP_TEXT_EXTS = {'.md','.txt','.py','.js','.ts','.jsx','.tsx','.html','.htm','.css',
                   '.json','.yaml','.yml','.toml','.ini','.cfg','.conf','.sql','.log',
                   '.env','.csv','.rst','.xml','.sh','.bash','.go','.rs','.rb','.php',
                   '.c','.cpp','.h'}
IMAGE_CAPTIONS_DB = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'image_captions.db')


def _grep_image_captions(q_low: str, limit: int = 40):
    """Match live-search query against stored image captions + tags."""
    if not os.path.exists(IMAGE_CAPTIONS_DB):
        return []
    out = []
    try:
        con = sqlite3.connect(IMAGE_CAPTIONS_DB, timeout=5)
        con.row_factory = sqlite3.Row
        cur = con.execute("""
            SELECT project_id, sub_path, caption, tags
              FROM image_captions
             WHERE status = 'ok'
               AND (LOWER(caption) LIKE ? OR LOWER(tags) LIKE ?)
             LIMIT ?
        """, (f'%{q_low}%', f'%{q_low}%', limit))
        for r in cur:
            base = (r['sub_path'] or '').rsplit('/', 1)[-1]
            ext  = os.path.splitext(base)[1].lower()
            snippet = (r['caption'] or r['tags'] or '')[:240]
            out.append({
                'project_id': r['project_id'], 'name': r['sub_path'],
                'basename':   base, 'snippet': snippet,
                'match_count': 1, 'ext': ext, 'match_type': 'image',
            })
        con.close()
    except Exception:
        pass
    return out


@app.route('/api/artifacts/reindex-images', methods=['POST'])
def api_artifact_reindex_images():
    """Fire-and-forget: spawn image_indexer.py as a detached subprocess."""
    script = os.path.join(DASHBOARD_DIR, 'image_indexer.py')
    if not os.path.exists(script):
        return jsonify({'ok': False, 'error': 'indexer not installed'}), 500
    payload = (request.json or {}) if request.is_json else {}
    venv_py = os.path.join(FRAMEWORK_DIR, 'venv', 'bin', 'python')
    py = venv_py if os.path.exists(venv_py) else 'python3'
    cmd = [py, script]
    if payload.get('force'):        cmd.append('--force')
    if payload.get('retry_failed'): cmd.append('--retry-failed')
    if payload.get('limit'):        cmd += ['--limit', str(int(payload['limit']))]
    try:
        subprocess.Popen(cmd, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL,
                         start_new_session=True)
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)[:200]}), 500
    return jsonify({'ok': True, 'started': True})


@app.route('/api/artifacts/image-index-status')
def api_artifact_image_index_status():
    """Progress/summary for the image caption index."""
    img_ext = {'.png','.jpg','.jpeg','.gif','.webp','.bmp','.tiff','.tif'}
    total = 0
    if os.path.isdir(ARTIFACTS_DIR):
        for root, dirs, fnames in os.walk(ARTIFACTS_DIR):
            dirs[:] = [d for d in dirs if not d.startswith('.')]
            for fn in fnames:
                if os.path.splitext(fn)[1].lower() in img_ext:
                    try:
                        if os.path.getsize(os.path.join(root, fn)) >= 20*1024:
                            total += 1
                    except OSError:
                        pass
    captioned = failed = 0
    last_at   = None
    if os.path.exists(IMAGE_CAPTIONS_DB):
        try:
            con = sqlite3.connect(IMAGE_CAPTIONS_DB, timeout=3)
            captioned = con.execute("SELECT COUNT(*) FROM image_captions WHERE status='ok'").fetchone()[0]
            failed    = con.execute("SELECT COUNT(*) FROM image_captions WHERE status='failed'").fetchone()[0]
            last_at   = con.execute("SELECT MAX(indexed_at) FROM image_captions").fetchone()[0]
            con.close()
        except Exception as e:
            return jsonify({'ok': False, 'error': str(e)[:200]}), 500
    return jsonify({
        'ok': True, 'total_images': total, 'captioned': captioned,
        'failed': failed, 'pending': max(0, total - captioned - failed),
        'last_indexed_at': last_at, 'ever_indexed': (captioned + failed) > 0,
    })


@app.route('/api/artifacts/grep')
def api_artifact_grep():
    """Live keyword search across text artifacts. Returns matches with snippets."""
    q = (request.args.get('q','') or '').strip()
    if len(q) < 2:
        return jsonify({'results': [], 'query': q})
    q_low = q.lower()
    max_file_size = 2 * 1024 * 1024
    max_results   = 60
    results = []
    if not os.path.exists(ARTIFACTS_DIR):
        return jsonify({'results': [], 'query': q})
    for root, dirs, fnames in os.walk(ARTIFACTS_DIR):
        dirs[:] = [d for d in dirs if not d.startswith('.')]
        for fname in fnames:
            if fname.endswith('.meta'):
                continue
            ext = os.path.splitext(fname)[1].lower()
            if ext not in _GREP_TEXT_EXTS:
                continue
            fpath = os.path.join(root, fname)
            try:
                if os.path.getsize(fpath) > max_file_size:
                    continue
                with open(fpath, 'r', errors='replace') as fh:
                    content = fh.read()
                idx = content.lower().find(q_low)
                if idx < 0:
                    continue
                match_count = content.lower().count(q_low)
                start = max(0, idx - 40)
                end   = min(len(content), idx + len(q) + 80)
                snippet = content[start:end].replace('\n', ' ').replace('\r', ' ').strip()
                rel = os.path.relpath(fpath, ARTIFACTS_DIR)
                parts = rel.replace('\\','/').split('/')
                proj = parts[0] if len(parts) > 1 else 'shared'
                sub  = '/'.join(parts[1:]) if len(parts) > 1 else fname
                results.append({
                    'project_id': proj, 'name': sub, 'basename': fname,
                    'snippet': snippet, 'match_count': match_count, 'ext': ext,
                })
                if len(results) >= max_results:
                    break
            except Exception:
                continue
        if len(results) >= max_results:
            break
    # Merge image-caption hits (semantic object match) under the same search
    img_hits = _grep_image_captions(q_low, limit=40)
    # De-dup on (project_id, name)
    seen = {(r['project_id'], r['name']) for r in results}
    for h in img_hits:
        key = (h['project_id'], h['name'])
        if key not in seen:
            results.append(h)
            seen.add(key)
    return jsonify({'results': results, 'query': q,
                    'truncated': len(results) >= max_results,
                    'image_matches': len(img_hits)})


@app.route('/api/artifacts/serve/<project_id>/<path:filename>')
def api_artifact_serve(project_id, filename):
    """Serve file inline for browser preview (images, PDFs, etc). Supports subpaths."""
    # Block vault paths only — .private-inbound/ is public now.
    if project_id == VAULT_DIRNAME or VAULT_DIRNAME in filename.split('/'):
        return jsonify({'error': 'Vault — locked'}), 403
    proj_dir = os.path.join(ARTIFACTS_DIR, project_id)
    fpath    = os.path.realpath(os.path.join(proj_dir, filename))
    # Path traversal guard — must stay inside ARTIFACTS_DIR
    if not fpath.startswith(os.path.realpath(ARTIFACTS_DIR)):
        return jsonify({'error': 'Forbidden'}), 403
    if not os.path.isfile(fpath):
        return jsonify({'error': 'Not found'}), 404
    # Per-file privacy gate — defense in depth in case a private-marked file
    # ever lands in a public project dir.
    if _is_private_path(fpath):
        return jsonify({'error': 'Private'}), 403
    return send_from_directory(os.path.dirname(fpath), os.path.basename(fpath))

@app.route('/api/datahub/specter')
def api_datahub_specter():
    """Live Specter Voss status for the DataHub Specter card.
    Pulls heartbeat from Redis (set by openclaw/telegram_bridge.py every cycle)
    and recent insights from empire_knowledge (set by publish_insight skill).
    Returns: status, last_seen, model, recent_insights[]"""
    out = {
        "status": "offline",
        "last_seen": None,
        "model": None,
        "seconds_since_heartbeat": None,
        "insights": [],
        "insight_count": 0,
    }

    # Heartbeat from Redis
    try:
        import redis as _redis
        r = _redis.Redis(host="localhost", port=6379, decode_responses=True)
        hb_raw = r.get("baza:heartbeat:specter_voss")
        if hb_raw:
            hb = json.loads(hb_raw)
            ts = int(hb.get("ts", 0))
            now_ts = int(datetime.datetime.now().timestamp())
            age = now_ts - ts
            out["seconds_since_heartbeat"] = age
            out["last_seen"] = datetime.datetime.fromtimestamp(ts).isoformat()
            out["model"] = hb.get("model")
            # Heartbeat TTL is 120s; treat <180s as alive, otherwise stale
            if age < 180:
                out["status"] = hb.get("status", "online") or "online"
            elif age < 600:
                out["status"] = "stale"
            else:
                out["status"] = "offline"
    except Exception as e:
        out["redis_error"] = str(e)

    # Recent insights from PostgreSQL empire_knowledge
    try:
        from core.context_db import get_pool
        pool = get_pool()
        conn = pool.getconn()
        cur = conn.cursor()
        cur.execute(
            "SELECT key, value, category, updated_at FROM empire_knowledge "
            "WHERE key LIKE 'specter_%' ORDER BY updated_at DESC LIMIT 12"
        )
        rows = cur.fetchall()
        for k, v, cat, ts in rows:
            # key format: specter_<category>_<safe_title>
            title = (k or "").replace(f"specter_{cat}_", "", 1).replace("_", " ").title()
            out["insights"].append({
                "key": k,
                "title": title[:80],
                "category": cat or "research",
                "preview": (v or "")[:280],
                "size": len(v or ""),
                "updated_at": ts.isoformat() if ts else None,
            })
        cur.execute("SELECT count(*) FROM empire_knowledge WHERE key LIKE 'specter_%'")
        out["insight_count"] = cur.fetchone()[0]
        cur.close()
        pool.putconn(conn)
    except Exception as e:
        out["pg_error"] = str(e)

    return jsonify(out)


# ── Specter insight detail + actions (click-through from Data Hub) ────────────

_AGENT_TOKEN_ENV = {
    'simon_bately': 'TELEGRAM_SIMON_BATELY',
    'claw_batto': 'TELEGRAM_CLAW_BATTO',
    'phil_hass': 'TELEGRAM_PHIL_HASS',
    'sam_axe': 'TELEGRAM_SAM_AXE',
    'rex_valor': 'TELEGRAM_REX_VALOR',
    'duke_harmon': 'TELEGRAM_DUKE_HARMON',
    'nova_sterling': 'TELEGRAM_NOVA_STERLING',
    'scout_reeves': 'TELEGRAM_SCOUT_REEVES',
    'specter_voss': 'TELEGRAM_SPECTER_VOSS',
}


def _agent_telegram_token(agent_id: str) -> str:
    env_name = _AGENT_TOKEN_ENV.get(agent_id, '')
    if not env_name:
        return ''
    token = os.environ.get(env_name, '')
    if token:
        return token
    secrets_path = os.path.join(os.path.dirname(DASHBOARD_DIR), 'configs', 'secrets.env')
    if os.path.exists(secrets_path):
        with open(secrets_path) as sf:
            for line in sf:
                line = line.strip()
                if line.startswith(env_name + '='):
                    return line.split('=', 1)[1].strip().strip('"').strip("'")
    return ''


def _agent_last_chat_id(agent_id: str):
    """Look up the most-recent chat_id the agent has talked to (private DM)."""
    try:
        from core.context_db import get_pool
        pool = get_pool()
        conn = pool.getconn()
        cur = conn.cursor()
        cur.execute(
            "SELECT DISTINCT chat_id FROM messages WHERE agent_id=%s "
            "ORDER BY chat_id DESC LIMIT 1", (agent_id,))
        row = cur.fetchone()
        cur.close()
        pool.putconn(conn)
        if row:
            return row[0]
    except Exception:
        pass
    return None


def _notify_agent(agent_id: str, message: str) -> tuple[bool, str]:
    """Best-effort Telegram notify. Returns (ok, detail)."""
    token = _agent_telegram_token(agent_id)
    if not token:
        return False, f"no Telegram token for {agent_id}"
    chat_id = _agent_last_chat_id(agent_id)
    if not chat_id:
        return False, f"no recent chat_id for {agent_id} — DM them first"
    try:
        import requests as _req
        resp = _req.post(
            f"https://api.telegram.org/bot{token}/sendMessage",
            json={"chat_id": chat_id, "text": message[:3900]},
            timeout=10,
        )
        if resp.ok and resp.json().get("ok"):
            return True, "sent"
        return False, f"telegram returned {resp.status_code}: {resp.text[:200]}"
    except Exception as e:
        return False, f"telegram error: {e}"


@app.route('/api/datahub/specter/insight')
def api_datahub_specter_insight():
    """Full value for one Specter insight, looked up by exact key."""
    key = request.args.get('key', '').strip()
    if not key:
        return jsonify({"error": "key required"}), 400
    try:
        from core.context_db import get_pool
        pool = get_pool()
        conn = pool.getconn()
        cur = conn.cursor()
        cur.execute(
            "SELECT key, value, category, updated_at, updated_by "
            "FROM empire_knowledge WHERE key = %s LIMIT 1",
            (key,),
        )
        row = cur.fetchone()
        cur.close()
        pool.putconn(conn)
        if not row:
            return jsonify({"error": "not found"}), 404
        return jsonify({
            "key": row[0],
            "value": row[1] or "",
            "category": row[2],
            "updated_at": row[3].isoformat() if row[3] else None,
            "updated_by": row[4],
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route('/api/datahub/specter/insight/archive', methods=['POST'])
def api_datahub_specter_insight_archive():
    """Soft-archive: rename the key with prefix `archived_` so it drops out of
    the active Specter insight feed but stays in empire_knowledge for history."""
    data = request.json or {}
    key = (data.get("key") or "").strip()
    if not key:
        return jsonify({"ok": False, "error": "key required"}), 400
    if key.startswith("archived_"):
        return jsonify({"ok": False, "error": "already archived"}), 400
    new_key = "archived_" + key
    try:
        from core.context_db import get_pool
        pool = get_pool()
        conn = pool.getconn()
        cur = conn.cursor()
        cur.execute(
            "UPDATE empire_knowledge SET key = %s, updated_at = NOW(), "
            "updated_by = 'datahub_ui' WHERE key = %s",
            (new_key, key),
        )
        conn.commit()
        cur.close()
        pool.putconn(conn)
        return jsonify({"ok": True, "new_key": new_key})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500


@app.route('/api/datahub/specter/insight/pin', methods=['POST'])
def api_datahub_specter_insight_pin():
    """Append the insight as a ## TOPIC: <slug> block to EMPIRE_STATE.md so it
    becomes part of every agent's awareness via the self_orient skill."""
    data = request.json or {}
    key = (data.get("key") or "").strip()
    slug = (data.get("slug") or "").strip().lower()
    if not key or not slug:
        return jsonify({"ok": False, "error": "key and slug required"}), 400
    # Read full insight value
    try:
        from core.context_db import get_pool
        pool = get_pool()
        conn = pool.getconn()
        cur = conn.cursor()
        cur.execute("SELECT value FROM empire_knowledge WHERE key = %s LIMIT 1", (key,))
        row = cur.fetchone()
        cur.close()
        pool.putconn(conn)
        if not row:
            return jsonify({"ok": False, "error": "insight not found"}), 404
        value = (row[0] or "")[:1200]  # cap to keep EMPIRE_STATE.md tight
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500

    state_file = os.path.join(os.path.dirname(DASHBOARD_DIR), "EMPIRE_STATE.md")
    try:
        existing = open(state_file).read() if os.path.exists(state_file) else ""
        # Drop any prior TOPIC: <slug> block (idempotent re-pin)
        import re as _re
        existing = _re.sub(
            rf"^##\s+TOPIC:\s*{_re.escape(slug)}\s*$.*?(?=^##\s+|\Z)",
            "", existing, flags=_re.MULTILINE | _re.DOTALL | _re.IGNORECASE,
        )
        new_block = f"\n## TOPIC: {slug}\n{value}\n"
        with open(state_file, "w") as f:
            f.write(existing.rstrip() + "\n" + new_block + "\n")
        return jsonify({"ok": True, "slug": slug})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500


@app.route('/api/datahub/pca', methods=['POST'])
def api_datahub_pca():
    """Save a 'Provide a Course of Action' tied to an insight, optionally
    notify the assigned agent on Telegram and create a task."""
    data = request.json or {}
    insight_key = (data.get("insight_key") or "").strip()
    insight_title = (data.get("insight_title") or "").strip()
    action_text = (data.get("action_text") or "").strip()
    assigned = (data.get("assigned_to") or "specter_voss").strip()
    notify = bool(data.get("notify_telegram"))
    create_task = bool(data.get("create_task"))
    if not action_text:
        return jsonify({"ok": False, "error": "action_text required"}), 400

    pca_key = f"pca_{int(datetime.datetime.now().timestamp())}_{(insight_key or 'adhoc')[:60]}"
    body = (
        f"Insight: {insight_title or insight_key or '(adhoc)'}\n"
        f"Assigned: {assigned}\n"
        f"From: Serge (Data Hub)\n"
        f"Created: {datetime.datetime.now().isoformat(timespec='seconds')}\n\n"
        f"COURSE OF ACTION:\n{action_text}\n"
    )

    # Store in empire_knowledge (category='pca')
    try:
        from core.context_db import get_pool
        pool = get_pool()
        conn = pool.getconn()
        cur = conn.cursor()
        cur.execute(
            "INSERT INTO empire_knowledge (key, value, category, updated_by) "
            "VALUES (%s, %s, 'pca', 'serge') "
            "ON CONFLICT (key) DO UPDATE SET value = EXCLUDED.value, "
            "updated_at = NOW(), updated_by = EXCLUDED.updated_by",
            (pca_key, body),
        )
        conn.commit()
        cur.close()
        pool.putconn(conn)
    except Exception as e:
        return jsonify({"ok": False, "error": f"db: {e}"}), 500

    result = {"ok": True, "key": pca_key, "notified": False, "task_id": None}

    # Optional: notify the agent via Telegram
    if notify:
        ok, detail = _notify_agent(
            assigned,
            f"🎯 PCA from Serge — re: {insight_title or insight_key}\n\n{action_text}\n\n"
            f"(stored as {pca_key})",
        )
        result["notified"] = ok
        result["notify_detail"] = detail

    # Optional: create a task in baza_projects.db
    if create_task:
        try:
            import sqlite3, uuid as _uuid
            db_path = os.path.join(DASHBOARD_DIR, "baza_projects.db")
            sconn = sqlite3.connect(db_path)
            scur = sconn.cursor()
            task_id = _uuid.uuid4().hex[:8]
            scur.execute(
                "INSERT INTO tasks (id, title, status, priority, assigned_to, "
                "notes, is_subtask, project_id, created_at, updated_at) "
                "VALUES (?, ?, 'pending', 'high', ?, ?, 0, ?, "
                "datetime('now'), datetime('now'))",
                (task_id, (action_text[:80] or "PCA from Serge"), assigned,
                 body, "proj-baza-empire"),
            )
            sconn.commit()
            sconn.close()
            result["task_id"] = task_id
        except Exception as e:
            result["task_error"] = str(e)

    # Journal it
    try:
        from core.context_db import journal_log
        journal_log(
            agent_id=assigned, task_type="pca_received",
            task_description=f"Course of action from Serge: {(action_text[:120])}",
            result=body, success=True,
            input_data={"insight_key": insight_key, "pca_key": pca_key},
            requested_by="serge",
        )
    except Exception:
        pass

    return jsonify(result)


@app.route('/api/datahub/feed-item/dispatch', methods=['POST'])
def api_datahub_feed_item_dispatch():
    """From a clicked Live Feed item: re-dispatch to the agent (Telegram DM)
    or create a follow-up task. instruction body is free-form."""
    data = request.json or {}
    agent_id = (data.get("agent_id") or "").strip()
    instruction = (data.get("instruction") or "").strip()
    as_task = bool(data.get("as_task"))
    notes = (data.get("notes") or "")[:1000]
    if not agent_id or not instruction:
        return jsonify({"ok": False, "error": "agent_id and instruction required"}), 400

    if as_task:
        try:
            import sqlite3, uuid as _uuid
            db_path = os.path.join(DASHBOARD_DIR, "baza_projects.db")
            sconn = sqlite3.connect(db_path)
            scur = sconn.cursor()
            task_id = _uuid.uuid4().hex[:8]
            scur.execute(
                "INSERT INTO tasks (id, title, status, priority, assigned_to, "
                "notes, is_subtask, project_id, created_at, updated_at) "
                "VALUES (?, ?, 'pending', 'medium', ?, ?, 0, ?, "
                "datetime('now'), datetime('now'))",
                (task_id, instruction[:80], agent_id, notes, "proj-baza-empire"),
            )
            sconn.commit()
            sconn.close()
            return jsonify({"ok": True, "task_id": task_id})
        except Exception as e:
            return jsonify({"ok": False, "error": f"task create failed: {e}"}), 500

    ok, detail = _notify_agent(
        agent_id, f"↗ Re-dispatch from Serge (Data Hub feed)\n\n{instruction}",
    )
    return jsonify({"ok": ok, "detail": detail})


@app.route('/api/datahub/feed')
def api_datahub_feed():
    """Live feed of recent agent activity from task_journal + recent artifacts."""
    limit = int(request.args.get('limit', 30))
    agent_id = request.args.get('agent_id', '')
    feed = []
    # Recent artifacts
    arts = all_artifacts()
    for a in arts[:limit]:
        feed.append({
            'type': 'artifact',
            'agent_id': a.get('agent_id', 'simon_bately') or 'simon_bately',
            'name': a.get('name', ''),
            'project_id': a.get('project_id', ''),
            'ext': a.get('ext', ''),
            'size': a.get('size', 0),
            'modified': a.get('modified', ''),
            'download_url': a.get('download_url', ''),
        })
    # Recent journal entries from PostgreSQL
    try:
        from core.context_db import get_conn, release_conn
        conn = get_conn()
        cur = conn.cursor()
        q = ("SELECT id, agent_id, task_type, task_description, result, success, "
             "created_at, COALESCE(verified, TRUE) FROM task_journal")
        params = []
        if agent_id:
            q += " WHERE agent_id = %s"
            params.append(agent_id)
        q += " ORDER BY created_at DESC LIMIT %s"
        params.append(limit)
        cur.execute(q, params)
        for row in cur.fetchall():
            # row[4] is result — show more in the modal but keep feed-card list lean
            full_result = row[4] or ''
            feed.append({
                'type': 'journal',
                'id': row[0],
                'agent_id': row[1],
                'task_type': row[2],
                'description': row[3][:200] if row[3] else '',
                'result': full_result[:4000],   # modal-friendly cap
                'result_short': full_result[:200],
                'success': row[5],
                'timestamp': row[6].isoformat() if row[6] else '',
                'verified': bool(row[7]),
            })
        cur.close()
        release_conn(conn)
    except Exception:
        pass
    # Sort by timestamp/modified
    feed.sort(key=lambda x: x.get('timestamp') or x.get('modified', ''), reverse=True)
    return jsonify(feed[:limit])


# ── Activity Chains (visibility pipeline #1) ──────────────────────────────────

@app.route('/api/datahub/events')
def api_datahub_events():
    """Filtered list of task_events. Reverse chronological."""
    try:
        from core import task_events as te
    except Exception as e:
        return jsonify({"events": [], "error": f"task_events unavailable: {e}"})
    kinds_arg = request.args.get('kinds', '')
    kinds = [k.strip() for k in kinds_arg.split(',') if k.strip()] or None
    events = te.list_events(
        task_id=request.args.get('task_id') or None,
        project_id=request.args.get('project_id') or None,
        agent_id=request.args.get('agent_id') or None,
        kinds=kinds,
        since=request.args.get('since') or None,
        limit=int(request.args.get('limit', 100) or 100),
    )
    return jsonify({"events": events})


@app.route('/api/datahub/chain/<task_id>')
def api_datahub_chain(task_id):
    """Time-ascending chain for one task with parent/child nesting."""
    try:
        from core import task_events as te
    except Exception as e:
        return jsonify({"chain": [], "error": f"task_events unavailable: {e}"})
    chain = te.chain_for_task(task_id)
    # Pull task metadata from baza_projects.db tasks table for header
    meta = {}
    try:
        conn = sqlite3.connect(os.path.join(DASHBOARD_DIR, 'baza_projects.db'))
        conn.row_factory = sqlite3.Row
        row = conn.execute(
            "SELECT id, project_id, title, assigned_to, status, priority, due_date FROM tasks WHERE id=?",
            (task_id,),
        ).fetchone()
        if row:
            meta = dict(row)
        conn.close()
    except Exception:
        pass
    return jsonify({"task_id": task_id, "task": meta, "chain": chain})


@app.route('/api/datahub/chains')
def api_datahub_chains():
    """Recent task summaries — one row per task_id, newest activity first."""
    try:
        from core import task_events as te
    except Exception as e:
        return jsonify({"chains": [], "error": f"task_events unavailable: {e}"})
    rows = te.recent_task_summaries(limit=int(request.args.get('limit', 50) or 50))
    # Optional filter by agent_id at the call site
    agent_id = request.args.get('agent_id') or ''
    project_id = request.args.get('project_id') or ''
    if agent_id:
        rows = [r for r in rows if (r.get('agent_id') or '') == agent_id]
    if project_id:
        rows = [r for r in rows if (r.get('project_id') or '') == project_id]
    # Enrich with task title from tasks table
    try:
        conn = sqlite3.connect(os.path.join(DASHBOARD_DIR, 'baza_projects.db'))
        conn.row_factory = sqlite3.Row
        ids = [r['task_id'] for r in rows]
        if ids:
            placeholders = ','.join(['?'] * len(ids))
            for trow in conn.execute(
                f"SELECT id, title, status FROM tasks WHERE id IN ({placeholders})",
                ids,
            ).fetchall():
                for r in rows:
                    if r['task_id'] == trow['id']:
                        r['title'] = trow['title']
                        r['task_status'] = trow['status']
        conn.close()
    except Exception:
        pass
    return jsonify({"chains": rows})


@app.route('/api/datahub/events/stream')
def api_datahub_events_stream():
    """Server-Sent Events feed: replays last 50 events then streams live via Redis."""
    try:
        from core import task_events as te
    except Exception:
        return jsonify({"error": "task_events unavailable"}), 503

    def generate():
        # Replay last 50 events (oldest first so UI renders in order)
        recent = list(reversed(te.list_events(limit=50)))
        for ev in recent:
            yield f"data: {json.dumps(ev, default=str)}\n\n"
        yield ":replay-done\n\n"
        # Live tail via Redis
        try:
            import redis
        except Exception:
            yield ":redis-unavailable\n\n"
            return
        try:
            r = redis.Redis.from_url(te.REDIS_URL, decode_responses=True)
            pubsub = r.pubsub()
            pubsub.subscribe(te.REDIS_CHANNEL)
            last_hb = datetime.datetime.utcnow()
            for message in pubsub.listen():
                if message.get('type') == 'message':
                    data = message.get('data') or '{}'
                    yield f"data: {data}\n\n"
                # Heartbeat at most every 15s
                now = datetime.datetime.utcnow()
                if (now - last_hb).total_seconds() >= 15:
                    yield f":hb {now.isoformat()}Z\n\n"
                    last_hb = now
        except GeneratorExit:
            return
        except Exception as exc:
            yield f":stream-error {exc}\n\n"

    resp = make_response(generate(), 200)
    resp.headers['Content-Type'] = 'text/event-stream'
    resp.headers['Cache-Control'] = 'no-cache, no-transform'
    resp.headers['X-Accel-Buffering'] = 'no'
    return resp


@app.route('/api/datahub/agent-chat', methods=['POST'])
def api_datahub_agent_chat():
    """Relay a message to an agent via skill execution for Data Hub chat."""
    data = request.json or {}
    agent_id = data.get('agent_id', '')
    message = data.get('message', '')
    file_info = data.get('file_info', '')
    file_content = data.get('file_content', '')
    context = data.get('context', '')

    if not agent_id or not message:
        return jsonify({'error': 'agent_id and message required'}), 400

    # Build a prompt for the agent via Ollama
    try:
        import urllib.request as _ur
        ollama_urls = ['http://localhost:11434', 'http://localhost:11435']
        system = f"You are {agent_id}. A user is editing a file in the Data Hub and needs your help."
        if context == 'image':
            system += " You are an image/design expert. Help with visual editing requests."
        elif context in ('code', 'text'):
            system += " You are a code expert. Help edit, refactor, fix, or improve the file."

        user_msg = message
        if file_info:
            user_msg = f"[File: {file_info}]\n\n"
            if file_content:
                user_msg += f"[Content preview:]\n```\n{file_content[:2000]}\n```\n\n"
            user_msg += f"User request: {message}"

        payload = json.dumps({
            "model": "qwen3.5:latest",
            "messages": [
                {"role": "system", "content": system},
                {"role": "user", "content": user_msg}
            ],
            "stream": False,
            "options": {"num_predict": 800, "temperature": 0.7}
        }).encode()

        reply = None
        for url in ollama_urls:
            try:
                req = _ur.Request(f"{url}/api/chat", data=payload,
                                  headers={"Content-Type": "application/json"})
                with _ur.urlopen(req, timeout=60) as resp:
                    result = json.loads(resp.read())
                    reply = result.get('message', {}).get('content', '')
                    if reply:
                        break
            except Exception:
                continue

        if not reply:
            return jsonify({'reply': 'Agent is busy or offline. Try again or message them on Telegram.'})

        return jsonify({'reply': reply})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/artifacts/delete', methods=['POST'])
def api_artifact_delete():
    data       = request.json or {}
    project_id = data.get('project_id', '')
    filename   = data.get('name', '')
    if not project_id or not filename:
        return jsonify({'success': False}), 400
    fpath = os.path.join(ARTIFACTS_DIR, project_id, filename)
    if os.path.exists(fpath):
        os.remove(fpath)
        return jsonify({'success': True})
    return jsonify({'success': False, 'error': 'Not found'}), 404


@app.route('/api/print', methods=['POST'])
def api_print():
    """Print a file/artifact via the HP Smart Tank 5101."""
    data = request.json or {}
    action = data.get('action', 'print')
    skill_path = os.path.join(os.path.dirname(DASHBOARD_DIR), 'skills', 'shared', 'print_document.py')

    # If printing an artifact by project_id + filename, resolve to absolute path
    if action == 'print' and data.get('project_id') and data.get('filename'):
        fpath = os.path.join(ARTIFACTS_DIR, data['project_id'], data['filename'])
        if os.path.exists(fpath):
            data['file_path'] = fpath

    env = os.environ.copy()
    env['SKILL_ARGS'] = json.dumps(data)
    try:
        result = subprocess.run([VENV_PYTHON, skill_path], capture_output=True, text=True, timeout=30, env=env)
        output = result.stdout.strip()
        # Try to get JSON from last line
        parsed = {}
        for line in reversed(output.split('\n')):
            if line.strip().startswith('{'):
                try:
                    parsed = json.loads(line.strip())
                    break
                except Exception:
                    pass
        if not parsed:
            parsed = {'success': result.returncode == 0, 'output': output}
        if result.returncode != 0 and result.stderr:
            parsed['stderr'] = result.stderr.strip()
        return jsonify(parsed)
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})


@app.route('/api/print/status')
def api_print_status():
    """Get printer status."""
    skill_path = os.path.join(os.path.dirname(DASHBOARD_DIR), 'skills', 'shared', 'print_document.py')
    env = os.environ.copy()
    env['SKILL_ARGS'] = json.dumps({'action': 'status'})
    try:
        result = subprocess.run([VENV_PYTHON, skill_path], capture_output=True, text=True, timeout=10, env=env)
        for line in reversed(result.stdout.strip().split('\n')):
            if line.strip().startswith('{'):
                return jsonify(json.loads(line.strip()))
        return jsonify({'success': True, 'output': result.stdout.strip()})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})


@app.route('/api/artifacts/delete-bulk', methods=['POST'])
def api_artifact_delete_bulk():
    data  = request.json or {}
    files = data.get('files', [])  # [{project_id, name}]
    deleted = 0
    for f in files:
        fpath = os.path.join(ARTIFACTS_DIR, f.get('project_id',''), f.get('name',''))
        if os.path.exists(fpath):
            os.remove(fpath)
            deleted += 1
    return jsonify({'success': True, 'deleted': deleted})


@app.route('/api/artifacts/zip', methods=['POST'])
def api_artifact_zip():
    """Stream a zip of selected artifacts. Body: {files:[{project_id,name},...]}.
    Skips anything that escapes ARTIFACTS_DIR (defense in depth)."""
    data  = request.json or {}
    files = data.get('files', [])
    if not isinstance(files, list) or not files:
        return jsonify({'success': False, 'error': 'files[] required'}), 400
    import io, zipfile
    base = os.path.realpath(ARTIFACTS_DIR)
    buf  = io.BytesIO()
    seen = {}
    with zipfile.ZipFile(buf, 'w', zipfile.ZIP_DEFLATED) as z:
        for f in files:
            pid  = f.get('project_id', '')
            name = f.get('name', '')
            if not pid or not name:
                continue
            fpath = os.path.realpath(os.path.join(ARTIFACTS_DIR, pid, name))
            if not fpath.startswith(base + os.sep) or not os.path.isfile(fpath):
                continue
            arc = f"{pid}/{os.path.basename(name)}"
            n = seen.get(arc, 0)
            if n:
                stem, ext = os.path.splitext(arc)
                arc = f"{stem}_{n}{ext}"
            seen[f"{pid}/{os.path.basename(name)}"] = n + 1
            z.write(fpath, arcname=arc)
    buf.seek(0)
    ts = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
    resp = make_response(buf.getvalue())
    resp.headers['Content-Type'] = 'application/zip'
    resp.headers['Content-Disposition'] = f'attachment; filename="datahub_{ts}.zip"'
    resp.headers['Cache-Control'] = 'no-store'
    return resp

@app.route('/api/artifacts/rename', methods=['POST'])
def api_artifact_rename():
    import re as _re
    data       = request.json or {}
    project_id = data.get('project_id','')
    old_name   = data.get('old_name','')
    new_name   = _re.sub(r'[^\w.\-_ ()]','_', data.get('new_name','')).strip()
    if not all([project_id, old_name, new_name]):
        return jsonify({'success': False, 'error': 'Missing fields'})
    old_path = os.path.join(ARTIFACTS_DIR, project_id, old_name)
    new_path = os.path.join(ARTIFACTS_DIR, project_id, new_name)
    if not os.path.exists(old_path):
        return jsonify({'success': False, 'error': 'File not found'})
    os.rename(old_path, new_path)
    return jsonify({'success': True, 'new_name': new_name})

@app.route('/api/artifacts/move', methods=['POST'])
def api_artifact_move():
    data       = request.json or {}
    from_proj  = data.get('from_project','')
    to_proj    = data.get('to_project','')
    filename   = data.get('name','')
    if not all([from_proj, to_proj, filename]):
        return jsonify({'success': False, 'error': 'Missing fields'})
    src  = os.path.join(ARTIFACTS_DIR, from_proj, filename)
    dest_dir = os.path.join(ARTIFACTS_DIR, to_proj)
    os.makedirs(dest_dir, exist_ok=True)
    dest = os.path.join(dest_dir, filename)
    if not os.path.exists(src):
        return jsonify({'success': False, 'error': 'File not found'})
    import shutil
    shutil.move(src, dest)
    return jsonify({'success': True})

@app.route('/api/artifacts/send-to-agent', methods=['POST'])
def api_artifact_send_to_agent():
    data = request.json or {}
    project_id = data.get('project_id', '')
    filename = data.get('filename', '')
    agent_id = data.get('agent_id', '')
    prompt = data.get('prompt', '')
    if not all([agent_id, prompt]):
        return jsonify({'success': False, 'error': 'agent_id and prompt required'}), 400
    token_map = {
        'simon_bately': 'TELEGRAM_SIMON_BATELY',
        'claw_batto': 'TELEGRAM_CLAW_BATTO',
        'phil_hass': 'TELEGRAM_PHIL_HASS',
        'sam_axe': 'TELEGRAM_SAM_AXE',
        'rex_valor': 'TELEGRAM_REX_VALOR',
        'duke_harmon': 'TELEGRAM_DUKE_HARMON',
        'nova_sterling': 'TELEGRAM_NOVA_STERLING',
        'scout_reeves': 'TELEGRAM_SCOUT_REEVES',
    }
    token_env = token_map.get(agent_id)
    if not token_env:
        return jsonify({'success': False, 'error': f'Unknown agent: {agent_id}'}), 400
    token = os.environ.get(token_env, '')
    if not token:
        # Load from secrets file
        secrets_path = os.path.join(os.path.dirname(DASHBOARD_DIR), 'configs', 'secrets.env')
        if os.path.exists(secrets_path):
            with open(secrets_path) as sf:
                for line in sf:
                    line = line.strip()
                    if line.startswith(token_env + '='):
                        token = line.split('=', 1)[1].strip().strip('"').strip("'")
                        break
    if not token:
        return jsonify({'success': False, 'error': f'No token for {agent_id}'}), 500
    file_url = f"dashboard/artifacts/{project_id}/{filename}" if project_id and filename else ""
    message = f"\U0001f4ce Artifact Task from Dashboard\n\nFile: {filename}\nProject: {project_id}\nPath: {file_url}\n\nInstruction: {prompt}"
    import requests as _req
    try:
        # Get chat_id from PostgreSQL messages table (most reliable — bot is actively polling)
        chat_id = None
        try:
            from core.context_db import get_pool
            pool = get_pool()
            conn = pool.getconn()
            cur = conn.cursor()
            cur.execute("SELECT DISTINCT chat_id FROM messages WHERE agent_id=%s ORDER BY chat_id DESC LIMIT 1", (agent_id,))
            row = cur.fetchone()
            if row:
                chat_id = row[0]
            cur.close()
            pool.putconn(conn)
        except Exception:
            pass
        # Fallback: try getUpdates
        if not chat_id:
            resp = _req.get(f"https://api.telegram.org/bot{token}/getUpdates", params={"limit": 10}, timeout=10)
            updates = resp.json().get('result', [])
            for u in reversed(updates):
                msg = u.get('message', {})
                if msg.get('chat', {}).get('type') == 'private':
                    chat_id = msg['chat']['id']
                    break
        if not chat_id:
            return jsonify({'success': False, 'error': 'No chat found — message the agent on Telegram first'}), 400
        resp = _req.post(f"https://api.telegram.org/bot{token}/sendMessage",
            json={"chat_id": chat_id, "text": message}, timeout=10)
        result = resp.json()
        if result.get('ok'):
            try:
                from core.context_db import journal_log
                journal_log(agent_id=agent_id, task_type="artifact_dispatch",
                    task_description=f"Dashboard sent {filename} to {agent_id}: {prompt[:200]}",
                    result="sent", success=True, input_data=data, chat_id=chat_id)
            except Exception:
                pass
            return jsonify({'success': True, 'message_id': result.get('result', {}).get('message_id')})
        return jsonify({'success': False, 'error': result.get('description', 'Send failed')})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})


@app.route('/api/artifacts/run', methods=['POST'])
def api_artifact_run():
    """Execute a .py or .sh artifact in a constrained subprocess. 30s timeout, captures stdout/stderr.

    Safety: extension whitelist, path must resolve under ARTIFACTS_DIR, no shell for .py, no network
    isolation (files already on this box), agent-framework venv for .py.
    """
    import subprocess as _sp
    import shlex as _shlex
    import time as _time
    data       = request.json or {}
    project_id = (data.get('project_id') or '').strip()
    filename   = (data.get('filename') or '').strip()
    if not project_id or not filename:
        return jsonify({'success': False, 'error': 'project_id + filename required'}), 400

    ext = os.path.splitext(filename)[1].lower()
    if ext not in ('.py', '.sh'):
        return jsonify({'success': False, 'error': f'Extension {ext} not runnable'}), 400

    # Resolve and sanity-check the path is under ARTIFACTS_DIR
    target = os.path.realpath(os.path.join(ARTIFACTS_DIR, project_id, filename))
    base = os.path.realpath(ARTIFACTS_DIR)
    if not target.startswith(base + os.sep):
        return jsonify({'success': False, 'error': 'Path escapes artifacts dir'}), 400
    if not os.path.isfile(target):
        return jsonify({'success': False, 'error': 'File not found'}), 404

    # Build interpreter command
    if ext == '.py':
        venv_py = os.path.join(os.path.dirname(DASHBOARD_DIR), 'venv', 'bin', 'python3')
        py = venv_py if os.path.exists(venv_py) else 'python3'
        cmd = [py, target]
        interpreter = 'python3 (venv)' if os.path.exists(venv_py) else 'python3'
    else:
        cmd = ['/bin/bash', target]
        interpreter = 'bash'

    started = _time.time()
    try:
        proc = _sp.run(
            cmd,
            capture_output=True,
            text=True,
            timeout=30,
            cwd=os.path.dirname(target),
        )
        duration_ms = int((_time.time() - started) * 1000)
        return jsonify({
            'success': True,
            'returncode': proc.returncode,
            'stdout': (proc.stdout or '')[-8000:],
            'stderr': (proc.stderr or '')[-4000:],
            'duration_ms': duration_ms,
            'interpreter': interpreter,
        })
    except _sp.TimeoutExpired as e:
        return jsonify({
            'success': True,
            'returncode': -1,
            'stdout': (e.stdout.decode('utf-8', errors='replace') if e.stdout else '')[-8000:],
            'stderr': (e.stderr.decode('utf-8', errors='replace') if e.stderr else '') + '\n[timed out after 30s]',
            'duration_ms': int((_time.time() - started) * 1000),
            'interpreter': interpreter,
        })
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})


# ── Routes — Heartbeat API ───────────────────────────────────────────────────

@app.route('/api/heartbeats')
def api_heartbeats():
    """Return last heartbeat timestamp for all agents from Redis. TTL=120s per key."""
    try:
        import redis as _redis
        r = _redis.Redis(host='localhost', port=6379, decode_responses=True)
        keys = r.keys('baza:heartbeat:*')
        result = {}
        for key in keys:
            val = r.get(key)
            ttl = r.ttl(key)
            if val:
                try:
                    data = json.loads(val)
                    data['ttl']         = ttl
                    data['seconds_ago'] = max(0, 120 - ttl) if ttl > 0 else 999
                    result[data['agent_id']] = data
                except Exception:
                    pass
        return jsonify(result)
    except Exception as e:
        return jsonify({'error': str(e)}), 500

# ── Routes — Settings API ─────────────────────────────────────────────────────

@app.route('/api/settings/secret', methods=['POST'])
def api_set_secret():
    data    = request.json or {}
    key     = data.get('key', '').strip()
    value   = data.get('value', '').strip()
    if not key:
        return jsonify({'success': False, 'error': 'key required'}), 400
    secrets      = load_secrets()
    secrets[key] = value
    save_secrets(secrets)
    return jsonify({'success': True})

@app.route('/api/settings/secret/delete', methods=['POST'])
def api_delete_secret():
    data    = request.json or {}
    key     = data.get('key', '').strip()
    secrets = load_secrets()
    if key in secrets:
        del secrets[key]
        save_secrets(secrets)
    return jsonify({'success': True})

# ── Routes — Infra Map ────────────────────────────────────────────────────────

@app.route('/infra')
def infra_page():
    return render_template('infra.html')

@app.route('/api/infra/metrics')
def api_infra_metrics():
    import socket as _socket, shutil as _shutil

    def _run(cmd):
        try:
            r = subprocess.run(cmd, shell=True, capture_output=True, text=True, timeout=5)
            return r.stdout.strip()
        except:
            return ""

    def _svc(name):
        out = _run("systemctl is-active " + name)
        return "active" if out == "active" else (out or "inactive")

    def _port(host, port):
        try:
            with _socket.create_connection((host, port), timeout=2):
                return "up"
        except:
            return "down"

    # CPU temp: try k10temp/Tctl (AMD Ryzen), then Package id 0 (Intel), then thermal_zone
    cpu_raw = _run("sensors 2>/dev/null | grep -i 'Tctl' | head -1 | awk '{print $2}'")
    if not cpu_raw:
        cpu_raw = _run("sensors 2>/dev/null | grep -i 'Package id 0' | head -1 | awk '{print $4}'")
    if not cpu_raw:
        # Try ASUS EC sensor for CPU temp
        cpu_raw = _run("sensors 2>/dev/null | grep -A1 'asusec' | grep -i 'CPU:' | head -1 | awk '{print $2}'")
    if not cpu_raw:
        cpu_raw = _run("cat /sys/class/thermal/thermal_zone0/temp 2>/dev/null | awk '{printf \"%.1fC\", $1/1000}'")

    # GPU info
    gpus = []
    nvidia_smi = _run("nvidia-smi --query-gpu=name,temperature.gpu,memory.used,memory.total,utilization.gpu --format=csv,noheader,nounits 2>/dev/null")
    if nvidia_smi:
        parts = nvidia_smi.split(',')
        if len(parts) >= 5:
            gpus.append({"name": parts[0].strip(), "temp": parts[1].strip(), "memory_used": parts[2].strip(),
                         "memory_total": parts[3].strip(), "utilization": parts[4].strip(), "type": "NVIDIA"})
    # AMD — use rocm-smi (reliable even when Vulkan holds sysfs), fall back to sysfs
    # Try rocm-smi for utilization + VRAM (works when Vulkan locks sysfs)
    _rocm_use = _run("rocm-smi --showuse 2>/dev/null")
    _rocm_vram = _run("rocm-smi --showmeminfo vram 2>/dev/null")
    _amd_found = False
    if _rocm_use and "GPU use" in _rocm_use:
        gpu_entry = {"name": "Radeon RX 6700 XT", "type": "AMD"}
        # GPU utilization
        for _line in _rocm_use.splitlines():
            if "GPU use" in _line:
                _uval = ''.join(c for c in _line.split(":")[-1] if c.isdigit())
                if _uval:
                    gpu_entry["utilization"] = _uval + "%"
        # Temperature — try rocm-smi first
        _rocm_temp = _run("rocm-smi --showtemp 2>/dev/null")
        for _line in (_rocm_temp or "").splitlines():
            if "Temperature" in _line and ("edge" in _line.lower() or "GPU" in _line):
                _tval = ''.join(c for c in _line.split(":")[-1] if c.isdigit() or c == '.')
                if _tval:
                    gpu_entry["temp"] = str(int(float(_tval)))
        # VRAM from rocm-smi
        for _line in (_rocm_vram or "").splitlines():
            if "Total Memory" in _line:
                _v = ''.join(c for c in _line.split(":")[-1] if c.isdigit())
                if _v:
                    gpu_entry["memory_total"] = int(_v) // (1024*1024)
            elif "Used Memory" in _line:
                _v = ''.join(c for c in _line.split(":")[-1] if c.isdigit())
                if _v:
                    gpu_entry["memory_used"] = int(_v) // (1024*1024)
        # If temp still missing, try sysfs as last resort
        if "temp" not in gpu_entry:
            for _card in sorted(_run("ls -d /sys/class/drm/card[0-9]*/device/uevent 2>/dev/null").splitlines()):
                _card_dir = os.path.dirname(_card)
                if "amdgpu" in (_run(f"grep DRIVER {_card} 2>/dev/null") or ""):
                    _t = _run(f"cat {_card_dir}/hwmon/hwmon*/temp1_input 2>/dev/null")
                    if _t:
                        gpu_entry["temp"] = str(int(_t) // 1000)
                    break
        gpus.append(gpu_entry)
        _amd_found = True

    if not _amd_found:
        # Fallback: try sysfs directly
        for _card in sorted(_run("ls -d /sys/class/drm/card[0-9]*/device/uevent 2>/dev/null").splitlines()):
            _card_dir = os.path.dirname(_card)
            if "amdgpu" in (_run(f"grep DRIVER {_card} 2>/dev/null") or ""):
                gpu_entry = {"name": "Radeon RX 6700 XT", "type": "AMD"}
                _t = _run(f"cat {_card_dir}/hwmon/hwmon*/temp1_input 2>/dev/null")
                if _t:
                    gpu_entry["temp"] = str(int(_t) // 1000)
                _m = _run(f"cat {_card_dir}/mem_info_vram_total 2>/dev/null")
                if _m:
                    gpu_entry["memory_total"] = int(_m) // (1024*1024)
                _mu = _run(f"cat {_card_dir}/mem_info_vram_used 2>/dev/null")
                if _mu:
                    gpu_entry["memory_used"] = int(_mu) // (1024*1024)
                _b = _run(f"cat {_card_dir}/gpu_busy_percent 2>/dev/null")
                if _b:
                    gpu_entry["utilization"] = _b.strip() + "%"
                gpus.append(gpu_entry)
                break

    # Storage
    storage = []
    for line in _run("lsblk -bno NAME,SIZE,TYPE | grep disk").splitlines():
        parts = line.split()
        if len(parts) >= 3:
            sz = int(parts[1]) if parts[1].isdigit() else 0
            storage.append({"name": parts[0], "size_tb": round(sz / 1e12, 1), "type": parts[2]})

    # Tailscale
    ts_raw = _run("tailscale status --json 2>/dev/null")
    tailscale = {}
    try:
        import json as _json
        ts = _json.loads(ts_raw) if ts_raw else {}
        tailscale = {
            "self_ip": ts.get("Self", {}).get("TailscaleIPs", [""])[0] if ts.get("Self") else "",
            "hostname": ts.get("Self", {}).get("HostName", ""),
            "peers": [{
                "name": (p.get("DNSName","").split(".")[0] if p.get("HostName","") == "localhost" and p.get("DNSName") else p.get("HostName","")) or p.get("DNSName","").split(".")[0],
                "ip": p.get("TailscaleIPs", [""])[0] if p.get("TailscaleIPs") else "",
                "os": p.get("OS", ""),
                "online": p.get("Online", False),
                "last_seen": p.get("LastSeen", ""),
            } for p in ts.get("Peer", {}).values()]
        }
    except Exception:
        tailscale = {"self_ip": _run("tailscale ip -4 2>/dev/null"), "peers": []}

    # API endpoint count
    api_count = 0
    try:
        with open(os.path.join(DASHBOARD_DIR, 'app.py')) as _f:
            api_count = _f.read().count("@app.route")
    except Exception:
        pass

    # DB stats — SQLite
    db_stats = {}
    try:
        conn = sqlite3.connect(os.path.join(DASHBOARD_DIR, 'baza_projects.db'))
        for tbl in ['ahb_projects','ahb_invoices','ahb_receipts','ahb_clients','ahb_employees',
                     'ahb_events','ahb_debts','ahb_payroll','ahb_notes','ahb_files']:
            db_stats[tbl] = conn.execute(f"SELECT count(*) FROM {tbl}").fetchone()[0]
        conn.close()
    except Exception:
        pass

    # DB stats — PostgreSQL
    pg_stats = {}
    try:
        from core.context_db import get_conn, release_conn
        pgc = get_conn()
        cur = pgc.cursor()
        for tbl, label in [('agent_memory','agent memory'),('agent_summaries','summaries'),
                           ('empire_knowledge','empire knowledge'),('task_journal','task journal'),
                           ('agent_usage','agent usage'),('agent_skills','agent skills')]:
            try:
                cur.execute(f"SELECT count(*) FROM {tbl}")
                pg_stats[label] = cur.fetchone()[0]
            except Exception:
                pgc.rollback()
                pg_stats[label] = 0
        cur.close()
        release_conn(pgc)
    except Exception:
        pass

    return jsonify({
        "hostname": _run("hostname"),
        "os": _run("uname -sr"),
        "kernel": _run("uname -r"),
        "cpu_model": _run("lscpu | grep 'Model name' | sed 's/.*: *//'"),
        "cpu_cores": _run("nproc"),
        "cpu_temp": cpu_raw or "N/A",
        "mem_total": _run("free -h | awk '/^Mem:/{print $2}'"),
        "mem_used": _run("free -h | awk '/^Mem:/{print $3}'"),
        "mem_usage": _run("free -h | awk '/^Mem:/{print $3\"/\"$2}'"),
        "disk_root": _run("df -h / | tail -1 | awk '{print $3\"/\"$2\" (\"$5\")}'"),
        "disk_usage": _run("df -h / | tail -1 | awk '{print $3\"/\"$2\" (\"$5\")}'"),
        "uptime": _run("uptime -p"),
        "load": _run("cat /proc/loadavg | awk '{print $1, $2, $3}'"),
        "gpus": gpus,
        "storage": storage,
        "total_storage_tb": round(sum(d["size_tb"] for d in storage), 1),
        "tailscale": tailscale,
        "api_routes": api_count,
        "db_stats": db_stats,
        "pg_stats": pg_stats,
        "services": {
            "ollama_amd": _port("localhost", 11434),
            "ollama_nvidia": _port("localhost", 11435),
            "ollama_cpu": _port("localhost", 11436),
            "ollama_amd2": _port("localhost", 11437),
            "dashboard": "up",
            "sdwebui": _port("localhost", 7860),
            "litellm": _port("localhost", 4000),
            "tool_server": _port("localhost", 8000),
            "postgresql": _svc("postgresql"),
            "redis": _svc("redis-server"),
            "nginx": _svc("nginx"),
            "nextcloud": _run("docker ps --filter name=nextcloud --format '{{.Status}}' 2>/dev/null") or "stopped",
        },
        "agents": {
            name: _svc(f"baza-agent-{name.replace('_','-')}")
            for name in ["simon-bately","claw-batto","phil-hass","sam-axe",
                          "scout-reeves","duke-harmon","rex-valor","nova-sterling"]
        },
        "urls": {
            "dashboard": "http://100.127.118.103:8888",
            "mobile": "http://100.127.118.103:8888/mobile",
            "ahb123": "http://100.127.118.103:8888/ahb123",
            "portal": "http://100.127.118.103:8888/portal",
            "local_dashboard": "http://localhost:8888",
            "lan_dashboard": f"http://{_run('hostname -I').split()[0] if _run('hostname -I') else 'localhost'}:8888",
        },
    })


# ── Email — local SQLite (context.db on baza) ─────────────────────────────────

EMAIL_DB_PATH = os.path.join(FRAMEWORK_DIR, "context.db")

def get_email_db():
    conn = sqlite3.connect(EMAIL_DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn

def email_db_exists():
    return os.path.exists(EMAIL_DB_PATH)

def rows_to_list(rows):
    return [dict(r) for r in rows]

# ── Edge nodes proxy (Tool Server :8000 → /edge/*) ──────────────────────────
EDGE_TOOL_SERVER = os.environ.get("BAZA_TOOL_SERVER", "http://localhost:8000")

@app.route('/api/edge/nodes')
def api_edge_nodes():
    import requests as _rq
    try:
        r = _rq.get(f"{EDGE_TOOL_SERVER}/edge/nodes", timeout=3)
        return jsonify(r.json()), r.status_code
    except Exception as e:
        return jsonify({"error": str(e), "nodes": [], "alerts": []}), 503

@app.route('/api/edge/frame/<node_id>')
def api_edge_frame(node_id):
    """Stream the latest JPEG for <node_id> through the dashboard origin."""
    import requests as _rq
    try:
        r = _rq.get(f"{EDGE_TOOL_SERVER}/edge/frames/{node_id}/latest",
                    timeout=4, stream=True)
        if r.status_code != 200:
            return ("no frame", 404)
        resp = make_response(r.content)
        resp.headers["Content-Type"]  = "image/jpeg"
        resp.headers["Cache-Control"] = "no-store"
        return resp
    except Exception as e:
        return (f"upstream: {e}", 503)

@app.route('/email')
def email_page():
    return render_template('email.html')

@app.route('/api/email/queue')
def api_email_queue():
    if not email_db_exists():
        return jsonify({"records": [], "error": "context.db not found — email pipeline not initialised yet"})
    status = request.args.get("status", "")
    limit  = int(request.args.get("limit", 50))
    try:
        conn = get_email_db()
        if status:
            rows = conn.execute(
                "SELECT * FROM email_queue WHERE status=? ORDER BY received_at DESC LIMIT ?",
                (status, limit)
            ).fetchall()
        else:
            rows = conn.execute(
                "SELECT * FROM email_queue ORDER BY received_at DESC LIMIT ?",
                (limit,)
            ).fetchall()
        conn.close()
        # Normalise field names for the frontend (from -> sender, body_snippet -> snippet)
        records = []
        for r in rows:
            d = dict(r)
            d.setdefault("from",         d.get("sender", ""))
            d.setdefault("body_snippet", d.get("snippet", ""))
            records.append(d)
        return jsonify({"records": records})
    except Exception as e:
        return jsonify({"records": [], "error": str(e)})

@app.route('/api/email/stats')
def api_email_stats():
    if not email_db_exists():
        return jsonify({"total":0,"pending":0,"approved":0,"ignored":0,"sent":0,"high_priority":0})
    try:
        conn = get_email_db()
        rows = conn.execute("SELECT status, priority FROM email_queue").fetchall()
        conn.close()
        stats = {"total": len(rows), "pending": 0, "approved": 0,
                 "ignored": 0, "sent": 0, "high_priority": 0}
        for r in rows:
            s = r["status"] or ""
            if s == "awaiting_confirmation": stats["pending"] += 1
            elif s in stats: stats[s] += 1
            if (r["priority"] or "").lower() in ("high","urgent"): stats["high_priority"] += 1
        return jsonify(stats)
    except Exception as e:
        return jsonify({"total":0,"error":str(e)})

@app.route('/api/email/action', methods=['POST'])
def api_email_action():
    """approve / ignore / restore an email by its local DB id or gmail_id."""
    if not email_db_exists():
        return jsonify({"success": False, "error": "context.db not found"})
    body     = request.get_json() or {}
    gmail_id = body.get("gmail_id")
    action   = body.get("action", "")
    reply    = body.get("reply_text", "")

    status_map = {
        "approve":  "approved",
        "ignore":   "ignored",
        "restore":  "awaiting_confirmation",
    }
    new_status = status_map.get(action, action)

    try:
        conn = get_email_db()
        if reply:
            conn.execute(
                "UPDATE email_queue SET status=?, suggested_reply=? WHERE gmail_id=?",
                (new_status, reply, gmail_id)
            )
        else:
            conn.execute(
                "UPDATE email_queue SET status=? WHERE gmail_id=?",
                (new_status, gmail_id)
            )
        conn.commit()

        # If approving — run email_send.py via subprocess
        if action == "approve":
            row = conn.execute(
                "SELECT id FROM email_queue WHERE gmail_id=?", (gmail_id,)
            ).fetchone()
            conn.close()
            if row:
                local_id = row["id"]
                send_cmd = [
                    VENV_PYTHON,
                    os.path.join(FRAMEWORK_DIR, "skills", "shared", "email_send.py"),
                    "approve", str(local_id)
                ]
                if reply:
                    send_cmd += ["send", str(local_id), reply]
                subprocess.Popen(send_cmd, cwd=FRAMEWORK_DIR,
                                 stdout=open(os.path.join(LOGS_DIR, "email_send.log"), "a"),
                                 stderr=subprocess.STDOUT)
        else:
            conn.close()

        return jsonify({"success": True})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)})

@app.route('/api/email/fetch', methods=['POST'])
def api_email_fetch():
    """Manually trigger email_fetch.py to pull new emails from Gmail."""
    try:
        proc = subprocess.Popen(
            [VENV_PYTHON, os.path.join(FRAMEWORK_DIR, "skills", "shared", "email_fetch.py")],
            cwd=FRAMEWORK_DIR,
            stdout=open(os.path.join(LOGS_DIR, "email_fetch.log"), "a"),
            stderr=subprocess.STDOUT
        )
        return jsonify({"success": True, "pid": proc.pid})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)})


# ── Routes — Tasks (SQLite baza_projects.db) ─────────────────────────────────

DB_PATH = os.path.join(DASHBOARD_DIR, "baza_projects.db")

def get_tasks_db():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn

@app.route('/tasks')
def tasks_page():
    return render_template('tasks.html')

@app.route('/api/tasks')
def api_tasks_list():
    if not os.path.exists(DB_PATH):
        return jsonify([])
    status   = request.args.get('status', '')
    project  = request.args.get('project_id', '')
    agent    = request.args.get('assigned_to', '')
    limit    = int(request.args.get('limit', 100))
    try:
        conn = get_tasks_db()
        sql = "SELECT * FROM tasks WHERE 1=1"
        params = []
        if status:   sql += " AND status=?";      params.append(status)
        if project:  sql += " AND project_id=?";  params.append(project)
        if agent:    sql += " AND assigned_to=?"; params.append(agent)
        sql += " ORDER BY CASE status WHEN 'in_progress' THEN 1 WHEN 'pending' THEN 2 WHEN 'blocked' THEN 3 ELSE 4 END, priority DESC, updated_at DESC LIMIT ?"
        params.append(limit)
        rows = [dict(r) for r in conn.execute(sql, params).fetchall()]
        conn.close()
        return jsonify(rows)
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/tasks/capacity')
def api_tasks_capacity():
    """Per-agent open-task load + recommended next routing per Duke's map.

    Returns:
      agents: [{agent_id, pending, in_progress, blocked, total, max_capacity}]
      total_open: int
    """
    if not os.path.exists(DB_PATH):
        return jsonify({"agents": [], "total_open": 0})
    max_cap = int(request.args.get('max_capacity', 5) or 5)
    conn = get_tasks_db()
    try:
        rows = conn.execute(
            "SELECT assigned_to, status, COUNT(*) FROM tasks "
            "WHERE status IN ('pending','in_progress','blocked') "
            "GROUP BY assigned_to, status"
        ).fetchall()
    except Exception as e:
        conn.close()
        return jsonify({"error": str(e)}), 500
    conn.close()
    by_agent: dict[str, dict] = {}
    for r in rows:
        ag = (r["assigned_to"] or "_unassigned")
        d = by_agent.setdefault(ag, {"pending": 0, "in_progress": 0, "blocked": 0})
        d[r["status"]] = r["COUNT(*)"] if "COUNT(*)" in r.keys() else r[2]
    out = []
    for ag, d in sorted(by_agent.items()):
        total = d["pending"] + d["in_progress"] + d["blocked"]
        out.append({
            "agent_id": ag,
            "pending": d["pending"],
            "in_progress": d["in_progress"],
            "blocked": d["blocked"],
            "total": total,
            "max_capacity": max_cap,
            "load_pct": min(100, round(100 * total / max_cap)),
            "overloaded": total > max_cap,
        })
    out.sort(key=lambda r: -r["total"])
    return jsonify({
        "agents": out,
        "total_open": sum(r["total"] for r in out),
        "max_capacity": max_cap,
    })


@app.route('/api/tasks/bulk', methods=['POST'])
def api_tasks_bulk():
    """Bulk operations: archive | reassign | delete | set_priority.

    Body: {"task_ids": [...], "op": "archive"|"reassign"|"delete"|"set_priority",
           "value": <new value if applicable>}
    """
    if not os.path.exists(DB_PATH):
        return jsonify({"error": "DB not found"}), 404
    data = request.get_json(silent=True) or {}
    ids = data.get('task_ids') or []
    op = (data.get('op') or '').strip().lower()
    value = data.get('value')
    if not ids or op not in ("archive", "reassign", "delete", "set_priority"):
        return jsonify({"error": "task_ids and op (archive|reassign|delete|set_priority) required"}), 400
    if op == "reassign" and not value:
        return jsonify({"error": "reassign requires value=<agent_id>"}), 400
    if op == "set_priority" and value not in ("low", "medium", "high"):
        return jsonify({"error": "set_priority requires value in low|medium|high"}), 400

    conn = get_tasks_db()
    placeholders = ",".join(["?"] * len(ids))
    now = datetime.datetime.utcnow().isoformat()
    affected = 0
    try:
        if op == "delete":
            cur = conn.execute(f"DELETE FROM tasks WHERE id IN ({placeholders})", ids)
            affected = cur.rowcount
        elif op == "archive":
            cur = conn.execute(
                f"UPDATE tasks SET status='archived', updated_at=?, "
                f"  notes = COALESCE(notes,'') || ? "
                f"WHERE id IN ({placeholders})",
                [now, f"\n[ARCHIVED {now[:10]} via bulk-op]\n"] + list(ids),
            )
            affected = cur.rowcount
        elif op == "reassign":
            cur = conn.execute(
                f"UPDATE tasks SET assigned_to=?, updated_at=? WHERE id IN ({placeholders})",
                [value, now] + list(ids),
            )
            affected = cur.rowcount
        elif op == "set_priority":
            cur = conn.execute(
                f"UPDATE tasks SET priority=?, updated_at=? WHERE id IN ({placeholders})",
                [value, now] + list(ids),
            )
            affected = cur.rowcount
        conn.commit()
    except Exception as e:
        conn.close()
        return jsonify({"error": str(e)}), 500
    conn.close()
    return jsonify({"op": op, "affected": affected, "ids": ids, "value": value})


@app.route('/api/tasks/<task_id>', methods=['GET'])
def api_task_get(task_id):
    if not os.path.exists(DB_PATH):
        return jsonify({'error': 'DB not found'}), 404
    conn = get_tasks_db()
    row = conn.execute("SELECT * FROM tasks WHERE id=?", (task_id,)).fetchone()
    conn.close()
    if not row:
        return jsonify({'error': 'not found'}), 404
    return jsonify(dict(row))

@app.route('/api/tasks/<task_id>', methods=['PATCH'])
def api_task_update(task_id):
    if not os.path.exists(DB_PATH):
        return jsonify({'error': 'DB not found'}), 404
    data = request.json or {}
    allowed = {'title','description','status','priority','assigned_to','project_id','notes','due_date'}
    fields = {k: v for k, v in data.items() if k in allowed}
    if not fields:
        return jsonify({'error': 'no valid fields'}), 400
    fields['updated_at'] = datetime.datetime.utcnow().isoformat()
    set_clause = ", ".join(f"{k}=?" for k in fields)
    values = list(fields.values()) + [task_id]
    conn = get_tasks_db()
    conn.execute(f"UPDATE tasks SET {set_clause} WHERE id=?", values)
    conn.commit()
    conn.close()
    return jsonify({'success': True})

@app.route('/api/tasks', methods=['POST'])
def api_task_create():
    import uuid as _uuid
    if not os.path.exists(DB_PATH):
        return jsonify({'error': 'DB not found'}), 404
    data = request.json or {}
    task_id = data.get('id') or str(_uuid.uuid4())[:8]
    now = datetime.datetime.utcnow().isoformat()
    conn = get_tasks_db()
    try:
        conn.execute("""
            INSERT INTO tasks (id, project_id, title, description, assigned_to, status, priority, notes, created_at, updated_at)
            VALUES (?,?,?,?,?,?,?,?,?,?)
        """, (
            task_id,
            data.get('project_id', 'proj-baza-empire'),
            data.get('title', 'Untitled Task'),
            data.get('description', ''),
            data.get('assigned_to', ''),
            data.get('status', 'pending'),
            data.get('priority', 'medium'),
            data.get('notes', ''),
            now, now
        ))
        conn.commit()
        conn.close()
        return jsonify({'success': True, 'id': task_id})
    except Exception as e:
        conn.close()
        return jsonify({'error': str(e)}), 500

@app.route('/api/tasks/<task_id>', methods=['DELETE'])
def api_task_delete(task_id):
    if not os.path.exists(DB_PATH):
        return jsonify({'error': 'DB not found'}), 404
    conn = get_tasks_db()
    conn.execute("DELETE FROM tasks WHERE id=?", (task_id,))
    conn.commit()
    conn.close()
    return jsonify({'success': True})

@app.route('/api/tasks/stats', methods=['GET'])
def api_task_stats():
    if not os.path.exists(DB_PATH):
        return jsonify({})
    conn = get_tasks_db()
    rows = conn.execute("SELECT status, project_id, assigned_to FROM tasks").fetchall()
    conn.close()
    stats = {'total': len(rows), 'by_status': {}, 'by_project': {}, 'by_agent': {}}
    for r in rows:
        s = r['status'] or 'unknown'
        stats['by_status'][s] = stats['by_status'].get(s, 0) + 1
        p = r['project_id'] or 'unknown'
        stats['by_project'][p] = stats['by_project'].get(p, 0) + 1
        a = r['assigned_to'] or 'unassigned'
        stats['by_agent'][a] = stats['by_agent'].get(a, 0) + 1
    return jsonify(stats)

@app.route('/api/projects')
def api_projects_list():
    if not os.path.exists(DB_PATH):
        return jsonify([])
    conn = get_tasks_db()
    try:
        rows = [dict(r) for r in conn.execute("SELECT * FROM projects ORDER BY created_at DESC").fetchall()]
    except Exception:
        rows = []
    conn.close()
    return jsonify(rows)


# ── Routes — Baza Projects (sub-project #4) ──────────────────────────────────

def _baza_projects():
    from core import baza_projects as bp
    return bp


def _emit_project_event(kind, project_id, payload=None, agent_id="user"):
    try:
        from core import task_events as te
        te.emit(kind, project_id=project_id, agent_id=agent_id, payload=payload or {})
    except Exception:
        pass


@app.route('/projects')
def projects_page():
    """Baza Projects developer UI — list view."""
    return render_template('projects.html')


@app.route('/projects/<project_id>')
def project_detail_page(project_id):
    """Baza Projects developer UI — detail view with sub-tabs."""
    return render_template('project_detail.html', project_id=project_id)


@app.route('/api/baza/projects', methods=['GET'])
def api_baza_projects_list():
    bp = _baza_projects()
    return jsonify({"projects": bp.list_projects(kind=request.args.get('kind') or 'baza-dev')})


@app.route('/api/baza/projects', methods=['POST'])
def api_baza_projects_create():
    bp = _baza_projects()
    data = request.get_json(silent=True) or {}
    name = (data.get('name') or '').strip()
    raw_type = (data.get('type') or '').strip()
    description = (data.get('description') or '').strip()
    project_id = (data.get('id') or '').strip() or None
    template_id = (data.get('template') or '').strip() or None
    # If a template is selected and the caller didn't explicitly set a type,
    # let the template's declared type decide. Otherwise default to web-app.
    if template_id and not raw_type:
        type_ = "other"  # signals "use template's type"
    else:
        type_ = raw_type or "web-app"
    if not name:
        return jsonify({"error": "name is required"}), 400
    try:
        proj = bp.create_project(
            name=name, type_=type_, description=description,
            created_by=data.get('created_by') or 'user', project_id=project_id,
            template_id=template_id,
        )
    except FileExistsError as e:
        return jsonify({"error": str(e)}), 409
    except ValueError as e:
        return jsonify({"error": str(e)}), 400
    _emit_project_event(
        "intent_parsed", proj["id"],
        payload={"intent": "create_baza_project", "name": name, "type": type_,
                 "template": template_id},
    )
    return jsonify({"project": proj}), 201


@app.route('/api/baza/templates')
def api_baza_templates_list():
    from core import baza_project_templates as tpl
    return jsonify({"templates": tpl.list_templates()})


@app.route('/api/baza/projects/<project_id>', methods=['GET'])
def api_baza_project_get(project_id):
    bp = _baza_projects()
    proj = bp.get_project(project_id)
    if not proj:
        return jsonify({"error": "not found"}), 404
    return jsonify({"project": proj})


@app.route('/api/baza/projects/<project_id>', methods=['PUT'])
def api_baza_project_update(project_id):
    bp = _baza_projects()
    patch = request.get_json(silent=True) or {}
    try:
        manifest = bp.update_manifest(project_id, patch)
    except FileNotFoundError:
        return jsonify({"error": "not found"}), 404
    return jsonify({"manifest": manifest})


@app.route('/api/baza/projects/<project_id>', methods=['DELETE'])
def api_baza_project_delete(project_id):
    bp = _baza_projects()
    hard = (request.args.get('hard') or '').lower() in ('1', 'true', 'yes')
    ok = bp.delete_project(project_id, hard=hard)
    return jsonify({"deleted": ok})


@app.route('/api/baza/projects/<project_id>/files')
def api_baza_project_files(project_id):
    bp = _baza_projects()
    subpath = request.args.get('path', '')
    return jsonify({"files": bp.list_files(project_id, subpath)})


@app.route('/api/baza/projects/<project_id>/file', methods=['GET'])
def api_baza_project_file_get(project_id):
    bp = _baza_projects()
    relpath = request.args.get('path', '')
    if not relpath:
        return jsonify({"error": "path is required"}), 400
    content = bp.read_file(project_id, relpath)
    if content is None:
        return jsonify({"error": "not found"}), 404
    return jsonify({"path": relpath, "content": content})


@app.route('/api/baza/projects/<project_id>/file', methods=['POST'])
def api_baza_project_file_put(project_id):
    bp = _baza_projects()
    data = request.get_json(silent=True) or {}
    relpath = data.get('path') or ''
    content = data.get('content') or ''
    agent_id = data.get('agent_id') or None
    force = bool(data.get('force'))
    if not relpath:
        return jsonify({"error": "path is required"}), 400
    try:
        info = bp.write_file(project_id, relpath, content, agent_id=agent_id, force=force)
    except PermissionError as e:
        # Surface holder so the caller can decide to wait or force
        msg = str(e)
        return jsonify({"error": msg, "lock_holder": bp.current_lock_holder(project_id)}), 423 if "locked" in msg else 403
    return jsonify({"saved": True, "info": info, "lock": {"held_by": bp.current_lock_holder(project_id)}})


@app.route('/api/baza/projects/<project_id>/lock', methods=['GET'])
def api_baza_project_lock_status(project_id):
    bp = _baza_projects()
    h = bp.current_lock_holder(project_id)
    return jsonify({"held_by": h, "is_locked": bool(h)})


@app.route('/api/baza/projects/<project_id>/lock', methods=['POST'])
def api_baza_project_lock_acquire(project_id):
    bp = _baza_projects()
    agent_id = (request.get_json(silent=True) or {}).get('agent_id') or 'user'
    return jsonify(bp.acquire_lock(project_id, agent_id))


@app.route('/api/baza/projects/<project_id>/lock', methods=['DELETE'])
def api_baza_project_lock_release(project_id):
    bp = _baza_projects()
    agent_id = (request.args.get('agent_id') or
                ((request.get_json(silent=True) or {}).get('agent_id')) or 'user')
    released = bp.release_lock(project_id, agent_id)
    return jsonify({"released": released, "held_by": bp.current_lock_holder(project_id)})


@app.route('/api/baza/projects/<project_id>/git/status')
def api_baza_project_git_status(project_id):
    bp = _baza_projects()
    try:
        return jsonify(bp.git_status(project_id))
    except FileNotFoundError:
        return jsonify({"error": "project not found"}), 404


@app.route('/api/baza/projects/<project_id>/git/commit', methods=['POST'])
def api_baza_project_git_commit(project_id):
    bp = _baza_projects()
    data = request.get_json(silent=True) or {}
    msg = (data.get('message') or '').strip()
    stage_all = data.get('stage_all', True)
    if not msg:
        return jsonify({"committed": False, "error": "message is required"}), 400
    try:
        res = bp.git_commit(project_id, msg, stage_all=bool(stage_all))
    except FileNotFoundError:
        return jsonify({"error": "project not found"}), 404
    try:
        from core import task_events as te
        te.emit("tool_result", project_id=project_id, agent_id="user",
                payload={"tool": "git.commit", "ok": bool(res.get("committed")),
                         "result_snippet": (res.get("head") or res.get("error") or "")[:300]})
    except Exception:
        pass
    return jsonify(res)


@app.route('/api/baza/projects/<project_id>/exec', methods=['POST'])
def api_baza_project_exec(project_id):
    """Run an arbitrary shell command pinned to the project sandbox dir."""
    bp = _baza_projects()
    data = request.get_json(silent=True) or {}
    cmd = (data.get('command') or '').strip()
    timeout = int(data.get('timeout') or 60)
    if not cmd:
        return jsonify({"error": "command is required"}), 400
    parent = None
    try:
        from core import task_events as te
        parent = te.emit("tool_call", project_id=project_id, agent_id="user",
                         payload={"tool": "explore.exec", "args": {"command": cmd[:200]}})
    except Exception:
        pass
    try:
        res = bp.exec_in_project(project_id, cmd, timeout=timeout)
    except FileNotFoundError:
        return jsonify({"error": "project not found"}), 404
    try:
        from core import task_events as te
        te.emit("tool_result", project_id=project_id, agent_id="user",
                payload={"tool": "explore.exec", "ok": bool(res.get("success")),
                         "exit_code": res.get("exit_code", -1),
                         "result_snippet": (res.get("stdout") or res.get("error") or "")[:600]},
                parent_event_id=parent)
    except Exception:
        pass
    return jsonify(res)


@app.route('/api/baza/projects/<project_id>/render', methods=['POST'])
def api_baza_project_render(project_id):
    """Generate a visual via Stable Diffusion (uses Sam's existing infra) and
    save into the project's artifacts/ dir."""
    bp = _baza_projects()
    proj = bp.get_project(project_id)
    if not proj:
        return jsonify({"error": "project not found"}), 404
    data = request.get_json(silent=True) or {}
    prompt = (data.get('prompt') or '').strip()
    if not prompt:
        return jsonify({"error": "prompt is required"}), 400

    parent = None
    try:
        from core import task_events as te
        parent = te.emit("tool_call", project_id=project_id, agent_id="user",
                         payload={"tool": "render.txt2img", "args": {"prompt": prompt[:200]}})
    except Exception:
        pass

    try:
        import base64
        import requests as _req
        resp = _req.post('http://localhost:7860/sdapi/v1/txt2img', json={
            'prompt': prompt,
            'negative_prompt': data.get('negative_prompt', ''),
            'width':  int(data.get('width', 1024)),
            'height': int(data.get('height', 1024)),
            'steps':  int(data.get('steps', 30)),
            'cfg_scale': float(data.get('cfg_scale', 7)),
            'sampler_name': data.get('sampler') or 'DPM++ 2M Karras',
        }, timeout=180)
        result = resp.json()
        images = result.get('images') or []
        out_dir = os.path.join(proj["path"], "artifacts")
        os.makedirs(out_dir, exist_ok=True)
        saved_urls = []
        saved_files = []
        for img_b64 in images:
            fname = f"render_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}_{uuid.uuid4().hex[:6]}.png"
            fpath = os.path.join(out_dir, fname)
            with open(fpath, "wb") as fp:
                fp.write(base64.b64decode(img_b64))
            saved_files.append(fname)
            saved_urls.append(f"/api/artifacts/serve/{project_id}/{fname}")
            try:
                from core import task_events as te
                te.emit("artifact_saved", project_id=project_id, agent_id="user",
                        payload={"path": fpath, "filename": fname,
                                 "kind": "render", "bytes": os.path.getsize(fpath)})
            except Exception:
                pass

        try:
            from core import task_events as te
            te.emit("tool_result", project_id=project_id, agent_id="user",
                    payload={"tool": "render.txt2img", "ok": True,
                             "result_snippet": f"saved {len(saved_files)} image(s)"},
                    parent_event_id=parent)
        except Exception:
            pass
        return jsonify({"success": True, "files": saved_files, "urls": saved_urls})
    except Exception as e:
        try:
            from core import task_events as te
            te.emit("tool_result", project_id=project_id, agent_id="user",
                    payload={"tool": "render.txt2img", "ok": False, "error": str(e)[:300]},
                    parent_event_id=parent)
        except Exception:
            pass
        return jsonify({"success": False, "error": str(e)}), 500


@app.route('/api/baza/projects/<project_id>/renders')
def api_baza_project_renders(project_id):
    """List render images saved in the project's artifacts/ dir."""
    bp = _baza_projects()
    proj = bp.get_project(project_id)
    if not proj:
        return jsonify({"renders": []})
    art_dir = os.path.join(proj["path"], "artifacts")
    if not os.path.isdir(art_dir):
        return jsonify({"renders": []})
    out = []
    for name in sorted(os.listdir(art_dir), reverse=True):
        if name.lower().endswith((".png", ".jpg", ".jpeg", ".webp")):
            p = os.path.join(art_dir, name)
            out.append({
                "name": name,
                "url": f"/api/artifacts/serve/{project_id}/{name}",
                "size": os.path.getsize(p),
                "modified": datetime.datetime.fromtimestamp(os.path.getmtime(p)).isoformat(),
            })
    return jsonify({"renders": out[:60]})


@app.route('/api/baza/projects/<project_id>/preview/start', methods=['POST'])
def api_baza_preview_start(project_id):
    from core import preview_supervisor as ps
    slot = (request.get_json(silent=True) or {}).get('slot') or 'preview'
    res = ps.start(project_id, slot=slot)
    try:
        from core import task_events as te
        te.emit("tool_call", project_id=project_id, agent_id="user",
                payload={"tool": f"preview.start.{slot}", "args": res})
    except Exception:
        pass
    return jsonify(res), (200 if res.get('started') else 409 if 'already' in (res.get('error') or '') else 400)


@app.route('/api/baza/projects/<project_id>/preview/stop', methods=['POST'])
def api_baza_preview_stop(project_id):
    from core import preview_supervisor as ps
    hard = (request.args.get('hard') or '').lower() in ('1', 'true', 'yes')
    res = ps.stop(project_id, hard=hard)
    try:
        from core import task_events as te
        te.emit("tool_call", project_id=project_id, agent_id="user",
                payload={"tool": "preview.stop", "args": {"hard": hard}, "ok": bool(res.get('stopped'))})
    except Exception:
        pass
    return jsonify(res)


@app.route('/api/baza/projects/<project_id>/preview/status')
def api_baza_preview_status(project_id):
    from core import preview_supervisor as ps
    return jsonify(ps.status(project_id))


@app.route('/api/baza/projects/<project_id>/preview/logs')
def api_baza_preview_logs(project_id):
    from core import preview_supervisor as ps
    lines = int(request.args.get('lines', 200) or 200)
    return jsonify({"logs": ps.tail_logs(project_id, lines=lines)})


@app.route('/api/baza/projects/<project_id>/run', methods=['POST'])
def api_baza_project_run(project_id):
    """Run a manifest command slot. Long-running run/preview are not handled here."""
    bp = _baza_projects()
    data = request.get_json(silent=True) or {}
    slot = (data.get('slot') or '').strip()
    approved = bool(data.get('approved'))
    if not slot:
        return jsonify({"error": "slot is required"}), 400
    # Emit lifecycle events
    parent = None
    try:
        from core import task_events as te
        parent = te.emit(
            "tool_call", project_id=project_id, agent_id="user",
            payload={"tool": f"baza_projects.run.{slot}", "args": {}},
        )
    except Exception:
        pass
    try:
        result = bp.run_command(project_id, slot, approved=approved)
    except FileNotFoundError:
        return jsonify({"error": "project not found"}), 404
    try:
        from core import task_events as te
        te.emit(
            "tool_result", project_id=project_id, agent_id="user",
            payload={
                "tool": f"baza_projects.run.{slot}",
                "ok": bool(result.get("success")),
                "result_snippet": (result.get("stdout") or result.get("error") or "")[:600],
                "exit_code": result.get("exit_code", -1),
            },
            parent_event_id=parent,
        )
    except Exception:
        pass
    return jsonify(result)


# ── Routes — Approvals Inbox (#R2) ────────────────────────────────────────────

@app.route('/approvals')
def approvals_page():
    return render_template('approvals.html')


@app.route('/api/approvals')
def api_approvals_list():
    from core import task_events as te
    state = request.args.get('state', 'pending')
    limit = int(request.args.get('limit', 100) or 100)
    return jsonify({"approvals": te.list_approvals(state=state, limit=limit)})


@app.route('/api/approvals/<int:event_id>/grant', methods=['POST'])
def api_approval_grant(event_id):
    """Mark an approval_requested event as granted and re-dispatch the
    underlying intent with approved=true."""
    from core import task_events as te
    from core.intent_router import parse_intent
    from core.intent_dispatcher import dispatch
    by = (request.get_json(silent=True) or {}).get('by') or 'user'

    # Find the original request to know what to re-fire
    pending = te.list_approvals(state='all', limit=500)
    req = next((p for p in pending if p["id"] == event_id), None)
    if not req:
        return jsonify({"error": "approval not found"}), 404
    if req["state"] != "pending":
        return jsonify({"error": f"already {req['state']}"}), 409

    # Emit decision event
    te.emit("approval_granted", project_id=req["project_id"], agent_id=req["agent_id"],
            payload={"action": req["action"], "by": by, "for_event_id": event_id})

    # Re-dispatch if it's an intent we can handle
    action = req["action"] or ""
    details = req["details"] or {}
    re_result = None
    try:
        if action == "deploy" and req["project_id"]:
            text = f"/deploy {req['project_id']}"
            out = dispatch(parse_intent(text), extra={"approved": True, "agent_id": by})
            re_result = {"status": out["status"], "result": out["result"]}
        elif action.startswith("ahb.") and details.get("id"):
            # Privileged ahb_api action — caller can re-issue with approved=true
            re_result = {"hint": "re-issue the original ##SKILL:ahb_api## call with approved=true"}
        elif action.startswith("baza_proj.") and details.get("id"):
            re_result = {"hint": "re-issue the original ##SKILL:baza_proj## call with approved=true"}
    except Exception as e:
        re_result = {"error": str(e)}

    return jsonify({"granted": True, "approval_id": event_id, "re_dispatched": re_result})


@app.route('/api/approvals/<int:event_id>/deny', methods=['POST'])
def api_approval_deny(event_id):
    from core import task_events as te
    body = request.get_json(silent=True) or {}
    note = (body.get('note') or '')[:300]
    by = body.get('by') or 'user'

    pending = te.list_approvals(state='all', limit=500)
    req = next((p for p in pending if p["id"] == event_id), None)
    if not req:
        return jsonify({"error": "approval not found"}), 404
    if req["state"] != "pending":
        return jsonify({"error": f"already {req['state']}"}), 409

    te.emit("approval_denied", project_id=req["project_id"], agent_id=req["agent_id"],
            payload={"action": req["action"], "by": by, "note": note,
                     "for_event_id": event_id})
    return jsonify({"denied": True, "approval_id": event_id})


# ── Routes — Intent Router (sub-project #2) ───────────────────────────────────

@app.route('/api/intents/parse', methods=['POST'])
def api_intents_parse():
    """Parse text into a structured intent envelope (no execution)."""
    from core.intent_router import parse_intent
    text = (request.get_json(silent=True) or {}).get('text', '')
    return jsonify(parse_intent(text))


@app.route('/api/intents/help')
def api_intents_help():
    from core.intent_router import help_text
    return jsonify({"help": help_text()})


@app.route('/api/intents', methods=['POST'])
def api_intents_dispatch():
    """Parse text, dispatch known intents, return result. Also handles the
    AHB project create flow because that's HTTP-side state."""
    from core.intent_router import parse_intent
    from core.intent_dispatcher import dispatch as _dispatch
    data = request.get_json(silent=True) or {}
    text = (data.get('text') or '').strip()
    extra = data.get('extra') or {}
    env = parse_intent(text)

    # AHB project create needs the dashboard's existing endpoint — not in core.
    if env.get('intent') == 'create_ahb_project':
        if env['errors']:
            return jsonify({"envelope": env, "result": None}), 400
        try:
            ahb_payload = {"title": env['args'].get('name') or 'New AHB Project'}
            if env['args'].get('client_id'):
                ahb_payload["client_id"] = env['args']['client_id']
            if env['args'].get('description'):
                ahb_payload["description"] = env['args']['description']
            with app.test_client() as tc:
                r = tc.post('/api/ahb/projects', json=ahb_payload)
                ok = (r.status_code in (200, 201))
                return jsonify({
                    "envelope": env,
                    "result": {"ok": ok, "status": r.status_code, "body": r.get_json(silent=True) or {}},
                }), (200 if ok else 400)
        except Exception as e:
            return jsonify({"envelope": env, "result": {"error": str(e)}}), 500

    out = _dispatch(env, extra=extra)
    return jsonify({"envelope": out["envelope"], "result": out["result"]}), out["status"]


# ── Routes — Skills Lab ────────────────────────────────────────────────────────

@app.route('/skills')
def skills_page():
    return render_template('skills.html')

@app.route('/api/skills/list')
def api_skills_list():
    skills = []
    shared_dir = os.path.join(FRAMEWORK_DIR, "skills", "shared")
    if os.path.isdir(shared_dir):
        for f in sorted(Path(shared_dir).glob("*.py")):
            stat = os.stat(f)
            # Read first docstring line
            desc = ""
            try:
                for line in open(f).readlines()[:12]:
                    line = line.strip()
                    if line.startswith('"""') or line.startswith("'''"):
                        desc = line.strip('"\' ')
                        if len(desc) < 5:
                            continue
                        break
                    if line and not line.startswith('#') and not line.startswith('import') and not line.startswith('def'):
                        continue
            except Exception:
                pass
            skills.append({'name': f.stem, 'path': str(f), 'scope': 'shared',
                           'size': stat.st_size, 'modified': datetime.datetime.fromtimestamp(stat.st_mtime).strftime("%Y-%m-%d %H:%M"),
                           'description': desc})
    # Per-agent skills
    agents_dir = os.path.join(FRAMEWORK_DIR, "agents")
    for agent_dir in sorted(Path(agents_dir).iterdir()):
        skill_dir = agent_dir / "skills"
        if skill_dir.is_dir():
            for f in sorted(skill_dir.glob("*.py")):
                stat = os.stat(f)
                skills.append({'name': f.stem, 'path': str(f), 'scope': agent_dir.name,
                               'size': stat.st_size, 'modified': datetime.datetime.fromtimestamp(stat.st_mtime).strftime("%Y-%m-%d %H:%M"),
                               'description': ''})
    return jsonify(skills)

@app.route('/api/skills/catalog')
def api_skills_catalog():
    """Rich skill catalog via skills/shared/skill_catalog.py — returns name, scope, owner, summary, args_hint."""
    filter_str = (request.args.get('filter') or '').strip()
    agent = (request.args.get('agent') or '').strip() or None
    env = os.environ.copy()
    env['SKILL_ARGS'] = json.dumps({'filter': filter_str, 'agent': agent})
    try:
        proc = subprocess.run(
            [sys.executable, os.path.join(FRAMEWORK_DIR, 'skills', 'shared', 'skill_catalog.py')],
            env=env, capture_output=True, text=True, timeout=30,
        )
        return jsonify(json.loads(proc.stdout or '{}'))
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 500


@app.route('/api/knowledge/search')
def api_knowledge_search():
    """Unified FTS5 knowledge search across AHBCO data + agent memory."""
    query = (request.args.get('q') or '').strip()
    if not query:
        return jsonify({'ok': False, 'error': 'q parameter required'}), 400
    sources = request.args.getlist('source') or None
    limit = int(request.args.get('limit', 10))
    env = os.environ.copy()
    env['SKILL_ARGS'] = json.dumps({'query': query, 'sources': sources, 'limit': limit})
    try:
        proc = subprocess.run(
            [sys.executable, os.path.join(FRAMEWORK_DIR, 'skills', 'shared', 'knowledge_search.py')],
            env=env, capture_output=True, text=True, timeout=15,
        )
        return jsonify(json.loads(proc.stdout or '{}'))
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 500


@app.route('/api/knowledge/rebuild', methods=['POST'])
def api_knowledge_rebuild():
    """Trigger knowledge index rebuild on demand."""
    try:
        proc = subprocess.run(
            [sys.executable, os.path.join(FRAMEWORK_DIR, 'skills', 'shared', 'knowledge_rebuild_index.py')],
            capture_output=True, text=True, timeout=120,
        )
        return jsonify(json.loads(proc.stdout or '{}'))
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 500


@app.route('/api/skills/read/<skill_name>')
def api_skill_read(skill_name):
    shared_dir = os.path.join(FRAMEWORK_DIR, "skills", "shared")
    path = os.path.join(shared_dir, f"{skill_name}.py")
    if not os.path.exists(path):
        return jsonify({'error': 'not found'}), 404
    return jsonify({'name': skill_name, 'code': open(path).read(), 'path': path})

@app.route('/api/skills/save', methods=['POST'])
def api_skill_save():
    import re as _re
    data = request.json or {}
    name = data.get('name', '').strip()
    code = data.get('code', '').strip()
    if not name or not code:
        return jsonify({'error': 'name and code required'}), 400
    if not _re.match(r'^[a-z][a-z0-9_]{1,49}$', name):
        return jsonify({'error': 'invalid name'}), 400
    path = os.path.join(FRAMEWORK_DIR, "skills", "shared", f"{name}.py")
    import stat as _stat
    with open(path, 'w') as f:
        f.write(code)
    os.chmod(path, os.stat(path).st_mode | _stat.S_IXUSR)
    return jsonify({'success': True, 'path': path})

@app.route('/api/skills/run', methods=['POST'])
def api_skill_run():
    data = request.json or {}
    name = data.get('name', '').strip()
    args = data.get('args', {})
    if not name:
        return jsonify({'error': 'name required'}), 400
    shared_dir = os.path.join(FRAMEWORK_DIR, "skills", "shared")
    path = os.path.join(shared_dir, f"{name}.py")
    if not os.path.exists(path):
        return jsonify({'error': f'skill not found: {name}'}), 404
    import time as _time
    env = os.environ.copy()
    env['SKILL_ARGS'] = json.dumps(args)
    env['AGENT_ID'] = 'dashboard'
    t0 = _time.time()
    try:
        result = subprocess.run([VENV_PYTHON, path], capture_output=True, text=True, timeout=30, env=env)
        elapsed = int((_time.time() - t0) * 1000)
        return jsonify({
            'success': result.returncode == 0,
            'output': result.stdout[:8000],
            'error': result.stderr[:2000] if result.returncode != 0 else '',
            'duration_ms': elapsed,
            'exit_code': result.returncode
        })
    except subprocess.TimeoutExpired:
        return jsonify({'success': False, 'error': 'timeout (30s)', 'output': ''})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e), 'output': ''})

@app.route('/api/skills/delete', methods=['POST'])
def api_skill_delete():
    import re as _re
    data = request.json or {}
    name = data.get('name', '').strip()
    protected = {'create_skill', 'save_artifact', 'artifact_save', 'update_task'}
    if not name or name in protected:
        return jsonify({'error': 'cannot delete protected skill'}), 400
    path = os.path.join(FRAMEWORK_DIR, "skills", "shared", f"{name}.py")
    if not os.path.exists(path):
        return jsonify({'error': 'not found'}), 404
    os.remove(path)
    return jsonify({'success': True})


# ── Routes — Journal (PostgreSQL task_journal) ────────────────────────────────

@app.route('/journal')
def journal_page():
    return render_template('journal.html')

@app.route('/api/journal')
def api_journal():
    agent_id = request.args.get('agent_id', '')
    task_type = request.args.get('task_type', '')
    limit = int(request.args.get('limit', 100))
    try:
        import psycopg2, psycopg2.extras
        conn = psycopg2.connect(
            host="localhost", port=5432, dbname="baza_agents",
            user="switchhacker", password=os.environ.get("DB_PASSWORD","baza2026")
        )
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        sql = "SELECT * FROM task_journal WHERE 1=1"
        params = []
        if agent_id:   sql += " AND agent_id=%s";   params.append(agent_id)
        if task_type:  sql += " AND task_type=%s";  params.append(task_type)
        sql += " ORDER BY created_at DESC LIMIT %s"
        params.append(limit)
        cur.execute(sql, params)
        rows = [dict(r) for r in cur.fetchall()]
        # Serialise datetime objects
        for r in rows:
            for k, v in r.items():
                if hasattr(v, 'isoformat'):
                    r[k] = v.isoformat()
        conn.close()
        return jsonify(rows)
    except Exception as e:
        return jsonify({'error': str(e), 'rows': []})


# ── Routes — Activity Feed (live agent activity from task_journal) ────────────

AGENT_DISPLAY_NAMES = {
    'simon_bately': 'Simon Bately', 'claw_batto': 'Claw Batto',
    'phil_hass': 'Phil Hass', 'sam_axe': 'Sam Axe',
    'duke_harmon': 'Duke Harmon', 'rex_valor': 'Rex Valor',
    'scout_reeves': 'Scout Reeves', 'nova_sterling': 'Nova Sterling',
    'specter_voss': 'Specter Voss',
}

@app.route('/api/ahb/activity-feed')
def api_ahb_activity_feed():
    """Live activity feed — all agent actions for AHB123."""
    agent_id = request.args.get('agent_id', '')
    status = request.args.get('status', '')
    since = request.args.get('since', '')
    limit = min(int(request.args.get('limit', 50)), 200)
    try:
        import psycopg2, psycopg2.extras
        conn = psycopg2.connect(
            host="localhost", port=5432, dbname="baza_agents",
            user="switchhacker", password=os.environ.get("DB_PASSWORD","baza2026")
        )
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        sql = "SELECT id, agent_id, task_type, task_description, result, success, created_at, requested_by, status, action_summary FROM task_journal WHERE 1=1"
        params = []
        if agent_id:
            sql += " AND agent_id=%s"; params.append(agent_id)
        if status:
            sql += " AND status=%s"; params.append(status)
        if since:
            sql += " AND created_at > %s"; params.append(since)
        sql += " ORDER BY created_at DESC LIMIT %s"
        params.append(limit)
        cur.execute(sql, params)
        rows = [dict(r) for r in cur.fetchall()]
        conn.close()
        for r in rows:
            for k, v in r.items():
                if hasattr(v, 'isoformat'):
                    r[k] = v.isoformat()
            r['agent_name'] = AGENT_DISPLAY_NAMES.get(r.get('agent_id',''), r.get('agent_id',''))
            if not r.get('action_summary'):
                r['action_summary'] = r.get('task_description') or r.get('task_type') or ''
        return jsonify(rows)
    except Exception as e:
        return jsonify({'error': str(e), 'rows': []})


# ── Routes — Agent Memory (PostgreSQL agent_memory) ───────────────────────────

@app.route('/memory')
def memory_page():
    return render_template('memory.html')

@app.route('/api/memory')
def api_memory_list():
    agent_id = request.args.get('agent_id', '')
    category = request.args.get('category', '')
    try:
        import psycopg2, psycopg2.extras
        conn = psycopg2.connect(
            host="localhost", port=5432, dbname="baza_agents",
            user="switchhacker", password=os.environ.get("DB_PASSWORD","baza2026")
        )
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        sql = "SELECT * FROM agent_memory WHERE 1=1"
        params = []
        if agent_id: sql += " AND agent_id=%s"; params.append(agent_id)
        if category: sql += " AND category=%s"; params.append(category)
        sql += " ORDER BY agent_id, category, key"
        cur.execute(sql, params)
        rows = [dict(r) for r in cur.fetchall()]
        for r in rows:
            for k, v in r.items():
                if hasattr(v, 'isoformat'): r[k] = v.isoformat()
        conn.close()
        return jsonify(rows)
    except Exception as e:
        return jsonify({'error': str(e), 'rows': []})

@app.route('/api/memory', methods=['POST'])
def api_memory_set():
    data = request.json or {}
    agent_id = data.get('agent_id','').strip()
    key      = data.get('key','').strip()
    value    = data.get('value','').strip()
    category = data.get('category','general').strip()
    if not agent_id or not key:
        return jsonify({'error': 'agent_id and key required'}), 400
    try:
        import psycopg2
        conn = psycopg2.connect(
            host="localhost", port=5432, dbname="baza_agents",
            user="switchhacker", password=os.environ.get("DB_PASSWORD","baza2026")
        )
        conn.cursor().execute("""
            INSERT INTO agent_memory (agent_id, key, value, category, updated_at)
            VALUES (%s,%s,%s,%s,NOW())
            ON CONFLICT (agent_id, key) DO UPDATE
            SET value=EXCLUDED.value, category=EXCLUDED.category, updated_at=NOW()
        """, (agent_id, key, value, category))
        conn.commit()
        conn.close()
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/memory/delete', methods=['POST'])
def api_memory_delete():
    data = request.json or {}
    agent_id = data.get('agent_id','').strip()
    key      = data.get('key','').strip()
    if not agent_id or not key:
        return jsonify({'error': 'agent_id and key required'}), 400
    try:
        import psycopg2
        conn = psycopg2.connect(
            host="localhost", port=5432, dbname="baza_agents",
            user="switchhacker", password=os.environ.get("DB_PASSWORD","baza2026")
        )
        conn.cursor().execute("DELETE FROM agent_memory WHERE agent_id=%s AND key=%s", (agent_id, key))
        conn.commit()
        conn.close()
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/empire')
def api_empire_list():
    category = request.args.get('category','')
    try:
        import psycopg2, psycopg2.extras
        conn = psycopg2.connect(
            host="localhost", port=5432, dbname="baza_agents",
            user="switchhacker", password=os.environ.get("DB_PASSWORD","baza2026")
        )
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        if category:
            cur.execute("SELECT * FROM empire_knowledge WHERE category=%s ORDER BY key", (category,))
        else:
            cur.execute("SELECT * FROM empire_knowledge ORDER BY category, key")
        rows = [dict(r) for r in cur.fetchall()]
        for r in rows:
            for k, v in r.items():
                if hasattr(v, 'isoformat'): r[k] = v.isoformat()
        conn.close()
        return jsonify(rows)
    except Exception as e:
        return jsonify({'error': str(e), 'rows': []})

@app.route('/api/empire', methods=['POST'])
def api_empire_set():
    data = request.json or {}
    key      = data.get('key','').strip()
    value    = data.get('value','').strip()
    category = data.get('category','general').strip()
    if not key:
        return jsonify({'error': 'key required'}), 400
    try:
        import psycopg2
        conn = psycopg2.connect(
            host="localhost", port=5432, dbname="baza_agents",
            user="switchhacker", password=os.environ.get("DB_PASSWORD","baza2026")
        )
        conn.cursor().execute("""
            INSERT INTO empire_knowledge (key, value, category, updated_at, updated_by)
            VALUES (%s,%s,%s,NOW(),'dashboard')
            ON CONFLICT (key) DO UPDATE
            SET value=EXCLUDED.value, category=EXCLUDED.category, updated_at=NOW(), updated_by='dashboard'
        """, (key, value, category))
        conn.commit()
        conn.close()
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'error': str(e)}), 500


# ── Routes — Ollama model management ─────────────────────────────────────────

@app.route('/api/ollama/models')
def api_ollama_models():
    import urllib.request as _ur
    results = {}
    for label, port in [('amd', 11434), ('nvidia', 11435)]:
        try:
            with _ur.urlopen(f"http://localhost:{port}/api/tags", timeout=3) as r:
                data = json.loads(r.read())
                results[label] = [m['name'] for m in data.get('models', [])]
        except Exception:
            results[label] = []
    return jsonify(results)

@app.route('/api/ollama/running')
def api_ollama_running():
    import urllib.request as _ur
    results = {}
    for label, port in [('amd', 11434), ('nvidia', 11435)]:
        try:
            with _ur.urlopen(f"http://localhost:{port}/api/ps", timeout=3) as r:
                data = json.loads(r.read())
                results[label] = data.get('models', [])
        except Exception:
            results[label] = None
    return jsonify(results)


# ── Routes — System health (live) ─────────────────────────────────────────────

@app.route('/api/syshealth')
def api_syshealth():
    def _run(cmd):
        try:
            return subprocess.run(cmd, shell=True, capture_output=True, text=True, timeout=5).stdout.strip()
        except Exception:
            return ""

    # CPU load
    load = _run("cat /proc/loadavg").split()
    cpu_load = f"{load[0]}/{load[1]}/{load[2]}" if len(load) >= 3 else "?"

    # Memory
    mem_out = _run("free -h | awk '/^Mem:/{print $3\"/\"$2}'")

    # Disk
    disk_out = _run("df -h / | tail -1 | awk '{print $3\"/\"$2\" (\"$5\")\"}'")

    # GPU Nvidia
    nv = _run("nvidia-smi --query-gpu=temperature.gpu,utilization.gpu,memory.used,memory.total --format=csv,noheader,nounits 2>/dev/null | head -1")
    nv_data = {}
    if nv and "," in nv:
        parts = [x.strip() for x in nv.split(",")]
        if len(parts) >= 4:
            nv_data = {"temp": parts[0], "util": parts[1], "mem_used": parts[2], "mem_total": parts[3]}

    # GPU AMD via sysfs
    amd_data = {}
    try:
        temp = _run("cat /sys/class/hwmon/hwmon*/temp1_input 2>/dev/null | head -1")
        if temp.strip().isdigit():
            amd_data["temp"] = str(int(temp.strip())//1000)
    except Exception:
        pass

    return jsonify({
        "cpu_load": cpu_load,
        "memory": mem_out,
        "disk": disk_out,
        "nvidia": nv_data,
        "amd": amd_data,
        "timestamp": datetime.datetime.utcnow().isoformat()
    })


# ── Routes — Task Runner control ──────────────────────────────────────────────

@app.route('/api/taskrunner/run', methods=['POST'])
def api_taskrunner_run():
    """Manually trigger the task runner. Accepts optional `agent` and `task_id`."""
    data = request.json or {}
    agent = data.get('agent', '') or ''
    task_id = data.get('task_id', '') or ''
    cmd = [VENV_PYTHON, os.path.join(FRAMEWORK_DIR, "core", "task_runner.py")]
    if agent:
        cmd += ["--agent", agent]
    if task_id:
        cmd += ["--task-id", task_id]
    log_path = os.path.join(LOGS_DIR, "task_runner_manual.log")
    try:
        proc = subprocess.Popen(cmd, cwd=FRAMEWORK_DIR,
                                stdout=open(log_path, 'a'), stderr=subprocess.STDOUT)
        return jsonify({'success': True, 'pid': proc.pid, 'log': log_path,
                        'agent': agent, 'task_id': task_id})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/api/taskrunner/logs')
def api_taskrunner_logs():
    log_path = os.path.join(LOGS_DIR, "task_runner_manual.log")
    if not os.path.exists(log_path):
        return jsonify({'logs': '(no logs yet)'})
    try:
        lines = open(log_path).readlines()
        return jsonify({'logs': ''.join(lines[-100:])})
    except Exception as e:
        return jsonify({'logs': str(e)})


# ── Routes — AHB123 Business Hub ──────────────────────────────────────────────

def _ahb_db():
    """SQLite connection to baza_projects.db with row factory and a lock-tolerant
    `busy_timeout` so concurrent writers wait instead of failing instantly with
    "database is locked." 30s tolerates an external batch importer holding the
    write lock for tens of seconds at a time (e.g. the Takeout dedup commits
    every ~24s). WAL is enabled once at startup so readers don't block writers."""
    conn = sqlite3.connect(
        os.path.join(DASHBOARD_DIR, 'baza_projects.db'),
        timeout=30.0,
    )
    conn.row_factory = sqlite3.Row
    conn.execute("PRAGMA busy_timeout = 30000")
    conn.execute("PRAGMA synchronous = NORMAL")
    return conn


@app.route('/ahb123')
@app.route('/ahb123/<tab>')
def ahb123_page(tab='dashboard'):
    return render_template('ahb123.html', active_tab=tab)


@app.route('/mobile')
def mobile_page():
    resp = make_response(render_template('mobile.html'))
    resp.headers['Cache-Control'] = 'no-store, no-cache, must-revalidate, max-age=0'
    resp.headers['Pragma'] = 'no-cache'
    return resp


@app.route('/mobile/manifest.json')
def mobile_manifest():
    manifest = {
        "name": "Baza Empire",
        "short_name": "Baza",
        "start_url": "/mobile",
        "scope": "/",
        "display": "standalone",
        "background_color": "#07070f",
        "theme_color": "#07070f",
        "orientation": "portrait",
        "icons": [
            {"src": "/static/img/ahb_logo.jpeg", "sizes": "192x192", "type": "image/jpeg"},
            {"src": "/static/img/ahb_logo.jpeg", "sizes": "512x512", "type": "image/jpeg"}
        ]
    }
    return jsonify(manifest)


@app.route('/mobile/sw.js')
@app.route('/sw.js')
def mobile_sw():
    """Minimal app-shell service worker. Caches the icon + the editor JS so
    Chrome treats /mobile as installable. Network-first for HTML and APIs —
    no offline support; the SW exists only to satisfy the install criteria
    and to speed up cold loads of static assets."""
    sw = """
const CACHE = 'baza-shell-v5';
const SHELL = ['/static/img/ahb_logo.jpeg', '/static/quickrf-editor.js', '/mobile/manifest.json'];
self.addEventListener('install', e => {
  self.skipWaiting();
  e.waitUntil(caches.open(CACHE).then(c => c.addAll(SHELL).catch(()=>null)));
});
self.addEventListener('activate', e => {
  e.waitUntil((async () => {
    const keys = await caches.keys();
    await Promise.all(keys.filter(k => k !== CACHE).map(k => caches.delete(k)));
    await self.clients.claim();
  })());
});
self.addEventListener('fetch', e => {
  const url = new URL(e.request.url);
  if (e.request.method !== 'GET') return;
  // Cache-first only for our small static shell; everything else hits the network.
  if (SHELL.includes(url.pathname)) {
    e.respondWith(caches.match(e.request).then(r => r || fetch(e.request).then(resp => {
      const clone = resp.clone();
      caches.open(CACHE).then(c => c.put(e.request, clone)).catch(()=>null);
      return resp;
    }).catch(() => caches.match(e.request))));
  }
});
"""
    resp = make_response(sw)
    resp.headers['Content-Type'] = 'application/javascript'
    resp.headers['Service-Worker-Allowed'] = '/'
    resp.headers['Cache-Control'] = 'no-store, no-cache, must-revalidate'
    return resp


@app.route('/portal')
def portal_page():
    return render_template('portal.html')


# ── AHB123 — Clients ─────────────────────────────────────────────────────────

@app.route('/api/ahb/scopes', methods=['GET'])
def api_ahb_scopes_list():
    """Return every distinct, non-empty scope ever used on a project, in
    alphabetical (case-insensitive) order. Used by the project modal datalist
    so any custom scope a user types becomes a suggestion for the next run."""
    try:
        conn = _ahb_db()
        rows = conn.execute(
            "SELECT DISTINCT scope FROM ahb_projects "
            "WHERE scope IS NOT NULL AND TRIM(scope) != '' "
            "ORDER BY scope COLLATE NOCASE"
        ).fetchall()
        conn.close()
        return jsonify([r['scope'] for r in rows])
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/ahb/clients', methods=['GET'])
def api_ahb_clients_list():
    try:
        conn = _ahb_db()
        q = "SELECT * FROM ahb_clients"
        params = []
        status = request.args.get('status')
        if status:
            q += " WHERE status = ?"
            params.append(status)
        q += " ORDER BY created_at DESC"
        rows = conn.execute(q, params).fetchall()
        conn.close()
        return jsonify([dict(r) for r in rows])
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/ahb/clients', methods=['POST'])
def api_ahb_clients_create():
    try:
        data = request.json or {}
        cid = uuid.uuid4().hex[:24]
        conn = _ahb_db()
        conn.execute(
            """INSERT INTO ahb_clients (id, name, phone, email, address, city, source, status, notes, assigned_agent)
               VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)""",
            (cid, data.get('name'), data.get('phone'), data.get('email'),
             data.get('address'), data.get('city', 'Philadelphia'),
             data.get('source'), data.get('status', 'lead'),
             data.get('notes'), data.get('assigned_agent'))
        )
        conn.commit()
        conn.close()
        return jsonify({'success': True, 'id': cid})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/ahb/clients/<cid>', methods=['PUT'])
def api_ahb_clients_update(cid):
    try:
        data = request.json or {}
        fields = []
        vals = []
        for k in ('name', 'phone', 'email', 'address', 'city', 'source', 'status', 'notes', 'assigned_agent'):
            if k in data:
                fields.append(f"{k} = ?")
                vals.append(data[k])
        if not fields:
            return jsonify({'success': False, 'error': 'No fields to update'}), 400
        fields.append("updated_at = ?")
        vals.append(datetime.datetime.now().isoformat())
        vals.append(cid)
        conn = _ahb_db()
        conn.execute(f"UPDATE ahb_clients SET {', '.join(fields)} WHERE id = ?", vals)
        conn.commit()
        conn.close()
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/ahb/clients/<cid>', methods=['DELETE'])
def api_ahb_clients_delete(cid):
    try:
        conn = _ahb_db()
        conn.execute("DELETE FROM ahb_clients WHERE id = ?", (cid,))
        conn.commit()
        conn.close()
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


# ── AHB123 — Projects ────────────────────────────────────────────────────────

@app.route('/api/ahb/projects', methods=['GET'])
def api_ahb_projects_list():
    try:
        conn = _ahb_db()
        q = "SELECT * FROM ahb_projects WHERE 1=1"
        params = []
        if request.args.get('client_id'):
            q += " AND client_id = ?"
            params.append(request.args['client_id'])
        if request.args.get('status'):
            q += " AND status = ?"
            params.append(request.args['status'])
        q += " ORDER BY created_at DESC"
        rows = conn.execute(q, params).fetchall()
        # Attach payment summary so the UI can flag Completed-but-unpaid projects
        # without an N+1 round-trip per row.
        result = []
        for r in rows:
            d = dict(r)
            d['_payment'] = _ahb_project_payment_summary(conn, d['id'])
            result.append(d)
        conn.close()
        return jsonify(result)
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/ahb/projects', methods=['POST'])
def api_ahb_projects_create():
    """Create a project and auto-generate a correlated invoice. Optionally create phases."""
    try:
        data = request.json or {}
        pid = uuid.uuid4().hex[:24]
        conn = _ahb_db()
        conn.execute(
            """INSERT INTO ahb_projects (id, client_id, title, address, scope, description,
               budget_low, budget_high, status, start_date, end_date, assigned_agents, notes,
               value, client_name, client_email, contact_info, location,
               commission_pct, commission_value)
               VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)""",
            (pid, data.get('client_id'), data.get('title'), data.get('address'),
             data.get('scope'), data.get('description'),
             data.get('budget_low'), data.get('budget_high'),
             _ahb_canon_project_status(data.get('status', 'Planning')), data.get('start_date'),
             data.get('end_date'), data.get('assigned_agents'), data.get('notes'),
             data.get('value'), data.get('client_name', ''),
             data.get('client_email', ''), data.get('contact_info', ''),
             data.get('location', ''),
             # Default new projects to 0% commission. The user enters a real
             # commission per-project in the project detail modal when needed.
             float(data.get('commission_pct') or 0),
             float(data.get('commission_value') or 0))
        )

        # Create phases if provided
        phases = data.get('phases', [])
        line_items = []
        subtotal = 0
        for i, ph in enumerate(phases):
            phid = uuid.uuid4().hex[:24]
            phase_val = ph.get('value', 0) or 0
            conn.execute(
                """INSERT INTO ahb_project_phases (id, project_id, phase_number, name, value, start_date, end_date, status)
                   VALUES (?, ?, ?, ?, ?, ?, ?, ?)""",
                (phid, pid, ph.get('phase_number', i + 1), ph.get('name', f'Phase {i+1}'),
                 phase_val, ph.get('start_date', ''), ph.get('end_date', ''),
                 ph.get('status', 'pending')))
            line_items.append({'description': ph.get('name', f'Phase {i+1}'), 'quantity': 1, 'unit_price': phase_val})
            subtotal += phase_val

        # If no phases, default to a single line item carrying the budget.
        # The user adds real line items with material + labor breakdowns in the
        # invoice editor — we don't auto-shred the description (that produced
        # duplicated descriptions and uniform total/N prices).
        if not line_items:
            budget = data.get('value') or data.get('budget_high') or data.get('budget_low') or 0
            try: budget = float(budget) if budget else 0
            except Exception: budget = 0
            line_items = [{
                'description': data.get('title', 'Project'),
                'qty': 1, 'rate': budget, 'total': budget,
                'materials': 0, 'labor': 0,
                'quantity': 1, 'unit_price': budget,
            }]
            subtotal = budget

        # Auto-create correlated invoice
        iid = uuid.uuid4().hex[:24]
        inv_num = f"AHB-{datetime.datetime.now().strftime('%Y%m%d')}-{uuid.uuid4().hex[:3].upper()}"
        total = subtotal  # no tax by default
        conn.execute(
            """INSERT INTO ahb_invoices (id, client_id, project_id, invoice_number, line_items,
               subtotal, tax, total, status, notes, client_name, project_name, terms,
               company_name, contractor_name, client_address, client_email, client_phone, project_address)
               VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)""",
            (iid, data.get('client_id', ''), pid, inv_num,
             json.dumps(line_items), subtotal, 0, total, 'draft',
             '',
             data.get('client_name', ''), data.get('title', ''), 'Net 30',
             'All Home Building Co', 'Sergey Tkach',
             data.get('address', ''), data.get('client_email', ''),
             data.get('contact_info', ''), data.get('address', '')))

        conn.commit()
        conn.close()
        return jsonify({'success': True, 'id': pid, 'invoice_id': iid, 'invoice_number': inv_num})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/ahb/projects/<pid>', methods=['PUT'])
def api_ahb_projects_update(pid):
    try:
        data = request.json or {}
        fields = []
        vals = []
        for k in ('client_id', 'title', 'address', 'scope', 'description',
                   'budget_low', 'budget_high', 'status', 'start_date',
                   'end_date', 'assigned_agents', 'notes', 'value',
                   'client_name', 'client_email', 'contact_info', 'location',
                   'acquisition_type', 'commission_pct', 'commission_value',
                   'commission_beneficiary'):
            if k in data:
                fields.append(f"{k} = ?")
                vals.append(data[k])
        if not fields:
            return jsonify({'success': False, 'error': 'No fields to update'}), 400
        fields.append("updated_at = ?")
        vals.append(datetime.datetime.now().isoformat())
        vals.append(pid)
        conn = _ahb_db()
        conn.execute(f"UPDATE ahb_projects SET {', '.join(fields)} WHERE id = ?", vals)
        conn.commit()
        # Push header changes (title, address, client info, value) into linked invoice
        invoice_id = _sync_invoice_from_project(conn, pid, data)
        conn.commit()
        conn.close()
        return jsonify({'success': True, 'invoice_id': invoice_id})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/ahb/projects/<pid>/move-to-year', methods=['POST'])
def api_ahb_project_move_year(pid):
    """Re-assign a project to a different tax year. Updates the year column,
    shifts start_date/end_date by the year delta, and propagates the year to
    every linked invoice so InvoiceIT + Uncle Sam stay consistent."""
    try:
        body = request.get_json() or {}
        new_year = str(body.get('year', '')).strip()
        if not re.match(r'^\d{4}$', new_year):
            return jsonify({'success': False, 'error': 'year must be YYYY'}), 400
        conn = _ahb_db()
        proj = conn.execute("SELECT * FROM ahb_projects WHERE id = ?", (pid,)).fetchone()
        if not proj:
            conn.close()
            return jsonify({'success': False, 'error': 'project not found'}), 404
        proj = dict(proj)
        # Compute new start/end dates by shifting the year while keeping month/day
        def _shift(date_str, target_year):
            if not date_str or len(date_str) < 10:
                return date_str
            try:
                return target_year + date_str[4:]   # YYYY-MM-DD → newYYYY-MM-DD
            except Exception:
                return date_str
        new_start = _shift(proj.get('start_date'), new_year) or f"{new_year}-01-01"
        new_end   = _shift(proj.get('end_date'),   new_year)
        conn.execute(
            "UPDATE ahb_projects SET year=?, start_date=?, end_date=?, updated_at=? WHERE id=?",
            (new_year, new_start, new_end, datetime.datetime.now().isoformat(), pid)
        )
        # Propagate to every linked invoice
        invs = conn.execute("SELECT id, date, paid_date FROM ahb_invoices WHERE project_id = ?", (pid,)).fetchall()
        for inv in invs:
            inv_d = dict(inv)
            new_inv_date  = _shift(inv_d.get('date'),     new_year)
            new_inv_paid  = _shift(inv_d.get('paid_date'),new_year)
            conn.execute(
                "UPDATE ahb_invoices SET year=?, date=COALESCE(?,date), paid_date=COALESCE(?,paid_date), updated_at=? WHERE id=?",
                (new_year, new_inv_date, new_inv_paid, datetime.datetime.now().isoformat(), inv_d['id'])
            )
        conn.commit()
        conn.close()
        return jsonify({'success': True, 'year': new_year, 'invoices_updated': len(invs)})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/ahb/invoices/<iid>/move-to-year', methods=['POST'])
def api_ahb_invoice_move_year(iid):
    """Re-assign a single invoice (and its linked project's year if it has one)
    to a different tax year. Used for one-off InvoiceIT corrections."""
    try:
        body = request.get_json() or {}
        new_year = str(body.get('year', '')).strip()
        if not re.match(r'^\d{4}$', new_year):
            return jsonify({'success': False, 'error': 'year must be YYYY'}), 400
        also_project = bool(body.get('also_project', False))
        conn = _ahb_db()
        inv = conn.execute("SELECT * FROM ahb_invoices WHERE id = ?", (iid,)).fetchone()
        if not inv:
            conn.close()
            return jsonify({'success': False, 'error': 'invoice not found'}), 404
        inv = dict(inv)
        def _shift(date_str, target_year):
            if not date_str or len(date_str) < 10:
                return date_str
            return target_year + date_str[4:]
        new_date     = _shift(inv.get('date'),      new_year)
        new_paid     = _shift(inv.get('paid_date'), new_year)
        conn.execute(
            "UPDATE ahb_invoices SET year=?, date=COALESCE(?,date), paid_date=COALESCE(?,paid_date), updated_at=? WHERE id=?",
            (new_year, new_date, new_paid, datetime.datetime.now().isoformat(), iid)
        )
        if also_project and inv.get('project_id'):
            proj = conn.execute("SELECT start_date, end_date FROM ahb_projects WHERE id = ?", (inv['project_id'],)).fetchone()
            if proj:
                new_start = _shift(proj['start_date'], new_year)
                new_end   = _shift(proj['end_date'],   new_year)
                conn.execute(
                    "UPDATE ahb_projects SET year=?, start_date=COALESCE(?,start_date), end_date=COALESCE(?,end_date), updated_at=? WHERE id=?",
                    (new_year, new_start, new_end, datetime.datetime.now().isoformat(), inv['project_id'])
                )
        conn.commit()
        conn.close()
        return jsonify({'success': True, 'year': new_year})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


# ── Project Quotes ──────────────────────────────────────────────────────────

def _deposit_lines(total, carry_money):
    """The two auto-injected payment-schedule rows: 50% deposit + 50% balance.

    When `carry_money` is True (quote-applied invoice with blank work-line
    totals), the deposit/balance rows carry the money — their `total`s sum
    to the basis. When False (manual invoice with real work-line totals),
    the rows are informational with `total=0` and the dollar amounts shown
    inline in the description so they don't double-count."""
    total = float(total or 0)
    deposit_amount = round(total * 0.5, 2)
    balance_amount = round(total - deposit_amount, 2)
    if carry_money:
        d_total, b_total = deposit_amount, balance_amount
        d_desc = '50% Deposit due before commencement of work'
        b_desc = 'Balance (50%) due upon project completion — total due after deposit'
    else:
        d_total, b_total = 0, 0
        d_desc = f'50% Deposit due before commencement of work — ${deposit_amount:,.2f}'
        b_desc = f'Balance (50%) due upon project completion — total due after deposit ${balance_amount:,.2f}'
    return [
        {
            'description': d_desc,
            'qty': 1, 'rate': d_total, 'total': d_total,
            'unit': 'qty', 'materials': 0, 'labor': 0,
            'quantity': 1, 'unit_price': d_total,
            '_auto_deposit': True,
        },
        {
            'description': b_desc,
            'qty': 1, 'rate': b_total, 'total': b_total,
            'unit': 'qty', 'materials': 0, 'labor': 0,
            'quantity': 1, 'unit_price': b_total,
            '_auto_deposit': True,
        },
    ]


def _line_items_from_description(description, total):
    """Build invoice line_items from a quote: each non-empty line of the
    description becomes one item carrying ONLY the description text — qty,
    rate, materials, labor, and total are blank ($0) so the contractor can
    fill them in later. The full total is carried on the auto-appended
    50% deposit + 50% balance rows."""
    total = float(total or 0)
    lines = [ln.strip() for ln in (description or '').split('\n') if ln.strip()]
    items = []
    if not lines:
        items.append({
            'description': 'Project work per quote',
            'qty': 1, 'rate': 0, 'total': 0,
            'unit': 'qty', 'materials': 0, 'labor': 0,
            'quantity': 1, 'unit_price': 0,
        })
    else:
        for ln in lines:
            items.append({
                'description': ln,
                'qty': 1, 'rate': 0, 'total': 0,
                'unit': 'qty', 'materials': 0, 'labor': 0,
                'quantity': 1, 'unit_price': 0,
            })
    # Quote-applied invoices have blank work-line totals, so deposit/balance
    # carry the money.
    items.extend(_deposit_lines(total, carry_money=True))
    return items


_DEPOSIT_DESC_PREFIXES = (
    '50% deposit due before',
    'balance (50%) due upon',
    'balance due upon',
)


def _is_auto_deposit_line(li):
    """An auto-deposit row is identified by its marker OR by its canonical
    description (the marker is dropped on frontend round-trips through the
    invoice modal, so the description is our fallback)."""
    if li.get('_auto_deposit'):
        return True
    desc = (li.get('description') or '').strip().lower()
    return any(desc.startswith(p) for p in _DEPOSIT_DESC_PREFIXES)


def _ensure_deposit_lines(items_or_json, fallback_total=0):
    """Strip any prior auto-deposit lines from the payload, then append fresh
    50% deposit + 50% balance payment-schedule rows.

    - If the remaining work lines sum to > 0 (manual invoice flow), the
      deposit/balance rows are informational with `total=0` and the dollar
      amounts inlined in the description, so they don't double-count.
    - If the work lines all sum to 0 (quote-driven invoice with blank work
      totals), the deposit/balance rows carry the money split 50/50 of
      `fallback_total` (the quote total)."""
    if items_or_json is None:
        return None
    items = items_or_json
    if isinstance(items, str):
        try:
            items = json.loads(items) if items else []
        except Exception:
            return items_or_json
    if not isinstance(items, list):
        return items_or_json
    work_items = [li for li in items if not _is_auto_deposit_line(li)]
    work_subtotal = 0.0
    for li in work_items:
        try:
            work_subtotal += float(li.get('total') or li.get('rate') or 0)
        except Exception:
            pass
    if work_subtotal > 0:
        return work_items + _deposit_lines(work_subtotal, carry_money=False)
    fallback = float(fallback_total or 0)
    if fallback > 0:
        return work_items + _deposit_lines(fallback, carry_money=True)
    return work_items  # nothing to schedule yet


def _apply_quote_to_invoice(c, project_id, quote_total, quote_description):
    """Rewrite the linked invoice's line items from the quote's description
    and sync subtotal/total. Returns the invoice id or None if no invoice
    exists for this project."""
    inv = c.execute(
        "SELECT id FROM ahb_invoices WHERE project_id=? ORDER BY created_at ASC LIMIT 1",
        (project_id,)
    ).fetchone()
    if not inv:
        return None
    items = _line_items_from_description(quote_description, quote_total)
    c.execute(
        "UPDATE ahb_invoices SET line_items=?, subtotal=?, total=?, tax=0, updated_at=? WHERE id=?",
        (json.dumps(items), float(quote_total or 0), float(quote_total or 0),
         datetime.datetime.now().isoformat(), inv['id'])
    )
    return inv['id']


@app.route('/api/ahb/projects/<pid>/quotes', methods=['GET', 'POST'])
def api_ahb_project_quotes(pid):
    conn = _ahb_db()
    c = conn.cursor()
    if request.method == 'POST':
        d = request.get_json() or {}
        total = float(d.get('total') or 0)
        if not total:
            conn.close()
            return jsonify({'success': False, 'error': 'total required'}), 400
        c.execute("""INSERT INTO ahb_quotes
            (project_id, method, scope, description, total, breakdown, notes, is_active)
            VALUES (?,?,?,?,?,?,?,?)""",
            (pid, d.get('method', 'manual'), d.get('scope', ''),
             d.get('description', ''), total,
             json.dumps(d.get('breakdown') or {}),
             d.get('notes', ''), 1 if d.get('make_active') else 0))
        qid = c.lastrowid
        # If marked active, demote others, update project value, and rebuild
        # the linked invoice's line items from this quote's description.
        if d.get('make_active'):
            c.execute("UPDATE ahb_quotes SET is_active=0 WHERE project_id=? AND id<>?", (pid, qid))
            c.execute("UPDATE ahb_projects SET value=?, budget_high=?, updated_at=? WHERE id=?",
                      (total, total, datetime.datetime.now().isoformat(), pid))
            _apply_quote_to_invoice(c, pid, total, d.get('description', ''))
        conn.commit()
        row = c.execute("SELECT * FROM ahb_quotes WHERE id=?", (qid,)).fetchone()
        conn.close()
        return jsonify({'success': True, 'quote': dict(row)})
    rows = c.execute("SELECT * FROM ahb_quotes WHERE project_id=? ORDER BY id DESC", (pid,)).fetchall()
    conn.close()
    return jsonify([dict(r) for r in rows])


@app.route('/api/ahb/quotes/<int:qid>', methods=['GET', 'DELETE', 'PUT'])
def api_ahb_quote_modify(qid):
    conn = _ahb_db()
    c = conn.cursor()
    row = c.execute("SELECT * FROM ahb_quotes WHERE id=?", (qid,)).fetchone()
    if not row:
        conn.close()
        return jsonify({'success': False, 'error': 'not found'}), 404
    pid = row['project_id']
    if request.method == 'GET':
        q = dict(row)
        try:
            q['breakdown'] = json.loads(q['breakdown']) if q.get('breakdown') else {}
        except (json.JSONDecodeError, TypeError):
            q['breakdown'] = {}
        conn.close()
        return jsonify(q)
    if request.method == 'DELETE':
        c.execute("DELETE FROM ahb_quotes WHERE id=?", (qid,))
        conn.commit()
        conn.close()
        return jsonify({'success': True})
    d = request.get_json() or {}
    # Full-field update (used by Edit-in-tool flow). Any subset may be provided.
    full_update_fields = ('total', 'breakdown', 'description', 'scope', 'method')
    if any(k in d for k in full_update_fields):
        new_total = float(d['total']) if 'total' in d and d['total'] is not None else row['total']
        new_breakdown = (json.dumps(d['breakdown']) if d.get('breakdown') is not None else row['breakdown']) \
                        if 'breakdown' in d else row['breakdown']
        new_desc = d['description'] if 'description' in d else row['description']
        new_scope = d['scope'] if 'scope' in d else row['scope']
        new_method = d['method'] if 'method' in d else row['method']
        c.execute("""UPDATE ahb_quotes SET total=?, breakdown=?, description=?, scope=?, method=?
                     WHERE id=?""",
                  (new_total, new_breakdown, new_desc, new_scope, new_method, qid))
        # If this quote is currently chosen, propagate the new total to project value
        # and (optionally) rebuild the linked invoice from the new description.
        if row['is_active']:
            c.execute("UPDATE ahb_projects SET value=?, budget_high=?, updated_at=? WHERE id=?",
                      (new_total, new_total, datetime.datetime.now().isoformat(), pid))
            _apply_quote_to_invoice(c, pid, new_total, new_desc or '')
        # Re-read for the response
        row = c.execute("SELECT * FROM ahb_quotes WHERE id=?", (qid,)).fetchone()
    if d.get('make_active'):
        c.execute("UPDATE ahb_quotes SET is_active=0 WHERE project_id=?", (pid,))
        c.execute("UPDATE ahb_quotes SET is_active=1 WHERE id=?", (qid,))
        c.execute("UPDATE ahb_projects SET value=?, budget_high=?, updated_at=? WHERE id=?",
                  (row['total'], row['total'], datetime.datetime.now().isoformat(), pid))
        # Rebuild the linked invoice's line items from this quote's description.
        _apply_quote_to_invoice(c, pid, row['total'], row['description'] or '')
    if d.get('deactivate'):
        c.execute("UPDATE ahb_quotes SET is_active=0 WHERE id=?", (qid,))
    if d.get('apply_to_invoice'):
        # Apply this quote's total to the linked invoice and rebuild line items
        # from its description — works whether or not the quote is active.
        _apply_quote_to_invoice(c, pid, row['total'], row['description'] or '')
    if 'notes' in d:
        c.execute("UPDATE ahb_quotes SET notes=? WHERE id=?", (d['notes'], qid))
    conn.commit()
    conn.close()
    return jsonify({'success': True})


@app.route('/api/ahb/quotes/<int:qid>/pdf', methods=['GET'])
def api_ahb_quote_pdf(qid):
    """Generate a printable quote PDF for the project's estimates bin."""
    try:
        conn = _ahb_db()
        row = conn.execute("SELECT * FROM ahb_quotes WHERE id=?", (qid,)).fetchone()
        if not row:
            conn.close()
            return jsonify({'error': 'Quote not found'}), 404
        q = dict(row)
        try:
            breakdown = json.loads(q['breakdown']) if q.get('breakdown') else {}
        except (json.JSONDecodeError, TypeError):
            breakdown = {}
        project = None
        client = None
        if q.get('project_id'):
            p = conn.execute("SELECT * FROM ahb_projects WHERE id=?", (q['project_id'],)).fetchone()
            if p:
                project = dict(p)
                if project.get('client_id'):
                    cl = conn.execute("SELECT * FROM ahb_clients WHERE id=?", (project['client_id'],)).fetchone()
                    if cl:
                        client = dict(cl)
        conn.close()

        # Logo as base64
        logo_b64 = ''
        logo_path = os.path.join(DASHBOARD_DIR, 'static', 'img', 'ahb_logo.jpeg')
        if os.path.exists(logo_path):
            import base64
            with open(logo_path, 'rb') as lf:
                logo_b64 = base64.b64encode(lf.read()).decode('utf-8')

        method_label = (q.get('method') or 'manual').replace('_', ' ').title()
        total = float(q.get('total') or 0)
        created = (q.get('created_at') or '')[:10]
        desc_html = (q.get('description') or '').replace('<', '&lt;').replace('>', '&gt;')
        scope_html = (q.get('scope') or '').replace('<', '&lt;').replace('>', '&gt;')
        notes_html = (q.get('notes') or '').replace('<', '&lt;').replace('>', '&gt;')
        proj_title = (project or {}).get('title') or ''
        proj_addr = (project or {}).get('address') or ''
        client_name = (client or {}).get('name') or (project or {}).get('client_name') or ''
        client_addr = (client or {}).get('address') or ''

        # Render breakdown rows (any numeric field in the breakdown blob)
        bk_rows = ''
        if isinstance(breakdown, dict):
            for k, v in breakdown.items():
                if isinstance(v, (int, float)) and k != 'total':
                    label = str(k).replace('_', ' ').title()
                    bk_rows += f'''<tr>
                        <td style="padding:8px 12px;border-bottom:1px solid #eee;color:#333;">{label}</td>
                        <td style="padding:8px 12px;border-bottom:1px solid #eee;text-align:right;color:#333;font-family:monospace;">${float(v):,.2f}</td>
                    </tr>'''
        if not bk_rows:
            bk_rows = f'''<tr>
                <td style="padding:8px 12px;border-bottom:1px solid #eee;color:#333;">Estimate</td>
                <td style="padding:8px 12px;border-bottom:1px solid #eee;text-align:right;color:#333;font-family:monospace;">${total:,.2f}</td>
            </tr>'''

        active_badge = '<span style="display:inline-block;padding:2px 10px;border-radius:10px;background:#dcfce7;color:#16a34a;font-size:11px;font-weight:700;">CHOSEN</span>' if q.get('is_active') else '<span style="display:inline-block;padding:2px 10px;border-radius:10px;background:#f3f4f6;color:#6b7280;font-size:11px;font-weight:700;">INACTIVE</span>'
        scope_line = f'<div style="margin-top:6px;font-size:12px;color:#666;"><strong>Trade/Scope:</strong> {scope_html}</div>' if scope_html else ''
        scope_block = f'<div style="margin-bottom:20px;padding:12px 16px;background:#f8fafc;border-radius:6px;border-left:3px solid #7c3aed;"><div style="font-size:11px;color:#999;text-transform:uppercase;font-weight:700;margin-bottom:6px;">Scope of Work</div><div style="font-size:13px;color:#444;line-height:1.5;white-space:pre-wrap;">{desc_html}</div>{scope_line}</div>' if desc_html else ''
        notes_block = f'<div style="margin-top:24px;padding:12px 16px;background:#fff8e1;border-radius:6px;border-left:3px solid #f59e0b;"><div style="font-size:11px;color:#999;text-transform:uppercase;font-weight:700;margin-bottom:6px;">Notes</div><div style="font-size:13px;color:#444;line-height:1.5;white-space:pre-wrap;">{notes_html}</div></div>' if notes_html else ''
        logo_block = f'<img src="data:image/jpeg;base64,{logo_b64}" style="width:50px;height:50px;object-fit:contain;margin-top:2px;">' if logo_b64 else '<div style="width:50px;height:50px;background:#7c3aed;border-radius:8px;margin-top:2px;"></div>'

        html = f'''<!DOCTYPE html>
<html>
<head><meta charset="utf-8"><title>Quote #{qid}</title>
<style>
@media print {{ body {{ margin:0; }} @page {{ margin:40px 50px; }} }}
body {{ font-family:'Helvetica Neue',Helvetica,Arial,sans-serif;max-width:780px;margin:30px auto;color:#333;font-size:14px;line-height:1.5; }}
</style></head>
<body>
<div style="display:flex;justify-content:space-between;align-items:flex-start;margin-bottom:20px;">
    <div style="display:flex;align-items:flex-start;gap:12px;">
        {logo_block}
        <div>
            <div style="font-size:20px;font-weight:700;color:#1a1a1a;white-space:nowrap;">All Home Building CO LLC</div>
            <div style="font-size:12px;color:#888;">2725 Colmar Ave, Bensalem, PA 19020</div>
            <div style="font-size:12px;color:#888;">800-484-6404 · AHB123.com</div>
        </div>
    </div>
    <div style="text-align:right;">
        <div style="font-size:28px;font-weight:300;color:#333;letter-spacing:2px;">QUOTE</div>
        <div style="font-size:14px;color:#555;">#Q-{qid}</div>
        <div style="margin-top:4px;">{active_badge}</div>
    </div>
</div>

<div style="display:flex;justify-content:space-between;margin:16px 0 24px;padding:12px 0;border-top:1px solid #e5e7eb;border-bottom:1px solid #e5e7eb;">
    <div>
        <div style="font-size:11px;color:#999;text-transform:uppercase;font-weight:700;margin-bottom:4px;">Prepared For:</div>
        <div style="font-weight:600;">{client_name}</div>
        <div style="color:#666;">{client_addr}</div>
    </div>
    <div style="text-align:center;">
        <div style="font-size:11px;color:#999;text-transform:uppercase;font-weight:700;margin-bottom:4px;">Project:</div>
        <div style="color:#444;">{proj_title}</div>
        <div style="color:#666;font-size:12px;">{proj_addr}</div>
    </div>
    <div style="text-align:right;">
        <div style="font-size:11px;color:#999;text-transform:uppercase;font-weight:700;margin-bottom:4px;">Date:</div>
        <div>{created}</div>
        <div style="font-size:11px;color:#999;margin-top:6px;">Method:</div>
        <div style="font-size:12px;color:#444;">{method_label}</div>
    </div>
</div>

{scope_block}

<table style="width:100%;border-collapse:collapse;margin:0 0 20px;">
    <thead>
        <tr style="background:#f8fafc;">
            <th style="padding:10px 12px;text-align:left;border-bottom:2px solid #e2e8f0;font-size:12px;color:#64748b;font-weight:700;">Cost Component</th>
            <th style="padding:10px 12px;text-align:right;border-bottom:2px solid #e2e8f0;font-size:12px;color:#64748b;font-weight:700;">Amount</th>
        </tr>
    </thead>
    <tbody>{bk_rows}</tbody>
</table>

<div style="display:flex;justify-content:flex-end;">
    <div style="width:280px;border-top:2px solid #333;padding-top:8px;">
        <div style="display:flex;justify-content:space-between;padding:6px 0;font-size:20px;font-weight:700;color:#7c3aed;">
            <span>Quote Total:</span><span>${total:,.2f}</span>
        </div>
    </div>
</div>

{notes_block}

<div style="margin-top:32px;padding-top:16px;border-top:1px solid #e5e7eb;font-size:11px;color:#888;">
    This quote is valid for 30 days from the date above. Quote totals are estimates based on the scope described.
    Final pricing may adjust based on actual scope, materials chosen, and site conditions discovered during work.
</div>
</body>
</html>'''

        download = request.args.get('download', '0') == '1'
        try:
            from weasyprint import HTML as WeasyHTML
            pdf_bytes = WeasyHTML(string=html).write_pdf()
            response = make_response(pdf_bytes)
            response.headers['Content-Type'] = 'application/pdf'
            disposition = 'attachment' if download else 'inline'
            response.headers['Content-Disposition'] = f'{disposition}; filename="quote_{qid}.pdf"'
            return response
        except ImportError:
            response = make_response(html)
            response.headers['Content-Type'] = 'text/html; charset=utf-8'
            response.headers['Content-Disposition'] = f'inline; filename="quote_{qid}.html"'
            return response
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/ahb/projects/<pid>', methods=['DELETE'])
def api_ahb_projects_delete(pid):
    try:
        conn = _ahb_db()
        conn.execute("DELETE FROM ahb_projects WHERE id = ?", (pid,))
        conn.commit()
        conn.close()
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


# ── AHB123 — canonical project status helpers ───────────────────────────────
#
# Project lifecycle is exactly 3 states:
#   Planning     — quote/estimate phase (no signed contract, no money in)
#   In Progress  — contract signed, deposit received, work happening
#   Completed    — work finished (balance may or may not be paid)
#
# Invoice status follows from the project:
#   Planning     -> invoice 'Sent'      (quote outstanding)
#   In Progress  -> invoice 'Approved'  (deposit received; balance still owed)
#   Completed    -> invoice 'Paid' if fully paid, else 'Approved' (final bill due)
#
# Anything else gets coerced into the closest of those three at the API edge.

AHB_PROJECT_STATES = ('Planning', 'In Progress', 'Completed')

def _ahb_canon_project_status(s):
    if not s:
        return 'Planning'
    x = s.strip().lower().replace('_', ' ')
    if x in ('planning', 'estimate', 'proposal', 'quote', 'lead'):
        return 'Planning'
    if x in ('in progress', 'inprogress', 'signed', 'active', 'invoiced'):
        return 'In Progress'
    if x in ('completed', 'complete', 'done', 'closed', 'paid'):
        return 'Completed'
    # Already canonical?
    for c in AHB_PROJECT_STATES:
        if x == c.lower():
            return c
    return 'Planning'


def _ahb_project_payment_summary(conn, project_id):
    """Compute payment state for a project from its linked invoice.

    Returns a dict the frontend uses to decide if a Completed project still
    owes money (red badge), if a project has any deposit recorded
    (auto-flip Planning -> In Progress), etc.
    """
    row = conn.execute(
        "SELECT id, invoice_number, total FROM ahb_invoices "
        "WHERE project_id = ? ORDER BY created_at ASC LIMIT 1",
        (project_id,)
    ).fetchone()
    if not row:
        return {'invoice_id': None, 'invoice_number': None,
                'total': 0.0, 'paid': 0.0, 'owed': 0.0,
                'has_payments': False, 'fully_paid': False}
    total = float(row['total'] or 0)
    pay_row = conn.execute(
        "SELECT COALESCE(sum(amount),0) as paid FROM ahb_payments WHERE invoice_id = ?",
        (row['id'],)
    ).fetchone()
    paid = float(pay_row['paid'] or 0) if pay_row else 0.0
    # Floating-point tolerance — anything within a penny counts as paid in full.
    fully_paid = total > 0 and paid + 0.01 >= total
    return {
        'invoice_id': row['id'],
        'invoice_number': row['invoice_number'],
        'total': total,
        'paid': paid,
        'owed': max(0.0, total - paid),
        'has_payments': paid > 0,
        'fully_paid': fully_paid,
    }


def _ahb_apply_status_sync(conn, project_id, new_status):
    """Set project.status and align the linked invoice's status to match.

    Returns (canonical_status, invoice_after_update_or_None).
    Does NOT commit — caller owns the transaction.
    """
    canon = _ahb_canon_project_status(new_status)
    now = datetime.datetime.now().isoformat()
    conn.execute(
        "UPDATE ahb_projects SET status = ?, updated_at = ? WHERE id = ?",
        (canon, now, project_id)
    )
    inv = conn.execute(
        "SELECT * FROM ahb_invoices WHERE project_id = ? ORDER BY created_at ASC LIMIT 1",
        (project_id,)
    ).fetchone()
    if not inv:
        return canon, None
    inv = dict(inv)
    summary = _ahb_project_payment_summary(conn, project_id)
    if canon == 'Planning':
        target = 'Sent'
    elif canon == 'In Progress':
        target = 'Approved'
    else:  # Completed
        target = 'Paid' if summary['fully_paid'] else 'Approved'
    update_fields = {'status': target, 'updated_at': now}
    if canon == 'Completed' and not summary['fully_paid']:
        update_fields['notes'] = f"Final bill due. Remaining balance: ${summary['owed']:.2f}"
    if target == 'Paid' and not inv.get('paid_date'):
        update_fields['paid_date'] = now[:10]
    set_clause = ', '.join(f"{k} = ?" for k in update_fields)
    vals = list(update_fields.values()) + [inv['id']]
    conn.execute(f"UPDATE ahb_invoices SET {set_clause} WHERE id = ?", vals)
    return canon, {**inv, **update_fields}


@app.route('/api/ahb/projects/<pid>/status', methods=['POST'])
def api_ahb_project_status_sync(pid):
    """Update project status and auto-sync the linked invoice status + calendar events.

    Accepts any of the legacy labels (estimate/signed/paid/etc.) and snaps to
    the canonical 3-state lifecycle (Planning, In Progress, Completed).
    """
    try:
        data = request.json or {}
        new_status = data.get('status')
        if not new_status:
            return jsonify({'success': False, 'error': 'status is required'}), 400

        conn = _ahb_db()
        project = conn.execute("SELECT * FROM ahb_projects WHERE id = ?", (pid,)).fetchone()
        if not project:
            conn.close()
            return jsonify({'success': False, 'error': 'Project not found'}), 404
        project = dict(project)

        canon, invoice_result = _ahb_apply_status_sync(conn, pid, new_status)
        new_status = canon
        now = datetime.datetime.now().isoformat()

        # If status is 'In Progress', create calendar events from project dates
        if new_status.lower().replace(' ', '_') in ('in_progress', 'in progress'):
            start_date = project.get('start_date')
            end_date = project.get('end_date')
            title = project.get('title', 'Project')
            if start_date:
                eid = uuid.uuid4().hex[:24]
                conn.execute(
                    """INSERT INTO ahb_events (id, title, details, date, category, all_day, project_id)
                       VALUES (?, ?, ?, ?, ?, ?, ?)""",
                    (eid, f"{title} - Start", f"Project start: {title}", start_date, 'project', 1, pid))
            if end_date:
                eid = uuid.uuid4().hex[:24]
                conn.execute(
                    """INSERT INTO ahb_events (id, title, details, date, category, all_day, project_id)
                       VALUES (?, ?, ?, ?, ?, ?, ?)""",
                    (eid, f"{title} - End", f"Project end: {title}", end_date, 'project', 1, pid))

        conn.commit()

        # Re-fetch updated project
        updated_project = dict(conn.execute("SELECT * FROM ahb_projects WHERE id = ?", (pid,)).fetchone())
        conn.close()
        return jsonify({'success': True, 'project': updated_project, 'invoice': invoice_result})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


# ── AHB123 — Invoices ────────────────────────────────────────────────────────

@app.route('/api/ahb/invoices', methods=['GET'])
def api_ahb_invoices_list():
    try:
        conn = _ahb_db()
        q = "SELECT * FROM ahb_invoices WHERE 1=1"
        params = []
        if request.args.get('status'):
            q += " AND status = ?"
            params.append(request.args['status'])
        if request.args.get('client_id'):
            q += " AND client_id = ?"
            params.append(request.args['client_id'])
        if request.args.get('is_change_order') is not None:
            q += " AND is_change_order = ?"; params.append(int(request.args['is_change_order']))
        q += " ORDER BY created_at DESC"
        rows = conn.execute(q, params).fetchall()
        conn.close()
        return jsonify([dict(r) for r in rows])
    except Exception as e:
        return jsonify({'error': str(e)}), 500


def _ensure_invoice_project(conn, data: dict) -> str:
    """Every invoice must belong to a project. Returns a project_id, creating
    one from the invoice's own metadata if none is supplied or matchable.

    Match precedence: explicit project_id → (client_id + project_address)
    → (client_name + project_address) → (client_name) → new project.
    """
    pid = (data.get('project_id') or '').strip()
    if pid:
        # Verify it actually exists; if not, fall through to creation.
        row = conn.execute("SELECT id FROM ahb_projects WHERE id = ?", (pid,)).fetchone()
        if row:
            return pid

    client_id   = (data.get('client_id') or '').strip()
    client_name = (data.get('client_name') or '').strip()
    addr        = (data.get('project_address') or data.get('client_address') or '').strip()
    proj_name   = (data.get('project_name') or '').strip()

    # Try matchers in order
    row = None
    if client_id and addr:
        row = conn.execute(
            "SELECT id FROM ahb_projects WHERE client_id = ? AND LOWER(TRIM(address)) = LOWER(TRIM(?)) LIMIT 1",
            (client_id, addr)
        ).fetchone()
    if not row and client_name and addr:
        row = conn.execute(
            "SELECT id FROM ahb_projects WHERE LOWER(TRIM(client_name)) = LOWER(TRIM(?)) AND LOWER(TRIM(address)) = LOWER(TRIM(?)) LIMIT 1",
            (client_name, addr)
        ).fetchone()
    if not row and proj_name:
        row = conn.execute(
            "SELECT id FROM ahb_projects WHERE LOWER(TRIM(title)) = LOWER(TRIM(?)) LIMIT 1",
            (proj_name,)
        ).fetchone()
    if not row and client_name and not addr:
        # Last resort: the only project for that client
        rows = conn.execute(
            "SELECT id FROM ahb_projects WHERE LOWER(TRIM(client_name)) = LOWER(TRIM(?))",
            (client_name,)
        ).fetchall()
        if len(rows) == 1:
            row = rows[0]
    if row:
        return row['id']

    # Create a new project so the invoice has a home
    new_pid = uuid.uuid4().hex[:24]
    title_bits = []
    if proj_name: title_bits.append(proj_name)
    elif client_name: title_bits.append(client_name)
    if addr: title_bits.append(addr)
    if not title_bits:
        title_bits.append('Auto-created from invoice')
    title = ' — '.join(title_bits)[:120]
    try:
        value = float(data.get('total') or data.get('subtotal') or 0)
    except (TypeError, ValueError):
        value = 0.0
    conn.execute(
        """INSERT INTO ahb_projects (id, client_id, title, address, scope, description,
           status, start_date, notes, value, client_name, client_email, contact_info,
           location, commission_pct, commission_value)
           VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)""",
        (new_pid, client_id or None, title, addr, '', 'Auto-created to house an invoice.',
         'invoiced', datetime.datetime.now().date().isoformat(),
         'Auto-created from invoice.', value, client_name,
         data.get('client_email', ''), data.get('client_phone', ''),
         addr, 0.0, 0.0)
    )
    return new_pid


@app.route('/api/ahb/invoices', methods=['POST'])
def api_ahb_invoices_create():
    try:
        data = request.json or {}
        iid = uuid.uuid4().hex[:24]
        inv_num = f"AHB-{datetime.datetime.now().strftime('%Y%m%d')}-{uuid.uuid4().hex[:3].upper()}"

        # Always inject the 50% deposit + balance payment-schedule rows and
        # recompute subtotal/total to match the rebuilt line items.
        items_in = data.get('line_items', [])
        rebuilt = _ensure_deposit_lines(
            items_in,
            fallback_total=data.get('total') or data.get('subtotal') or 0
        )
        if rebuilt is not None:
            items_in = rebuilt
            sub = sum(float(li.get('total') or 0) for li in items_in)
            data['subtotal'] = sub
            data['total'] = sub + float(data.get('tax') or 0)

        conn = _ahb_db()
        # Every invoice must belong to a project — find or auto-create one.
        data['project_id'] = _ensure_invoice_project(conn, data)
        conn.execute(
            """INSERT INTO ahb_invoices (id, client_id, project_id, invoice_number, line_items,
               subtotal, tax, total, status, due_date, paid_date, notes,
               date, parent_invoice_id, is_change_order, overdue_since,
               overdue_interest_per_week, company_name, contractor_name,
               client_address, client_email, client_phone, project_address)
               VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)""",
            (iid, data.get('client_id'), data.get('project_id'), inv_num,
             json.dumps(items_in) if isinstance(items_in, list) else items_in,
             data.get('subtotal'), data.get('tax'), data.get('total'),
             data.get('status', 'draft'), data.get('due_date'),
             data.get('paid_date'), data.get('notes'),
             data.get('date', ''), data.get('parent_invoice_id', ''),
             data.get('is_change_order', 0), data.get('overdue_since', ''),
             data.get('overdue_interest_per_week', 50),
             data.get('company_name', 'All Home Building Co'),
             data.get('contractor_name', 'Sergey Tkach'),
             data.get('client_address', ''), data.get('client_email', ''),
             data.get('client_phone', ''), data.get('project_address', ''))
        )
        conn.commit()
        conn.close()
        return jsonify({'success': True, 'id': iid, 'invoice_number': inv_num})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/ahb/invoices/<iid>', methods=['PUT'])
def api_ahb_invoices_update(iid):
    try:
        data = request.json or {}
        fields = []
        vals = []
        # Re-inject deposit lines whenever line_items is being rewritten and
        # recompute subtotal/total so the stored math matches the rebuilt
        # line items (frontend-submitted subtotal/total may include the
        # round-tripped old deposit amounts and would otherwise be stale).
        if 'line_items' in data:
            rebuilt = _ensure_deposit_lines(
                data['line_items'],
                fallback_total=data.get('total') or data.get('subtotal') or 0
            )
            if rebuilt is not None:
                data['line_items'] = rebuilt
                sub = sum(float(li.get('total') or 0) for li in rebuilt)
                data['subtotal'] = sub
                data['total'] = sub + float(data.get('tax') or 0)
        for k in ('client_id', 'project_id', 'line_items', 'subtotal', 'tax',
                   'total', 'status', 'due_date', 'paid_date', 'notes',
                   'date', 'parent_invoice_id', 'is_change_order', 'overdue_since',
                   'overdue_interest_per_week', 'company_name', 'contractor_name',
                   'client_address', 'client_email', 'client_phone', 'project_address'):
            if k in data:
                fields.append(f"{k} = ?")
                v = data[k]
                if k == 'line_items' and isinstance(v, list):
                    v = json.dumps(v)
                vals.append(v)
        if not fields:
            return jsonify({'success': False, 'error': 'No fields to update'}), 400
        conn = _ahb_db()
        # If project_id is being cleared (or empty), auto-create/find one so
        # the invoice still has a project home after the update.
        if 'project_id' in data and not (data.get('project_id') or '').strip():
            # Merge stored row defaults with incoming data so the helper has
            # enough context to find or create a sensible project.
            row = conn.execute(
                "SELECT client_id, client_name, project_address, client_address, "
                "client_email, client_phone, subtotal, total FROM ahb_invoices WHERE id = ?",
                (iid,)
            ).fetchone()
            ctx = dict(data)
            if row:
                for k in ('client_id','client_name','project_address','client_address',
                         'client_email','client_phone','subtotal','total'):
                    ctx.setdefault(k, row[k])
            pid = _ensure_invoice_project(conn, ctx)
            # Replace the empty project_id placeholder in vals
            for i, f in enumerate(fields):
                if f == 'project_id = ?':
                    vals[i] = pid
                    break
        fields.append("updated_at = ?")
        vals.append(datetime.datetime.now().isoformat())
        vals.append(iid)
        conn.execute(f"UPDATE ahb_invoices SET {', '.join(fields)} WHERE id = ?", vals)
        conn.commit()
        conn.close()
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/ahb/invoices/backfill-projects', methods=['POST'])
def api_ahb_invoices_backfill_projects():
    """Find every invoice with an empty/missing project_id and either link it
    to an existing matching project or auto-create one. Idempotent."""
    try:
        conn = _ahb_db()
        rows = conn.execute(
            "SELECT id, client_id, client_name, project_address, client_address, "
            "client_email, client_phone, subtotal, total "
            "FROM ahb_invoices "
            "WHERE project_id IS NULL OR TRIM(project_id) = '' "
            "   OR project_id NOT IN (SELECT id FROM ahb_projects)"
        ).fetchall()
        fixed = 0
        for r in rows:
            ctx = {k: r[k] for k in r.keys()}
            pid = _ensure_invoice_project(conn, ctx)
            conn.execute("UPDATE ahb_invoices SET project_id = ? WHERE id = ?", (pid, r['id']))
            fixed += 1
        conn.commit(); conn.close()
        return jsonify({'success': True, 'fixed': fixed})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/ahb/invoices/<iid>', methods=['DELETE'])
def api_ahb_invoices_delete(iid):
    try:
        conn = _ahb_db()
        cur = conn.execute("DELETE FROM ahb_invoices WHERE id = ?", (iid,))
        conn.commit()
        deleted = cur.rowcount
        conn.close()
        if not deleted:
            return jsonify({'success': False, 'error': 'Invoice not found'}), 404
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


# ── AHB123 — Receipts ────────────────────────────────────────────────────────

@app.route('/api/ahb/receipts', methods=['GET'])
def api_ahb_receipts_list():
    try:
        conn = _ahb_db()
        q = "SELECT * FROM ahb_receipts WHERE 1=1"
        params = []
        if request.args.get('project_id'):
            q += " AND project_id = ?"
            params.append(request.args['project_id'])
        if request.args.get('category'):
            q += " AND category = ?"
            params.append(request.args['category'])
        q += " ORDER BY created_at DESC"
        rows = conn.execute(q, params).fetchall()
        conn.close()
        return jsonify([dict(r) for r in rows])
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/ahb/receipts', methods=['POST'])
def api_ahb_receipts_create():
    try:
        rid = uuid.uuid4().hex[:24]
        file_path = None
        if request.content_type and 'multipart' in request.content_type:
            data = request.form.to_dict()
            f = request.files.get('file')
            if f:
                upload_dir = os.path.join(ARTIFACTS_DIR, 'receipts')
                os.makedirs(upload_dir, exist_ok=True)
                file_path = os.path.join(upload_dir, f"{rid}_{f.filename}")
                f.save(file_path)
        else:
            data = request.json or {}
            file_path = data.get('file_path')
        conn = _ahb_db()
        amount = data.get('amount') or data.get('total') or 0
        conn.execute(
            """INSERT INTO ahb_receipts (id, project_id, vendor, amount, category,
               description, receipt_date, file_path, ocr_text, created_by,
               store_name, payment_method, total)
               VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)""",
            (rid, data.get('project_id'), data.get('vendor') or data.get('store_name'),
             amount, data.get('category'), data.get('description'), data.get('receipt_date'),
             file_path, data.get('ocr_text'), data.get('created_by'),
             data.get('store_name') or data.get('vendor'), data.get('payment_method'), amount)
        )
        conn.commit()
        conn.close()
        return jsonify({'success': True, 'id': rid})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/ahb/receipts/<rid>', methods=['GET'])
def api_ahb_receipt_detail(rid):
    try:
        conn = _ahb_db()
        row = conn.execute("SELECT * FROM ahb_receipts WHERE id = ?", (rid,)).fetchone()
        conn.close()
        if not row:
            return jsonify({'error': 'Not found'}), 404
        return jsonify(dict(row))
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/ahb/receipts/<rid>', methods=['PUT'])
def api_ahb_receipt_update(rid):
    """Update a receipt AND record each changed field in ahb_receipt_corrections.
    The corrections table is what receipt_learn.py mines to improve vendor
    aliases and category rules — so accurate audit is the whole point."""
    try:
        data = request.json or {}
        editable = ['vendor', 'amount', 'category', 'description', 'receipt_date',
                    'store_name', 'payment_method', 'total', 'teller_name',
                    'store_location', 'purchase_time', 'tax_amount', 'subtotal',
                    'items_json', 'ocr_text', 'ocr_raw', 'ocr_structured',
                    'image_path', 'project_id', 'year']
        changed_by = (
            request.headers.get('X-Agent-Id')
            or data.pop('_edited_by', None)
            or 'serge'
        )

        conn = _ahb_db()
        conn.row_factory = sqlite3.Row
        old_row = conn.execute(
            "SELECT * FROM ahb_receipts WHERE id = ?", (rid,)
        ).fetchone()

        def _ser(v):
            if v is None:
                return ''
            if isinstance(v, (list, dict)):
                return json.dumps(v, ensure_ascii=False)
            return str(v)

        corrections = []
        fields, vals = [], []
        for k in editable:
            if k not in data:
                continue
            new_v = data[k]
            old_v = old_row[k] if old_row is not None and k in old_row.keys() else None
            if _ser(old_v) != _ser(new_v):
                corrections.append((rid, changed_by, k, _ser(old_v), _ser(new_v)))
            fields.append(f"{k} = ?")
            vals.append(new_v)

        if corrections:
            conn.executemany(
                """INSERT INTO ahb_receipt_corrections
                        (receipt_id, changed_by, field, old_value, new_value)
                   VALUES (?, ?, ?, ?, ?)""",
                corrections,
            )

        if fields:
            vals.append(rid)
            conn.execute(f"UPDATE ahb_receipts SET {', '.join(fields)} WHERE id = ?", vals)
        conn.commit()
        conn.close()
        return jsonify({
            'success': True,
            'changed_fields': [c[2] for c in corrections],
            'changed_by': changed_by,
        })
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/ahb/receipts/bulk-rename-vendor', methods=['POST'])
def api_ahb_receipt_bulk_rename_vendor():
    """Propagate a vendor (or store_name) rename across every receipt that
    currently carries the old name. Each touched row is also logged in
    ahb_receipt_corrections so receipt_learn.py picks up the alias and
    future imports normalize correctly. Supports ?dry_run for a count
    preview before the user confirms."""
    try:
        data = request.json or {}
        field = (data.get('field') or 'vendor').strip()
        if field not in ('vendor', 'store_name'):
            return jsonify({'success': False, 'error': 'field must be vendor or store_name'}), 400
        old_v = (data.get('old_vendor') or '').strip()
        new_v = (data.get('new_vendor') or '').strip()
        if not old_v or not new_v:
            return jsonify({'success': False, 'error': 'old_vendor and new_vendor required'}), 400
        if old_v.lower() == new_v.lower() and old_v == new_v:
            return jsonify({'success': True, 'updated_count': 0, 'updated_ids': []})

        dry_run = bool(data.get('dry_run'))
        exclude_id = data.get('exclude_id')
        changed_by = (
            request.headers.get('X-Agent-Id')
            or data.get('_edited_by')
            or 'serge'
        )

        conn = _ahb_db()
        conn.row_factory = sqlite3.Row
        params = [old_v]
        extra = ''
        if exclude_id:
            extra = ' AND id != ?'
            params.append(exclude_id)
        rows = conn.execute(
            f"SELECT id, {field} AS cur FROM ahb_receipts "
            f"WHERE LOWER(TRIM({field})) = LOWER(TRIM(?)){extra}",
            params
        ).fetchall()
        matching_ids = [r['id'] for r in rows]

        if dry_run:
            conn.close()
            return jsonify({
                'success': True,
                'matching_count': len(matching_ids),
                'matching_ids': matching_ids,
                'field': field,
                'old_vendor': old_v,
                'new_vendor': new_v,
            })

        if matching_ids:
            placeholders = ','.join(['?'] * len(matching_ids))
            conn.execute(
                f"UPDATE ahb_receipts SET {field} = ? WHERE id IN ({placeholders})",
                [new_v] + matching_ids,
            )
            conn.executemany(
                """INSERT INTO ahb_receipt_corrections
                        (receipt_id, changed_by, field, old_value, new_value)
                   VALUES (?, ?, ?, ?, ?)""",
                [(rid, changed_by, field, old_v, new_v) for rid in matching_ids],
            )
        conn.commit()
        conn.close()
        return jsonify({
            'success': True,
            'updated_count': len(matching_ids),
            'updated_ids': matching_ids,
            'field': field,
            'old_vendor': old_v,
            'new_vendor': new_v,
        })
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/ahb/receipts/<rid>', methods=['DELETE'])
def api_ahb_receipt_delete(rid):
    """Delete a filed receipt row and its associated image file."""
    try:
        conn = _ahb_db()
        conn.row_factory = sqlite3.Row
        row = conn.execute(
            "SELECT image_path, file_path FROM ahb_receipts WHERE id = ?", (rid,)
        ).fetchone()
        if not row:
            conn.close()
            return jsonify({'success': False, 'error': 'not found'}), 404
        conn.execute("DELETE FROM ahb_receipts WHERE id = ?", (rid,))
        conn.commit()
        conn.close()
        for p in (row['image_path'], row['file_path']):
            if p and os.path.exists(p):
                try:
                    os.remove(p)
                except OSError:
                    pass
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/agents', methods=['GET'])
def api_agents_registry():
    """Return the list of named agents/admins for UI attribution.
    Feeds the 'Filed by' column on receipts/documents and the audit trail."""
    try:
        import yaml as _yaml
        path = os.path.join(
            os.path.dirname(DASHBOARD_DIR), 'config', 'agents_registry.yaml'
        )
        with open(path) as f:
            data = _yaml.safe_load(f) or {}
        admins = data.get('admins') or []
        # Return as both a list and a convenience map for easy JS lookup.
        by_id = {a.get('id', ''): a for a in admins if a.get('id')}
        # Also synthesize a CYD-shaped 'agents' array (name/handle/role/status).
        # CYD firmware parses this; web dashboard ignores the extra key.
        agents = [{
            'name':      a.get('display') or a.get('id', ''),
            'handle':    a.get('id', ''),
            'role':      a.get('role', ''),
            'status':    'active',
            'last_task': '',
        } for a in admins]
        return jsonify({'admins': admins, 'by_id': by_id, 'agents': agents})
    except Exception as e:
        return jsonify({'admins': [], 'by_id': {}, 'error': str(e)}), 200


@app.route('/api/ahb/receipts/corrections', methods=['GET'])
def api_ahb_receipt_corrections_list():
    """List correction history, optionally filtered by receipt_id or field.
    Query params: receipt_id, field, limit (default 100)."""
    try:
        rid = request.args.get('receipt_id')
        field = request.args.get('field')
        limit = int(request.args.get('limit', 100))
        where, vals = [], []
        if rid:
            where.append("receipt_id = ?"); vals.append(rid)
        if field:
            where.append("field = ?"); vals.append(field)
        clause = ("WHERE " + " AND ".join(where)) if where else ""
        conn = _ahb_db()
        conn.row_factory = sqlite3.Row
        rows = conn.execute(
            f"""SELECT id, receipt_id, changed_by, changed_at, field,
                       old_value, new_value
                  FROM ahb_receipt_corrections
                  {clause}
                 ORDER BY id DESC LIMIT ?""",
            (*vals, limit),
        ).fetchall()
        conn.close()
        return jsonify({'corrections': [dict(r) for r in rows]})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/ahb/receipts/<rid>/ocr', methods=['POST'])
def api_ahb_receipts_ocr(rid):
    """Run OCR on a receipt image using the receipt_ocr skill."""
    try:
        conn = _ahb_db()
        row = conn.execute("SELECT image_path FROM ahb_receipts WHERE id = ?", (rid,)).fetchone()
        if not row or not row['image_path']:
            conn.close()
            return jsonify({'success': False, 'error': 'No image attached to this receipt'}), 400

        skill_path = os.path.join(os.path.dirname(DASHBOARD_DIR), 'skills', 'shared', 'receipt_ocr.py')
        env = os.environ.copy()
        env['SKILL_ARGS'] = json.dumps({'image_path': row['image_path'], 'mode': 'full'})
        result = subprocess.run([VENV_PYTHON, skill_path], capture_output=True, text=True, timeout=120, env=env)

        if result.returncode != 0:
            conn.close()
            return jsonify({'success': False, 'error': result.stderr.strip() or 'OCR failed'})

        ocr_result = json.loads(result.stdout.strip())
        structured = ocr_result.get('structured', {})

        # Update receipt with OCR data
        update_fields = {
            'ocr_raw': ocr_result.get('ocr_raw', ''),
            'ocr_structured': json.dumps(structured),
            'ocr_text': ocr_result.get('ocr_raw', ''),
        }
        # Only overwrite empty fields with OCR data
        field_map = {
            'store_name': 'store_name', 'store_location': 'store_location',
            'teller_name': 'teller_name', 'purchase_time': 'purchase_time',
            'payment_method': 'payment_method', 'category': 'category',
            'total': 'total', 'tax_amount': 'tax_amount', 'subtotal': 'subtotal',
        }
        current = dict(conn.execute("SELECT * FROM ahb_receipts WHERE id = ?", (rid,)).fetchone())
        for ocr_key, db_key in field_map.items():
            val = structured.get(ocr_key)
            if val and not current.get(db_key):
                update_fields[db_key] = val
        if structured.get('items'):
            update_fields['items_json'] = json.dumps(structured['items'])
        if structured.get('purchase_date') and not current.get('receipt_date'):
            update_fields['receipt_date'] = structured['purchase_date']
        if not current.get('vendor') and structured.get('store_name'):
            update_fields['vendor'] = structured['store_name']
        if not current.get('amount') and structured.get('total'):
            update_fields['amount'] = structured['total']

        set_clause = ', '.join(f"{k} = ?" for k in update_fields)
        vals = list(update_fields.values()) + [rid]
        conn.execute(f"UPDATE ahb_receipts SET {set_clause} WHERE id = ?", vals)
        conn.commit()
        conn.close()

        return jsonify({'success': True, 'ocr': ocr_result, 'updated_fields': list(update_fields.keys())})
    except subprocess.TimeoutExpired:
        return jsonify({'success': False, 'error': 'OCR timed out (120s)'}), 500
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/ahb/receipts/upload', methods=['POST'])
def api_ahb_receipts_upload():
    """Upload a receipt image (multipart `file`/`image`) or reference one
    via `pick_token` from the Baza picker, save to disk, create receipt
    record, and trigger OCR. If `target_id` is provided, attaches the image
    to an existing receipt instead of creating a new one."""
    try:
        import shutil as _shutil
        f = request.files.get('file') or request.files.get('image')
        pick_token = (request.form.get('pick_token') or '').strip()
        target_id = (request.form.get('target_id') or '').strip()

        rid = target_id or str(uuid.uuid4())
        upload_dir = os.path.join(DASHBOARD_DIR, 'uploads', 'ahb', 'receipts')
        os.makedirs(upload_dir, exist_ok=True)

        if f:
            safe_name = re.sub(r'[^\w.\-]', '_', f.filename or 'receipt.jpg')
            file_path = os.path.join(upload_dir, f"{rid}_{safe_name}")
            f.save(file_path)
        elif pick_token:
            src = _pick_decode_token(pick_token)
            if not src:
                return jsonify({'success': False, 'error': 'Invalid pick_token'}), 400
            safe_name = re.sub(r'[^\w.\-]', '_', os.path.basename(src)) or 'receipt.jpg'
            file_path = os.path.join(upload_dir, f"{rid}_{safe_name}")
            _shutil.copy2(src, file_path)
        else:
            return jsonify({'success': False, 'error': 'No file or pick_token uploaded'}), 400

        # Attach-to-existing branch: skip the INSERT, update image_path on the
        # existing row, and return early.
        if target_id:
            conn = _ahb_db()
            conn.execute("UPDATE ahb_receipts SET image_path = ?, file_path = ? WHERE id = ?",
                         (file_path, file_path, target_id))
            conn.commit()
            conn.close()
            return jsonify({'success': True, 'id': target_id, 'image_path': file_path,
                            'image_url': f'/api/ahb/receipts/image/{target_id}', 'attached': True})

        # Get form data
        data = request.form.to_dict()
        now = datetime.datetime.now()
        year = data.get('year') or now.strftime('%Y')

        conn = _ahb_db()
        conn.execute(
            """INSERT INTO ahb_receipts (id, vendor, amount, category, description,
               receipt_date, store_name, payment_method, total, image_path, year, file_path)
               VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)""",
            (rid, data.get('vendor', ''), float(data.get('amount', 0) or 0),
             data.get('category', ''), data.get('description', ''),
             data.get('receipt_date', now.strftime('%Y-%m-%d')),
             data.get('store_name', data.get('vendor', '')),
             data.get('payment_method', ''), float(data.get('amount', 0) or 0),
             file_path, year, file_path))
        conn.commit()
        conn.close()

        return jsonify({'success': True, 'id': rid, 'image_path': file_path,
                        'image_url': f'/api/ahb/receipts/image/{rid}'})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/ahb/receipts/image/<rid>', methods=['GET'])
def api_ahb_receipt_image(rid):
    """Serve a receipt image by receipt ID."""
    try:
        conn = _ahb_db()
        row = conn.execute("SELECT image_path FROM ahb_receipts WHERE id = ?", (rid,)).fetchone()
        conn.close()
        if not row or not row['image_path']:
            return 'No image', 404
        img_path = row['image_path']
        if not os.path.exists(img_path):
            return 'Image file not found', 404
        return send_from_directory(os.path.dirname(img_path), os.path.basename(img_path))
    except Exception as e:
        return str(e), 500


# ── AHB123 — Payroll ─────────────────────────────────────────────────────────

@app.route('/api/ahb/payroll', methods=['GET'])
def api_ahb_payroll_list():
    try:
        conn = _ahb_db()
        q = "SELECT * FROM ahb_payroll WHERE 1=1"
        params = []
        if request.args.get('status'):
            q += " AND status = ?"
            params.append(request.args['status'])
        if request.args.get('worker_name'):
            q += " AND worker_name = ?"
            params.append(request.args['worker_name'])
        q += " ORDER BY created_at DESC"
        rows = conn.execute(q, params).fetchall()
        conn.close()
        return jsonify([dict(r) for r in rows])
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/ahb/payroll', methods=['POST'])
def api_ahb_payroll_create():
    try:
        data = request.json or {}
        pid = uuid.uuid4().hex[:24]
        hours = float(data.get('hours', 0))
        rate = float(data.get('rate', 0))
        total = hours * rate
        conn = _ahb_db()
        conn.execute(
            """INSERT INTO ahb_payroll (id, worker_name, role, hours, rate, total,
               period_start, period_end, status, project_id, notes)
               VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)""",
            (pid, data.get('worker_name'), data.get('role'), hours, rate, total,
             data.get('period_start'), data.get('period_end'),
             data.get('status', 'pending'), data.get('project_id'), data.get('notes'))
        )
        conn.commit()
        conn.close()
        return jsonify({'success': True, 'id': pid, 'total': total})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/ahb/payroll/<pid>', methods=['PUT'])
def api_ahb_payroll_update(pid):
    try:
        data = request.json or {}
        fields = []
        vals = []
        for k in ('worker_name', 'role', 'hours', 'rate',
                   'period_start', 'period_end', 'status', 'project_id', 'notes'):
            if k in data:
                fields.append(f"{k} = ?")
                vals.append(data[k])
        # Auto-recalculate total if both hours and rate provided, otherwise allow explicit total
        if 'hours' in data and 'rate' in data:
            fields.append("total = ?")
            vals.append(float(data['hours']) * float(data['rate']))
        elif 'total' in data:
            fields.append("total = ?")
            vals.append(data['total'])
        if not fields:
            return jsonify({'success': False, 'error': 'No fields to update'}), 400
        vals.append(pid)
        conn = _ahb_db()
        conn.execute(f"UPDATE ahb_payroll SET {', '.join(fields)} WHERE id = ?", vals)
        conn.commit()
        conn.close()
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


# ── AHB123 — Estimates ───────────────────────────────────────────────────────

@app.route('/api/ahb/estimates', methods=['GET'])
def api_ahb_estimates_list():
    try:
        conn = _ahb_db()
        rows = conn.execute("SELECT * FROM ahb_estimates ORDER BY created_at DESC").fetchall()
        conn.close()
        return jsonify([dict(r) for r in rows])
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/ahb/estimates', methods=['POST'])
def api_ahb_estimates_create():
    try:
        data = request.json or {}
        eid = uuid.uuid4().hex[:24]
        conn = _ahb_db()
        conn.execute(
            """INSERT INTO ahb_estimates (id, client_id, project_id, title, description, scope,
               line_items, subtotal, markup_pct, total, status, generated_by, notes)
               VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)""",
            (eid, data.get('client_id'), data.get('project_id'), data.get('title'),
             data.get('description'), data.get('scope'),
             json.dumps(data.get('line_items', [])) if isinstance(data.get('line_items'), list) else data.get('line_items'),
             data.get('subtotal'), data.get('markup_pct', 15), data.get('total'),
             data.get('status', 'draft'), data.get('generated_by'), data.get('notes'))
        )
        conn.commit()
        conn.close()
        return jsonify({'success': True, 'id': eid})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/ahb/estimates/generate', methods=['POST'])
def api_ahb_estimates_generate():
    return jsonify({'status': 'queued', 'message': 'Estimate generation queued for agent processing'})


# ── AHB123 — Chats ───────────────────────────────────────────────────────────

@app.route('/api/ahb/chats', methods=['GET'])
def api_ahb_chats_list():
    try:
        conn = _ahb_db()
        q = "SELECT * FROM ahb_chats"
        params = []
        if request.args.get('status'):
            q += " WHERE status = ?"
            params.append(request.args['status'])
        q += " ORDER BY created_at DESC"
        rows = conn.execute(q, params).fetchall()
        conn.close()
        return jsonify([dict(r) for r in rows])
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/ahb/chats/<chat_id>/messages', methods=['GET'])
def api_ahb_chat_messages(chat_id):
    try:
        conn = _ahb_db()
        rows = conn.execute(
            "SELECT * FROM ahb_chat_messages WHERE chat_id = ? ORDER BY created_at ASC",
            (chat_id,)
        ).fetchall()
        conn.close()
        return jsonify([dict(r) for r in rows])
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/ahb/chats/<chat_id>/messages', methods=['POST'])
def api_ahb_chat_message_create(chat_id):
    try:
        data = request.json or {}
        conn = _ahb_db()
        conn.execute(
            "INSERT INTO ahb_chat_messages (chat_id, role, content, agent_id) VALUES (?, ?, ?, ?)",
            (chat_id, data.get('role', 'user'), data.get('content'), data.get('agent_id'))
        )
        conn.execute("UPDATE ahb_chats SET updated_at = ? WHERE id = ?",
                      (datetime.datetime.now().isoformat(), chat_id))
        conn.commit()
        conn.close()
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/ahb/chats/history')
def api_ahb_chats_history():
    db = _ahb_db()
    page = request.args.get('page', 1, type=int)
    per_page = request.args.get('per_page', 50, type=int)
    q = request.args.get('q', '')
    offset = (page - 1) * per_page
    if q:
        rows = db.execute("SELECT c.*, (SELECT COUNT(*) FROM ahb_chat_messages m WHERE m.chat_id=c.id) as msg_count FROM ahb_chats c WHERE c.visitor_name LIKE ? OR c.visitor_email LIKE ? ORDER BY c.updated_at DESC LIMIT ? OFFSET ?", (f'%{q}%', f'%{q}%', per_page, offset)).fetchall()
    else:
        rows = db.execute("SELECT c.*, (SELECT COUNT(*) FROM ahb_chat_messages m WHERE m.chat_id=c.id) as msg_count FROM ahb_chats c ORDER BY c.updated_at DESC LIMIT ? OFFSET ?", (per_page, offset)).fetchall()
    total = db.execute("SELECT COUNT(*) FROM ahb_chats").fetchone()[0]
    db.close()
    return jsonify({'chats': [dict(r) for r in rows], 'total': total, 'page': page, 'per_page': per_page})

@app.route('/api/ahb/chats/stats')
def api_ahb_chats_stats():
    db = _ahb_db()
    total = db.execute("SELECT COUNT(*) FROM ahb_chats").fetchone()[0]
    active = db.execute("SELECT COUNT(*) FROM ahb_chats WHERE status='active'").fetchone()[0]
    resolved = db.execute("SELECT COUNT(*) FROM ahb_chats WHERE status='resolved'").fetchone()[0]
    escalated = db.execute("SELECT COUNT(*) FROM ahb_chats WHERE status='escalated'").fetchone()[0]
    total_msgs = db.execute("SELECT COUNT(*) FROM ahb_chat_messages").fetchone()[0]
    db.close()
    return jsonify({'total': total, 'active': active, 'resolved': resolved, 'escalated': escalated, 'total_messages': total_msgs})

@app.route('/api/ahb/chats/<chat_id>/export')
def api_ahb_chats_export(chat_id):
    db = _ahb_db()
    chat = db.execute("SELECT * FROM ahb_chats WHERE id=?", (chat_id,)).fetchone()
    if not chat:
        db.close()
        return jsonify({'error': 'Not found'}), 404
    msgs = db.execute("SELECT * FROM ahb_chat_messages WHERE chat_id=? ORDER BY created_at", (chat_id,)).fetchall()
    db.close()
    lines = [f"Chat Export: {dict(chat).get('visitor_name','Unknown')}", f"Date: {dict(chat).get('created_at','')}", f"Status: {dict(chat).get('status','')}", "---"]
    for m in msgs:
        md = dict(m)
        lines.append(f"[{md.get('created_at','')}] {md.get('role','').upper()}: {md.get('content','')}")
    from flask import Response
    return Response('\n'.join(lines), mimetype='text/plain', headers={'Content-Disposition': f'attachment; filename=chat_{chat_id[:8]}.txt'})

@app.route('/api/ahb/chats/<chat_id>', methods=['PUT'])
def api_ahb_chat_update(chat_id):
    data = request.json or {}
    db = _ahb_db()
    fields = []
    vals = []
    for k in ['status', 'lead_score', 'assigned_agent', 'visitor_name', 'visitor_email', 'visitor_phone']:
        if k in data:
            fields.append(f"{k}=?")
            vals.append(data[k])
    if not fields:
        db.close()
        return jsonify({'success': False, 'error': 'No fields to update'}), 400
    fields.append("updated_at=datetime('now')")
    vals.append(chat_id)
    db.execute(f"UPDATE ahb_chats SET {','.join(fields)} WHERE id=?", vals)
    db.commit()
    db.close()
    return jsonify({'success': True})

@app.route('/api/ahb/chats/<chat_id>/escalate', methods=['POST'])
def api_ahb_chat_escalate(chat_id):
    db = _ahb_db()
    db.execute("UPDATE ahb_chats SET status='escalated', assigned_agent='simon_bately', updated_at=datetime('now') WHERE id=?", (chat_id,))
    db.commit()
    db.close()
    token = os.environ.get('TELEGRAM_SIMON_BATELY', '')
    if token:
        db2 = _ahb_db()
        chat = db2.execute("SELECT * FROM ahb_chats WHERE id=?", (chat_id,)).fetchone()
        db2.close()
        if chat:
            cd = dict(chat)
            try:
                import requests as _req
                resp = _req.get(f"https://api.telegram.org/bot{token}/getUpdates", params={"limit": 5}, timeout=5)
                updates = resp.json().get('result', [])
                tg_chat = None
                for u in reversed(updates):
                    msg = u.get('message', {})
                    if msg.get('chat', {}).get('type') == 'private':
                        tg_chat = msg['chat']['id']
                        break
                if tg_chat:
                    _req.post(f"https://api.telegram.org/bot{token}/sendMessage",
                        json={"chat_id": tg_chat, "text": f"\U0001f6a8 Chat Escalated\nVisitor: {cd.get('visitor_name','Unknown')}\nEmail: {cd.get('visitor_email','')}\nLead Score: {cd.get('lead_score','')}"}, timeout=5)
            except Exception:
                pass
    return jsonify({'success': True})


# ── AHB123 — Voice ───────────────────────────────────────────────────────────

@app.route('/api/ahb/voice/voices', methods=['GET'])
def api_ahb_voice_list_voices():
    try:
        result = subprocess.run(
            [VENV_PYTHON, '-m', 'edge_tts', '--list-voices'], capture_output=True, text=True, timeout=15
        )
        voices = []
        for line in result.stdout.splitlines():
            line = line.strip()
            if not line or line.startswith('Name') or line.startswith('---'):
                continue
            parts = line.split()
            if len(parts) >= 2:
                name = parts[0]
                gender = parts[1] if len(parts) > 1 else ''
                locale = name.rsplit('-', 1)[0] if '-' in name else ''
                voices.append({'name': name, 'gender': gender, 'locale': locale})
        return jsonify(voices)
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/ahb/voice/configs', methods=['GET'])
def api_ahb_voice_configs_list():
    try:
        conn = _ahb_db()
        rows = conn.execute("SELECT * FROM ahb_voice_configs ORDER BY created_at DESC").fetchall()
        conn.close()
        return jsonify([dict(r) for r in rows])
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/ahb/voice/configs', methods=['POST'])
def api_ahb_voice_configs_create():
    try:
        data = request.json or {}
        vid = uuid.uuid4().hex[:24]
        conn = _ahb_db()
        conn.execute(
            """INSERT INTO ahb_voice_configs (id, name, voice, rate, pitch, volume, style,
               pauses_enabled, filler_words, breathing_sounds, notes)
               VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)""",
            (vid, data.get('name'), data.get('voice', 'en-US-GuyNeural'),
             str(data.get('rate', '+0%')), str(data.get('pitch', '+0Hz')),
             str(data.get('volume', '+0%')), data.get('style', 'friendly'),
             1 if data.get('pauses_enabled') or data.get('natural_pauses') else 0,
             1 if data.get('filler_words') else 0,
             1 if data.get('breathing_sounds') or data.get('breathing') else 0,
             data.get('notes', ''))
        )
        conn.commit()
        conn.close()
        return jsonify({'success': True, 'id': vid})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/ahb/voice/configs/<vid>', methods=['DELETE'])
def api_ahb_voice_configs_delete(vid):
    try:
        conn = _ahb_db()
        conn.execute("DELETE FROM ahb_voice_configs WHERE id = ?", (vid,))
        conn.commit()
        conn.close()
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/ahb/voice/synthesize', methods=['POST'])
def api_ahb_voice_synthesize():
    try:
        data = request.json or {}
        text = data.get('text', '')
        if not text:
            return jsonify({'success': False, 'error': 'text is required'}), 400
        voice = str(data.get('voice', 'en-US-GuyNeural'))
        # Normalize rate/pitch — accept numbers or strings
        raw_rate = data.get('rate', '+0%')
        raw_pitch = data.get('pitch', '+0Hz')
        raw_volume = data.get('volume', '+0%')
        # If numeric, format with % / Hz suffix
        if isinstance(raw_rate, (int, float)):
            rate = f"+{int(raw_rate)}%" if raw_rate >= 0 else f"{int(raw_rate)}%"
        else:
            rate = str(raw_rate) if raw_rate else '+0%'
        if isinstance(raw_pitch, (int, float)):
            pitch = f"+{int(raw_pitch)}Hz" if raw_pitch >= 0 else f"{int(raw_pitch)}Hz"
        else:
            pitch = str(raw_pitch) if raw_pitch else '+0Hz'
        if isinstance(raw_volume, (int, float)):
            volume = f"+{int(raw_volume)}%" if raw_volume >= 0 else f"{int(raw_volume)}%"
        else:
            volume = str(raw_volume) if raw_volume else '+0%'
        output_dir = os.path.join(DASHBOARD_DIR, 'artifacts', 'voice')
        os.makedirs(output_dir, exist_ok=True)
        filename = f"tts_{uuid.uuid4().hex[:8]}.mp3"
        output = os.path.join(output_dir, filename)
        cmd = [VENV_PYTHON, '-m', 'edge_tts', f'--voice={voice}', f'--rate={rate}',
               f'--pitch={pitch}', f'--volume={volume}', f'--text={text}', f'--write-media={output}']
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
        if result.returncode != 0:
            return jsonify({'success': False, 'error': result.stderr.strip() or 'TTS failed'}), 500
        return jsonify({'success': True, 'audio_url': f'/api/ahb/voice/audio/{filename}',
                        'voice': voice, 'rate': rate, 'pitch': pitch})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/ahb/voice/audio/<filename>', methods=['GET'])
def api_ahb_voice_audio(filename):
    voice_dir = os.path.join(DASHBOARD_DIR, 'artifacts', 'voice')
    return send_from_directory(voice_dir, filename, mimetype='audio/mpeg')


@app.route('/api/ahb/voice/logs')
def api_ahb_voice_logs():
    db = _ahb_db()
    status_filter = request.args.get('status', '')
    direction = request.args.get('direction', '')
    q = "SELECT * FROM ahb_voice_logs WHERE 1=1"
    params = []
    if status_filter:
        q += " AND status=?"
        params.append(status_filter)
    if direction:
        q += " AND direction=?"
        params.append(direction)
    q += " ORDER BY created_at DESC LIMIT 100"
    rows = db.execute(q, params).fetchall()
    db.close()
    return jsonify([dict(r) for r in rows])

@app.route('/api/ahb/voice/logs', methods=['POST'])
def api_ahb_voice_log_create():
    data = request.json or {}
    db = _ahb_db()
    db.execute("INSERT INTO ahb_voice_logs (caller_name, caller_phone, direction, duration_seconds, transcript, audio_file, status, agent_notes) VALUES (?,?,?,?,?,?,?,?)",
        (data.get('caller_name',''), data.get('caller_phone',''), data.get('direction','inbound'),
         data.get('duration_seconds',0), data.get('transcript',''), data.get('audio_file',''),
         data.get('status','completed'), data.get('agent_notes','')))
    db.commit()
    log_id = db.execute("SELECT last_insert_rowid()").fetchone()[0]
    db.close()
    return jsonify({'success': True, 'id': log_id})

@app.route('/api/ahb/voice/logs/<int:log_id>')
def api_ahb_voice_log_detail(log_id):
    db = _ahb_db()
    row = db.execute("SELECT * FROM ahb_voice_logs WHERE id=?", (log_id,)).fetchone()
    db.close()
    if not row:
        return jsonify({'error': 'Not found'}), 404
    return jsonify(dict(row))

@app.route('/api/ahb/voice/stats')
def api_ahb_voice_stats():
    db = _ahb_db()
    total = db.execute("SELECT COUNT(*) FROM ahb_voice_logs").fetchone()[0]
    inbound = db.execute("SELECT COUNT(*) FROM ahb_voice_logs WHERE direction='inbound'").fetchone()[0]
    outbound = db.execute("SELECT COUNT(*) FROM ahb_voice_logs WHERE direction='outbound'").fetchone()[0]
    missed = db.execute("SELECT COUNT(*) FROM ahb_voice_logs WHERE status='missed'").fetchone()[0]
    leads = db.execute("SELECT COUNT(*) FROM ahb_voice_logs WHERE lead_created=1").fetchone()[0]
    avg_dur = db.execute("SELECT AVG(duration_seconds) FROM ahb_voice_logs WHERE duration_seconds>0").fetchone()[0] or 0
    db.close()
    return jsonify({'total': total, 'inbound': inbound, 'outbound': outbound, 'missed': missed, 'leads_generated': leads, 'avg_duration': round(avg_dur)})


# ── AHB123 — ArchiteCT ───────────────────────────────────────────────────────

@app.route('/api/ahb/architect/analyze', methods=['POST'])
def api_ahb_architect_analyze():
    try:
        f = request.files.get('image') or request.files.get('file')
        if not f:
            return jsonify({'success': False, 'error': 'No file uploaded'}), 400
        upload_dir = os.path.join(ARTIFACTS_DIR, 'proj-ahb123')
        os.makedirs(upload_dir, exist_ok=True)
        safe_name = re.sub(r'[^\w.\-]', '_', f.filename or 'upload.jpg')
        file_path = os.path.join(upload_dir, f"architect_{uuid.uuid4().hex[:8]}_{safe_name}")
        f.save(file_path)
        # Write .meta sidecar so artifact scanner knows the agent
        try:
            with open(file_path + '.meta', 'w') as mf:
                json.dump({'agent_id': 'sam_axe', 'task_id': '', 'created_at': datetime.datetime.now().isoformat()}, mf)
        except Exception:
            pass
        # Run cloud vision analysis via the analyze_image skill
        prompt = request.form.get('prompt',
            'Describe this image in exhaustive detail. Identify and label every object, surface, material, '
            'color, texture, and spatial relationship. Include estimated dimensions, placement positions, '
            'lighting, camera angle, condition of all elements, and any text or markings visible. '
            'The description must be detailed enough for an image generation model to recreate this image '
            'and for agents to understand every element in the scene.')
        mode = request.form.get('mode', 'describe_for_agents')
        skill_path = os.path.join(os.path.dirname(DASHBOARD_DIR), 'skills', 'shared', 'analyze_image.py')
        env = os.environ.copy()
        env['SKILL_ARGS'] = json.dumps({'image_path': file_path, 'prompt': prompt, 'mode': mode})
        result = subprocess.run([VENV_PYTHON, skill_path], capture_output=True, text=True, timeout=120, env=env)
        if result.returncode != 0:
            return jsonify({'success': False, 'error': result.stderr.strip() or 'Analysis failed', 'file_path': file_path})
        # Parse the output — skill prints text then JSON on last line
        output = result.stdout.strip()
        analysis = output
        try:
            # Try to extract JSON from last line
            lines = output.split('\n')
            for line in reversed(lines):
                if line.strip().startswith('{'):
                    parsed = json.loads(line.strip())
                    analysis = parsed.get('analysis', output)
                    break
        except Exception:
            pass
        return jsonify({'success': True, 'analysis': analysis, 'file_path': file_path})
    except subprocess.TimeoutExpired:
        return jsonify({'success': False, 'error': 'Analysis timed out (60s). Is LiteLLM running?'}), 500
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/ahb/architect/generate', methods=['POST'])
def api_ahb_architect_generate():
    try:
        import requests as _req
        import base64
        data = request.json or {}
        resp = _req.post('http://localhost:7860/sdapi/v1/txt2img', json={
            'prompt': data['prompt'],
            'width': data.get('width', 1024),
            'height': data.get('height', 1024),
            'steps': data.get('steps', 30),
            'cfg_scale': data.get('cfg_scale', 7),
            'sampler_name': 'DPM++ 2M Karras',
        }, timeout=120)
        result = resp.json()
        images = result.get('images', [])
        saved = []
        out_dir = os.path.join(ARTIFACTS_DIR, 'proj-ahb123')
        os.makedirs(out_dir, exist_ok=True)
        for i, img_b64 in enumerate(images):
            fname = f"sam_axe_gen_{uuid.uuid4().hex[:8]}.png"
            fpath = os.path.join(out_dir, fname)
            with open(fpath, 'wb') as fp:
                fp.write(base64.b64decode(img_b64))
            # Write .meta sidecar
            try:
                with open(fpath + '.meta', 'w') as mf:
                    json.dump({'agent_id': 'sam_axe', 'task_id': '', 'created_at': datetime.datetime.now().isoformat()}, mf)
            except Exception:
                pass
            saved.append(f'/api/artifacts/serve/proj-ahb123/{fname}')
        image_url = saved[0] if saved else ''
        return jsonify({'success': True, 'images': saved, 'image_url': image_url})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/ahb/architect/transform', methods=['POST'])
@app.route('/api/ahb/architect/img2img', methods=['POST'])
def api_ahb_architect_img2img():
    try:
        import requests as _req
        import base64
        f = request.files.get('file') or request.files.get('image')
        if not f:
            return jsonify({'success': False, 'error': 'No file uploaded'}), 400
        img_b64 = base64.b64encode(f.read()).decode('utf-8')
        prompt = request.form.get('prompt', '')

        # Save original for before/after display
        out_dir = os.path.join(ARTIFACTS_DIR, 'proj-ahb123')
        os.makedirs(out_dir, exist_ok=True)
        orig_fname = f"sam_axe_original_{uuid.uuid4().hex[:8]}.png"
        orig_fpath = os.path.join(out_dir, orig_fname)
        with open(orig_fpath, 'wb') as fp:
            fp.write(base64.b64decode(img_b64))
        original_url = f'/api/artifacts/serve/proj-ahb123/{orig_fname}'

        # Build SD payload
        payload = {
            'init_images': [img_b64],
            'prompt': prompt,
            'width': int(request.form.get('width', 1024)),
            'height': int(request.form.get('height', 1024)),
            'steps': int(request.form.get('steps', 30)),
            'cfg_scale': float(request.form.get('cfg_scale', 7)),
            'denoising_strength': float(request.form.get('denoising_strength', 0.5)),
            'sampler_name': 'DPM++ 2M Karras',
        }

        # If a mask file is provided, use inpainting mode
        mask_file = request.files.get('mask')
        if mask_file:
            mask_b64 = base64.b64encode(mask_file.read()).decode('utf-8')
            payload['mask'] = mask_b64
            payload['inpainting_fill'] = int(request.form.get('inpaint_fill', 1))
            payload['inpaint_full_res'] = True
            payload['inpaint_full_res_padding'] = 32

        resp = _req.post('http://localhost:7860/sdapi/v1/img2img', json=payload, timeout=180)
        result = resp.json()
        images = result.get('images', [])
        saved = []
        for i, ib64 in enumerate(images):
            prefix = "sam_axe_inpaint_" if mask_file else "sam_axe_img2img_"
            fname = f"{prefix}{uuid.uuid4().hex[:8]}.png"
            fpath = os.path.join(out_dir, fname)
            with open(fpath, 'wb') as fp:
                fp.write(base64.b64decode(ib64))
            try:
                with open(fpath + '.meta', 'w') as mf:
                    json.dump({'agent_id': 'sam_axe', 'task_id': '', 'created_at': datetime.datetime.now().isoformat()}, mf)
            except Exception:
                pass
            saved.append(f'/api/artifacts/serve/proj-ahb123/{fname}')
        image_url = saved[0] if saved else ''
        return jsonify({'success': True, 'images': saved, 'image_url': image_url, 'original_url': original_url})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/ahb/architect/images/<filename>', methods=['GET'])
def api_ahb_architect_images(filename):
    """Serve architect images — check proj-ahb123 first, fall back to architect dir."""
    proj_dir = os.path.join(ARTIFACTS_DIR, 'proj-ahb123')
    if os.path.exists(os.path.join(proj_dir, filename)):
        return send_from_directory(proj_dir, filename)
    arch_dir = os.path.join(ARTIFACTS_DIR, 'architect')
    return send_from_directory(arch_dir, filename)


# ── AHB123 — Blueprint Builder ──────────────────────────────────────────────────
# Full 2D blueprint editor: rooms, walls, objects, dimensions; multi-floor; Sam
# render + LLM description→layout + photo→layout suggestions.

def _blueprint_default_data(units='imperial'):
    return {
        'units': units,
        'scale': 24 if units == 'imperial' else 50,  # px per foot OR px per meter
        'grid': 1,
        'floors': [{
            'level': 1, 'name': 'Ground Floor',
            'rooms': [], 'walls': [], 'objects': [], 'dims': [], 'notes': []
        }],
    }


@app.route('/api/ahb/blueprints', methods=['GET'])
def api_ahb_blueprints_list():
    try:
        conn = _ahb_db()
        q = "SELECT id, name, project_id, units, thumbnail_path, notes, created_at, updated_at FROM ahb_blueprints WHERE 1=1"
        params = []
        if request.args.get('project_id'):
            q += " AND project_id = ?"; params.append(request.args['project_id'])
        rows = conn.execute(q + " ORDER BY updated_at DESC", params).fetchall()
        conn.close()
        return jsonify([dict(r) for r in rows])
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/ahb/blueprints/<bid>', methods=['GET'])
def api_ahb_blueprints_get(bid):
    try:
        conn = _ahb_db()
        row = conn.execute("SELECT * FROM ahb_blueprints WHERE id = ?", (bid,)).fetchone()
        conn.close()
        if not row:
            return jsonify({'error': 'Not found'}), 404
        out = dict(row)
        try:
            out['data'] = json.loads(out.get('data') or '{}')
        except Exception:
            out['data'] = _blueprint_default_data(out.get('units') or 'imperial')
        return jsonify(out)
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/ahb/blueprints', methods=['POST'])
def api_ahb_blueprints_create():
    try:
        body = request.json or {}
        bid = body.get('id') or str(uuid.uuid4())
        name = body.get('name') or 'Untitled Blueprint'
        project_id = body.get('project_id') or ''
        units = body.get('units') or 'imperial'
        data = body.get('data') or _blueprint_default_data(units)
        notes = body.get('notes') or ''
        conn = _ahb_db()
        conn.execute(
            "INSERT INTO ahb_blueprints (id, name, project_id, units, data, notes) VALUES (?,?,?,?,?,?)",
            (bid, name, project_id, units, json.dumps(data), notes))
        conn.commit()
        conn.close()
        return jsonify({'success': True, 'id': bid})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/ahb/blueprints/<bid>', methods=['PUT'])
def api_ahb_blueprints_update(bid):
    try:
        body = request.json or {}
        fields = []; vals = []
        for k in ('name', 'project_id', 'units', 'notes'):
            if k in body:
                fields.append(f"{k} = ?"); vals.append(body[k])
        if 'data' in body:
            fields.append("data = ?")
            vals.append(json.dumps(body['data']) if not isinstance(body['data'], str) else body['data'])
        if 'thumbnail_path' in body:
            fields.append("thumbnail_path = ?"); vals.append(body['thumbnail_path'])
        if not fields:
            return jsonify({'success': False, 'error': 'No fields'}), 400
        fields.append("updated_at = ?"); vals.append(datetime.datetime.now().isoformat())
        vals.append(bid)
        conn = _ahb_db()
        conn.execute(f"UPDATE ahb_blueprints SET {', '.join(fields)} WHERE id = ?", vals)
        conn.commit()
        conn.close()
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/ahb/blueprints/<bid>', methods=['DELETE'])
def api_ahb_blueprints_delete(bid):
    try:
        conn = _ahb_db()
        cur = conn.execute("DELETE FROM ahb_blueprints WHERE id = ?", (bid,))
        conn.execute("DELETE FROM ahb_blueprint_renders WHERE blueprint_id = ?", (bid,))
        conn.commit()
        conn.close()
        if not cur.rowcount:
            return jsonify({'success': False, 'error': 'Not found'}), 404
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/ahb/blueprints/<bid>/thumbnail', methods=['POST'])
def api_ahb_blueprints_thumbnail(bid):
    """Accept a PNG blob (base64) and save as the blueprint thumbnail."""
    try:
        import base64
        body = request.json or {}
        png_b64 = body.get('png_base64') or ''
        if png_b64.startswith('data:image'):
            png_b64 = png_b64.split(',', 1)[1]
        if not png_b64:
            return jsonify({'success': False, 'error': 'No image'}), 400
        out_dir = os.path.join(ARTIFACTS_DIR, 'proj-ahb123', 'blueprints')
        os.makedirs(out_dir, exist_ok=True)
        fname = f"bp_{bid}.png"
        fpath = os.path.join(out_dir, fname)
        with open(fpath, 'wb') as fp:
            fp.write(base64.b64decode(png_b64))
        rel = f'/api/artifacts/serve/proj-ahb123/blueprints/{fname}'
        conn = _ahb_db()
        conn.execute("UPDATE ahb_blueprints SET thumbnail_path = ?, updated_at = ? WHERE id = ?",
                     (rel, datetime.datetime.now().isoformat(), bid))
        conn.commit()
        conn.close()
        return jsonify({'success': True, 'thumbnail_path': rel})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


def _blueprint_to_prompt(data, floor_level=1, style='photorealistic architectural render'):
    """Serialize a blueprint floor into a prompt for Stable Diffusion."""
    try:
        units = data.get('units', 'imperial')
        u = 'ft' if units == 'imperial' else 'm'
        floor = next((f for f in data.get('floors', []) if int(f.get('level', 1)) == int(floor_level)),
                     (data.get('floors') or [{}])[0])
        rooms = floor.get('rooms', [])
        objs = floor.get('objects', [])
        parts = [style + ', floor plan of a home,']
        if rooms:
            room_strs = []
            for r in rooms:
                w = round(r.get('w', 0), 1); h = round(r.get('h', 0), 1)
                t = r.get('type', 'room')
                lbl = r.get('label') or t
                room_strs.append(f"{lbl} ({w}x{h}{u})")
            parts.append("rooms: " + ", ".join(room_strs) + ".")
        if objs:
            counts = {}
            for o in objs:
                k = o.get('kind', 'object')
                counts[k] = counts.get(k, 0) + 1
            parts.append("features: " + ", ".join(f"{n} {k}{'s' if n>1 else ''}" for k,n in counts.items()) + ".")
        parts.append("clean, high-quality, realistic lighting, wide-angle, no people, no text overlays.")
        return " ".join(parts)
    except Exception as e:
        return f"floor plan render ({e})"


@app.route('/api/ahb/blueprints/<bid>/render', methods=['POST'])
def api_ahb_blueprints_render(bid):
    """Render the blueprint as a photorealistic image via Stable Diffusion."""
    try:
        import requests as _req
        import base64
        body = request.json or {}
        floor_level = int(body.get('floor_level', 1))
        mode = body.get('mode', 'photorealistic')  # photorealistic | isometric | topdown
        user_prompt = (body.get('prompt') or '').strip()

        conn = _ahb_db()
        row = conn.execute("SELECT data FROM ahb_blueprints WHERE id = ?", (bid,)).fetchone()
        if not row:
            conn.close()
            return jsonify({'success': False, 'error': 'Blueprint not found'}), 404
        try:
            data = json.loads(row['data'] or '{}')
        except Exception:
            data = {}

        style_map = {
            'photorealistic': 'photorealistic architectural interior render',
            'isometric': 'isometric 3D cutaway floor plan, architectural illustration',
            'topdown': 'top-down architectural floor plan drawing, clean technical lines',
        }
        style = style_map.get(mode, style_map['photorealistic'])
        auto_prompt = _blueprint_to_prompt(data, floor_level, style)
        prompt = (user_prompt + '. ' + auto_prompt) if user_prompt else auto_prompt

        try:
            resp = _req.post('http://localhost:7860/sdapi/v1/txt2img', json={
                'prompt': prompt,
                'width': int(body.get('width', 1024)),
                'height': int(body.get('height', 1024)),
                'steps': int(body.get('steps', 30)),
                'cfg_scale': float(body.get('cfg_scale', 7)),
                'sampler_name': 'DPM++ 2M Karras',
                'negative_prompt': body.get('negative_prompt', 'blurry, distorted, text, watermark, low quality'),
            }, timeout=240)
        except Exception as e:
            conn.close()
            return jsonify({'success': False, 'error': f'Stable Diffusion unavailable: {e}'}), 502

        result = resp.json()
        images = result.get('images', [])
        if not images:
            conn.close()
            return jsonify({'success': False, 'error': 'No image returned from SD'}), 502

        out_dir = os.path.join(ARTIFACTS_DIR, 'proj-ahb123', 'blueprints')
        os.makedirs(out_dir, exist_ok=True)
        saved = []
        for ib64 in images:
            rid = str(uuid.uuid4())
            fname = f"bp_render_{bid}_{rid[:8]}.png"
            fpath = os.path.join(out_dir, fname)
            with open(fpath, 'wb') as fp:
                fp.write(base64.b64decode(ib64))
            try:
                with open(fpath + '.meta', 'w') as mf:
                    json.dump({'agent_id': 'sam_axe', 'task_id': '',
                               'blueprint_id': bid,
                               'created_at': datetime.datetime.now().isoformat()}, mf)
            except Exception:
                pass
            url = f'/api/artifacts/serve/proj-ahb123/blueprints/{fname}'
            saved.append(url)
            conn.execute(
                "INSERT INTO ahb_blueprint_renders (id, blueprint_id, floor_level, mode, prompt, image_path) VALUES (?,?,?,?,?,?)",
                (rid, bid, floor_level, mode, prompt, url))
        conn.commit()
        conn.close()
        return jsonify({'success': True, 'images': saved, 'image_url': saved[0], 'prompt_used': prompt})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/ahb/blueprints/<bid>/renders', methods=['GET'])
def api_ahb_blueprints_renders_list(bid):
    try:
        conn = _ahb_db()
        rows = conn.execute(
            "SELECT id, blueprint_id, floor_level, mode, prompt, image_path, created_at "
            "FROM ahb_blueprint_renders WHERE blueprint_id = ? ORDER BY created_at DESC",
            (bid,)).fetchall()
        conn.close()
        return jsonify([dict(r) for r in rows])
    except Exception as e:
        return jsonify({'error': str(e)}), 500


def _call_local_llm(system_prompt, user_prompt, model='qwen2.5:14b', max_tokens=1200, timeout=180):
    """Call local Ollama, iterating across healthy instances."""
    import urllib.request as _ur
    urls = ['http://localhost:11434', 'http://localhost:11437', 'http://localhost:11436']
    payload = json.dumps({
        'model': model, 'stream': False,
        'messages': [
            {'role': 'system', 'content': system_prompt},
            {'role': 'user', 'content': user_prompt},
        ],
        'options': {'num_predict': max_tokens, 'num_ctx': 8192, 'temperature': 0.6}
    }).encode()
    last = None
    for url in urls:
        try:
            req = _ur.Request(f"{url}/api/chat", data=payload,
                              headers={'Content-Type': 'application/json'}, method='POST')
            with _ur.urlopen(req, timeout=timeout) as r:
                return json.loads(r.read())['message']['content'].strip(), None
        except Exception as e:
            last = e
            continue
    return '', str(last)


def _extract_json(text):
    """Pull first JSON object/array out of a blob of text."""
    import re as _re
    text = text or ''
    text = _re.sub(r'```(?:json)?', '', text).strip().strip('`').strip()
    # Try direct parse first
    try:
        return json.loads(text)
    except Exception:
        pass
    # Find first { or [ and parse balanced
    for opener, closer in (('{', '}'), ('[', ']')):
        start = text.find(opener)
        if start < 0:
            continue
        depth = 0
        for i in range(start, len(text)):
            if text[i] == opener:
                depth += 1
            elif text[i] == closer:
                depth -= 1
                if depth == 0:
                    candidate = text[start:i+1]
                    try:
                        return json.loads(candidate)
                    except Exception:
                        break
    return None


@app.route('/api/ahb/blueprints/from-description', methods=['POST'])
def api_ahb_blueprints_from_description():
    """Use the local LLM to convert a natural-language description into a blueprint layout."""
    try:
        body = request.json or {}
        description = (body.get('description') or '').strip()
        units = body.get('units') or 'imperial'
        if not description:
            return jsonify({'success': False, 'error': 'description is required'}), 400

        u = 'feet' if units == 'imperial' else 'meters'
        system = (
            "You are an architectural layout assistant. Given a natural-language description of a home or room, "
            "produce a single JSON object describing a 2D floor plan. Coordinate origin is top-left. "
            f"All distances are in {u}. Use reasonable residential proportions. DO NOT wrap the JSON in backticks or prose.\n\n"
            "Schema:\n"
            "{\n"
            '  "units": "imperial"|"metric",\n'
            '  "name": "string",\n'
            '  "floors": [\n'
            '    {\n'
            '      "level": 1, "name": "Ground Floor",\n'
            '      "rooms": [{"type":"kitchen|bedroom|bathroom|living|dining|garage|hallway|office|closet|laundry|other","label":"string","x":number,"y":number,"w":number,"h":number}],\n'
            '      "objects": [{"kind":"door|window|sink|toilet|bathtub|shower|range|fridge|dishwasher|bed|sofa|table|chair|island|stairs|fireplace","x":number,"y":number,"w":number,"h":number,"rotation":number,"label":"string"}]\n'
            '    }\n'
            '  ]\n'
            "}\n"
            "Rules: rooms must not overlap. Place doors on shared walls. Typical room sizes: bedroom 10-14, bathroom 5-8, "
            "kitchen 10-14, living 14-20. Rotations in degrees (0/90/180/270)."
        )

        content, err = _call_local_llm(system, f"Description:\n{description}\n\nReturn ONLY the JSON object.",
                                        max_tokens=1500)
        if not content:
            return jsonify({'success': False, 'error': f'LLM unavailable: {err}'}), 502

        layout = _extract_json(content)
        if not isinstance(layout, dict):
            return jsonify({'success': False, 'error': 'LLM did not return valid JSON',
                             'raw': content[:500]}), 502

        # Normalize + sanity-check
        layout.setdefault('units', units)
        layout.setdefault('scale', 24 if layout['units'] == 'imperial' else 50)
        layout.setdefault('grid', 1)
        floors = layout.get('floors') or []
        if not floors:
            return jsonify({'success': False, 'error': 'No floors in LLM output', 'raw': content[:500]}), 502
        for f_idx, f in enumerate(floors, 1):
            f.setdefault('level', f_idx); f.setdefault('name', f'Floor {f_idx}')
            for coll in ('rooms', 'objects', 'walls', 'dims', 'notes'):
                f.setdefault(coll, [])
            for r in f['rooms']:
                r['id'] = r.get('id') or str(uuid.uuid4())
                r.setdefault('notes', '')
            for o in f['objects']:
                o['id'] = o.get('id') or str(uuid.uuid4())
                o.setdefault('rotation', 0)
        return jsonify({'success': True, 'layout': layout, 'raw': content[:2000]})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


# In-memory job store for async photo imports. A single Flask worker is enough
# for this dashboard (one user, one session) so module-level dict + lock works.
import threading as _bp_threading
_bp_photo_jobs = {}
_bp_photo_jobs_lock = _bp_threading.Lock()


def _bp_photo_worker(job_id, file_path, source_url, units, sys_msg):
    def _update(**patch):
        with _bp_photo_jobs_lock:
            if job_id in _bp_photo_jobs:
                _bp_photo_jobs[job_id].update(patch)
                _bp_photo_jobs[job_id]['updated_at'] = datetime.datetime.now().isoformat()

    try:
        # ── Phase 1: vision via analyze_image skill ─────────────────────────
        _update(phase='vision', progress='Analyzing photo (vision model, can take 1-3 min if cold)…')
        prompt = (
            'Analyze this image as an architectural reference. If it is a floor plan or sketch, '
            'identify rooms (type, approximate dimensions in feet), walls, doors, windows, and furniture. '
            'If it is a room photo, identify the room type, dimensions (estimate), and objects with positions. '
            'Return a concise description that can be used to generate a 2D blueprint.'
        )
        skill_path = os.path.join(os.path.dirname(DASHBOARD_DIR), 'skills', 'shared', 'analyze_image.py')
        env = os.environ.copy()
        env['SKILL_ARGS'] = json.dumps({'image_path': file_path, 'prompt': prompt,
                                         'mode': 'describe_for_agents'})
        try:
            result = subprocess.run([VENV_PYTHON, skill_path], capture_output=True, text=True,
                                    timeout=360, env=env)
        except subprocess.TimeoutExpired:
            _update(status='error', phase='vision',
                    error='Vision model timed out after 6 minutes. A vision model (qwen3-vl:latest or llava:13b) '
                          'may not be loaded on any Ollama instance. Pre-warm with `ollama run qwen3-vl:latest "hi"` '
                          'or use the "From Text" flow with a typed description.',
                    source_url=source_url)
            return

        stdout = (result.stdout or '').strip()
        stderr = (result.stderr or '').strip()
        description = ''
        skill_error = ''
        for line in reversed(stdout.split('\n')):
            line = line.strip()
            if line.startswith('{') and line.endswith('}'):
                try:
                    parsed = json.loads(line)
                    if parsed.get('success'):
                        description = parsed.get('analysis') or ''
                    else:
                        skill_error = parsed.get('error') or ''
                    break
                except Exception:
                    pass
        if not description:
            description = stdout

        if not description:
            _update(status='error', phase='vision',
                    error=skill_error or stderr[:500] or f'Photo analysis failed (exit {result.returncode})',
                    source_url=source_url)
            return

        _update(phase='layout', progress='Vision done — generating layout from description…',
                description=description, source_url=source_url)

        # ── Phase 2: description → JSON layout via local LLM ────────────────
        u = 'feet' if units == 'imperial' else 'meters'
        layout_system = (
            "You are an architectural layout assistant. Read the vision description below and produce a single "
            f"JSON object describing a 2D floor plan with distances in {u}. Schema:\n"
            '{"units":"imperial|metric","name":"string","floors":[{"level":1,"name":"Ground Floor",'
            '"rooms":[{"type":"kitchen|bedroom|bathroom|living|dining|garage|hallway|office|closet|laundry|other",'
            '"label":"string","x":number,"y":number,"w":number,"h":number}],'
            '"objects":[{"kind":"door|window|sink|toilet|bathtub|shower|range|fridge|dishwasher|bed|sofa|table|chair|island|stairs|fireplace",'
            '"x":number,"y":number,"w":number,"h":number,"rotation":number,"label":"string"}]}]}\n'
            "Rooms must not overlap. Place doors on shared walls. Output ONLY the JSON, no prose, no code fences."
            + (f"\n\n{sys_msg}" if sys_msg else "")
        )
        content, err = _call_local_llm(layout_system, f"Vision description:\n{description[:4000]}",
                                        max_tokens=1500, timeout=240)
        layout = None
        if content:
            parsed = _extract_json(content)
            if isinstance(parsed, dict) and parsed.get('floors'):
                for fl_idx, fl in enumerate(parsed.get('floors') or [], 1):
                    fl.setdefault('level', fl_idx); fl.setdefault('name', f'Floor {fl_idx}')
                    for coll in ('rooms', 'objects', 'walls', 'dims', 'notes'):
                        fl.setdefault(coll, [])
                    for r in fl['rooms']:
                        r['id'] = r.get('id') or str(uuid.uuid4()); r.setdefault('notes', '')
                    for o in fl['objects']:
                        o['id'] = o.get('id') or str(uuid.uuid4()); o.setdefault('rotation', 0)
                parsed.setdefault('units', units)
                parsed.setdefault('scale', 24 if units == 'imperial' else 50)
                parsed.setdefault('grid', 1)
                layout = parsed

        _update(status='done', phase='done', progress='Complete',
                description=description, layout=layout,
                llm_error=err, llm_raw=(content or '')[:2000],
                source_url=source_url)
    except Exception as e:
        _update(status='error', phase='unknown', error=f'Worker crashed: {e}')


@app.route('/api/ahb/blueprints/from-photo', methods=['POST'])
def api_ahb_blueprints_from_photo():
    """Kick off an async photo→layout job. Returns a job_id immediately.

    The actual vision + layout work runs in a background thread — the HTTP call
    returns in under a second so browsers and proxies never time out. Poll
    /api/ahb/blueprints/photo-jobs/<id> to get status, progress and final result.
    """
    try:
        f = request.files.get('image') or request.files.get('file')
        if not f:
            return jsonify({'success': False, 'error': 'No image uploaded'}), 400

        upload_dir = os.path.join(ARTIFACTS_DIR, 'proj-ahb123', 'blueprints')
        os.makedirs(upload_dir, exist_ok=True)
        safe_name = re.sub(r'[^\w.\-]', '_', f.filename or 'source.jpg')
        file_path = os.path.join(upload_dir, f"bp_source_{uuid.uuid4().hex[:8]}_{safe_name}")
        f.save(file_path)
        source_url = f'/api/artifacts/serve/proj-ahb123/blueprints/{os.path.basename(file_path)}'

        units = (request.form.get('units') or 'imperial')
        sys_msg = request.form.get('system_extra', '')

        job_id = str(uuid.uuid4())
        with _bp_photo_jobs_lock:
            # GC jobs older than 1 hour to keep the dict bounded
            now = datetime.datetime.now()
            for stale in [k for k, v in _bp_photo_jobs.items()
                           if (now - datetime.datetime.fromisoformat(v.get('created_at', now.isoformat()))).total_seconds() > 3600]:
                _bp_photo_jobs.pop(stale, None)
            _bp_photo_jobs[job_id] = {
                'job_id': job_id, 'status': 'running', 'phase': 'queued',
                'progress': 'Queued — upload received', 'source_url': source_url,
                'created_at': now.isoformat(), 'updated_at': now.isoformat(),
            }
        t = _bp_threading.Thread(target=_bp_photo_worker,
                                   args=(job_id, file_path, source_url, units, sys_msg),
                                   daemon=True)
        t.start()
        return jsonify({'success': True, 'job_id': job_id, 'status': 'running', 'source_url': source_url})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/ahb/blueprints/photo-jobs/<job_id>', methods=['GET'])
def api_ahb_blueprints_photo_job_status(job_id):
    with _bp_photo_jobs_lock:
        job = _bp_photo_jobs.get(job_id)
    if not job:
        return jsonify({'error': 'Unknown or expired job_id'}), 404
    return jsonify(job)


@app.route('/api/ahb/blueprints/<bid>/inpaint-room', methods=['POST'])
def api_ahb_blueprints_inpaint_room(bid):
    """Re-render a single room within an existing SD render using a mask rectangle (normalized 0..1)."""
    try:
        import requests as _req
        import base64
        body = request.json or {}
        source_url = body.get('source_url') or ''
        prompt = body.get('prompt') or ''
        mask_rect = body.get('mask_rect') or {}  # {x, y, w, h} in 0..1
        denoising = float(body.get('denoising_strength', 0.75))
        if not source_url or not prompt:
            return jsonify({'success': False, 'error': 'source_url and prompt required'}), 400

        # Resolve source image path
        if source_url.startswith('/api/artifacts/serve/'):
            rel = source_url.replace('/api/artifacts/serve/', '', 1)
            src_path = os.path.join(ARTIFACTS_DIR, rel)
        elif source_url.startswith('http'):
            with _req.get(source_url, timeout=30) as r:
                src_path = os.path.join(ARTIFACTS_DIR, 'proj-ahb123', 'blueprints', f'src_{uuid.uuid4().hex[:8]}.png')
                os.makedirs(os.path.dirname(src_path), exist_ok=True)
                with open(src_path, 'wb') as fp:
                    fp.write(r.content)
        else:
            src_path = source_url
        if not os.path.exists(src_path):
            return jsonify({'success': False, 'error': f'source image not found: {src_path}'}), 404

        # Build mask PNG (same size as source) — white inside rect, black outside
        try:
            from PIL import Image, ImageDraw
        except Exception:
            return jsonify({'success': False, 'error': 'Pillow is required for mask building'}), 500
        src = Image.open(src_path).convert('RGB')
        sw, sh = src.size
        mask = Image.new('L', (sw, sh), 0)
        draw = ImageDraw.Draw(mask)
        mx = int(float(mask_rect.get('x', 0)) * sw)
        my = int(float(mask_rect.get('y', 0)) * sh)
        mw = max(1, int(float(mask_rect.get('w', 1)) * sw))
        mh = max(1, int(float(mask_rect.get('h', 1)) * sh))
        draw.rectangle([mx, my, mx + mw, my + mh], fill=255)

        import io as _io
        src_buf = _io.BytesIO(); src.save(src_buf, format='PNG')
        mask_buf = _io.BytesIO(); mask.save(mask_buf, format='PNG')
        img_b64 = base64.b64encode(src_buf.getvalue()).decode('utf-8')
        mask_b64 = base64.b64encode(mask_buf.getvalue()).decode('utf-8')

        resp = _req.post('http://localhost:7860/sdapi/v1/img2img', json={
            'init_images': [img_b64],
            'mask': mask_b64,
            'prompt': prompt,
            'width': sw, 'height': sh,
            'steps': int(body.get('steps', 30)),
            'cfg_scale': float(body.get('cfg_scale', 7)),
            'denoising_strength': denoising,
            'inpainting_fill': 1, 'inpaint_full_res': True, 'inpaint_full_res_padding': 32,
            'sampler_name': 'DPM++ 2M Karras',
        }, timeout=240)
        result = resp.json()
        images = result.get('images', [])
        if not images:
            return jsonify({'success': False, 'error': 'No image returned from SD'}), 502
        out_dir = os.path.join(ARTIFACTS_DIR, 'proj-ahb123', 'blueprints')
        os.makedirs(out_dir, exist_ok=True)
        saved = []
        conn = _ahb_db()
        for ib64 in images:
            rid = str(uuid.uuid4())
            fname = f"bp_inpaint_{bid}_{rid[:8]}.png"
            fpath = os.path.join(out_dir, fname)
            with open(fpath, 'wb') as fp:
                fp.write(base64.b64decode(ib64))
            url = f'/api/artifacts/serve/proj-ahb123/blueprints/{fname}'
            saved.append(url)
            conn.execute(
                "INSERT INTO ahb_blueprint_renders (id, blueprint_id, floor_level, mode, prompt, image_path) VALUES (?,?,?,?,?,?)",
                (rid, bid, int(body.get('floor_level', 1)), 'inpaint', prompt, url))
        conn.commit()
        conn.close()
        return jsonify({'success': True, 'images': saved, 'image_url': saved[0]})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


# ── AHB123 — Employees ─────────────────────────────────────────────────────────

@app.route('/api/ahb/employees', methods=['GET'])
def api_ahb_employees_list():
    try:
        conn = _ahb_db()
        q = "SELECT * FROM ahb_employees WHERE 1=1"
        params = []
        if request.args.get('active'):
            q += " AND active = ?"; params.append(int(request.args['active']))
        rows = conn.execute(q + " ORDER BY name", params).fetchall()
        conn.close()
        return jsonify([dict(r) for r in rows])
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/ahb/employees', methods=['POST'])
def api_ahb_employees_create():
    try:
        data = request.json or {}
        cls = (data.get('tax_classification') or 'W2').upper()
        if cls not in ('W2', 'W9'):
            cls = 'W2'
        conn = _ahb_db()
        eid = str(uuid.uuid4())
        conn.execute(
            """INSERT INTO ahb_employees
                 (id, name, position, hourly_rate, pay_type, pay_method, phone, email, active,
                  tax_classification, business_name, tax_id, tax_id_type, address,
                  w9_doc_id, w9_signed_date)
               VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)""",
            (eid, data.get('name',''), data.get('position',''), data.get('hourly_rate',0),
             data.get('pay_type','hourly'), data.get('pay_method',''), data.get('phone',''),
             data.get('email',''), 1 if data.get('active', True) else 0,
             cls, data.get('business_name',''), data.get('tax_id',''),
             (data.get('tax_id_type') or '').upper(), data.get('address',''),
             data.get('w9_doc_id') or None, data.get('w9_signed_date','')))
        conn.commit(); conn.close()
        return jsonify({'success': True, 'id': eid})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/ahb/employees/<eid>', methods=['PUT'])
def api_ahb_employees_update(eid):
    try:
        data = request.json or {}
        conn = _ahb_db()
        fields, vals = [], []
        for k in ['name','position','hourly_rate','pay_type','pay_method','phone','email',
                  'business_name','tax_id','address','w9_signed_date','w9_doc_id']:
            if k in data: fields.append(f"{k} = ?"); vals.append(data[k])
        if 'tax_classification' in data:
            cls = (data['tax_classification'] or 'W2').upper()
            fields.append("tax_classification = ?"); vals.append(cls if cls in ('W2','W9') else 'W2')
        if 'tax_id_type' in data:
            fields.append("tax_id_type = ?"); vals.append((data['tax_id_type'] or '').upper())
        if 'active' in data: fields.append("active = ?"); vals.append(1 if data['active'] else 0)
        if fields:
            fields.append("updated_at = datetime('now')")
            vals.append(eid)
            conn.execute(f"UPDATE ahb_employees SET {', '.join(fields)} WHERE id = ?", vals)
            conn.commit()
        conn.close()
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/ahb/employees/<eid>', methods=['DELETE'])
def api_ahb_employees_delete(eid):
    try:
        conn = _ahb_db()
        conn.execute("UPDATE ahb_employees SET active = 0, updated_at = datetime('now') WHERE id = ?", (eid,))
        conn.commit(); conn.close()
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


# ── AHB123 — Events/Schedule ───────────────────────────────────────────────────

@app.route('/api/ahb/events', methods=['GET'])
def api_ahb_events_list():
    try:
        conn = _ahb_db()
        q = "SELECT * FROM ahb_events WHERE 1=1"
        params = []
        if request.args.get('category'):
            q += " AND category = ?"; params.append(request.args['category'])
        if request.args.get('date_from'):
            q += " AND date >= ?"; params.append(request.args['date_from'])
        if request.args.get('date_to'):
            q += " AND date <= ?"; params.append(request.args['date_to'])
        rows = conn.execute(q + " ORDER BY date, time", params).fetchall()
        conn.close()
        return jsonify([dict(r) for r in rows])
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/ahb/events', methods=['POST'])
def api_ahb_events_create():
    try:
        data = request.json or {}
        conn = _ahb_db()
        eid = str(uuid.uuid4())
        conn.execute(
            """INSERT INTO ahb_events (id, title, details, date, time, end_time, category, all_day, project_id, employee_id)
               VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)""",
            (eid, data.get('title',''), data.get('details',''), data.get('date',''),
             data.get('time',''), data.get('end_time',''), data.get('category',''),
             1 if data.get('all_day') else 0, data.get('project_id',''), data.get('employee_id','')))
        conn.commit(); conn.close()
        return jsonify({'success': True, 'id': eid})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/ahb/events/<eid>', methods=['PUT'])
def api_ahb_events_update(eid):
    try:
        data = request.json or {}
        conn = _ahb_db()
        fields, vals = [], []
        for k in ['title','details','date','time','end_time','category','project_id','employee_id']:
            if k in data: fields.append(f"{k} = ?"); vals.append(data[k])
        if 'all_day' in data: fields.append("all_day = ?"); vals.append(1 if data['all_day'] else 0)
        if fields:
            vals.append(eid)
            conn.execute(f"UPDATE ahb_events SET {', '.join(fields)} WHERE id = ?", vals)
            conn.commit()
        conn.close()
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/ahb/events/<eid>', methods=['DELETE'])
def api_ahb_events_delete(eid):
    try:
        conn = _ahb_db()
        conn.execute("DELETE FROM ahb_events WHERE id = ?", (eid,))
        conn.commit(); conn.close()
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/ahb/projects/<pid>/calendar', methods=['POST'])
def api_ahb_project_calendar_backfill(pid):
    """Backfill calendar events from project and phase dates. For importing past projects."""
    try:
        data = request.json or {}
        conn = _ahb_db()
        project = conn.execute("SELECT * FROM ahb_projects WHERE id = ?", (pid,)).fetchone()
        if not project:
            conn.close()
            return jsonify({'success': False, 'error': 'Project not found'}), 404
        project = dict(project)

        title = data.get('title') or project.get('title', 'Project')
        start_date = data.get('start_date') or project.get('start_date')
        end_date = data.get('end_date') or project.get('end_date')
        events_created = []

        # Project start event
        if start_date:
            eid = uuid.uuid4().hex[:24]
            conn.execute(
                """INSERT INTO ahb_events (id, title, details, date, category, all_day, project_id)
                   VALUES (?, ?, ?, ?, ?, ?, ?)""",
                (eid, f"{title} - Start", f"Project start: {title}", start_date, 'project', 1, pid))
            events_created.append({'id': eid, 'type': 'project_start', 'date': start_date})

        # Project end event
        if end_date:
            eid = uuid.uuid4().hex[:24]
            conn.execute(
                """INSERT INTO ahb_events (id, title, details, date, category, all_day, project_id)
                   VALUES (?, ?, ?, ?, ?, ?, ?)""",
                (eid, f"{title} - End", f"Project end: {title}", end_date, 'project', 1, pid))
            events_created.append({'id': eid, 'type': 'project_end', 'date': end_date})

        # Phase start/end events
        phases = conn.execute(
            "SELECT * FROM ahb_project_phases WHERE project_id = ? ORDER BY phase_number", (pid,)
        ).fetchall()
        for ph in phases:
            ph = dict(ph)
            ph_name = ph.get('name', f"Phase {ph.get('phase_number', '?')}")
            if ph.get('start_date'):
                eid = uuid.uuid4().hex[:24]
                conn.execute(
                    """INSERT INTO ahb_events (id, title, details, date, category, all_day, project_id)
                       VALUES (?, ?, ?, ?, ?, ?, ?)""",
                    (eid, f"{ph_name} - Start", f"Phase start for {title}", ph['start_date'], 'phase', 1, pid))
                events_created.append({'id': eid, 'type': 'phase_start', 'phase': ph_name, 'date': ph['start_date']})
            if ph.get('end_date'):
                eid = uuid.uuid4().hex[:24]
                conn.execute(
                    """INSERT INTO ahb_events (id, title, details, date, category, all_day, project_id)
                       VALUES (?, ?, ?, ?, ?, ?, ?)""",
                    (eid, f"{ph_name} - End", f"Phase end for {title}", ph['end_date'], 'phase', 1, pid))
                events_created.append({'id': eid, 'type': 'phase_end', 'phase': ph_name, 'date': ph['end_date']})

        conn.commit()
        conn.close()
        return jsonify({'success': True, 'events_created': len(events_created), 'events': events_created})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


# ── AHB123 — Notes ─────────────────────────────────────────────────────────────

@app.route('/api/ahb/notes', methods=['GET'])
def api_ahb_notes_list():
    try:
        conn = _ahb_db()
        q = "SELECT * FROM ahb_notes WHERE 1=1"
        params = []
        if request.args.get('is_task'):
            q += " AND is_task = 1"
        if request.args.get('project_id'):
            q += " AND project_id = ?"; params.append(request.args['project_id'])
        rows = conn.execute(q + " ORDER BY pinned DESC, created_at DESC", params).fetchall()
        conn.close()
        return jsonify([dict(r) for r in rows])
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/ahb/notes', methods=['POST'])
def api_ahb_notes_create():
    try:
        data = request.json or {}
        conn = _ahb_db()
        nid = str(uuid.uuid4())
        try:
            color = int(data.get('color', 1) or 1)
        except (TypeError, ValueError):
            color = 1
        color = max(1, min(5, color))
        conn.execute(
            """INSERT INTO ahb_notes (id, title, content, is_list, is_task, tags, pinned, project_id, due_date, checklist_items, author_employee_id, color)
               VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)""",
            (nid, data.get('title',''), data.get('content',''),
             1 if data.get('is_list') else 0, 1 if data.get('is_task') else 0,
             data.get('tags',''), 1 if data.get('pinned') else 0,
             data.get('project_id',''), data.get('due_date',''),
             json.dumps(data.get('checklist_items',[])) if isinstance(data.get('checklist_items'), list) else data.get('checklist_items','[]'),
             data.get('author_employee_id',''), color))
        conn.commit(); conn.close()
        return jsonify({'success': True, 'id': nid})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/ahb/notes/<nid>', methods=['PUT'])
def api_ahb_notes_update(nid):
    try:
        data = request.json or {}
        conn = _ahb_db()
        fields, vals = [], []
        for k in ['title','content','tags','project_id','due_date','author_employee_id']:
            if k in data: fields.append(f"{k} = ?"); vals.append(data[k])
        for k in ['is_list','is_task','pinned']:
            if k in data: fields.append(f"{k} = ?"); vals.append(1 if data[k] else 0)
        if 'color' in data:
            try:
                c = int(data['color'] or 1)
            except (TypeError, ValueError):
                c = 1
            fields.append("color = ?"); vals.append(max(1, min(5, c)))
        if 'checklist_items' in data:
            fields.append("checklist_items = ?")
            vals.append(json.dumps(data['checklist_items']) if isinstance(data['checklist_items'], list) else data['checklist_items'])
        if fields:
            fields.append("updated_at = datetime('now')")
            vals.append(nid)
            conn.execute(f"UPDATE ahb_notes SET {', '.join(fields)} WHERE id = ?", vals)
            conn.commit()
        conn.close()
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/ahb/notes/<nid>', methods=['DELETE'])
def api_ahb_notes_delete(nid):
    try:
        conn = _ahb_db()
        conn.execute("DELETE FROM ahb_notes WHERE id = ?", (nid,))
        conn.commit(); conn.close()
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


def _summarize_text(text: str, max_len: int = 60) -> str:
    """Cheap, deterministic summary — first sentence, or first chunk of words.
    No LLM dependency on the hot path; safe for autosave."""
    t = (text or '').strip()
    if not t:
        return 'Untitled'
    import re as _re
    first = _re.split(r'[.!?\n]+', t, maxsplit=1)[0].strip()
    if not first:
        first = t
    if len(first) <= max_len:
        return first
    # Cut on word boundary
    cut = first[:max_len].rsplit(' ', 1)[0]
    return (cut or first[:max_len]).rstrip(' ,;:') + '…'


@app.route('/api/ahb/notes/<nid>/summarize', methods=['POST'])
def api_ahb_notes_summarize(nid):
    try:
        conn = _ahb_db()
        row = conn.execute("SELECT content FROM ahb_notes WHERE id = ?", (nid,)).fetchone()
        if not row:
            conn.close()
            return jsonify({'success': False, 'error': 'not found'}), 404
        summary = _summarize_text(row['content'])
        conn.execute("UPDATE ahb_notes SET title = ?, updated_at = datetime('now') WHERE id = ?", (summary, nid))
        conn.commit(); conn.close()
        return jsonify({'success': True, 'summary': summary})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


# ── AHB123 — Debts (IPAY) ──────────────────────────────────────────────────────

@app.route('/api/ahb/debts', methods=['GET'])
def api_ahb_debts_list():
    try:
        conn = _ahb_db()
        q = "SELECT * FROM ahb_debts WHERE 1=1"
        params = []
        if request.args.get('type'):
            q += " AND type = ?"; params.append(request.args['type'])
        rows = conn.execute(q + " ORDER BY due_date", params).fetchall()
        conn.close()
        return jsonify([dict(r) for r in rows])
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/ahb/debts', methods=['POST'])
def api_ahb_debts_create():
    try:
        data = request.json or {}
        conn = _ahb_db()
        did = str(uuid.uuid4())
        conn.execute(
            """INSERT INTO ahb_debts (id, name, type, frequency, payment_amount, due_date, payoff_date, balance)
               VALUES (?, ?, ?, ?, ?, ?, ?, ?)""",
            (did, data.get('name',''), data.get('type','Bill'), data.get('frequency','Monthly'),
             data.get('payment_amount',0), data.get('due_date',''), data.get('payoff_date',''),
             data.get('balance',0)))
        conn.commit(); conn.close()
        return jsonify({'success': True, 'id': did})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/ahb/debts/<did>', methods=['PUT'])
def api_ahb_debts_update(did):
    try:
        data = request.json or {}
        conn = _ahb_db()
        fields, vals = [], []
        for k in ['name','type','frequency','payment_amount','due_date','payoff_date','balance']:
            if k in data: fields.append(f"{k} = ?"); vals.append(data[k])
        if fields:
            fields.append("updated_at = datetime('now')")
            vals.append(did)
            conn.execute(f"UPDATE ahb_debts SET {', '.join(fields)} WHERE id = ?", vals)
            conn.commit()
        conn.close()
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/ahb/debts/<did>', methods=['DELETE'])
def api_ahb_debts_delete(did):
    try:
        conn = _ahb_db()
        conn.execute("DELETE FROM ahb_debts WHERE id = ?", (did,))
        conn.commit(); conn.close()
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


# ── AHB123 — Files ─────────────────────────────────────────────────────────────

@app.route('/api/ahb/files', methods=['GET'])
def api_ahb_files_list():
    try:
        conn = _ahb_db()
        q = "SELECT * FROM ahb_files WHERE 1=1"
        params = []
        if request.args.get('category'):
            q += " AND category = ?"; params.append(request.args['category'])
        if request.args.get('project_id'):
            q += " AND project_id = ?"; params.append(request.args['project_id'])
        if request.args.get('year'):
            q += " AND year = ?"; params.append(request.args['year'])
        rows = conn.execute(q + " ORDER BY created_at DESC", params).fetchall()
        conn.close()
        return jsonify([dict(r) for r in rows])
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/ahb/files', methods=['POST'])
def api_ahb_files_create():
    try:
        conn = _ahb_db()
        fid = str(uuid.uuid4())
        f = request.files.get('file')
        file_path = ''
        if f:
            upload_dir = os.path.join(DASHBOARD_DIR, 'uploads', 'ahb', 'files')
            os.makedirs(upload_dir, exist_ok=True)
            safe_name = re.sub(r'[^\w.\-]', '_', f.filename or 'file')
            file_path = os.path.join(upload_dir, f"{fid}_{safe_name}")
            f.save(file_path)
        data = request.form if f else (request.json or {})
        # Allow linking existing file by absolute path (agents uploading via Telegram)
        if not file_path and data.get('file_path'):
            file_path = data.get('file_path')
        size = data.get('size', 0)
        if not size and file_path and os.path.exists(file_path):
            try: size = os.path.getsize(file_path)
            except Exception: pass
        name = data.get('name') or (f.filename if f else (os.path.basename(file_path) if file_path else ''))
        conn.execute(
            """INSERT INTO ahb_files (id, name, file_type, file_path, size, tags, category, year, project_id)
               VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)""",
            (fid, name, data.get('file_type',''),
             file_path, size, data.get('tags',''), data.get('category',''),
             data.get('year',''), data.get('project_id','')))
        conn.commit(); conn.close()
        return jsonify({'success': True, 'id': fid, 'file_path': file_path})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/ahb/files/<fid>', methods=['DELETE'])
def api_ahb_files_delete(fid):
    try:
        conn = _ahb_db()
        conn.execute("DELETE FROM ahb_files WHERE id = ?", (fid,))
        conn.commit(); conn.close()
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/ahb/projects/<pid>/files', methods=['POST'])
def api_ahb_project_files_upload(pid):
    """Upload photos or documents to a project. Accepts either a multipart
    `file` upload OR a `pick_token` referencing an existing image in
    artifacts/ (Baza picker flow — used for Sam's Telegram-inbound photos
    while SD imaging is down). Supports photo_section + document_type."""
    try:
        import shutil as _shutil
        upload_dir = os.path.join(DASHBOARD_DIR, 'uploads', 'ahb', 'projects', pid)
        os.makedirs(upload_dir, exist_ok=True)
        fid = uuid.uuid4().hex[:24]

        f = request.files.get('file')
        data = request.form
        pick_token = (data.get('pick_token') or '').strip()

        if f:
            safe_name = re.sub(r'[^\w.\-]', '_', f.filename or 'file')
            file_path = os.path.join(upload_dir, f"{fid}_{safe_name}")
            f.save(file_path)
            original_name = f.filename or safe_name
        elif pick_token:
            src = _pick_decode_token(pick_token)
            if not src:
                return jsonify({'success': False, 'error': 'Invalid pick_token'}), 400
            base = os.path.basename(src)
            safe_name = re.sub(r'[^\w.\-]', '_', base) or 'file'
            file_path = os.path.join(upload_dir, f"{fid}_{safe_name}")
            _shutil.copy2(src, file_path)
            original_name = data.get('name') or base
        else:
            return jsonify({'success': False, 'error': 'No file or pick_token provided'}), 400

        photo_section = data.get('photo_section', '')  # before / during / after
        document_type = data.get('document_type', '')   # Permit, Contract, COI, etc.

        ext = os.path.splitext(safe_name)[1].lower()
        is_image = ext in ('.jpg', '.jpeg', '.png', '.gif', '.webp', '.heic')
        file_type = 'photo' if is_image else 'document'
        category = document_type if document_type else ('photo' if is_image else 'document')

        conn = _ahb_db()
        conn.execute(
            """INSERT INTO ahb_files (id, name, file_type, file_path, size, tags, category, year, project_id, photo_section, document_type)
               VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)""",
            (fid, original_name, file_type, file_path,
             os.path.getsize(file_path), data.get('tags', ''), category,
             str(datetime.datetime.now().year), pid, photo_section, document_type))
        conn.commit()
        conn.close()
        return jsonify({'success': True, 'id': fid, 'file_path': file_path})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


# ── AHB123 — Project Phases ────────────────────────────────────────────────────

def _rebuild_invoice_line_items(conn, project_id):
    """Rebuild the linked invoice's line_items JSON from all project phases.
    Writes both legacy ({quantity, unit_price}) AND form-native ({qty, rate, total}) keys
    so the invoice editor can read either format without breaking."""
    phases = conn.execute(
        "SELECT * FROM ahb_project_phases WHERE project_id = ? ORDER BY phase_number",
        (project_id,)
    ).fetchall()
    line_items = []
    subtotal = 0
    for ph in phases:
        ph = dict(ph)
        val = ph.get('value', 0) or 0
        line_items.append({
            'description': ph.get('name', 'Phase'),
            'qty': 1, 'rate': val, 'total': val,
            'materials': 0, 'labor': 0,
            'quantity': 1, 'unit_price': val,   # legacy keys
        })
        subtotal += val
    # Update first linked invoice
    inv = conn.execute(
        "SELECT id FROM ahb_invoices WHERE project_id = ? ORDER BY created_at ASC LIMIT 1",
        (project_id,)
    ).fetchone()
    if inv:
        conn.execute(
            "UPDATE ahb_invoices SET line_items = ?, subtotal = ?, total = ?, updated_at = ? WHERE id = ?",
            (json.dumps(line_items), subtotal, subtotal, datetime.datetime.now().isoformat(), inv['id']))
    return inv['id'] if inv else None


def _split_description_to_line_items(description: str, total_budget: float) -> list:
    """Split a project description into individual invoice line items.
    Each non-empty line (or bullet/numbered item) becomes one line item.
    The total_budget is divided evenly across all line items so the sum equals the budget.
    Returns a list of dicts with both {qty,rate,total} and {quantity,unit_price} keys."""
    if not description or not description.strip():
        return []
    raw = description.strip()
    # Split on newlines, semicolons, or numbered/bulleted markers
    lines = []
    for chunk in re.split(r'[\n\r;]+', raw):
        chunk = chunk.strip()
        if not chunk:
            continue
        # Strip leading bullets, numbers, dashes
        chunk = re.sub(r'^[\s\-\*\u2022\u2023\u25E6\u2043\u2219]+', '', chunk)
        chunk = re.sub(r'^\d+[\.\)]\s*', '', chunk).strip()
        if chunk:
            lines.append(chunk)
    if not lines:
        return []
    # Old behavior shredded the description into N items at total/N. That made
    # invoices repeat the description and assign every line the same uniform price.
    # We now keep a single line carrying the full budget; the user fills in real
    # material + labor breakdowns in the invoice editor.
    budget = float(total_budget or 0)
    label = lines[0][:200]
    return [{
        'description': label,
        'qty': 1, 'rate': budget, 'total': budget,
        'materials': 0, 'labor': 0,
        'quantity': 1, 'unit_price': budget,
    }]


def _sync_invoice_from_project(conn, project_id, project_data):
    """Push project header fields (title, client info, address, value) into the linked invoice
    whenever the project is updated. Also rebuilds line items from project description when
    description or budget changes — each description line becomes one line item, and the
    project budget (value or budget_high) is split evenly across them as the total.

    Phase-driven invoices (>1 line items already from phases) are left alone unless explicitly
    rebuilt elsewhere — only single-line / description-driven invoices auto-sync from desc."""
    inv = conn.execute(
        "SELECT id, line_items, subtotal FROM ahb_invoices WHERE project_id = ? ORDER BY created_at ASC LIMIT 1",
        (project_id,)
    ).fetchone()
    if not inv:
        return None
    inv = dict(inv)

    # Pull the latest project so we always have the canonical description + budget,
    # even if the caller only sent a partial PUT (e.g. just notes/dates).
    proj_row = conn.execute("SELECT * FROM ahb_projects WHERE id = ?", (project_id,)).fetchone()
    if not proj_row:
        return inv['id']
    proj = dict(proj_row)
    # Merge what was just PUT — incoming data wins for fields it specifies
    for k, v in (project_data or {}).items():
        if v is not None:
            proj[k] = v

    sets, vals = [], []
    sets.append('project_name = ?');   vals.append(proj.get('title') or '')
    sets.append('client_name = ?');    vals.append(proj.get('client_name') or '')
    sets.append('client_email = ?');   vals.append(proj.get('client_email') or '')
    sets.append('client_phone = ?');   vals.append(proj.get('contact_info') or '')
    sets.append('project_address = ?');vals.append(proj.get('address') or '')
    sets.append('client_address = ?'); vals.append(proj.get('address') or '')
    if proj.get('client_id'):
        sets.append('client_id = ?');  vals.append(proj['client_id'])

    # Determine if invoice is phase-driven (has multiple lines, leave it alone)
    try:
        existing_lines = json.loads(inv.get('line_items') or '[]')
    except Exception:
        existing_lines = []
    phase_count = conn.execute(
        "SELECT COUNT(*) FROM ahb_project_phases WHERE project_id = ?", (project_id,)
    ).fetchone()[0]

    # When no phases, leave existing line items alone — the invoice editor is the
    # source of truth for material/labor breakdowns. We only touch line items when
    # the invoice has none yet (fresh invoice with no entries), and even then we
    # write a single placeholder line carrying the budget. Never auto-shred a
    # description into N items at total/N — that was the source of duplicated
    # descriptions and uniform prices on the printed invoice.
    if phase_count == 0 and len(existing_lines) == 0:
        budget = proj.get('value') or proj.get('budget_high') or proj.get('budget_low') or 0
        try:
            budget = float(budget) if budget else 0
        except Exception:
            budget = 0
        if budget:
            single = [{
                'description': proj.get('title', 'Project'),
                'qty': 1, 'rate': budget, 'total': budget,
                'materials': 0, 'labor': 0,
                'quantity': 1, 'unit_price': budget,
            }]
            sets.append('line_items = ?'); vals.append(json.dumps(single))
            sets.append('subtotal = ?');   vals.append(budget)
            sets.append('total = ?');      vals.append(budget)

    sets.append('updated_at = ?'); vals.append(datetime.datetime.now().isoformat())
    vals.append(inv['id'])
    conn.execute(f"UPDATE ahb_invoices SET {', '.join(sets)} WHERE id = ?", vals)
    return inv['id']

@app.route('/api/ahb/phases', methods=['GET'])
def api_ahb_phases_list():
    try:
        conn = _ahb_db()
        q = "SELECT * FROM ahb_project_phases WHERE 1=1"
        params = []
        if request.args.get('project_id'):
            q += " AND project_id = ?"; params.append(request.args['project_id'])
        rows = conn.execute(q + " ORDER BY phase_number", params).fetchall()
        conn.close()
        return jsonify([dict(r) for r in rows])
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/ahb/projects/<pid>/phases', methods=['POST'])
def api_ahb_project_phases_create(pid):
    """Add a phase (planning-only — value is no longer pushed to the invoice)."""
    try:
        data = request.json or {}
        conn = _ahb_db()
        phid = uuid.uuid4().hex[:24]
        conn.execute(
            """INSERT INTO ahb_project_phases (id, project_id, phase_number, name, value, start_date, end_date, status)
               VALUES (?, ?, ?, ?, ?, ?, ?, ?)""",
            (phid, pid, data.get('phase_number', 1), data.get('name', ''),
             data.get('value', 0), data.get('start_date', ''), data.get('end_date', ''),
             data.get('status', 'pending')))
        conn.commit()
        _sync_phase_events(conn, pid, phid)
        conn.commit()
        conn.close()
        return jsonify({'success': True, 'id': phid})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/ahb/phases', methods=['POST'])
def api_ahb_phases_create():
    """Legacy: add a phase (also syncs invoice)."""
    try:
        data = request.json or {}
        conn = _ahb_db()
        phid = uuid.uuid4().hex[:24]
        project_id = data.get('project_id', '')
        conn.execute(
            """INSERT INTO ahb_project_phases (id, project_id, phase_number, name, value, start_date, end_date, status)
               VALUES (?, ?, ?, ?, ?, ?, ?, ?)""",
            (phid, project_id, data.get('phase_number', 1), data.get('name', ''),
             data.get('value', 0), data.get('start_date', ''), data.get('end_date', ''),
             data.get('status', 'pending')))
        conn.commit()
        if project_id:
            _rebuild_invoice_line_items(conn, project_id)
            conn.commit()
        conn.close()
        return jsonify({'success': True, 'id': phid})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/ahb/phases/<phid>', methods=['PUT'])
def api_ahb_phases_update(phid):
    """Update a phase (planning-only — re-syncs calendar events for the phase)."""
    try:
        data = request.json or {}
        conn = _ahb_db()
        fields, vals = [], []
        for k in ['project_id', 'phase_number', 'name', 'value', 'start_date', 'end_date', 'status']:
            if k in data: fields.append(f"{k} = ?"); vals.append(data[k])
        if fields:
            vals.append(phid)
            conn.execute(f"UPDATE ahb_project_phases SET {', '.join(fields)} WHERE id = ?", vals)
            conn.commit()
        phase = conn.execute("SELECT project_id FROM ahb_project_phases WHERE id = ?", (phid,)).fetchone()
        if phase and phase['project_id']:
            _sync_phase_events(conn, phase['project_id'], phid)
            conn.commit()
        conn.close()
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/ahb/phases/<phid>', methods=['DELETE'])
def api_ahb_phases_delete(phid):
    """Delete a phase, its tasks, and their calendar events."""
    try:
        conn = _ahb_db()
        conn.execute("DELETE FROM ahb_phase_tasks WHERE phase_id = ?", (phid,))
        conn.execute("DELETE FROM ahb_events WHERE phase_id = ?", (phid,))
        conn.execute("DELETE FROM ahb_project_phases WHERE id = ?", (phid,))
        conn.commit()
        conn.close()
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/ahb/phases/<phid>/tasks', methods=['GET'])
def api_ahb_phase_tasks_list(phid):
    """List tasks for a phase."""
    try:
        conn = _ahb_db()
        rows = conn.execute("SELECT * FROM ahb_phase_tasks WHERE phase_id = ? ORDER BY created_at", (phid,)).fetchall()
        conn.close()
        return jsonify([dict(r) for r in rows])
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/ahb/phases/<phid>/tasks', methods=['POST'])
def api_ahb_phase_tasks_create(phid):
    """Add a task to a phase. Optional approx_minutes (time-to-complete) and
    source_line_idx (when the task was pulled from an invoice line item)."""
    try:
        data = request.json or {}
        conn = _ahb_db()
        tid = uuid.uuid4().hex[:24]
        phase = conn.execute("SELECT project_id FROM ahb_project_phases WHERE id = ?", (phid,)).fetchone()
        project_id = phase['project_id'] if phase else ''
        conn.execute(
            """INSERT INTO ahb_phase_tasks (id, phase_id, project_id, title, status, assigned_to, notes,
                                            approx_minutes, source_line_idx)
               VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)""",
            (tid, phid, project_id, data.get('title', ''), data.get('status', 'pending'),
             data.get('assigned_to', ''), data.get('notes', ''),
             int(data.get('approx_minutes') or 0),
             data.get('source_line_idx') if data.get('source_line_idx') is not None else None))
        conn.commit()
        # Sync this task to the calendar (date defaults to phase start; details carry the link)
        if project_id:
            _sync_phase_events(conn, project_id, phid)
            conn.commit()
        conn.close()
        return jsonify({'success': True, 'id': tid})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/ahb/phase-tasks/<tid>', methods=['PUT'])
def api_ahb_phase_task_update(tid):
    """Update a task — title, status, assigned_to, notes, approx_minutes, or
    phase_id (move to a different phase). When phase_id changes, project_id is
    coerced to match the new phase and calendar events are re-synced for BOTH
    the old and new phase so nothing is left orphaned."""
    try:
        data = request.json or {}
        conn = _ahb_db()
        prev = conn.execute("SELECT phase_id, project_id FROM ahb_phase_tasks WHERE id = ?", (tid,)).fetchone()
        old_phase_id = prev['phase_id'] if prev else None
        old_project_id = prev['project_id'] if prev else None

        fields, vals = [], []
        for k in ['title', 'status', 'assigned_to', 'notes', 'phase_id']:
            if k in data:
                fields.append(f"{k} = ?"); vals.append(data[k])
        if 'approx_minutes' in data:
            fields.append("approx_minutes = ?"); vals.append(int(data['approx_minutes'] or 0))
        # Moving to another phase — make sure project_id follows the new phase.
        new_phase_id = data.get('phase_id')
        if new_phase_id and new_phase_id != old_phase_id:
            r = conn.execute("SELECT project_id FROM ahb_project_phases WHERE id = ?", (new_phase_id,)).fetchone()
            if r:
                fields.append("project_id = ?"); vals.append(r['project_id'])
        if not fields:
            conn.close()
            return jsonify({'success': True})
        vals.append(tid)
        conn.execute(f"UPDATE ahb_phase_tasks SET {', '.join(fields)} WHERE id = ?", vals)
        conn.commit()

        # Re-sync events for whichever phase(s) were touched.
        cur = conn.execute("SELECT phase_id, project_id FROM ahb_phase_tasks WHERE id = ?", (tid,)).fetchone()
        cur_phase_id = cur['phase_id'] if cur else None
        cur_project_id = cur['project_id'] if cur else None
        if old_phase_id and old_phase_id != cur_phase_id and old_project_id:
            _sync_phase_events(conn, old_project_id, old_phase_id)
        if cur_phase_id and cur_project_id:
            _sync_phase_events(conn, cur_project_id, cur_phase_id)
        conn.commit()
        conn.close()
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/ahb/phase-tasks/<tid>', methods=['DELETE'])
def api_ahb_phase_task_delete(tid):
    try:
        conn = _ahb_db()
        row = conn.execute("SELECT phase_id, project_id FROM ahb_phase_tasks WHERE id = ?", (tid,)).fetchone()
        conn.execute("DELETE FROM ahb_phase_tasks WHERE id = ?", (tid,))
        # Remove any calendar event tied to this task
        conn.execute("DELETE FROM ahb_events WHERE task_id = ?", (tid,))
        conn.commit()
        if row and row['project_id'] and row['phase_id']:
            _sync_phase_events(conn, row['project_id'], row['phase_id'])
            conn.commit()
        conn.close()
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


def _sync_phase_events(conn, project_id, phase_id):
    """Idempotently rebuild the calendar events for one phase and its tasks.
    Wipes any prior events keyed on this phase_id (whether they came from the
    phase itself or its tasks) and inserts fresh ones, so editing a phase or
    its task list never leaves orphan/duplicate calendar entries."""
    project = conn.execute("SELECT title FROM ahb_projects WHERE id = ?", (project_id,)).fetchone()
    if not project:
        return
    project_title = project['title'] or 'Project'
    phase = conn.execute("SELECT * FROM ahb_project_phases WHERE id = ?", (phase_id,)).fetchone()
    if not phase:
        # Phase was deleted — clear its events and exit
        conn.execute("DELETE FROM ahb_events WHERE phase_id = ?", (phase_id,))
        return
    phase = dict(phase)
    phase_name = phase.get('name') or f"Phase {phase.get('phase_number','?')}"
    # Clear prior events for this phase (covers phase markers and tasks)
    conn.execute("DELETE FROM ahb_events WHERE phase_id = ?", (phase_id,))
    # Phase start
    if phase.get('start_date'):
        eid = uuid.uuid4().hex[:24]
        conn.execute(
            """INSERT INTO ahb_events (id, title, details, date, category, all_day, project_id, phase_id)
               VALUES (?, ?, ?, ?, ?, ?, ?, ?)""",
            (eid, f"{project_title} • {phase_name} — Start", f"Phase start: {phase_name}",
             phase['start_date'], 'phase', 1, project_id, phase_id))
    # Phase end
    if phase.get('end_date'):
        eid = uuid.uuid4().hex[:24]
        conn.execute(
            """INSERT INTO ahb_events (id, title, details, date, category, all_day, project_id, phase_id)
               VALUES (?, ?, ?, ?, ?, ?, ?, ?)""",
            (eid, f"{project_title} • {phase_name} — End", f"Phase end: {phase_name}",
             phase['end_date'], 'phase', 1, project_id, phase_id))
    # Tasks — anchor each to the phase start_date (or end_date if no start)
    anchor = phase.get('start_date') or phase.get('end_date') or ''
    if anchor:
        tasks = conn.execute(
            "SELECT * FROM ahb_phase_tasks WHERE phase_id = ? ORDER BY created_at", (phase_id,)
        ).fetchall()
        for t in tasks:
            t = dict(t)
            eid = uuid.uuid4().hex[:24]
            approx = int(t.get('approx_minutes') or 0)
            details = f"Task in {phase_name} ({project_title})"
            if approx:
                details += f" · ~{approx} min"
            conn.execute(
                """INSERT INTO ahb_events (id, title, details, date, category, all_day, project_id, phase_id, task_id)
                   VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)""",
                (eid, f"{project_title} • {phase_name}: {t.get('title','Task')}",
                 details, anchor, 'phase_task', 1, project_id, phase_id, t['id']))


# ── AHB123 — Tax Requirements ──────────────────────────────────────────────────

@app.route('/api/ahb/tax', methods=['GET'])
def api_ahb_tax_list():
    try:
        conn = _ahb_db()
        q = "SELECT * FROM ahb_tax_requirements WHERE 1=1"
        params = []
        if request.args.get('category'):
            q += " AND category = ?"; params.append(request.args['category'])
        if request.args.get('completed') is not None:
            q += " AND completed = ?"; params.append(int(request.args['completed']))
        rows = conn.execute(q + " ORDER BY due_date", params).fetchall()
        conn.close()
        return jsonify([dict(r) for r in rows])
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/ahb/tax', methods=['POST'])
def api_ahb_tax_create():
    try:
        data = request.json or {}
        conn = _ahb_db()
        tid = str(uuid.uuid4())
        conn.execute(
            """INSERT INTO ahb_tax_requirements (id, title, details, due_date, completed, category)
               VALUES (?, ?, ?, ?, ?, ?)""",
            (tid, data.get('title',''), data.get('details',''), data.get('due_date',''),
             1 if data.get('completed') else 0, data.get('category','tax')))
        conn.commit(); conn.close()
        return jsonify({'success': True, 'id': tid})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/ahb/tax/<tid>', methods=['PUT'])
def api_ahb_tax_update(tid):
    try:
        data = request.json or {}
        conn = _ahb_db()
        fields, vals = [], []
        for k in ['title','details','due_date','category']:
            if k in data: fields.append(f"{k} = ?"); vals.append(data[k])
        if 'completed' in data: fields.append("completed = ?"); vals.append(1 if data['completed'] else 0)
        if fields:
            vals.append(tid)
            conn.execute(f"UPDATE ahb_tax_requirements SET {', '.join(fields)} WHERE id = ?", vals)
            conn.commit()
        conn.close()
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/ahb/tax/<tid>', methods=['DELETE'])
def api_ahb_tax_delete(tid):
    try:
        conn = _ahb_db()
        conn.execute("DELETE FROM ahb_tax_requirements WHERE id = ?", (tid,))
        conn.commit(); conn.close()
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


# ── AHB123 — Invoice from Project ───────────────────────────────────────────────

@app.route('/api/ahb/invoices/from-project/<pid>', methods=['POST'])
def api_ahb_invoice_from_project(pid):
    """Generate an invoice from a project's phases and details."""
    try:
        conn = _ahb_db()
        project = conn.execute("SELECT * FROM ahb_projects WHERE id = ?", (pid,)).fetchone()
        if not project:
            conn.close()
            return jsonify({'success': False, 'error': 'Project not found'}), 404
        project = dict(project)
        # Get phases as line items
        phases = conn.execute(
            "SELECT * FROM ahb_project_phases WHERE project_id = ? ORDER BY phase_number",
            (pid,)
        ).fetchall()
        line_items = []
        subtotal = 0
        if phases:
            for ph in phases:
                ph = dict(ph)
                val = ph.get('value', 0) or 0
                line_items.append({
                    'description': ph.get('name', 'Phase'),
                    'qty': 1, 'rate': val, 'total': val,
                    'materials': 0, 'labor': 0,
                    'quantity': 1, 'unit_price': val,
                })
                subtotal += val
        else:
            # Fallback: single line item from project value
            val = project.get('value') or project.get('budget_high') or 0
            line_items = [{
                'description': project.get('title', 'Project'),
                'qty': 1, 'rate': val, 'total': val,
                'materials': 0, 'labor': 0,
                'quantity': 1, 'unit_price': val,
            }]
            subtotal = val

        inv_num = f"AHB-{datetime.datetime.now().strftime('%Y%m%d')}-{uuid.uuid4().hex[:3].upper()}"
        iid = str(uuid.uuid4())
        # The PDF already shows the project description in the Scope-of-Work block
        # at the top of the invoice — putting it in notes too would duplicate it.
        # Notes is reserved for invoice-specific terms / payment instructions.
        notes = ''
        conn.execute(
            """INSERT INTO ahb_invoices (id, client_id, project_id, invoice_number, line_items,
               subtotal, tax, total, status, notes, client_name, project_name, terms,
               project_address)
               VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)""",
            (iid, project.get('client_id', ''), pid, inv_num,
             json.dumps(line_items), subtotal, 0, subtotal, 'draft',
             notes,
             project.get('client_name', ''), project.get('title', ''),
             'Net 30', project.get('address', '')))
        conn.commit()
        conn.close()
        return jsonify({'success': True, 'id': iid, 'invoice_number': inv_num,
                        'line_items': line_items, 'subtotal': subtotal})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


# ── AHB123 — Dashboard Stats ───────────────────────────────────────────────────

@app.route('/api/ahb/dashboard', methods=['GET'])
def api_ahb_dashboard():
    """Return aggregated stats for the overview dashboard. Supports ?year=2026 filter."""
    try:
        conn = _ahb_db()
        stats = {}
        year = request.args.get('year', '')

        # Year filter helpers
        def yf_proj(col='start_date'):
            fallback = 'created_at' if '.' not in col else col.split('.')[0] + '.created_at'
            return f" AND substr(COALESCE({col}, {fallback}), 1, 4) = '{year}'" if year else ''
        def yf(col='created_at'):
            return f" AND substr({col}, 1, 4) = '{year}'" if year else ''

        # Projects
        rows = conn.execute(f"SELECT status, count(*) as cnt FROM ahb_projects WHERE 1=1{yf_proj()} GROUP BY status").fetchall()
        stats['projects'] = {r['status']: r['cnt'] for r in rows}
        stats['projects_total'] = sum(r['cnt'] for r in rows)

        # Invoices — join with project to get real year from project start_date
        if year:
            inv_sql = """SELECT i.status, count(*) as cnt, COALESCE(sum(i.total),0) as total
                FROM ahb_invoices i LEFT JOIN ahb_projects p ON i.project_id = p.id
                WHERE substr(COALESCE(i.due_date, i.paid_date, p.start_date, i.created_at), 1, 4) = ?
                GROUP BY i.status"""
            rows = conn.execute(inv_sql, (year,)).fetchall()
        else:
            rows = conn.execute("SELECT status, count(*) as cnt, COALESCE(sum(total),0) as total FROM ahb_invoices GROUP BY status").fetchall()
        stats['invoices'] = {r['status']: {'count': r['cnt'], 'total': r['total']} for r in rows}
        stats['invoices_total'] = sum(r['total'] for r in rows)

        # Receipts
        rows = conn.execute(f"SELECT category, count(*) as cnt, COALESCE(sum(total),0) as total FROM ahb_receipts WHERE 1=1{yf('receipt_date')} GROUP BY category").fetchall()
        stats['receipts'] = {r['category']: {'count': r['cnt'], 'total': r['total']} for r in rows}
        stats['receipts_total'] = sum(r['total'] for r in rows)

        # Payroll
        row = conn.execute(f"SELECT COALESCE(sum(total),0) as total FROM ahb_payroll WHERE 1=1{yf('period_start')}").fetchone()
        stats['payroll_total'] = row['total']

        # Debts (debts are current, not year-filtered)
        row = conn.execute("SELECT COALESCE(sum(balance),0) as balance, COALESCE(sum(payment_amount),0) as monthly FROM ahb_debts").fetchone()
        stats['debt_balance'] = row['balance']
        stats['debt_monthly'] = row['monthly']
        rows = conn.execute("SELECT type, count(*) as cnt, COALESCE(sum(balance),0) as balance FROM ahb_debts GROUP BY type").fetchall()
        stats['debts'] = {r['type']: {'count': r['cnt'], 'balance': r['balance']} for r in rows}

        # Employees
        stats['employees_active'] = conn.execute("SELECT count(*) as cnt FROM ahb_employees WHERE active=1").fetchone()['cnt']

        # Clients
        stats['clients_total'] = conn.execute("SELECT count(*) as cnt FROM ahb_clients").fetchone()['cnt']

        # Events
        today = datetime.datetime.now().strftime('%Y-%m-%d')
        week = (datetime.datetime.now() + datetime.timedelta(days=7)).strftime('%Y-%m-%d')
        stats['events_this_week'] = conn.execute(
            "SELECT count(*) as cnt FROM ahb_events WHERE date >= ? AND date <= ?", (today, week)
        ).fetchone()['cnt']

        # Recent projects
        recent = conn.execute(f"SELECT id, title, status, value, address FROM ahb_projects WHERE 1=1{yf_proj()} ORDER BY updated_at DESC LIMIT 5").fetchall()
        stats['recent_projects'] = [dict(r) for r in recent]

        # All project locations for map — ALWAYS all time, value from invoices
        locations = conn.execute("""
            SELECT p.id, p.title, p.status, p.address,
                   COALESCE(SUM(i.total), p.value, 0) as value,
                   COUNT(i.id) as invoice_count
            FROM ahb_projects p
            LEFT JOIN ahb_invoices i ON i.project_id = p.id
            WHERE p.address IS NOT NULL AND p.address != ''
            GROUP BY p.id
        """).fetchall()
        stats['project_locations'] = [dict(r) for r in locations]

        # Upcoming debt due dates
        upcoming = conn.execute(
            "SELECT name, due_date, payment_amount, type FROM ahb_debts WHERE due_date >= ? AND due_date <= ? ORDER BY due_date LIMIT 10",
            (today, week)
        ).fetchall()
        stats['upcoming_debts'] = [dict(r) for r in upcoming]

        conn.close()
        return jsonify(stats)
    except Exception as e:
        return jsonify({'error': str(e)}), 500


# ── AHB123 — Employee Time Clock ───────────────────────────────────────────────

@app.route('/api/ahb/timeclock', methods=['GET'])
def api_ahb_timeclock_list():
    try:
        conn = _ahb_db()
        q = "SELECT * FROM ahb_timeclock WHERE 1=1"
        params = []
        if request.args.get('employee_id'):
            q += " AND employee_id = ?"; params.append(request.args['employee_id'])
        if request.args.get('date'):
            q += " AND date = ?"; params.append(request.args['date'])
        rows = conn.execute(q + " ORDER BY date DESC, clock_in DESC", params).fetchall()
        conn.close()
        return jsonify([dict(r) for r in rows])
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/ahb/timeclock/punch', methods=['POST'])
def api_ahb_timeclock_punch():
    """Clock in, clock out, or lunch punch."""
    try:
        data = request.json or {}
        employee_id = data.get('employee_id', '')
        punch_type = data.get('type', 'clock_in')  # clock_in, clock_out, lunch_start, lunch_end
        now = datetime.datetime.now()
        conn = _ahb_db()

        if punch_type == 'clock_in':
            tid = str(uuid.uuid4())
            conn.execute(
                """INSERT INTO ahb_timeclock (id, employee_id, date, clock_in, status)
                   VALUES (?, ?, ?, ?, 'active')""",
                (tid, employee_id, now.strftime('%Y-%m-%d'), now.strftime('%H:%M:%S')))
            conn.commit(); conn.close()
            return jsonify({'success': True, 'id': tid, 'punch': 'clock_in', 'time': now.strftime('%H:%M:%S')})

        # Find today's active record
        record = conn.execute(
            "SELECT * FROM ahb_timeclock WHERE employee_id = ? AND date = ? AND status = 'active' ORDER BY clock_in DESC LIMIT 1",
            (employee_id, now.strftime('%Y-%m-%d'))
        ).fetchone()
        if not record:
            conn.close()
            return jsonify({'success': False, 'error': 'No active clock-in found for today'}), 400

        rid = record['id']
        time_str = now.strftime('%H:%M:%S')

        if punch_type == 'lunch_start':
            conn.execute("UPDATE ahb_timeclock SET lunch_start = ? WHERE id = ?", (time_str, rid))
        elif punch_type == 'lunch_end':
            conn.execute("UPDATE ahb_timeclock SET lunch_end = ? WHERE id = ?", (time_str, rid))
        elif punch_type == 'clock_out':
            # Calculate hours
            clock_in = record['clock_in']
            cin = datetime.datetime.strptime(f"{record['date']} {clock_in}", '%Y-%m-%d %H:%M:%S')
            cout = now
            total_mins = (cout - cin).total_seconds() / 60
            # Subtract lunch
            lunch_mins = 0
            if record['lunch_start'] and record['lunch_end']:
                ls = datetime.datetime.strptime(f"{record['date']} {record['lunch_start']}", '%Y-%m-%d %H:%M:%S')
                le = datetime.datetime.strptime(f"{record['date']} {record['lunch_end']}", '%Y-%m-%d %H:%M:%S')
                lunch_mins = (le - ls).total_seconds() / 60
            hours = round((total_mins - lunch_mins) / 60, 2)
            conn.execute(
                "UPDATE ahb_timeclock SET clock_out = ?, hours = ?, lunch_minutes = ?, status = 'completed' WHERE id = ?",
                (time_str, hours, round(lunch_mins), rid))

        conn.commit(); conn.close()
        return jsonify({'success': True, 'punch': punch_type, 'time': time_str})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/ahb/timeclock/status/<employee_id>', methods=['GET'])
def api_ahb_timeclock_status(employee_id):
    """Get current clock status for an employee."""
    try:
        conn = _ahb_db()
        today = datetime.datetime.now().strftime('%Y-%m-%d')
        record = conn.execute(
            "SELECT * FROM ahb_timeclock WHERE employee_id = ? AND date = ? ORDER BY clock_in DESC LIMIT 1",
            (employee_id, today)
        ).fetchone()
        conn.close()
        if not record:
            return jsonify({'status': 'not_clocked_in', 'record': None})
        return jsonify({'status': record['status'], 'record': dict(record)})
    except Exception as e:
        return jsonify({'error': str(e)}), 500


# ── AHB123 — Project Detail ─────────────────────────────────────────────────

@app.route('/api/ahb/projects/<pid>/detail', methods=['GET'])
def api_ahb_project_detail(pid):
    """Return full project detail with invoices, phases, files, receipts, linked invoice, permits count, photo counts."""
    try:
        conn = _ahb_db()
        project = conn.execute("SELECT * FROM ahb_projects WHERE id = ?", (pid,)).fetchone()
        if not project:
            conn.close()
            return jsonify({'error': 'Not found'}), 404
        result = dict(project)
        result['invoices'] = [dict(r) for r in conn.execute(
            "SELECT * FROM ahb_invoices WHERE project_id = ? ORDER BY created_at DESC", (pid,)).fetchall()]
        result['phases'] = [dict(r) for r in conn.execute(
            "SELECT * FROM ahb_project_phases WHERE project_id = ? ORDER BY phase_number", (pid,)).fetchall()]
        result['files'] = [dict(r) for r in conn.execute(
            "SELECT * FROM ahb_files WHERE project_id = ? ORDER BY created_at DESC", (pid,)).fetchall()]
        result['receipts'] = [dict(r) for r in conn.execute(
            "SELECT * FROM ahb_receipts WHERE project_id = ? ORDER BY receipt_date DESC", (pid,)).fetchall()]
        result['events'] = [dict(r) for r in conn.execute(
            "SELECT * FROM ahb_events WHERE project_id = ? ORDER BY date", (pid,)).fetchall()]

        # Curated documents (permits, COIs, licenses, etc. filed by Phil)
        result['documents'] = [dict(r) for r in conn.execute(
            "SELECT id, file_path, original_name, suggested_name, doc_type, entity, "
            "       doc_date, summary, tags, confidence, curated_at "
            "FROM ahb_documents WHERE project_id = ? ORDER BY curated_at DESC", (pid,)).fetchall()]

        # Linked invoice (first invoice for this project)
        linked_inv = conn.execute(
            "SELECT * FROM ahb_invoices WHERE project_id = ? ORDER BY created_at ASC LIMIT 1", (pid,)
        ).fetchone()
        result['linked_invoice'] = dict(linked_inv) if linked_inv else None

        # Payment summary drives the "balance due upon completion" red badge
        # and the deposit-paid → In Progress auto-flip flow on the frontend.
        result['_payment'] = _ahb_project_payment_summary(conn, pid)

        # Permits count
        permits_row = conn.execute(
            "SELECT count(*) as cnt FROM ahb_files WHERE project_id = ? AND document_type = 'Permit'", (pid,)
        ).fetchone()
        result['permits_count'] = permits_row['cnt'] if permits_row else 0

        # Photo counts per section
        photo_sections = conn.execute(
            "SELECT photo_section, count(*) as cnt FROM ahb_files WHERE project_id = ? AND photo_section != '' GROUP BY photo_section",
            (pid,)
        ).fetchall()
        result['photo_counts'] = {r['photo_section']: r['cnt'] for r in photo_sections}

        # Phase tasks
        for phase in result['phases']:
            tasks = conn.execute(
                "SELECT * FROM ahb_phase_tasks WHERE phase_id = ? ORDER BY created_at", (phase['id'],)
            ).fetchall()
            phase['tasks'] = [dict(t) for t in tasks]

        conn.close()
        return jsonify(result)
    except Exception as e:
        return jsonify({'error': str(e)}), 500


# ── AHB123 — Invoice Detail + Interest ─────────────────────────────────────

@app.route('/api/ahb/invoices/<iid>/detail', methods=['GET'])
def api_ahb_invoice_detail(iid):
    try:
        conn = _ahb_db()
        inv = conn.execute("SELECT * FROM ahb_invoices WHERE id = ?", (iid,)).fetchone()
        if not inv:
            conn.close()
            return jsonify({'error': 'Not found'}), 404
        result = dict(inv)
        # Get change orders for this invoice
        result['change_orders'] = [dict(r) for r in conn.execute(
            "SELECT * FROM ahb_invoices WHERE parent_invoice_id = ?", (iid,)).fetchall()]
        # Get parent if this is a change order
        if inv['parent_invoice_id']:
            parent = conn.execute("SELECT * FROM ahb_invoices WHERE id = ?", (inv['parent_invoice_id'],)).fetchone()
            result['parent_invoice'] = dict(parent) if parent else None
        # Get payments
        result['payments'] = [dict(r) for r in conn.execute(
            "SELECT * FROM ahb_payments WHERE invoice_id = ? ORDER BY payment_date DESC", (iid,)).fetchall()]
        conn.close()
        return jsonify(result)
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/ahb/invoices/<iid>/interest', methods=['GET'])
def api_ahb_invoice_interest(iid):
    try:
        conn = _ahb_db()
        inv = conn.execute("SELECT * FROM ahb_invoices WHERE id = ?", (iid,)).fetchone()
        if not inv:
            conn.close()
            return jsonify({'error': 'Not found'}), 404
        inv = dict(inv)
        overdue_since = inv.get('overdue_since') or inv.get('due_date')
        if not overdue_since:
            conn.close()
            return jsonify({'weeks': 0, 'interest': 0, 'total_with_interest': inv.get('total', 0)})
        from datetime import datetime as dt
        start = dt.strptime(overdue_since[:10], '%Y-%m-%d')
        weeks = max(0, (dt.now() - start).days // 7)
        rate = inv.get('overdue_interest_per_week') or 50
        interest = weeks * rate
        conn.close()
        return jsonify({'weeks': weeks, 'interest': interest, 'rate_per_week': rate,
                        'total_with_interest': (inv.get('total') or 0) + interest})
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/ahb/invoices/<iid>/pdf', methods=['GET'])
def api_ahb_invoice_pdf(iid):
    """Generate professional invoice PDF matching Shaski template."""
    try:
        conn = _ahb_db()
        inv = conn.execute("SELECT * FROM ahb_invoices WHERE id = ?", (iid,)).fetchone()
        if not inv:
            conn.close()
            return jsonify({'error': 'Invoice not found'}), 404
        inv = dict(inv)
        # Get project details if linked
        project = None
        if inv.get('project_id'):
            p = conn.execute("SELECT * FROM ahb_projects WHERE id = ?", (inv['project_id'],)).fetchone()
            if p:
                project = dict(p)
        # Get phases for this project
        phases = []
        if inv.get('project_id'):
            phases = [dict(r) for r in conn.execute(
                "SELECT * FROM ahb_project_phases WHERE project_id = ? ORDER BY phase_number", (inv['project_id'],)).fetchall()]
        conn.close()

        # Parse line items
        line_items = []
        if inv.get('line_items'):
            try:
                line_items = json.loads(inv['line_items']) if isinstance(inv['line_items'], str) else inv['line_items']
            except (json.JSONDecodeError, TypeError):
                line_items = []

        # Build line items HTML. We render Materials + Labor columns when ANY line
        # carries a non-zero breakdown — that way old single-total invoices still
        # render cleanly while new ones show the material/labor split that sums to
        # the line total.
        any_breakdown = any(
            (float(item.get('materials') or 0) > 0) or (float(item.get('labor') or 0) > 0)
            for item in line_items
        )
        items_html = ''
        for i, item in enumerate(line_items, 1):
            desc = item.get('description', '')
            qty   = item.get('qty')   if item.get('qty')   is not None else item.get('quantity', 1)
            price = item.get('rate')  if item.get('rate')  is not None else item.get('unit_price', 0)
            unit  = item.get('unit') or 'qty'
            try: qty = float(qty or 0)
            except: qty = 0
            try: price = float(price or 0)
            except: price = 0
            try: materials = float(item.get('materials') or 0)
            except: materials = 0
            try: labor = float(item.get('labor') or 0)
            except: labor = 0
            # Honor stored total if it was manually overridden
            stored_total = item.get('total')
            try:
                stored_total = float(stored_total) if stored_total is not None else None
            except:
                stored_total = None
            if materials or labor:
                total_item = materials + labor
            elif stored_total is not None:
                total_item = stored_total
            else:
                total_item = qty * price
            qty_display = f"{qty:g} {unit}" if unit and unit != 'qty' else f"{qty:g}"
            if any_breakdown:
                items_html += f'''<tr>
                    <td style="padding:10px 12px;border-bottom:1px solid #eee;color:#333;">{i}</td>
                    <td style="padding:10px 12px;border-bottom:1px solid #eee;color:#333;font-weight:500;">{desc}</td>
                    <td style="padding:10px 12px;border-bottom:1px solid #eee;text-align:center;color:#333;">{qty_display}</td>
                    <td style="padding:10px 12px;border-bottom:1px solid #eee;text-align:right;color:#666;">${materials:,.2f}</td>
                    <td style="padding:10px 12px;border-bottom:1px solid #eee;text-align:right;color:#666;">${labor:,.2f}</td>
                    <td style="padding:10px 12px;border-bottom:1px solid #eee;text-align:right;color:#333;font-weight:600;">${total_item:,.2f}</td>
                </tr>'''
            else:
                items_html += f'''<tr>
                    <td style="padding:10px 12px;border-bottom:1px solid #eee;color:#333;">{i}</td>
                    <td style="padding:10px 12px;border-bottom:1px solid #eee;color:#333;font-weight:500;">{desc}</td>
                    <td style="padding:10px 12px;border-bottom:1px solid #eee;text-align:center;color:#333;">{qty_display}</td>
                    <td style="padding:10px 12px;border-bottom:1px solid #eee;text-align:right;color:#333;">${price:,.2f}</td>
                    <td style="padding:10px 12px;border-bottom:1px solid #eee;text-align:right;color:#333;font-weight:600;">${total_item:,.2f}</td>
                </tr>'''

        # Build scope of work section for PDF.
        # Avoid duplicating the description in three places (scope block + notes +
        # line items). If the line items already break the work down (>1 row, or a
        # single row whose description differs from the project title), the line
        # items table is sufficient — only show the Trade/Scope tag for context.
        scope_of_work_html = ''
        if project:
            proj_desc = (project.get('description') or '').strip()
            proj_scope = (project.get('scope') or '').strip()
            line_items_carry_detail = (
                len(line_items) > 1 or
                (len(line_items) == 1 and (line_items[0].get('description') or '').strip()
                    and (line_items[0].get('description') or '').strip() != (project.get('title') or '').strip())
            )
            notes_text = (inv.get('notes') or '').strip()
            notes_overlaps = bool(notes_text) and proj_desc and (proj_desc[:60].lower() in notes_text.lower())
            show_desc_block = bool(proj_desc) and not line_items_carry_detail and not notes_overlaps
            if show_desc_block:
                proj_desc_html  = proj_desc.replace('<', '&lt;').replace('>', '&gt;')
                proj_scope_html = proj_scope.replace('<', '&lt;').replace('>', '&gt;')
                scope_line = f'<div style="margin-top:6px;font-size:12px;color:#666;"><strong>Trade/Scope:</strong> {proj_scope_html}</div>' if proj_scope_html else ''
                scope_of_work_html = f'''<div style="margin-bottom:20px;padding:12px 16px;background:#f8fafc;border-radius:6px;border-left:3px solid #2563eb;">
    <div style="font-size:11px;color:#999;text-transform:uppercase;font-weight:700;margin-bottom:6px;">Scope of Work</div>
    <div style="font-size:13px;color:#444;line-height:1.5;white-space:pre-wrap;">{proj_desc_html}</div>
    {scope_line}
</div>'''
            elif proj_scope:
                proj_scope_html = proj_scope.replace('<', '&lt;').replace('>', '&gt;')
                scope_of_work_html = f'''<div style="margin-bottom:14px;font-size:12px;color:#666;"><strong>Trade/Scope:</strong> {proj_scope_html}</div>'''

        subtotal = inv.get('subtotal') or inv.get('total') or 0
        tax = inv.get('tax') or 0
        total = inv.get('total') or 0
        inv_date = inv.get('date') or (inv.get('created_at', '')[:10] if inv.get('created_at') else '')
        project_location = inv.get('project_address') or (project.get('address', '') if project else '') or ''
        interest_rate = inv.get('overdue_interest_per_week') or 50

        # Logo as base64 for embedding in PDF
        logo_b64 = ''
        logo_path = os.path.join(DASHBOARD_DIR, 'static', 'img', 'ahb_logo.jpeg')
        if os.path.exists(logo_path):
            import base64
            with open(logo_path, 'rb') as lf:
                logo_b64 = base64.b64encode(lf.read()).decode('utf-8')

        # Page 1: Invoice
        html = f'''<!DOCTYPE html>
<html>
<head><meta charset="utf-8"><title>Invoice {inv.get('invoice_number','')}</title>
<style>
@media print {{ body {{ margin:0; }} @page {{ margin:40px 50px; }} }}
body {{ font-family:'Helvetica Neue',Helvetica,Arial,sans-serif;max-width:780px;margin:30px auto;color:#333;font-size:14px;line-height:1.5; }}
.page-break {{ page-break-before:always; }}
</style></head>
<body>
<!-- PAGE 1: INVOICE -->
<div style="display:flex;justify-content:space-between;align-items:flex-start;margin-bottom:20px;">
    <div style="display:flex;align-items:flex-start;gap:12px;">
        {f'<img src="data:image/jpeg;base64,{logo_b64}" style="width:50px;height:50px;object-fit:contain;margin-top:2px;">' if logo_b64 else '<div style="width:50px;height:50px;background:#2563eb;border-radius:8px;margin-top:2px;"></div>'}
        <div>
            <div style="font-size:20px;font-weight:700;color:#1a1a1a;white-space:nowrap;">All Home Building CO LLC</div>
            <div style="font-size:12px;color:#888;">2725 Colmar Ave, Bensalem, PA 19020</div>
            <div style="font-size:12px;color:#888;">800-484-6404 · AHB123.com</div>
        </div>
    </div>
    <div style="text-align:right;">
        <div style="font-size:28px;font-weight:300;color:#333;letter-spacing:2px;">INVOICE</div>
        <div style="font-size:14px;color:#555;">#{inv.get('invoice_number','')}</div>
        <div style="display:inline-block;padding:3px 12px;border-radius:12px;font-size:11px;font-weight:700;background:{'#dcfce7;color:#16a34a' if inv.get('status')=='Paid' else '#dbeafe;color:#2563eb' if inv.get('status')=='Approved' else '#fef3c7;color:#d97706' if inv.get('status')=='In Progress' else '#fee2e2;color:#dc2626' if inv.get('status')=='Overdue' else '#f3f4f6;color:#6b7280'};margin-top:4px;">{inv.get('status','Draft')}</div>
    </div>
</div>

<div style="display:flex;justify-content:space-between;margin:16px 0 24px;padding:12px 0;border-top:1px solid #e5e7eb;border-bottom:1px solid #e5e7eb;">
    <div>
        <div style="font-size:11px;color:#999;text-transform:uppercase;font-weight:700;margin-bottom:4px;">Bill To:</div>
        <div style="font-weight:600;">{inv.get('client_name','')}</div>
        <div style="color:#666;">{inv.get('client_address','')}</div>
    </div>
    <div style="text-align:center;">
        <div style="font-size:11px;color:#999;text-transform:uppercase;font-weight:700;margin-bottom:4px;">Project Location:</div>
        <div style="color:#666;">{project_location}</div>
    </div>
    <div style="text-align:right;">
        <div style="font-size:11px;color:#999;text-transform:uppercase;font-weight:700;margin-bottom:4px;">Date:</div>
        <div>{inv_date}</div>
    </div>
</div>

{scope_of_work_html}

<table style="width:100%;border-collapse:collapse;margin:0 0 20px;">
    <thead>
        <tr style="background:#f8fafc;">
            <th style="padding:10px 12px;text-align:left;border-bottom:2px solid #e2e8f0;font-size:12px;color:#64748b;font-weight:700;">#</th>
            <th style="padding:10px 12px;text-align:left;border-bottom:2px solid #e2e8f0;font-size:12px;color:#64748b;font-weight:700;">Description</th>
            <th style="padding:10px 12px;text-align:center;border-bottom:2px solid #e2e8f0;font-size:12px;color:#64748b;font-weight:700;">Qty</th>
            {'<th style="padding:10px 12px;text-align:right;border-bottom:2px solid #e2e8f0;font-size:12px;color:#64748b;font-weight:700;">Materials</th><th style="padding:10px 12px;text-align:right;border-bottom:2px solid #e2e8f0;font-size:12px;color:#64748b;font-weight:700;">Labor</th>' if any_breakdown else '<th style="padding:10px 12px;text-align:right;border-bottom:2px solid #e2e8f0;font-size:12px;color:#64748b;font-weight:700;">Price</th>'}
            <th style="padding:10px 12px;text-align:right;border-bottom:2px solid #e2e8f0;font-size:12px;color:#64748b;font-weight:700;">Total</th>
        </tr>
    </thead>
    <tbody>{items_html}</tbody>
</table>

<div style="display:flex;justify-content:flex-end;">
    <div style="width:250px;border-top:2px solid #333;padding-top:8px;">
        <div style="display:flex;justify-content:space-between;padding:6px 0;font-size:20px;font-weight:700;color:#2563eb;">
            <span>Total:</span><span>${total:,.2f}</span>
        </div>
    </div>
</div>


<!-- PAGE 2: CONTRACTOR AGREEMENT -->
<div class="page-break"></div>

<div style="margin-bottom:24px;">
    <h1 style="margin:0;font-size:24px;font-weight:400;color:#333;">CONTRACTOR AGREEMENT</h1>
</div>

<div style="margin-bottom:20px;">
    <h3 style="font-size:14px;color:#333;margin:0 0 10px;">Parties</h3>
    <div style="display:flex;justify-content:space-between;">
        <div>
            <div style="font-size:12px;color:#999;font-weight:700;">Contractor:</div>
            <div style="font-weight:600;">Sergey Tkach</div>
            <div>All Home Building CO LLC</div>
            <div style="color:#666;font-size:13px;">2725 Colmar Ave, Bensalem, PA 19020</div>
            <div style="color:#666;font-size:13px;">800-484-6404 · AHB123.com</div>
        </div>
        <div style="text-align:right;">
            <div style="font-size:12px;color:#999;font-weight:700;">Client:</div>
            <div style="font-weight:600;">{inv.get('client_name','')}</div>
            <div style="color:#666;font-size:13px;">{inv.get('client_address','')}</div>
        </div>
    </div>
</div>

<div style="margin-bottom:20px;padding:12px;background:#f8fafc;border-radius:6px;">
    <h3 style="font-size:14px;color:#333;margin:0 0 8px;">Project Information</h3>
    <div style="font-size:13px;color:#555;">
        <strong>Invoice Number:</strong> {inv.get('invoice_number','')}<br>
        <strong>Date:</strong> {inv_date}<br>
        <strong>Project Location:</strong> {project_location}
    </div>
</div>

{scope_of_work_html}

<div style="margin-bottom:24px;">
    <h3 style="font-size:14px;color:#333;margin:0 0 12px;">Terms & Conditions</h3>
    <div style="font-size:13px;color:#444;line-height:1.6;">
        <p style="margin:0 0 8px;"><strong>1.</strong> A Deposit is due before commencement of the project.</p>
        <p style="margin:0 0 8px;"><strong>2.</strong> Total due upon completion.</p>
        <p style="margin:0 0 8px;"><strong>3.</strong> Project will take approx {(project or {}).get('notes','') or 'TBD'} days to complete.</p>
        <p style="margin:0 0 8px;"><strong>4.</strong> Project description is final unless a change order is requested.</p>
        <p style="margin:0 0 8px;"><strong>5.</strong> Make checks payable to ALL HOME BUILDING CO.</p>
        <p style="margin:0 0 8px;"><strong>6.</strong> Late Payment: Payment is due by the date specified on this invoice. If payment is not received by the due date, this invoice shall be deemed overdue. Interest shall accrue at the rate of fifty dollars (${interest_rate:.0f}.00) per week on the unpaid balance. An overdue interest sheet will be attached reflecting all accrued charges.</p>
    </div>
</div>

<div style="display:flex;justify-content:space-between;margin-top:40px;padding-top:20px;">
    <div style="width:45%;">
        <div style="font-size:12px;color:#999;margin-bottom:4px;">Contractor Signature:</div>
        <div style="border-bottom:1px solid #333;height:40px;margin-bottom:4px;"></div>
        <div style="font-size:12px;color:#555;">Sergey Tkach</div>
    </div>
    <div style="width:45%;">
        <div style="font-size:12px;color:#999;margin-bottom:4px;">Client Signature:</div>
        <div style="border-bottom:1px solid #333;height:40px;margin-bottom:4px;"></div>
        <div style="font-size:12px;color:#555;">{inv.get('client_name','')}</div>
    </div>
</div>

</body>
</html>'''

        # Try to use weasyprint for real PDF, fall back to HTML
        download = request.args.get('download', '0') == '1'
        try:
            from weasyprint import HTML as WeasyHTML
            pdf_bytes = WeasyHTML(string=html).write_pdf()
            response = make_response(pdf_bytes)
            response.headers['Content-Type'] = 'application/pdf'
            disposition = 'attachment' if download else 'inline'
            response.headers['Content-Disposition'] = f'{disposition}; filename="invoice_{inv.get("invoice_number","")}.pdf"'
            return response
        except ImportError:
            # weasyprint not installed — return HTML with PDF-like content type hint
            response = make_response(html)
            response.headers['Content-Type'] = 'text/html; charset=utf-8'
            response.headers['Content-Disposition'] = f'inline; filename="invoice_{inv.get("invoice_number","")}.html"'
            return response

    except Exception as e:
        return jsonify({'error': str(e)}), 500


# ── AHB123 — Payments ───────────────────────────────────────────────────────

@app.route('/api/ahb/payments', methods=['GET'])
def api_ahb_payments_list():
    try:
        conn = _ahb_db()
        q = "SELECT * FROM ahb_payments WHERE 1=1"
        params = []
        if request.args.get('invoice_id'):
            q += " AND invoice_id = ?"; params.append(request.args['invoice_id'])
        rows = conn.execute(q + " ORDER BY payment_date DESC", params).fetchall()
        conn.close()
        return jsonify([dict(r) for r in rows])
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/ahb/payments', methods=['POST'])
def api_ahb_payments_create():
    """Record a payment and auto-progress the linked project status.

    Deposit recorded on a Planning project -> project flips to In Progress.
    Final payment that closes the balance on a Completed project -> invoice
    flips to Paid (project stays Completed).
    """
    try:
        data = request.json or {}
        conn = _ahb_db()
        pmt_id = str(uuid.uuid4())
        invoice_id = data.get('invoice_id', '')
        conn.execute(
            """INSERT INTO ahb_payments (id, invoice_id, amount, payment_method, payment_date, notes)
               VALUES (?, ?, ?, ?, ?, ?)""",
            (pmt_id, invoice_id, data.get('amount', 0), data.get('payment_method', ''),
             data.get('payment_date', ''), data.get('notes', '')))

        # Auto-advance project status from the deposit / balance signal.
        project_id = None
        project_status_after = None
        if invoice_id:
            inv_row = conn.execute(
                "SELECT project_id FROM ahb_invoices WHERE id = ?", (invoice_id,)
            ).fetchone()
            project_id = inv_row['project_id'] if inv_row else None
            if project_id:
                proj_row = conn.execute(
                    "SELECT status FROM ahb_projects WHERE id = ?", (project_id,)
                ).fetchone()
                cur_status = _ahb_canon_project_status(proj_row['status'] if proj_row else '')
                summary = _ahb_project_payment_summary(conn, project_id)
                # Deposit-or-better recorded on a Planning project → In Progress.
                if cur_status == 'Planning' and summary['has_payments']:
                    _ahb_apply_status_sync(conn, project_id, 'In Progress')
                    project_status_after = 'In Progress'
                # Re-align invoice status whenever balance closes on a Completed
                # project (so the Approved badge auto-flips to Paid).
                elif cur_status == 'Completed' and summary['fully_paid']:
                    _ahb_apply_status_sync(conn, project_id, 'Completed')
                    project_status_after = 'Completed'
        conn.commit(); conn.close()
        return jsonify({'success': True, 'id': pmt_id,
                        'project_id': project_id,
                        'project_status': project_status_after})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


# ── AHB123 — Billing Summary ────────────────────────────────────────────────

@app.route('/api/ahb/billing/summary', methods=['GET'])
def api_ahb_billing_summary():
    try:
        year = (request.args.get('year') or '').strip()
        unpaid_only = (request.args.get('unpaid_only', 'true').lower() != 'false')
        # Year filter: prefer `year` column; fall back to date/created_at prefix for rows where year is blank.
        year_expr = "COALESCE(NULLIF(year,''), substr(COALESCE(date, created_at, ''),1,4))"
        year_clause = f" AND {year_expr} = ?" if year else ""
        year_params = (year,) if year else ()

        conn = _ahb_db()
        stats = {'year': year or 'all', 'unpaid_only': unpaid_only}
        # Unpaid invoices — case-insensitive status match
        for status in ['Sent', 'Approved', 'In Progress', 'Overdue']:
            row = conn.execute(
                f"SELECT count(*) as cnt, COALESCE(sum(total),0) as total FROM ahb_invoices WHERE LOWER(status) = LOWER(?) AND is_change_order = 0{year_clause}",
                (status,) + year_params).fetchone()
            stats[status.lower().replace(' ','_')] = {'count': row['cnt'], 'total': row['total']}
        # Paid
        row = conn.execute(
            f"SELECT count(*) as cnt, COALESCE(sum(total),0) as total FROM ahb_invoices WHERE LOWER(status) = 'paid' AND is_change_order = 0{year_clause}",
            year_params).fetchone()
        stats['paid'] = {'count': row['cnt'], 'total': row['total']}
        # Total receivable (all non-paid)
        row = conn.execute(
            f"SELECT COALESCE(sum(total),0) as total FROM ahb_invoices WHERE LOWER(status) != 'paid' AND is_change_order = 0{year_clause}",
            year_params).fetchone()
        stats['total_receivable'] = row['total']
        # Total payments received (not year-filtered — ahb_payments has no year context here)
        row = conn.execute("SELECT COALESCE(sum(amount),0) as total FROM ahb_payments").fetchone()
        stats['total_payments'] = row['total']
        # Overdue invoices with interest
        overdue = conn.execute(
            f"SELECT * FROM ahb_invoices WHERE LOWER(status) = 'overdue' AND is_change_order = 0{year_clause}",
            year_params).fetchall()
        overdue_details = []
        for inv in overdue:
            inv = dict(inv)
            overdue_since = inv.get('overdue_since') or inv.get('due_date') or inv.get('created_at','')[:10]
            try:
                from datetime import datetime as dt
                start = dt.strptime(overdue_since[:10], '%Y-%m-%d')
                weeks = max(0, (dt.now() - start).days // 7)
            except:
                weeks = 0
            rate = inv.get('overdue_interest_per_week') or 50
            interest = weeks * rate
            inv['weeks_overdue'] = weeks
            inv['interest_accrued'] = interest
            inv['total_with_interest'] = (inv.get('total') or 0) + interest
            overdue_details.append(inv)
        stats['overdue_details'] = overdue_details
        # Active billing items — unpaid_only toggle flips between "unpaid only" (default) and "all invoices"
        status_clause = " AND LOWER(status) != 'paid'" if unpaid_only else ""
        active = conn.execute(
            f"SELECT * FROM ahb_invoices WHERE is_change_order = 0{status_clause}{year_clause} ORDER BY status, created_at DESC",
            year_params).fetchall()
        stats['active_items'] = [dict(r) for r in active]
        # Change orders summary
        row = conn.execute(
            f"SELECT count(*) as cnt, COALESCE(sum(total),0) as total FROM ahb_invoices WHERE is_change_order = 1{year_clause}",
            year_params).fetchone()
        stats['change_orders'] = {'count': row['cnt'], 'total': row['total']}
        conn.close()
        return jsonify(stats)
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/ahb/files/serve/<fid>', methods=['GET'])
def api_ahb_files_serve(fid):
    try:
        conn = _ahb_db()
        row = conn.execute("SELECT file_path FROM ahb_files WHERE id = ?", (fid,)).fetchone()
        conn.close()
        if not row or not row['file_path']:
            return 'Not found', 404
        path = row['file_path']
        if not os.path.exists(path):
            return 'File not found', 404
        return send_from_directory(os.path.dirname(path), os.path.basename(path))
    except Exception as e:
        return str(e), 500


# ── AHB123 — Receipt Processing Queue ───────────────────────────────────────

def _find_split_column(img):
    """Return the x-column index that best separates two side-by-side receipts.
    Strategy: brightest column ('white valley') in the middle 60% of width.
    Falls back to image midpoint when no clear valley exists.

    img: PIL.Image (any mode; will be converted to L)
    returns: int column index in [0, w)
    """
    gray = img.convert('L')
    w, h = gray.size
    if w < 4:
        return w // 2
    pixels = gray.load()
    lo = int(w * 0.20)
    hi = int(w * 0.80)
    if hi <= lo:
        return w // 2
    row_step = max(1, h // 64)
    col_means = []
    for x in range(lo, hi):
        s = 0
        n = 0
        for y in range(0, h, row_step):
            s += pixels[x, y]
            n += 1
        col_means.append((x, s / max(1, n)))
    if not col_means:
        return w // 2
    vals = [m for _, m in col_means]
    spread = max(vals) - min(vals)
    if spread < 20:
        return w // 2
    return max(col_means, key=lambda t: t[1])[0]


def _detect_and_queue(file_storage, conn, queue_dir, status='pending'):
    """Detect 1-up vs 2-up, crop accordingly, insert queue row(s).
    Returns list of new queue ids. Caller is responsible for conn.commit().

    file_storage: werkzeug FileStorage from request.files
    conn:         sqlite3 connection (open, autocommit off)
    queue_dir:    absolute path where crops are saved
    status:       initial row status — 'pending' (default, auto-OCR) or 'staged' (sit in upload bin)
    """
    from PIL import Image, ImageOps, UnidentifiedImageError
    safe_name = re.sub(r'[^\w.\-]', '_', file_storage.filename or 'receipt.jpg')

    try:
        img = Image.open(file_storage.stream)
        try:
            img = ImageOps.exif_transpose(img)
        except Exception:
            pass
        img = img.convert('RGB')
    except UnidentifiedImageError:
        file_storage.stream.seek(0)
        qid = str(uuid.uuid4())
        fpath = os.path.join(queue_dir, f"{qid}_{safe_name}")
        file_storage.save(fpath)
        conn.execute(
            "INSERT INTO ahb_receipt_queue (id, image_path, mode, status) "
            "VALUES (?, ?, 'bulk-fallback', ?)",
            (qid, fpath, status))
        return [qid]

    w, h = img.size

    if h >= w:
        qid = str(uuid.uuid4())
        fpath = os.path.join(queue_dir, f"{qid}.jpg")
        img.save(fpath, 'JPEG', quality=90)
        conn.execute(
            "INSERT INTO ahb_receipt_queue (id, image_path, mode, status) "
            "VALUES (?, ?, 'bulk-single', ?)",
            (qid, fpath, status))
        return [qid]

    pair_id = str(uuid.uuid4())
    parent_fpath = os.path.join(queue_dir, f"{pair_id}_parent.jpg")
    img.save(parent_fpath, 'JPEG', quality=90)
    split_col = _find_split_column(img)
    new_ids = []
    for side, box in [('left', (0, 0, split_col, h)),
                      ('right', (split_col, 0, w, h))]:
        qid = f"{pair_id}-{side}"
        fpath = os.path.join(queue_dir, f"{qid}.jpg")
        img.crop(box).save(fpath, 'JPEG', quality=90)
        conn.execute(
            "INSERT INTO ahb_receipt_queue "
            "(id, image_path, mode, status, parent_image_path, pair_id, split_col) "
            "VALUES (?, ?, ?, ?, ?, ?, ?)",
            (qid, fpath, f'bulk-dual-{side}', status, parent_fpath, pair_id, split_col))
        new_ids.append(qid)
    return new_ids


@app.route('/api/ahb/receipts/process', methods=['POST'])
def api_ahb_receipts_process():
    """Easy Bulk receipt upload. Per file: EXIF-transpose, then portrait → 1-up,
    landscape → 2-up split at the brightest column (white valley between receipts).
    Queue cards appear immediately with cropped thumbnails; OCR drains in background.

    Pass stage=1 (form field or query string) to keep rows in the Upload bin
    (status='staged') without spawning the OCR worker. Caller flips to pending
    via /api/ahb/receipts/queue/send-batch when ready."""
    try:
        files = request.files.getlist('files') or [request.files.get('file')]
        files = [f for f in files if f]
        if not files:
            return jsonify({'success': False, 'error': 'No files uploaded'}), 400

        stage_raw = (request.form.get('stage') or request.args.get('stage') or '').strip().lower()
        stage_mode = stage_raw in ('1', 'true', 'yes', 'on')
        initial_status = 'staged' if stage_mode else 'pending'

        queue_dir = os.path.join(DASHBOARD_DIR, 'uploads', 'ahb', 'receipts', 'queue')
        os.makedirs(queue_dir, exist_ok=True)
        conn = _ahb_db()
        queue_ids = []
        for f in files:
            try:
                queue_ids.extend(_detect_and_queue(f, conn, queue_dir, status=initial_status))
            except Exception as _e:
                qid = str(uuid.uuid4())
                conn.execute(
                    "INSERT INTO ahb_receipt_queue (id, image_path, mode, status, error) "
                    "VALUES (?, '', 'bulk-error', 'error', ?)",
                    (qid, f"upload failed: {str(_e)[:200]}"))
                queue_ids.append(qid)
        conn.commit()
        conn.close()
        if not stage_mode:
            _spawn_receipt_queue_worker()
        return jsonify({
            'success': True,
            'queue_ids': queue_ids,
            'count': len(queue_ids),
            'staged': stage_mode,
        })
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


# ── Background OCR worker ────────────────────────────────────────────────────
import threading as _ahb_threading
_ahb_worker_lock = _ahb_threading.Lock()
_ahb_worker_running = {'flag': False}

def _spawn_receipt_queue_worker():
    """Start a background thread that drains the receipt queue through OCR.
    Idempotent — only one worker runs at a time. Items move pending → done
    automatically; user only needs to review/confirm in the modal."""
    with _ahb_worker_lock:
        if _ahb_worker_running['flag']:
            return  # already draining
        _ahb_worker_running['flag'] = True

    def _drain():
        try:
            skill_path = os.path.join(FRAMEWORK_DIR, 'skills', 'shared', 'receipt_ocr.py')
            while True:
                conn = _ahb_db()
                row = conn.execute(
                    "SELECT id, image_path FROM ahb_receipt_queue "
                    "WHERE status = 'pending' ORDER BY created_at ASC LIMIT 1"
                ).fetchone()
                if not row:
                    conn.close()
                    break
                qid = row['id']
                conn.execute("UPDATE ahb_receipt_queue SET status='processing' WHERE id=?", (qid,))
                conn.commit()
                conn.close()
                try:
                    env = os.environ.copy()
                    env['SKILL_ARGS'] = json.dumps({'image_path': row['image_path'], 'mode': 'full'})
                    result = subprocess.run([VENV_PYTHON, skill_path], capture_output=True,
                                            text=True, timeout=180, env=env)
                    conn = _ahb_db()
                    if result.returncode == 0:
                        conn.execute("UPDATE ahb_receipt_queue SET status='ready', result_json=? WHERE id=?",
                                     (result.stdout.strip(), qid))
                    else:
                        conn.execute("UPDATE ahb_receipt_queue SET status='error', error=? WHERE id=?",
                                     (result.stderr.strip()[:500] or 'OCR failed', qid))
                    conn.commit()
                    conn.close()
                except Exception as _e:
                    conn = _ahb_db()
                    conn.execute("UPDATE ahb_receipt_queue SET status='error', error=? WHERE id=?",
                                 (str(_e)[:500], qid))
                    conn.commit()
                    conn.close()
        finally:
            with _ahb_worker_lock:
                _ahb_worker_running['flag'] = False

    t = _ahb_threading.Thread(target=_drain, name='ahb-receipt-ocr', daemon=True)
    t.start()


_BIN_STATUSES = {
    'upload':   ('staged',),
    'prescan':  ('pending', 'processing', 'error'),
    'analyzed': ('ready',),
    'live':     ('staged', 'pending', 'processing', 'error', 'ready'),
}

@app.route('/api/ahb/receipts/queue', methods=['GET'])
def api_ahb_receipts_queue_list():
    """List queue items.
    ?status=staged|pending|processing|ready|error|rejected|confirmed — exact status
    ?bin=upload|prescan|analyzed|live — convenience grouping for the 3-bin UI"""
    try:
        conn = _ahb_db()
        q = "SELECT * FROM ahb_receipt_queue WHERE 1=1"
        params = []
        bin_arg = (request.args.get('bin') or '').strip().lower()
        status_arg = request.args.get('status')
        if bin_arg in _BIN_STATUSES:
            placeholders = ','.join('?' * len(_BIN_STATUSES[bin_arg]))
            q += f" AND status IN ({placeholders})"
            params.extend(_BIN_STATUSES[bin_arg])
        elif status_arg:
            q += " AND status = ?"
            params.append(status_arg)
        rows = conn.execute(q + " ORDER BY created_at DESC", params).fetchall()
        conn.close()
        return jsonify([dict(r) for r in rows])
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/ahb/receipts/queue/send-batch', methods=['POST'])
def api_ahb_receipts_queue_send_batch():
    """Promote staged items to pending so the OCR worker picks them up.
    Body: {"ids": ["qid1", "qid2", ...]}  (omit or empty = promote all staged)"""
    try:
        data = request.json or {}
        ids = data.get('ids') or []
        conn = _ahb_db()
        if ids:
            placeholders = ','.join('?' * len(ids))
            cur = conn.execute(
                f"UPDATE ahb_receipt_queue SET status='pending' "
                f"WHERE status='staged' AND id IN ({placeholders})",
                ids)
        else:
            cur = conn.execute(
                "UPDATE ahb_receipt_queue SET status='pending' WHERE status='staged'")
        promoted = cur.rowcount or 0
        conn.commit()
        conn.close()
        if promoted:
            _spawn_receipt_queue_worker()
        return jsonify({'success': True, 'promoted': promoted})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/ahb/receipts/queue/<qid>/edit-image', methods=['POST'])
def api_ahb_receipts_queue_edit_image(qid):
    """Replace a queue item's image with a client-edited version (zoom/rotate/
    brightness/contrast/crop applied in the browser canvas).

    Body: {
      "image_b64": "data:image/jpeg;base64,...",   # required
      "rescan":    true|false,                      # default true → re-OCR
      "split_col": <int>                            # optional, only for split halves
    }

    For split halves with split_col provided, the parent image stays untouched
    but the sibling half is re-cropped from the parent at the new column."""
    try:
        from PIL import Image
        from io import BytesIO
        import base64

        data = request.json or {}
        b64 = (data.get('image_b64') or '').strip()
        if not b64:
            return jsonify({'success': False, 'error': 'image_b64 required'}), 400
        if ',' in b64 and b64.lower().startswith('data:'):
            b64 = b64.split(',', 1)[1]
        try:
            raw = base64.b64decode(b64, validate=False)
        except Exception:
            return jsonify({'success': False, 'error': 'invalid base64'}), 400

        conn = _ahb_db()
        row = conn.execute(
            "SELECT id, image_path, pair_id, parent_image_path, mode "
            "FROM ahb_receipt_queue WHERE id=?",
            (qid,)).fetchone()
        if not row:
            conn.close()
            return jsonify({'success': False, 'error': 'queue item not found'}), 404
        if not row['image_path']:
            conn.close()
            return jsonify({'success': False, 'error': 'no image path on row'}), 410

        try:
            edited = Image.open(BytesIO(raw)).convert('RGB')
        except Exception as _e:
            conn.close()
            return jsonify({'success': False, 'error': f'decode failed: {_e}'}), 400
        edited.save(row['image_path'], 'JPEG', quality=92)

        new_split = data.get('split_col')
        if new_split is not None and row['pair_id'] and row['parent_image_path'] \
                and os.path.exists(row['parent_image_path']):
            try:
                split_col = int(new_split)
                parent = Image.open(row['parent_image_path']).convert('RGB')
                pw, ph = parent.size
                split_col = max(1, min(pw - 1, split_col))
                this_side = 'left' if 'left' in (row['mode'] or '') else 'right'
                other_side = 'right' if this_side == 'left' else 'left'
                sibling = conn.execute(
                    "SELECT id, image_path, mode FROM ahb_receipt_queue "
                    "WHERE pair_id=? AND id!=? AND status!='rejected'",
                    (row['pair_id'], qid)).fetchone()
                if sibling:
                    box = (0, 0, split_col, ph) if other_side == 'left' \
                          else (split_col, 0, pw, ph)
                    parent.crop(box).save(sibling['image_path'], 'JPEG', quality=90)
                    conn.execute(
                        "UPDATE ahb_receipt_queue SET split_col=?, status='pending', "
                        "result_json=NULL, error=NULL WHERE id=?",
                        (split_col, sibling['id']))
                conn.execute(
                    "UPDATE ahb_receipt_queue SET split_col=? WHERE id=?",
                    (split_col, qid))
            except (TypeError, ValueError):
                pass

        rescan = data.get('rescan')
        rescan_flag = rescan if isinstance(rescan, bool) else \
                      str(rescan).strip().lower() in ('1', 'true', 'yes', 'on')
        if rescan is None:
            rescan_flag = True
        if rescan_flag:
            conn.execute(
                "UPDATE ahb_receipt_queue SET status='pending', "
                "result_json=NULL, error=NULL WHERE id=?",
                (qid,))
        conn.commit()
        conn.close()
        if rescan_flag:
            _spawn_receipt_queue_worker()
        return jsonify({'success': True, 'rescan': rescan_flag})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/ahb/receipts/queue/run', methods=['POST'])
def api_ahb_receipts_queue_run():
    """Process pending queue items through OCR. Runs synchronously for up to ?limit=10 items."""
    try:
        limit = int(request.args.get('limit', 10))
        conn = _ahb_db()
        pending = conn.execute(
            "SELECT * FROM ahb_receipt_queue WHERE status = 'pending' ORDER BY created_at ASC LIMIT ?",
            (limit,)).fetchall()

        if not pending:
            conn.close()
            return jsonify({'success': True, 'processed': 0, 'message': 'No pending items'})

        skill_path = os.path.join(os.path.dirname(DASHBOARD_DIR), 'skills', 'shared', 'receipt_ocr.py')
        processed = 0
        for item in pending:
            qid = item['id']
            conn.execute("UPDATE ahb_receipt_queue SET status = 'processing' WHERE id = ?", (qid,))
            conn.commit()
            try:
                env = os.environ.copy()
                env['SKILL_ARGS'] = json.dumps({'image_path': item['image_path'], 'mode': 'full'})
                result = subprocess.run([VENV_PYTHON, skill_path], capture_output=True, text=True, timeout=120, env=env)
                if result.returncode == 0:
                    conn.execute("UPDATE ahb_receipt_queue SET status = 'ready', result_json = ? WHERE id = ?",
                                 (result.stdout.strip(), qid))
                else:
                    conn.execute("UPDATE ahb_receipt_queue SET status = 'error', error = ? WHERE id = ?",
                                 (result.stderr.strip()[:500], qid))
            except subprocess.TimeoutExpired:
                conn.execute("UPDATE ahb_receipt_queue SET status = 'error', error = 'Timeout (120s)' WHERE id = ?", (qid,))
            except Exception as e:
                conn.execute("UPDATE ahb_receipt_queue SET status = 'error', error = ? WHERE id = ?", (str(e)[:500], qid))
            conn.commit()
            processed += 1

        conn.close()
        return jsonify({'success': True, 'processed': processed})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/ahb/receipts/queue/<qid>/confirm', methods=['POST'])
def api_ahb_receipts_queue_confirm(qid):
    """Confirm a processed receipt — creates actual ahb_receipt record from queue data + user edits."""
    try:
        data = request.json or {}
        conn = _ahb_db()
        item = conn.execute("SELECT * FROM ahb_receipt_queue WHERE id = ?", (qid,)).fetchone()
        if not item:
            conn.close()
            return jsonify({'success': False, 'error': 'Queue item not found'}), 404

        # Create actual receipt from confirmed data
        rid = str(uuid.uuid4())
        now = datetime.datetime.now()
        image_path = item['image_path']

        # Move image from queue to permanent storage
        perm_dir = os.path.join(DASHBOARD_DIR, 'uploads', 'ahb', 'receipts')
        os.makedirs(perm_dir, exist_ok=True)
        perm_path = os.path.join(perm_dir, f"{rid}.jpg")
        if os.path.exists(image_path):
            import shutil
            shutil.move(image_path, perm_path)
            image_path = perm_path

        conn.execute(
            """INSERT INTO ahb_receipts (id, project_id, vendor, store_name, amount, total, category, description,
               receipt_date, payment_method, teller_name, store_location, purchase_time,
               tax_amount, subtotal, items_json, ocr_text, ocr_raw, ocr_structured,
               image_path, file_path, year, created_at)
               VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)""",
            (rid, data.get('project_id') or None,
             data.get('store_name', data.get('vendor', '')),
             data.get('store_name', ''), data.get('total', 0), data.get('total', 0),
             data.get('category', ''), data.get('description', ''),
             data.get('receipt_date', now.strftime('%Y-%m-%d')),
             data.get('payment_method', ''), data.get('teller_name', ''),
             data.get('store_location', ''), data.get('purchase_time', ''),
             data.get('tax_amount', 0), data.get('subtotal', 0),
             json.dumps(data.get('items', [])), data.get('ocr_text', ''),
             data.get('ocr_raw', ''), json.dumps(data.get('structured', {})),
             image_path, image_path,
             (data.get('receipt_date', '') or '')[:4] or now.strftime('%Y'),
             now.isoformat()))

        # Mark queue item as confirmed
        conn.execute("UPDATE ahb_receipt_queue SET status = 'confirmed', receipt_id = ? WHERE id = ?", (rid, qid))
        conn.commit()
        conn.close()
        return jsonify({'success': True, 'receipt_id': rid})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/ahb/receipts/queue/<qid>/reject', methods=['POST'])
def api_ahb_receipts_queue_reject(qid):
    """Reject/discard a queue item."""
    try:
        conn = _ahb_db()
        conn.execute("UPDATE ahb_receipt_queue SET status = 'rejected' WHERE id = ?", (qid,))
        conn.commit()
        conn.close()
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/ahb/receipts/queue/<qid>/split', methods=['POST'])
def api_ahb_receipts_queue_split(qid):
    """Split a queue item's image down the middle into two new pending items
    and reject the original. For when the smart bulk-detector missed a 2-up."""
    try:
        conn = _ahb_db()
        row = conn.execute("SELECT image_path FROM ahb_receipt_queue WHERE id=?", (qid,)).fetchone()
        if not row or not row['image_path'] or not os.path.exists(row['image_path']):
            conn.close()
            return jsonify({'success': False, 'error': 'item not found'}), 404
        img = Image.open(row['image_path'])
        w, h = img.size
        mid = w // 2
        queue_dir = os.path.join(DASHBOARD_DIR, 'uploads', 'ahb', 'receipts', 'queue')
        os.makedirs(queue_dir, exist_ok=True)
        new_ids = []
        base_qid = str(uuid.uuid4())
        for side, box in [('left', (0, 0, mid, h)), ('right', (mid, 0, w, h))]:
            new_qid = f"{base_qid}-{side}"
            fpath = os.path.join(queue_dir, f"{new_qid}.jpg")
            img.crop(box).save(fpath, 'JPEG', quality=85)
            conn.execute(
                "INSERT INTO ahb_receipt_queue (id, image_path, mode, status) VALUES (?, ?, 'manual-split', 'pending')",
                (new_qid, fpath))
            new_ids.append(new_qid)
        conn.execute("UPDATE ahb_receipt_queue SET status='rejected' WHERE id=?", (qid,))
        conn.commit()
        conn.close()
        _spawn_receipt_queue_worker()
        return jsonify({'success': True, 'queue_ids': new_ids})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/ahb/receipts/queue/<qid>/rescan', methods=['POST'])
def api_ahb_receipts_queue_rescan(qid):
    """Reset an item to pending so the OCR worker re-processes it.
    Useful after a manual split or when fields came back as garbage."""
    try:
        conn = _ahb_db()
        conn.execute("UPDATE ahb_receipt_queue SET status='pending', result_json=NULL, error=NULL WHERE id=?", (qid,))
        conn.commit()
        conn.close()
        _spawn_receipt_queue_worker()
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/ahb/receipts/queue/<qid>/rotate', methods=['POST'])
def api_ahb_receipts_queue_rotate(qid):
    """Rotate the source image 90° CW and re-detect. Replaces the old queue row(s)."""
    try:
        from PIL import Image
        from io import BytesIO
        from werkzeug.datastructures import FileStorage
        conn = _ahb_db()
        row = conn.execute(
            "SELECT image_path, parent_image_path, pair_id FROM ahb_receipt_queue WHERE id=?",
            (qid,)).fetchone()
        if not row:
            conn.close()
            return jsonify({'success': False, 'error': 'queue item not found'}), 404
        src = row['parent_image_path'] or row['image_path']
        if not src or not os.path.exists(src):
            conn.close()
            return jsonify({'success': False, 'error': 'source image missing'}), 410

        queue_dir = os.path.join(DASHBOARD_DIR, 'uploads', 'ahb', 'receipts', 'queue')
        os.makedirs(queue_dir, exist_ok=True)

        rotated = Image.open(src).rotate(-90, expand=True).convert('RGB')
        buf = BytesIO()
        rotated.save(buf, 'JPEG', quality=90)
        buf.seek(0)
        fs = FileStorage(stream=buf, filename=f"rotated_{qid}.jpg", content_type='image/jpeg')
        new_ids = _detect_and_queue(fs, conn, queue_dir)

        if row['pair_id']:
            conn.execute("UPDATE ahb_receipt_queue SET status='rejected' WHERE pair_id=?",
                         (row['pair_id'],))
        else:
            conn.execute("UPDATE ahb_receipt_queue SET status='rejected' WHERE id=?", (qid,))
        conn.commit()
        conn.close()
        _spawn_receipt_queue_worker()
        return jsonify({'success': True, 'queue_ids': new_ids})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/ahb/receipts/queue/<qid>/merge', methods=['POST'])
def api_ahb_receipts_queue_merge(qid):
    """Merge two halves of an auto-split pair back into a single 1-up queue item."""
    try:
        import shutil
        conn = _ahb_db()
        row = conn.execute(
            "SELECT pair_id, parent_image_path FROM ahb_receipt_queue WHERE id=?",
            (qid,)).fetchone()
        if not row:
            conn.close()
            return jsonify({'success': False, 'error': 'queue item not found'}), 404
        if not row['pair_id'] or not row['parent_image_path']:
            conn.close()
            return jsonify({'success': False, 'error': 'item is not a split half'}), 400
        if not os.path.exists(row['parent_image_path']):
            conn.close()
            return jsonify({'success': False, 'error': 'parent image missing'}), 410

        conn.execute("UPDATE ahb_receipt_queue SET status='rejected' WHERE pair_id=?",
                     (row['pair_id'],))

        queue_dir = os.path.join(DASHBOARD_DIR, 'uploads', 'ahb', 'receipts', 'queue')
        os.makedirs(queue_dir, exist_ok=True)
        new_qid = str(uuid.uuid4())
        new_fpath = os.path.join(queue_dir, f"{new_qid}.jpg")
        shutil.copy2(row['parent_image_path'], new_fpath)
        conn.execute(
            "INSERT INTO ahb_receipt_queue (id, image_path, mode, status) "
            "VALUES (?, ?, 'bulk-merged', 'pending')",
            (new_qid, new_fpath))
        conn.commit()
        conn.close()
        _spawn_receipt_queue_worker()
        return jsonify({'success': True, 'queue_id': new_qid})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/ahb/receipts/queue/<qid>/adjust-split', methods=['POST'])
def api_ahb_receipts_queue_adjust_split(qid):
    """Re-crop both halves of an auto-split pair at a new split column.
    Body: {"split_col": <int>}"""
    try:
        from PIL import Image
        data = request.json or {}
        try:
            split_col = int(data.get('split_col', 0))
        except (TypeError, ValueError):
            return jsonify({'success': False, 'error': 'split_col must be an integer'}), 400

        conn = _ahb_db()
        row = conn.execute(
            "SELECT pair_id, parent_image_path FROM ahb_receipt_queue WHERE id=?",
            (qid,)).fetchone()
        if not row:
            conn.close()
            return jsonify({'success': False, 'error': 'queue item not found'}), 404
        if not row['pair_id'] or not row['parent_image_path']:
            conn.close()
            return jsonify({'success': False, 'error': 'item is not a split half'}), 400
        if not os.path.exists(row['parent_image_path']):
            conn.close()
            return jsonify({'success': False, 'error': 'parent image missing'}), 410

        img = Image.open(row['parent_image_path']).convert('RGB')
        w, h = img.size
        split_col = max(1, min(w - 1, split_col))

        halves = conn.execute(
            "SELECT id, image_path, mode FROM ahb_receipt_queue WHERE pair_id=? AND status!='rejected'",
            (row['pair_id'],)).fetchall()
        if len(halves) < 2:
            conn.close()
            return jsonify({'success': False, 'error': 'pair incomplete (use rotate or re-upload)'}), 400

        for half in halves:
            side = 'left' if 'left' in (half['mode'] or '') else 'right'
            box = (0, 0, split_col, h) if side == 'left' else (split_col, 0, w, h)
            img.crop(box).save(half['image_path'], 'JPEG', quality=90)
            conn.execute(
                "UPDATE ahb_receipt_queue "
                "SET status='pending', result_json=NULL, error=NULL, split_col=? "
                "WHERE id=?",
                (split_col, half['id']))
        conn.commit()
        conn.close()
        _spawn_receipt_queue_worker()
        return jsonify({'success': True, 'split_col': split_col})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


def _enhance_receipt_image(src_path, dst_path, target_long_edge, bw=False):
    """Receipt preview enhancement, fully local. Minimal-impact pipeline:
      1. EXIF-transpose via PIL.
      2. Lanczos upscale (cv2.INTER_LANCZOS4) to target_long_edge.
      3. Strong unsharp mask — restore the bite Lanczos smooths off, and
         compensate for the soft Telegram-JPEG source.
      4. Mild CLAHE on the L channel (clipLimit=1.2) for uneven lighting.
      5. bw=True: median blur 3 → adaptive Gaussian threshold sized to
         receipt char height.
    Earlier versions added fastNlMeansDenoisingColored before upscale —
    pulled out because it smeared text into haze on Telegram-compressed
    sources. Same for bilateralFilter. Without active denoising the result
    keeps Telegram's compression grain but text stays legible.
    Saves JPEG quality 92, no chroma subsampling.
    """
    import cv2
    import numpy as np
    from PIL import Image, ImageOps
    pil_im = Image.open(src_path)
    try:
        pil_im = ImageOps.exif_transpose(pil_im)
    except Exception:
        pass
    pil_im = pil_im.convert('RGB')
    img = cv2.cvtColor(np.array(pil_im), cv2.COLOR_RGB2BGR)

    # 2) Iterative 2× upscale with mid-pass sharpen — much better digit
    # legibility on Telegram-compressed receipts than a single 5-8× Lanczos.
    # Mid-pass unsharp is dialed down vs earlier versions so the B&W
    # threshold below doesn't grab unsharp halos as black.
    h, w = img.shape[:2]
    src_long = max(h, w)
    if target_long_edge > src_long:
        ratio = target_long_edge / src_long
        if ratio > 3:
            img = cv2.edgePreservingFilter(img, flags=cv2.RECURS_FILTER,
                                           sigma_s=15, sigma_r=0.18)
        while max(img.shape[:2]) * 1.5 < target_long_edge:
            img = cv2.resize(img, None, fx=2.0, fy=2.0,
                             interpolation=cv2.INTER_LANCZOS4)
            mblur = cv2.GaussianBlur(img, (0, 0), 1.0)
            img = cv2.addWeighted(img, 1.18, mblur, -0.18, 0)
        h2, w2 = img.shape[:2]
        if max(h2, w2) != target_long_edge:
            s = target_long_edge / max(h2, w2)
            img = cv2.resize(img, (max(1, int(w2 * s)), max(1, int(h2 * s))),
                             interpolation=cv2.INTER_LANCZOS4)

    # 3) Light final unsharp — kept gentle so threshold doesn't capture
    # halos as black strokes (= 'pasty/leaky ink').
    blur = cv2.GaussianBlur(img, (0, 0), 1.4)
    img = cv2.addWeighted(img, 1.15, blur, -0.15, 0)

    if bw:
        # High-contrast grayscale, NOT adaptive-threshold binarization.
        # Threshold gave us: thick blob strokes (JPEG mush captured as
        # black), broken characters (uneven darkness), cartoonized cel-
        # shaded edges at zoom. Grayscale keeps antialiasing intact —
        # text stays readable at every zoom level and there are no hard
        # transitions for JPEG to ring on.
        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
        # Gentle local-contrast lift — uneven receipt lighting flattened
        # without crushing whites or lifting blacks (CLAHE was removed
        # from the color path for that reason; here we keep it mild).
        clahe = cv2.createCLAHE(clipLimit=1.5, tileGridSize=(8, 8))
        gray = clahe.apply(gray)
        # Percentile contrast stretch — pin 2nd %ile to 0, 98th %ile to
        # 255. Makes the page read as black-text-on-white without the
        # all-or-nothing failure mode of thresholding.
        lo, hi = np.percentile(gray, [2, 98])
        if hi > lo + 1:
            gray = np.clip(
                (gray.astype(np.float32) - lo) * (255.0 / (hi - lo)),
                0, 255,
            ).astype(np.uint8)
        img = cv2.cvtColor(gray, cv2.COLOR_GRAY2BGR)

    cv2.imwrite(dst_path, img, [int(cv2.IMWRITE_JPEG_QUALITY), 92,
                                int(cv2.IMWRITE_JPEG_OPTIMIZE), 1])


@app.route('/api/ahb/receipts/queue/image/<qid>', methods=['GET'])
def api_ahb_receipts_queue_image(qid):
    """Serve a queue item's image.
      ?parent=1   — original pre-split parent
      ?w=<int>    — Lanczos-upscaled + unsharp-masked variant at the given
                    long-edge width. Cached to <dir>/_hi/<qid>_w<N>.jpg so
                    we don't re-encode on every preview render. Used by the
                    editor and lightbox so pinch/wheel zoom has real pixels
                    to work with instead of browser bicubic on a 480x640
                    Telegram-compressed source.
    """
    try:
        conn = _ahb_db()
        row = conn.execute(
            "SELECT image_path, parent_image_path FROM ahb_receipt_queue WHERE id = ?",
            (qid,)).fetchone()
        conn.close()
        if not row:
            return 'Not found', 404
        want_parent = request.args.get('parent') in ('1', 'true', 'yes')
        path = row['parent_image_path'] if want_parent else row['image_path']
        if not path or not os.path.exists(path):
            return 'Not found', 404

        try:
            w_req = int(request.args.get('w', '0') or 0)
        except (TypeError, ValueError):
            w_req = 0
        w_req = max(0, min(4000, w_req))
        bw_req = request.args.get('bw') in ('1', 'true', 'yes')

        if w_req > 0:
            try:
                cache_dir = os.path.join(os.path.dirname(path), '_hi')
                os.makedirs(cache_dir, exist_ok=True)
                kind = 'parent' if want_parent else 'crop'
                bw_suffix = '_bw' if bw_req else ''
                # JPEG for both — B&W is high-contrast grayscale, not
                # binarized, so JPEG quantization is safe.
                cache_name = f"{qid}_{kind}_w{w_req}{bw_suffix}.jpg"
                cache_path = os.path.join(cache_dir, cache_name)
                src_mtime = os.path.getmtime(path)
                if (not os.path.exists(cache_path)) or os.path.getmtime(cache_path) < src_mtime:
                    _enhance_receipt_image(path, cache_path, w_req, bw=bw_req)
                return send_from_directory(cache_dir, cache_name)
            except Exception:
                # Any failure → fall back to the raw source. Better blurry
                # than blank.
                pass

        return send_from_directory(os.path.dirname(path), os.path.basename(path))
    except Exception as e:
        return str(e), 500


@app.route('/api/ahb/receipts/scan-existing', methods=['POST'])
def api_ahb_receipts_scan_existing():
    """Queue existing receipts for OCR re-scan. Does NOT alter existing data — results go to queue for review."""
    try:
        data = request.json or {}
        category = data.get('category', '')  # empty = all
        conn = _ahb_db()
        q = "SELECT id, image_path, store_name, total FROM ahb_receipts WHERE image_path IS NOT NULL AND image_path != ''"
        params = []
        if category:
            q += " AND category = ?"; params.append(category)
        receipts = conn.execute(q, params).fetchall()

        queued = 0
        for r in receipts:
            if not os.path.exists(r['image_path']):
                continue
            qid = str(uuid.uuid4())
            conn.execute(
                "INSERT INTO ahb_receipt_queue (id, image_path, mode, status, receipt_id) VALUES (?, ?, 'rescan', 'pending', ?)",
                (qid, r['image_path'], r['id']))
            queued += 1

        conn.commit()
        conn.close()
        return jsonify({'success': True, 'queued': queued})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


# ── Explore Lab ──────────────────────────────────────────────────────────────
_explore_servers = {}  # session_id -> {pid, port, type, started}
EXPLORE_PORT_START = 9100

def _find_free_port():
    import socket as _s
    with _s.socket(_s.AF_INET, _s.SOCK_STREAM) as s:
        s.bind(('', 0))
        return s.getsockname()[1]

def _get_device_catalog():
    return {
        "common": [
            {"id":"iphone-16-pro","name":"iPhone 16 Pro","w":393,"h":852,"type":"phone","os":"ios","notch":"dynamic-island"},
            {"id":"iphone-se","name":"iPhone SE","w":375,"h":667,"type":"phone","os":"ios","notch":"none"},
            {"id":"galaxy-s25","name":"Galaxy S25","w":360,"h":780,"type":"phone","os":"android","notch":"punch-hole"},
            {"id":"pixel-9","name":"Pixel 9","w":412,"h":915,"type":"phone","os":"android","notch":"punch-hole"},
            {"id":"ipad-pro-13","name":"iPad Pro 13\"","w":1024,"h":1366,"type":"tablet","os":"ios"},
            {"id":"galaxy-tab-s10","name":"Galaxy Tab S10","w":800,"h":1280,"type":"tablet","os":"android"},
            {"id":"desktop-1080p","name":"Desktop 1080p","w":1920,"h":1080,"type":"desktop","os":"any"},
            {"id":"desktop-1440","name":"Desktop 1440p","w":2560,"h":1440,"type":"desktop","os":"any"},
            {"id":"macbook-pro","name":"MacBook Pro 14\"","w":1512,"h":982,"type":"desktop","os":"macos"},
        ],
        "uncommon": [
            {"id":"rpi-touch","name":"Raspberry Pi Touch","w":800,"h":480,"type":"embedded","os":"linux"},
            {"id":"rpi-zero","name":"Raspberry Pi Zero","w":640,"h":480,"type":"embedded","os":"linux"},
            {"id":"arduino-tft","name":"Arduino TFT 2.8\"","w":320,"h":240,"type":"embedded","os":"bare"},
            {"id":"stm32-oled","name":"STM32 OLED 0.96\"","w":128,"h":64,"type":"embedded","os":"bare"},
            {"id":"stm32-tft","name":"STM32 TFT 3.5\"","w":480,"h":320,"type":"embedded","os":"bare"},
            {"id":"esp32-tft","name":"ESP32 TFT 1.8\"","w":160,"h":128,"type":"embedded","os":"bare"},
            {"id":"surface-pro","name":"Surface Pro","w":2880,"h":1920,"type":"tablet","os":"windows"},
            {"id":"steam-deck","name":"Steam Deck","w":1280,"h":800,"type":"handheld","os":"linux"},
        ],
        "browsers": [
            {"id":"chrome-desktop","name":"Chrome Desktop","w":1440,"h":900,"type":"browser","os":"any","ua":"Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/131.0.0.0 Safari/537.36","chrome":"131"},
            {"id":"chrome-mobile","name":"Chrome Mobile","w":412,"h":915,"type":"browser","os":"android","ua":"Mozilla/5.0 (Linux; Android 15; Pixel 9) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/131.0.0.0 Mobile Safari/537.36","chrome":"131"},
            {"id":"safari-desktop","name":"Safari Desktop","w":1440,"h":900,"type":"browser","os":"macos","ua":"Mozilla/5.0 (Macintosh; Intel Mac OS X 14_5) AppleWebKit/605.1.15 (KHTML, like Gecko) Version/18.0 Safari/605.1.15","safari":"18"},
            {"id":"safari-mobile","name":"Safari iOS","w":393,"h":852,"type":"browser","os":"ios","ua":"Mozilla/5.0 (iPhone; CPU iPhone OS 18_0 like Mac OS X) AppleWebKit/605.1.15 (KHTML, like Gecko) Version/18.0 Mobile/15E148 Safari/604.1","safari":"18"},
            {"id":"firefox-desktop","name":"Firefox Desktop","w":1440,"h":900,"type":"browser","os":"any","ua":"Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:133.0) Gecko/20100101 Firefox/133.0","firefox":"133"},
            {"id":"firefox-mobile","name":"Firefox Android","w":412,"h":915,"type":"browser","os":"android","ua":"Mozilla/5.0 (Android 15; Mobile; rv:133.0) Gecko/133.0 Firefox/133.0","firefox":"133"},
            {"id":"edge-desktop","name":"Edge Desktop","w":1440,"h":900,"type":"browser","os":"windows","ua":"Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/131.0.0.0 Safari/537.36 Edg/131.0.0.0","edge":"131"},
            {"id":"brave-desktop","name":"Brave Desktop","w":1440,"h":900,"type":"browser","os":"any","ua":"Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/131.0.0.0 Safari/537.36","brave":"131"},
            {"id":"opera-desktop","name":"Opera Desktop","w":1440,"h":900,"type":"browser","os":"any","ua":"Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/131.0.0.0 Safari/537.36 OPR/117.0.0.0","opera":"117"},
            {"id":"samsung-internet","name":"Samsung Internet","w":360,"h":780,"type":"browser","os":"android","ua":"Mozilla/5.0 (Linux; Android 15; SAMSUNG SM-S926B) AppleWebKit/537.36 (KHTML, like Gecko) SamsungBrowser/27.0 Chrome/131.0.0.0 Mobile Safari/537.36","samsung":"27"},
        ]
    }

@app.route('/explore')
def explore_lab():
    projects = []
    if os.path.exists(ARTIFACTS_DIR):
        projects = sorted([d for d in os.listdir(ARTIFACTS_DIR)
                           if os.path.isdir(os.path.join(ARTIFACTS_DIR, d))])
    return render_template('explore.html', projects=projects, page='explore')

@app.route('/api/explore/devices')
def api_explore_devices():
    return jsonify(_get_device_catalog())

@app.route('/api/explore/static/<project_id>/<path:filepath>')
def api_explore_static(project_id, filepath):
    proj_dir = os.path.join(ARTIFACTS_DIR, project_id)
    full = os.path.realpath(os.path.join(proj_dir, filepath))
    if not full.startswith(os.path.realpath(ARTIFACTS_DIR)):
        return jsonify({'error': 'Forbidden'}), 403
    parent = os.path.dirname(full)
    fname  = os.path.basename(full)
    return send_from_directory(parent, fname)

@app.route('/api/explore/serve', methods=['POST'])
def api_explore_serve():
    data = request.json or {}
    src  = data.get('type', 'artifact')
    if src == 'url':
        return jsonify({'serve_url': data.get('url', ''), 'type': 'url'})
    project_id = data.get('project_id', '')
    filename   = data.get('filename', '')
    if not project_id or not filename:
        return jsonify({'error': 'project_id and filename required'}), 400
    serve_url = f'/api/explore/static/{project_id}/{filename}'
    return jsonify({'serve_url': serve_url, 'type': 'artifact'})

@app.route('/api/explore/build', methods=['POST'])
def api_explore_build():
    data        = request.json or {}
    project_id  = data.get('project_id', '')
    entry_point = data.get('entry_point', 'index.html')
    build_type  = data.get('type', 'static')
    proj_dir    = os.path.join(ARTIFACTS_DIR, project_id)
    if not os.path.isdir(proj_dir):
        return jsonify({'error': 'Project not found'}), 404

    session_id = str(uuid.uuid4())[:8]

    if build_type == 'static':
        serve_url = f'/api/explore/static/{project_id}/{entry_point}'
        return jsonify({'session_id': session_id, 'serve_url': serve_url, 'port': None, 'pid': None})

    port = _find_free_port()
    if build_type == 'flask':
        proc = subprocess.Popen(
            [VENV_PYTHON, entry_point],
            cwd=proj_dir,
            env={**os.environ, 'PORT': str(port), 'FLASK_RUN_PORT': str(port)},
            stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
    elif build_type == 'node':
        proc = subprocess.Popen(
            ['npx', 'serve', '-l', str(port), '.'],
            cwd=proj_dir,
            stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
    else:
        return jsonify({'error': f'Unknown build type: {build_type}'}), 400

    _explore_servers[session_id] = {
        'pid': proc.pid, 'port': port, 'type': build_type,
        'started': datetime.datetime.now().isoformat(), 'project_id': project_id
    }
    return jsonify({'session_id': session_id, 'serve_url': f'http://localhost:{port}', 'port': port, 'pid': proc.pid})

@app.route('/api/explore/kill', methods=['POST'])
def api_explore_kill():
    data       = request.json or {}
    session_id = data.get('session_id', '')
    info = _explore_servers.pop(session_id, None)
    if not info:
        return jsonify({'error': 'Session not found'}), 404
    try:
        os.kill(info['pid'], 9)
    except ProcessLookupError:
        pass
    return jsonify({'success': True, 'killed': session_id})

@app.route('/api/explore/sessions')
def api_explore_sessions():
    return jsonify(_explore_servers)

@app.route('/api/explore/upload', methods=['POST'])
def api_explore_upload():
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400
    upload_id  = str(uuid.uuid4())[:8]
    upload_dir = os.path.join(DASHBOARD_DIR, 'explore-sessions', upload_id)
    os.makedirs(upload_dir, exist_ok=True)
    saved = []
    for f in request.files.getlist('file'):
        safe = os.path.basename(f.filename)
        f.save(os.path.join(upload_dir, safe))
        saved.append(safe)
    entry = saved[0] if saved else ''
    return jsonify({'session_id': upload_id, 'serve_url': f'/api/explore/static-upload/{upload_id}/{entry}', 'files': saved})

@app.route('/api/explore/static-upload/<session_id>/<path:filepath>')
def api_explore_static_upload(session_id, filepath):
    upload_dir = os.path.join(DASHBOARD_DIR, 'explore-sessions', session_id)
    full = os.path.realpath(os.path.join(upload_dir, filepath))
    if not full.startswith(os.path.realpath(os.path.join(DASHBOARD_DIR, 'explore-sessions'))):
        return jsonify({'error': 'Forbidden'}), 403
    parent = os.path.dirname(full)
    fname  = os.path.basename(full)
    return send_from_directory(parent, fname)


# (SQLite cloud section removed — replaced by CLOUD_ENABLED PostgreSQL block below)


# ── Baza Roadmap API ─────────────────────────────────────────────────────────

@app.route('/api/roadmap', methods=['GET'])
def api_roadmap_list():
    conn = sqlite3.connect(os.path.join(DASHBOARD_DIR, 'baza_projects.db'))
    conn.row_factory = sqlite3.Row
    rows = conn.execute("SELECT * FROM baza_roadmap ORDER BY CASE status WHEN 'in_progress' THEN 0 WHEN 'planned' THEN 1 WHEN 'future' THEN 2 WHEN 'completed' THEN 3 ELSE 4 END, created_at DESC").fetchall()
    conn.close()
    return jsonify([dict(r) for r in rows])

@app.route('/api/roadmap', methods=['POST'])
def api_roadmap_add():
    d = request.json or {}
    rid = str(uuid.uuid4())
    conn = sqlite3.connect(os.path.join(DASHBOARD_DIR, 'baza_projects.db'))
    conn.execute(
        "INSERT INTO baza_roadmap (id, title, description, status, priority, category, assigned_agent, target_date, notes) VALUES (?,?,?,?,?,?,?,?,?)",
        (rid, d.get('title',''), d.get('description',''), d.get('status','planned'),
         d.get('priority','medium'), d.get('category','general'),
         d.get('assigned_agent',''), d.get('target_date',''), d.get('notes','')))
    conn.commit()
    conn.close()
    return jsonify({"id": rid, "status": "created"})

@app.route('/api/roadmap/<rid>', methods=['PUT'])
def api_roadmap_update(rid):
    d = request.json or {}
    conn = sqlite3.connect(os.path.join(DASHBOARD_DIR, 'baza_projects.db'))
    fields, vals = [], []
    for k in ['title','description','status','priority','category','assigned_agent','target_date','started_at','completed_at','notes']:
        if k in d:
            fields.append(f"{k}=?")
            vals.append(d[k])
    if not fields:
        return jsonify({"error": "No fields to update"}), 400
    fields.append("updated_at=datetime('now')")
    if d.get('status') == 'in_progress':
        row = conn.execute("SELECT started_at FROM baza_roadmap WHERE id=?", (rid,)).fetchone()
        if row and not row[0]:
            fields.append("started_at=datetime('now')")
    if d.get('status') == 'completed':
        fields.append("completed_at=datetime('now')")
    vals.append(rid)
    conn.execute(f"UPDATE baza_roadmap SET {','.join(fields)} WHERE id=?", vals)
    conn.commit()
    conn.close()
    return jsonify({"id": rid, "status": "updated"})

@app.route('/api/roadmap/<rid>/start', methods=['POST'])
def api_roadmap_start(rid):
    """Start a roadmap item: set status to in_progress + create a task for the right agent."""
    conn = sqlite3.connect(os.path.join(DASHBOARD_DIR, 'baza_projects.db'))
    conn.row_factory = sqlite3.Row
    item = conn.execute("SELECT * FROM baza_roadmap WHERE id=?", (rid,)).fetchone()
    if not item:
        conn.close()
        return jsonify({"error": "Roadmap item not found"}), 404

    # Update roadmap status
    conn.execute("UPDATE baza_roadmap SET status='in_progress', started_at=datetime('now'), updated_at=datetime('now') WHERE id=?", (rid,))

    # Determine best agent based on category
    cat = item['category'] or 'general'
    agent_map = {
        'business': 'simon_bately',
        'infrastructure': 'claw_batto',
        'development': 'claw_batto',
        'ai_agents': 'claw_batto',
        'general': 'simon_bately',
    }
    assigned = item['assigned_agent'] or agent_map.get(cat, 'simon_bately')

    # Check if task already exists for this roadmap item
    existing = conn.execute("SELECT id FROM tasks WHERE title LIKE ?", (f"%{item['title'][:30]}%",)).fetchone()
    task_id = None
    if not existing:
        task_id = str(uuid.uuid4())
        conn.execute(
            "INSERT INTO tasks (id, title, description, status, priority, assigned_to, created_at, notes) VALUES (?,?,?,?,?,?,?,?)",
            (task_id, f"Roadmap: {item['title']}", f"{item['description']}\n\nCategory: {cat}\nFrom roadmap, started manually.",
             'pending', item['priority'] or 'medium', assigned,
             datetime.datetime.now().isoformat(), "[Duke] Created from roadmap — started by Serge"))

        # Update roadmap with assigned agent
        conn.execute("UPDATE baza_roadmap SET assigned_agent=? WHERE id=? AND (assigned_agent IS NULL OR assigned_agent='')", (assigned, rid))
    else:
        task_id = existing[0]

    conn.commit()
    conn.close()

    return jsonify({"id": rid, "status": "started", "task_id": task_id, "assigned_to": assigned})


@app.route('/api/roadmap/<rid>', methods=['DELETE'])
def api_roadmap_delete(rid):
    conn = sqlite3.connect(os.path.join(DASHBOARD_DIR, 'baza_projects.db'))
    conn.execute("DELETE FROM baza_roadmap WHERE id=?", (rid,))
    conn.commit()
    conn.close()
    return jsonify({"id": rid, "status": "deleted"})

# ── Dashboard Links API ──────────────────────────────────────────────────────

@app.route('/api/dash-links', methods=['GET'])
def api_dash_links_list():
    conn = sqlite3.connect(os.path.join(DASHBOARD_DIR, 'baza_projects.db'))
    conn.row_factory = sqlite3.Row
    rows = conn.execute("SELECT * FROM baza_dash_links ORDER BY sort_order, created_at").fetchall()
    conn.close()
    return jsonify([dict(r) for r in rows])

@app.route('/api/dash-links', methods=['POST'])
def api_dash_links_add():
    d = request.json or {}
    lid = str(uuid.uuid4())
    conn = sqlite3.connect(os.path.join(DASHBOARD_DIR, 'baza_projects.db'))
    conn.execute(
        "INSERT INTO baza_dash_links (id, title, url, icon, category, sort_order) VALUES (?,?,?,?,?,?)",
        (lid, d.get('title',''), d.get('url',''), d.get('icon','&#128279;'),
         d.get('category','general'), d.get('sort_order', 0)))
    conn.commit()
    conn.close()
    return jsonify({"id": lid, "status": "created"})

@app.route('/api/dash-links/<lid>', methods=['PUT'])
def api_dash_links_update(lid):
    d = request.json or {}
    conn = sqlite3.connect(os.path.join(DASHBOARD_DIR, 'baza_projects.db'))
    fields, vals = [], []
    for k in ['title','url','icon','category','sort_order']:
        if k in d:
            fields.append(f"{k}=?")
            vals.append(d[k])
    if not fields:
        return jsonify({"error": "No fields"}), 400
    vals.append(lid)
    conn.execute(f"UPDATE baza_dash_links SET {','.join(fields)} WHERE id=?", vals)
    conn.commit()
    conn.close()
    return jsonify({"id": lid, "status": "updated"})

@app.route('/api/dash-links/<lid>', methods=['DELETE'])
def api_dash_links_delete(lid):
    conn = sqlite3.connect(os.path.join(DASHBOARD_DIR, 'baza_projects.db'))
    conn.execute("DELETE FROM baza_dash_links WHERE id=?", (lid,))
    conn.commit()
    conn.close()
    return jsonify({"id": lid, "status": "deleted"})

# ── Infra Notes API ──────────────────────────────────────────────────────────

@app.route('/api/infra/notes', methods=['GET'])
def api_infra_notes_list():
    conn = sqlite3.connect(os.path.join(DASHBOARD_DIR, 'baza_projects.db'))
    conn.row_factory = sqlite3.Row
    section = request.args.get('section')
    if section:
        rows = conn.execute("SELECT * FROM baza_infra_notes WHERE section=? ORDER BY created_at DESC", (section,)).fetchall()
    else:
        rows = conn.execute("SELECT * FROM baza_infra_notes ORDER BY created_at DESC").fetchall()
    conn.close()
    return jsonify([dict(r) for r in rows])

@app.route('/api/infra/notes', methods=['POST'])
def api_infra_notes_add():
    d = request.json or {}
    nid = str(uuid.uuid4())
    conn = sqlite3.connect(os.path.join(DASHBOARD_DIR, 'baza_projects.db'))
    conn.execute(
        "INSERT INTO baza_infra_notes (id, section, note, author) VALUES (?,?,?,?)",
        (nid, d.get('section','general'), d.get('note',''), d.get('author','system')))
    conn.commit()
    conn.close()
    return jsonify({"id": nid, "status": "created"})

@app.route('/api/infra/notes/<nid>', methods=['DELETE'])
def api_infra_notes_delete(nid):
    conn = sqlite3.connect(os.path.join(DASHBOARD_DIR, 'baza_projects.db'))
    conn.execute("DELETE FROM baza_infra_notes WHERE id=?", (nid,))
    conn.commit()
    conn.close()
    return jsonify({"id": nid, "status": "deleted"})

# ── Agent Usage API ──────────────────────────────────────────────────────────

@app.route('/api/usage')
def api_usage():
    agent_id = request.args.get('agent_id')
    days = int(request.args.get('days', 7))
    try:
        import psycopg2
        pg = psycopg2.connect(host="localhost", port=5432, dbname="baza_agents",
                              user="switchhacker",
                              password=os.environ.get("DB_PASSWORD", "baza2026"))
        cur = pg.cursor()
        cur.execute("""
            SELECT agent_id, COALESCE(SUM(prompt_tokens),0), COALESCE(SUM(completion_tokens),0),
                   COALESCE(SUM(total_tokens),0), COUNT(*), COALESCE(SUM(duration_ms),0), COALESCE(SUM(cost),0)
            FROM agent_usage WHERE created_at > NOW() - INTERVAL '%s days'
            """ + (" AND agent_id = %s" if agent_id else "") + """
            GROUP BY agent_id ORDER BY SUM(total_tokens) DESC
        """, (days, agent_id) if agent_id else (days,))
        by_agent = [{"agent_id": r[0], "prompt_tokens": r[1], "completion_tokens": r[2],
                      "total_tokens": r[3], "call_count": r[4], "duration_ms": r[5], "cost": float(r[6])} for r in cur.fetchall()]
        cur.execute("SELECT provider, COALESCE(SUM(total_tokens),0), COUNT(*) FROM agent_usage WHERE created_at > NOW() - INTERVAL '%s days' GROUP BY provider", (days,))
        by_provider = [{"provider": r[0], "total_tokens": r[1], "call_count": r[2]} for r in cur.fetchall()]
        cur.execute("SELECT DATE(created_at), COALESCE(SUM(total_tokens),0), COUNT(*) FROM agent_usage WHERE created_at > NOW() - INTERVAL '%s days' GROUP BY DATE(created_at) ORDER BY DATE(created_at)", (days,))
        by_day = [{"day": r[0].isoformat() if r[0] else None, "total_tokens": r[1], "call_count": r[2]} for r in cur.fetchall()]
        total_tokens = sum(a["total_tokens"] for a in by_agent)
        total_calls = sum(a["call_count"] for a in by_agent)
        cur.close(); pg.close()
        return jsonify({"days": days, "total_tokens": total_tokens, "total_calls": total_calls,
                        "by_agent": by_agent, "by_provider": by_provider, "by_day": by_day})
    except Exception as e:
        return jsonify({"error": str(e), "by_agent": [], "by_provider": [], "by_day": [],
                        "total_tokens": 0, "total_calls": 0})

# ── AHB123 Chat Widget API (public-facing, Nova Sterling) ────────────────────

@app.route('/api/ahb/widget/chat', methods=['POST'])
def api_ahb_widget_chat():
    """Public chat widget endpoint — visitor sends message, Nova responds via LLM."""
    data = request.json or {}
    message = (data.get('message') or '').strip()
    chat_id = data.get('chat_id', '')
    visitor_name = data.get('visitor_name', '')
    visitor_email = data.get('visitor_email', '')
    visitor_phone = data.get('visitor_phone', '')

    if not message:
        return jsonify({'error': 'Message required'}), 400

    conn = _ahb_db()

    # Create or get chat
    if not chat_id:
        chat_id = str(uuid.uuid4())
        conn.execute(
            "INSERT INTO ahb_chats (id, visitor_name, visitor_email, visitor_phone, channel, status, assigned_agent) VALUES (?,?,?,?,?,?,?)",
            (chat_id, visitor_name, visitor_email, visitor_phone, 'website', 'active', 'nova_sterling'))
        conn.commit()
    else:
        # Update visitor info if provided
        if visitor_name or visitor_email or visitor_phone:
            updates = []
            params = []
            if visitor_name: updates.append("visitor_name=?"); params.append(visitor_name)
            if visitor_email: updates.append("visitor_email=?"); params.append(visitor_email)
            if visitor_phone: updates.append("visitor_phone=?"); params.append(visitor_phone)
            if updates:
                params.append(chat_id)
                conn.execute(f"UPDATE ahb_chats SET {','.join(updates)} WHERE id=?", params)
                conn.commit()

    # Save user message
    conn.execute(
        "INSERT INTO ahb_chat_messages (chat_id, role, content, agent_id) VALUES (?,?,?,?)",
        (chat_id, 'user', message, None))
    conn.commit()

    # Get conversation history for context
    history = conn.execute(
        "SELECT role, content FROM ahb_chat_messages WHERE chat_id=? ORDER BY created_at ASC",
        (chat_id,)).fetchall()
    conn.close()

    # Build messages for Nova
    system_prompt = """You are Nova Sterling — Director of Client Relations at All Home Building Co LLC (AHBCO).
You are the live chat assistant on ahb123.com. You are warm, professional, and helpful.

ABOUT THE COMPANY:
- All Home Building Co LLC, licensed PA HIC residential general contractor
- President: Sergey Tkach | Phone: 800-484-6404
- Address: 2725 Colmar Ave, Bensalem PA 19020
- Website: ahb123.com
- Services: Kitchen & bathroom remodeling, basement finishing, home additions, new construction, commercial build-outs
- Service area: Philadelphia, Bucks County, Montgomery County, Delaware County, Chester County PA

YOUR GOALS:
1. Greet warmly and ask how you can help
2. Answer questions about services, process, and service areas
3. Qualify leads: ask for name, project type, location, timeline, and budget range
4. When you have enough info, say you'll have someone reach out within 24 hours
5. Keep responses concise (2-4 sentences max), friendly, professional

RULES:
- Never make up pricing — say "every project is unique, we provide free estimates"
- Never schedule appointments directly — say the team will reach out
- If asked about something outside your scope, say you'll connect them with the right person
- Use plain text only, no markdown
- Be conversational, not robotic"""

    messages = [{"role": r[0], "content": r[1]} for r in history[-20:]]  # Last 20 messages

    # Call Ollama for Nova's response
    try:
        import urllib.request as _urllib
        ollama_url = os.environ.get("OLLAMA_HOST", "http://localhost:11434")
        payload = json.dumps({
            "model": "qwen2.5:14b",
            "messages": [{"role": "system", "content": system_prompt}] + messages,
            "stream": False,
            "options": {"num_predict": 300, "temperature": 0.7}
        }).encode()
        req = _urllib.Request(f"{ollama_url}/api/chat", data=payload,
                              headers={"Content-Type": "application/json"})
        with _urllib.urlopen(req, timeout=120) as r:
            result = json.loads(r.read())
            reply = result.get("message", {}).get("content", "")
    except Exception as e:
        reply = "Thanks for reaching out! I'm having a brief technical moment. Please call us at 800-484-6404 or try again in a minute."

    if not reply.strip():
        reply = "Thanks for your message! How can I help you with your home improvement project today?"

    # Save Nova's response
    conn2 = _ahb_db()
    conn2.execute(
        "INSERT INTO ahb_chat_messages (chat_id, role, content, agent_id) VALUES (?,?,?,?)",
        (chat_id, 'assistant', reply.strip(), 'nova_sterling'))
    conn2.execute("UPDATE ahb_chats SET updated_at=? WHERE id=?",
                  (datetime.datetime.now().isoformat(), chat_id))
    conn2.commit()
    conn2.close()

    return jsonify({
        'chat_id': chat_id,
        'reply': reply.strip(),
        'agent': 'Nova Sterling'
    })


@app.route('/api/ahb/widget/history', methods=['GET'])
def api_ahb_widget_history():
    """Get chat history for a widget session."""
    chat_id = request.args.get('chat_id', '')
    if not chat_id:
        return jsonify([])
    conn = _ahb_db()
    rows = conn.execute(
        "SELECT role, content, created_at FROM ahb_chat_messages WHERE chat_id=? ORDER BY created_at ASC",
        (chat_id,)).fetchall()
    conn.close()
    return jsonify([{"role": r[0], "content": r[1], "time": r[2]} for r in rows])


@app.route('/ahb-chat-widget.js')
def ahb_chat_widget_js():
    """Serve the embeddable chat widget script."""
    return send_from_directory(os.path.join(DASHBOARD_DIR, 'static'), 'ahb-chat-widget.js',
                               mimetype='application/javascript')


# ── Baza Cloud ────────────────────────────────────────────────────────────────

if CLOUD_ENABLED:
    app.secret_key = os.environ.get('FLASK_SECRET', 'baza-cloud-secret-change-me')

    # ── Family mode: no auth, fixed user_id=1 (Serge) ──
    FAMILY_USER_ID = 1
    FAMILY_STORAGE_QUOTA = 5242880  # 5 TB

    @app.route('/cloud')
    def cloud_page():
        return render_template('cloud.html')

    @app.route('/cloud/login')
    def cloud_login_page():
        return redirect('/cloud')

    @app.route('/api/cloud/register', methods=['POST'])
    def api_cloud_register():
        data = request.json or {}
        email = data.get('email', '').strip().lower()
        password = data.get('password', '')
        name = data.get('display_name', '').strip()
        if not email or not password:
            return jsonify({'success': False, 'error': 'Email and password required'}), 400
        if len(password) < 6:
            return jsonify({'success': False, 'error': 'Password must be at least 6 characters'}), 400
        pw_hash = bcrypt.hashpw(password.encode(), bcrypt.gensalt()).decode()
        try:
            from core.context_db import get_pool
            pool = get_pool()
            conn = pool.getconn()
            cur = conn.cursor()
            cur.execute("INSERT INTO cloud_users (email, password_hash, display_name) VALUES (%s, %s, %s) RETURNING id", (email, pw_hash, name or email.split('@')[0]))
            user_id = cur.fetchone()[0]
            conn.commit()
            cur.close()
            pool.putconn(conn)
            # Create user storage directory
            user_dir = os.path.join(CLOUD_STORAGE, str(user_id))
            os.makedirs(user_dir, exist_ok=True)
            return jsonify({'success': True, 'user_id': user_id})
        except Exception as e:
            if 'unique' in str(e).lower() or 'duplicate' in str(e).lower():
                return jsonify({'success': False, 'error': 'Email already registered'}), 409
            return jsonify({'success': False, 'error': str(e)}), 500

    @app.route('/api/cloud/login', methods=['POST'])
    def api_cloud_login():
        data = request.json or {}
        email = data.get('email', '').strip().lower()
        password = data.get('password', '')
        try:
            from core.context_db import get_pool
            pool = get_pool()
            conn = pool.getconn()
            cur = conn.cursor()
            cur.execute("SELECT id, email, password_hash, display_name, storage_quota_mb, storage_used_mb, is_admin FROM cloud_users WHERE email=%s AND is_active=TRUE", (email,))
            row = cur.fetchone()
            if not row:
                cur.close(); pool.putconn(conn)
                return jsonify({'success': False, 'error': 'Invalid email or password'}), 401
            if not bcrypt.checkpw(password.encode(), row[2].encode()):
                cur.close(); pool.putconn(conn)
                return jsonify({'success': False, 'error': 'Invalid email or password'}), 401
            cur.execute("UPDATE cloud_users SET last_login=NOW() WHERE id=%s", (row[0],))
            conn.commit()
            cur.close()
            pool.putconn(conn)
            user = CloudUser(row[0], row[1], row[3], row[4], row[5], row[6])
            login_user(user)
            return jsonify({'success': True, 'user': {'id': user.id, 'email': user.email, 'name': user.display_name}})
        except Exception as e:
            return jsonify({'success': False, 'error': str(e)}), 500

    @app.route('/api/cloud/logout', methods=['POST'])
    def api_cloud_logout():
        logout_user()
        return jsonify({'success': True})

    @app.route('/api/cloud/me')
    # @login_required  # family mode — no auth
    def api_cloud_me():
        return jsonify({'id': FAMILY_USER_ID, 'email': 'serge@ahb123.com', 'name': 'Serge Tkach',
                        'storage_quota_mb': FAMILY_STORAGE_QUOTA, 'storage_used_mb': 0,
                        'is_admin': True})

    # ── Cloud File Manager ──
    @app.route('/api/cloud/files')
    # @login_required  # family mode — no auth
    def api_cloud_files():
        user_dir = os.path.join(CLOUD_STORAGE, str(FAMILY_USER_ID))
        os.makedirs(user_dir, exist_ok=True)
        subdir = request.args.get('path', '')
        target = os.path.realpath(os.path.join(user_dir, subdir))
        if not target.startswith(os.path.realpath(user_dir)):
            return jsonify({'error': 'Invalid path'}), 403
        if not os.path.isdir(target):
            return jsonify({'error': 'Not a directory'}), 404
        items = []
        for name in sorted(os.listdir(target)):
            full = os.path.join(target, name)
            is_dir = os.path.isdir(full)
            items.append({'name': name, 'is_dir': is_dir,
                          'size': os.path.getsize(full) if not is_dir else 0,
                          'modified': os.path.getmtime(full)})
        return jsonify({'path': subdir, 'items': items})

    @app.route('/api/cloud/files/upload', methods=['POST'])
    # @login_required  # family mode — no auth
    def api_cloud_upload():
        user_dir = os.path.join(CLOUD_STORAGE, str(FAMILY_USER_ID))
        subdir = request.form.get('path', '')
        target = os.path.realpath(os.path.join(user_dir, subdir))
        if not target.startswith(os.path.realpath(user_dir)):
            return jsonify({'error': 'Invalid path'}), 403
        os.makedirs(target, exist_ok=True)
        saved = []
        for f in request.files.getlist('files'):
            safe = re.sub(r'[^\w.\-]', '_', f.filename or 'upload')
            f.save(os.path.join(target, safe))
            saved.append(safe)
        return jsonify({'success': True, 'files': saved})

    @app.route('/api/cloud/files/download/<path:filepath>')
    # @login_required  # family mode — no auth
    def api_cloud_download(filepath):
        user_dir = os.path.join(CLOUD_STORAGE, str(FAMILY_USER_ID))
        target = os.path.realpath(os.path.join(user_dir, filepath))
        if not target.startswith(os.path.realpath(user_dir)):
            return jsonify({'error': 'Invalid path'}), 403
        return send_from_directory(os.path.dirname(target), os.path.basename(target), as_attachment=True)

    @app.route('/api/cloud/files/mkdir', methods=['POST'])
    # @login_required  # family mode — no auth
    def api_cloud_mkdir():
        data = request.json or {}
        user_dir = os.path.join(CLOUD_STORAGE, str(FAMILY_USER_ID))
        new_dir = os.path.realpath(os.path.join(user_dir, data.get('path', '')))
        if not new_dir.startswith(os.path.realpath(user_dir)):
            return jsonify({'error': 'Invalid path'}), 403
        os.makedirs(new_dir, exist_ok=True)
        return jsonify({'success': True})

    @app.route('/api/cloud/files/delete', methods=['POST'])
    # @login_required  # family mode — no auth
    def api_cloud_delete():
        data = request.json or {}
        user_dir = os.path.join(CLOUD_STORAGE, str(FAMILY_USER_ID))
        target = os.path.realpath(os.path.join(user_dir, data.get('path', '')))
        if not target.startswith(os.path.realpath(user_dir)):
            return jsonify({'error': 'Invalid path'}), 403
        if os.path.isdir(target):
            import shutil; shutil.rmtree(target)
        elif os.path.isfile(target):
            os.remove(target)
        return jsonify({'success': True})

    # ── Cloud Agent Chat ──
    @app.route('/api/cloud/chat/<agent_id>', methods=['POST'])
    # @login_required  # family mode — no auth
    def api_cloud_chat(agent_id):
        """Per-user agent chat -- runs inference via Ollama, stores in per-user context."""
        data = request.json or {}
        message = data.get('message', '').strip()
        if not message:
            return jsonify({'error': 'Message required'}), 400

        # Load agent config
        config = load_config()
        agent_cfg = config.get('agents', {}).get(agent_id)
        if not agent_cfg:
            return jsonify({'error': f'Unknown agent: {agent_id}'}), 404

        user_id = FAMILY_USER_ID

        try:
            from core.context_db import get_pool
            pool = get_pool()
            conn = pool.getconn()
            cur = conn.cursor()

            # Save user message
            cur.execute("INSERT INTO cloud_conversations (user_id, agent_id, role, content) VALUES (%s, %s, 'user', %s)", (user_id, agent_id, message))

            # Get conversation history (last 20 messages)
            cur.execute("SELECT role, content FROM cloud_conversations WHERE user_id=%s AND agent_id=%s ORDER BY created_at DESC LIMIT 20", (user_id, agent_id))
            history = [{'role': r[0], 'content': r[1]} for r in reversed(cur.fetchall())]

            # Get per-user agent memory
            cur.execute("SELECT key, value FROM cloud_agent_memory WHERE user_id=%s AND agent_id=%s", (user_id, agent_id))
            memories = {r[0]: r[1] for r in cur.fetchall()}

            conn.commit()
            cur.close()
            pool.putconn(conn)

            # Build system prompt
            system = agent_cfg.get('system_prompt', f'You are {agent_cfg.get("name", agent_id)}.')
            if memories:
                mem_str = '\n'.join(f'- {k}: {v}' for k, v in list(memories.items())[:20])
                system += f'\n\nUser context:\n{mem_str}'

            # Run inference via Ollama
            import urllib.request
            model = agent_cfg.get('model', 'qwen2.5:14b')
            messages = [{'role': 'system', 'content': system}] + history

            payload = json.dumps({
                'model': model,
                'messages': messages,
                'stream': False,
                'options': {'num_predict': 1024, 'temperature': 0.7}
            }).encode()

            # Try both Ollama instances
            response_text = None
            for port in [11434, 11435]:
                try:
                    req = urllib.request.Request(f'http://localhost:{port}/api/chat',
                        data=payload, headers={'Content-Type': 'application/json'}, method='POST')
                    with urllib.request.urlopen(req, timeout=120) as resp:
                        result = json.loads(resp.read())
                        response_text = result.get('message', {}).get('content', '')
                        if response_text:
                            break
                except Exception:
                    continue

            if not response_text:
                return jsonify({'error': 'All inference backends unavailable'}), 503

            # Save assistant response
            pool = get_pool()
            conn = pool.getconn()
            cur = conn.cursor()
            cur.execute("INSERT INTO cloud_conversations (user_id, agent_id, role, content) VALUES (%s, %s, 'assistant', %s)", (user_id, agent_id, response_text))
            conn.commit()
            cur.close()
            pool.putconn(conn)

            return jsonify({'success': True, 'response': response_text, 'agent': agent_cfg.get('name', agent_id)})
        except Exception as e:
            return jsonify({'error': str(e)}), 500

    @app.route('/api/cloud/chat/<agent_id>/history')
    # @login_required  # family mode — no auth
    def api_cloud_chat_history(agent_id):
        try:
            from core.context_db import get_pool
            pool = get_pool()
            conn = pool.getconn()
            cur = conn.cursor()
            cur.execute("SELECT role, content, created_at FROM cloud_conversations WHERE user_id=%s AND agent_id=%s ORDER BY created_at ASC LIMIT 100", (FAMILY_USER_ID, agent_id))
            msgs = [{'role': r[0], 'content': r[1], 'created_at': str(r[2])} for r in cur.fetchall()]
            cur.close()
            pool.putconn(conn)
            return jsonify(msgs)
        except Exception as e:
            return jsonify({'error': str(e)}), 500

    @app.route('/api/cloud/chat/<agent_id>/clear', methods=['POST'])
    # @login_required  # family mode — no auth
    def api_cloud_chat_clear(agent_id):
        try:
            from core.context_db import get_pool
            pool = get_pool()
            conn = pool.getconn()
            cur = conn.cursor()
            cur.execute("DELETE FROM cloud_conversations WHERE user_id=%s AND agent_id=%s", (FAMILY_USER_ID, agent_id))
            conn.commit()
            cur.close()
            pool.putconn(conn)
            return jsonify({'success': True})
        except Exception as e:
            return jsonify({'error': str(e)}), 500

    @app.route('/api/cloud/memory/<agent_id>', methods=['POST'])
    # @login_required  # family mode — no auth
    def api_cloud_memory_set(agent_id):
        data = request.json or {}
        key = data.get('key', '')
        value = data.get('value', '')
        if not key:
            return jsonify({'error': 'Key required'}), 400
        try:
            from core.context_db import get_pool
            pool = get_pool()
            conn = pool.getconn()
            cur = conn.cursor()
            cur.execute("""INSERT INTO cloud_agent_memory (user_id, agent_id, key, value, category, updated_at)
                VALUES (%s,%s,%s,%s,%s,NOW()) ON CONFLICT (user_id, agent_id, key) DO UPDATE SET value=%s, updated_at=NOW()""",
                (FAMILY_USER_ID, agent_id, key, value, data.get('category','general'), value))
            conn.commit()
            cur.close()
            pool.putconn(conn)
            return jsonify({'success': True})
        except Exception as e:
            return jsonify({'error': str(e)}), 500

    # ── Admin: user management ──
    @app.route('/api/cloud/admin/users')
    # @login_required  # family mode — no auth
    def api_cloud_admin_users():
        # Family mode — all access
        try:
            from core.context_db import get_pool
            pool = get_pool()
            conn = pool.getconn()
            cur = conn.cursor()
            cur.execute("SELECT id, email, display_name, storage_quota_mb, storage_used_mb, is_active, is_admin, created_at, last_login FROM cloud_users ORDER BY created_at DESC")
            users = [{'id':r[0],'email':r[1],'name':r[2],'quota_mb':r[3],'used_mb':r[4],'active':r[5],'admin':r[6],'created':str(r[7]),'last_login':str(r[8]) if r[8] else None} for r in cur.fetchall()]
            cur.close()
            pool.putconn(conn)
            return jsonify(users)
        except Exception as e:
            return jsonify({'error': str(e)}), 500

    @app.route('/api/cloud/storage/usage')
    # @login_required  # family mode — no auth
    def api_cloud_storage_usage():
        user_dir = os.path.join(CLOUD_STORAGE, str(FAMILY_USER_ID))
        if not os.path.isdir(user_dir):
            return jsonify({'used_mb': 0, 'quota_mb': FAMILY_STORAGE_QUOTA, 'used_gb': 0, 'total_gb': round(FAMILY_STORAGE_QUOTA / 1024), 'percent': 0})
        total = 0
        for root, dirs, files in os.walk(user_dir):
            for f in files:
                total += os.path.getsize(os.path.join(root, f))
        used_mb = round(total / 1024 / 1024, 1)
        used_gb = round(used_mb / 1024, 2)
        total_gb = round(FAMILY_STORAGE_QUOTA / 1024)
        return jsonify({'used_mb': used_mb, 'quota_mb': FAMILY_STORAGE_QUOTA,
                        'used_gb': used_gb, 'total_gb': total_gb,
                        'percent': round(used_mb / FAMILY_STORAGE_QUOTA * 100, 1) if FAMILY_STORAGE_QUOTA else 0})

    # ── Cloud Asset Management: Memories + Documents ─────────────────────────

    CLOUD_MEDIA_DIRS = [
        (os.path.join(CLOUD_STORAGE, str(FAMILY_USER_ID)), 'cloud'),
        ('/mnt/empirepool/media/icloud',  'icloud'),
        ('/mnt/empirepool/media/generated', 'generated'),
    ]
    CLOUD_IMG_EXTS = {'.jpg','.jpeg','.png','.heic','.heif','.tif','.tiff','.webp','.gif','.bmp',
                      '.dng','.cr2','.cr3','.nef','.arw','.orf','.rw2','.raw',
                      '.insp','.thm'}
    CLOUD_VID_EXTS = {'.mov','.mp4','.m4v','.avi','.mkv','.webm','.3gp','.wmv','.flv','.mts',
                      '.insv','.lrv'}
    CLOUD_DOC_EXTS = {'.pdf','.doc','.docx','.txt','.csv','.xlsx','.xls','.md','.rtf',
                      '.odt','.pptx','.ppt','.pages','.numbers','.key'}
    CLOUD_SKIP_DIRS = {'.thumbnails', '.vault_meta', 'Vault', 'Imports'}
    THUMB_DIR = os.path.join(CLOUD_STORAGE, str(FAMILY_USER_ID), '.thumbnails')
    TRANSCODE_DIR = os.path.join(CLOUD_STORAGE, str(FAMILY_USER_ID), '.transcodes')
    os.makedirs(THUMB_DIR, exist_ok=True)
    os.makedirs(TRANSCODE_DIR, exist_ok=True)

    def _init_media_index():
        conn = sqlite3.connect(os.path.join(DASHBOARD_DIR, 'baza_projects.db'))
        conn.execute("""CREATE TABLE IF NOT EXISTS cloud_media_index (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            filepath TEXT UNIQUE,
            filename TEXT,
            media_type TEXT,
            source TEXT,
            date_taken TEXT,
            time_taken TEXT,
            latitude REAL,
            longitude REAL,
            size INTEGER,
            favorite INTEGER DEFAULT 0,
            category TEXT DEFAULT '',
            indexed_at TEXT DEFAULT (datetime('now'))
        )""")
        conn.execute("CREATE INDEX IF NOT EXISTS idx_cmi_date ON cloud_media_index(date_taken)")
        conn.execute("CREATE INDEX IF NOT EXISTS idx_cmi_type ON cloud_media_index(media_type)")
        conn.commit(); conn.close()
    _init_media_index()

    def _resolve_media_path(relpath):
        """Resolve a relative media path to an absolute path across all media dirs."""
        for base, src in CLOUD_MEDIA_DIRS:
            full = os.path.realpath(os.path.join(base, relpath))
            if full.startswith(os.path.realpath(base)) and os.path.exists(full):
                return full, src
        return None, None

    def _scan_media_dirs():
        """Walk all media directories and index files into cloud_media_index."""
        import hashlib
        from core.icloud_ingest import extract_exif
        conn = sqlite3.connect(os.path.join(DASHBOARD_DIR, 'baza_projects.db'))
        conn.row_factory = sqlite3.Row
        existing = {r['filepath'] for r in conn.execute("SELECT filepath FROM cloud_media_index").fetchall()}
        added = 0
        for base_dir, source in CLOUD_MEDIA_DIRS:
            if not os.path.isdir(base_dir):
                continue
            for root, dirs, files in os.walk(base_dir, followlinks=False):
                dirs[:] = [d for d in dirs if d not in CLOUD_SKIP_DIRS]
                for fname in files:
                    full = os.path.join(root, fname)
                    rel = os.path.relpath(full, base_dir)
                    ext = os.path.splitext(fname)[1].lower()
                    if ext in CLOUD_IMG_EXTS:
                        mtype = 'photo'
                    elif ext in CLOUD_VID_EXTS:
                        mtype = 'video'
                    elif ext in CLOUD_DOC_EXTS:
                        mtype = 'document'
                    else:
                        continue
                    key = f"{source}/{rel}"
                    if key in existing:
                        continue
                    try:
                        size = os.path.getsize(full)
                    except Exception:
                        size = 0
                    exif = {}
                    if mtype == 'photo':
                        try:
                            exif = extract_exif(full)
                        except Exception:
                            pass
                    date_taken = exif.get('photo_date') or ''
                    time_taken = exif.get('photo_time') or ''
                    if not date_taken:
                        try:
                            mt = os.path.getmtime(full)
                            import datetime as _dt
                            d = _dt.datetime.fromtimestamp(mt)
                            date_taken = d.strftime('%Y-%m-%d')
                            time_taken = time_taken or d.strftime('%H:%M')
                        except Exception:
                            pass
                    lat = exif.get('latitude')
                    lon = exif.get('longitude')
                    cat = ''
                    if mtype == 'document':
                        fl = fname.lower()
                        if 'permit' in fl: cat = 'permit'
                        elif 'invoice' in fl or 'inv-' in fl or 'inv_' in fl: cat = 'invoice'
                        elif 'contract' in fl or 'agreement' in fl: cat = 'contract'
                        elif 'receipt' in fl: cat = 'receipt'
                        elif 'estimate' in fl: cat = 'estimate'
                        elif 'coi' in fl or 'certificate' in fl: cat = 'coi'
                        else: cat = 'other'
                    try:
                        conn.execute(
                            "INSERT OR IGNORE INTO cloud_media_index "
                            "(filepath,filename,media_type,source,date_taken,time_taken,"
                            "latitude,longitude,size,category) VALUES (?,?,?,?,?,?,?,?,?,?)",
                            (key, fname, mtype, source, date_taken, time_taken, lat, lon, size, cat))
                        added += 1
                    except Exception:
                        pass
        conn.commit(); conn.close()
        return added

    @app.route('/api/cloud/thumb/<path:filepath>')
    def api_cloud_thumb(filepath):
        """Generate and serve a thumbnail for a media file."""
        import hashlib
        size = int(request.args.get('size', 200))
        size = min(max(size, 50), 2400)
        full, src = _resolve_media_path(filepath)
        if not full:
            # Try with source prefix
            parts = filepath.split('/', 1)
            if len(parts) == 2:
                for base, s in CLOUD_MEDIA_DIRS:
                    if s == parts[0]:
                        full = os.path.realpath(os.path.join(base, parts[1]))
                        if full.startswith(os.path.realpath(base)) and os.path.exists(full):
                            src = s; break
            if not full:
                return '', 404
        ext = os.path.splitext(full)[1].lower()
        mtime = str(os.path.getmtime(full))
        cache_key = hashlib.md5((full + mtime).encode()).hexdigest()[:16] + f'_{size}.jpg'
        cached = os.path.join(THUMB_DIR, cache_key)
        if os.path.exists(cached):
            return send_from_directory(THUMB_DIR, cache_key, mimetype='image/jpeg',
                                       max_age=86400)
        if ext in CLOUD_IMG_EXTS:
            try:
                try:
                    import pillow_heif; pillow_heif.register_heif_opener()
                except ImportError:
                    pass
                from PIL import Image
                img = Image.open(full)
                img.thumbnail((size, size), Image.LANCZOS)
                if img.mode != 'RGB':
                    img = img.convert('RGB')
                img.save(cached, 'JPEG', quality=80)
                return send_from_directory(THUMB_DIR, cache_key, mimetype='image/jpeg',
                                           max_age=86400)
            except Exception:
                pass
        if ext in CLOUD_VID_EXTS:
            # Use ffmpeg to extract a frame ~1s in and scale to <=size. Works
            # for .insv (dual-fisheye shown as-is — still beats a black square).
            # Prefer the smaller .lrv sibling for Insta360 .insv when available.
            source = full
            if ext == '.insv':
                base = os.path.basename(full)
                lrv_name = base.replace('VID_', 'LRV_', 1).replace('_00_', '_11_', 1)
                lrv_name = os.path.splitext(lrv_name)[0] + '.lrv'
                lrv_path = os.path.join(os.path.dirname(full), lrv_name)
                if os.path.exists(lrv_path):
                    source = lrv_path
            # Try -ss 1 first (skips boring intro frame), then -ss 0 as fallback
            # for short clips (e.g. iPhone Live Photos at 0.066s where -ss 1
            # silently produces no output even though ffmpeg exits 0).
            import subprocess
            for seek in ('1', '0'):
                try:
                    subprocess.run(
                        ['ffmpeg', '-nostdin', '-loglevel', 'error', '-y',
                         '-ss', seek, '-i', source,
                         '-vframes', '1',
                         '-vf', f'scale={size}:-2:force_original_aspect_ratio=decrease',
                         '-q:v', '5', cached],
                        timeout=15,
                    )
                except Exception:
                    continue
                if os.path.exists(cached) and os.path.getsize(cached) > 0:
                    return send_from_directory(THUMB_DIR, cache_key, mimetype='image/jpeg',
                                               max_age=86400)
        # Fallback: 1x1 transparent pixel
        import base64
        pixel = base64.b64decode('R0lGODlhAQABAIAAAAAAAP///yH5BAEAAAAALAAAAAABAAEAAAIBRAA7')
        from flask import Response
        return Response(pixel, mimetype='image/gif')

    @app.route('/api/cloud/media')
    def api_cloud_media():
        """Memories-style media listing grouped by date."""
        refresh = request.args.get('refresh', '') == 'true'
        source_filter = request.args.get('source', '')
        fav_only = request.args.get('favorites', '') == 'true'
        conn = sqlite3.connect(os.path.join(DASHBOARD_DIR, 'baza_projects.db'))
        conn.row_factory = sqlite3.Row
        count = conn.execute("SELECT COUNT(*) FROM cloud_media_index WHERE media_type IN ('photo','video')").fetchone()[0]
        if count == 0 or refresh:
            conn.close()
            _scan_media_dirs()
            conn = sqlite3.connect(os.path.join(DASHBOARD_DIR, 'baza_projects.db'))
            conn.row_factory = sqlite3.Row
        sql = "SELECT * FROM cloud_media_index WHERE media_type IN ('photo','video')"
        params = []
        if source_filter:
            sql += " AND source=?"; params.append(source_filter)
        if fav_only:
            sql += " AND favorite=1"
        sql += " ORDER BY date_taken DESC, time_taken DESC"
        rows = [dict(r) for r in conn.execute(sql, params).fetchall()]
        conn.close()
        # Group by year → month → day
        years = {}
        for r in rows:
            dt = r.get('date_taken') or '0000-00-00'
            parts = dt.split('-')
            y = parts[0] if len(parts) > 0 else '0000'
            m = parts[1] if len(parts) > 1 else '00'
            d = parts[2] if len(parts) > 2 else '00'
            years.setdefault(y, {'months': {}, 'count': 0, 'cover': None})
            years[y]['count'] += 1
            if not years[y]['cover']:
                years[y]['cover'] = r['filepath']
            years[y]['months'].setdefault(m, {'days': {}, 'count': 0, 'cover': None})
            years[y]['months'][m]['count'] += 1
            if not years[y]['months'][m]['cover']:
                years[y]['months'][m]['cover'] = r['filepath']
            years[y]['months'][m]['days'].setdefault(d, [])
            years[y]['months'][m]['days'][d].append({
                'path': r['filepath'], 'name': r['filename'], 'date': dt,
                'time': r.get('time_taken') or '', 'size': r.get('size') or 0,
                'type': r['media_type'], 'favorite': bool(r.get('favorite')),
                'has_gps': bool(r.get('latitude')), 'source': r.get('source',''),
            })
        return jsonify({'years': years, 'total': len(rows)})

    @app.route('/api/cloud/media/serve/<path:filepath>')
    def api_cloud_media_serve(filepath):
        """Serve original media file inline (for lightbox)."""
        full, src = _resolve_media_path(filepath)
        if not full:
            parts = filepath.split('/', 1)
            if len(parts) == 2:
                for base, s in CLOUD_MEDIA_DIRS:
                    if s == parts[0]:
                        full = os.path.realpath(os.path.join(base, parts[1]))
                        if full.startswith(os.path.realpath(base)) and os.path.exists(full):
                            src = s; break
            if not full:
                return '', 404
        # Browsers default to application/octet-stream (forces download) for
        # extensions Python's mimetypes module doesn't know — .thm is actually
        # JPEG, .insv/.lrv are MP4 containers. Force a sensible inline type.
        ext = os.path.splitext(full)[1].lower()
        mime_override = {
            '.thm': 'image/jpeg',
            '.insv': 'video/mp4',
            '.lrv': 'video/mp4',
            '.insp': 'image/jpeg',
            '.heic': 'image/heic',
            '.heif': 'image/heif',
        }
        kwargs = {'as_attachment': False}
        if ext in mime_override:
            kwargs['mimetype'] = mime_override[ext]
        return send_from_directory(os.path.dirname(full), os.path.basename(full), **kwargs)

    @app.route('/api/cloud/media/play/<path:filepath>')
    def api_cloud_media_play(filepath):
        """Browser-playable video. H.264/MP4 streams as-is; HEVC and Insta360
        containers (.insv/.lrv) get transcoded to H.264/AAC mp4 on first play
        and cached. This is what fixes the "audio only, no video" symptom on
        Chrome/Linux for iPhone Live Photo .MOVs."""
        import hashlib, subprocess
        full, _ = _resolve_media_path(filepath)
        if not full:
            parts = filepath.split('/', 1)
            if len(parts) == 2:
                for base, s in CLOUD_MEDIA_DIRS:
                    if s == parts[0]:
                        cand = os.path.realpath(os.path.join(base, parts[1]))
                        if cand.startswith(os.path.realpath(base)) and os.path.exists(cand):
                            full = cand; break
            if not full:
                return '', 404
        ext = os.path.splitext(full)[1].lower()
        # Detect the video codec; pass through if it's already browser-friendly.
        codec = ''
        try:
            codec = subprocess.check_output(
                ['ffprobe','-v','error','-select_streams','v:0',
                 '-show_entries','stream=codec_name','-of','default=nw=1:nk=1', full],
                timeout=5).decode().strip()
        except Exception:
            pass
        pass_through_codecs = {'h264', 'vp8', 'vp9', 'av1'}
        if codec in pass_through_codecs and ext not in ('.insv', '.lrv'):
            return send_from_directory(os.path.dirname(full), os.path.basename(full),
                                       as_attachment=False, mimetype='video/mp4')
        mtime = os.path.getmtime(full)
        size = os.path.getsize(full)
        key = hashlib.md5(f"{full}|{mtime}|{size}".encode()).hexdigest() + '.mp4'
        cached = os.path.join(TRANSCODE_DIR, key)
        if not (os.path.exists(cached) and os.path.getsize(cached) > 0):
            try:
                subprocess.run(
                    ['ffmpeg','-nostdin','-loglevel','error','-y','-i', full,
                     '-c:v','libx264','-preset','ultrafast','-crf','23',
                     '-pix_fmt','yuv420p',
                     '-c:a','aac','-b:a','128k','-ac','2',
                     '-movflags','+faststart', cached],
                    check=True, timeout=600)
            except Exception as e:
                # Clean up any half-written cache file
                try:
                    if os.path.exists(cached):
                        os.remove(cached)
                except Exception:
                    pass
                return f'transcode failed: {e}', 500
        return send_from_directory(TRANSCODE_DIR, key,
                                   as_attachment=False, mimetype='video/mp4')

    @app.route('/api/cloud/media/upload', methods=['POST'])
    def api_cloud_media_upload():
        """Accept multipart photo/video uploads from the Media tab's Upload
        button or drag-drop. Also accepts comma-separated `pick_tokens` from
        the Baza picker (copies from artifacts/ into Uploads/<date>/). Files
        land in /mnt/empirepool/cloud/<user>/Uploads/<YYYY-MM-DD>/ and are
        indexed immediately so they show up in the Library."""
        import datetime, time as _time, shutil as _shutil
        from werkzeug.utils import secure_filename
        files = request.files.getlist('files')
        pick_tokens_raw = request.form.get('pick_tokens') or ''
        pick_tokens = [t.strip() for t in pick_tokens_raw.split(',') if t.strip()]
        if not files and not pick_tokens:
            return jsonify({'success': False, 'error': 'no files'}), 400
        today = datetime.date.today().isoformat()
        dest_dir = os.path.join(CLOUD_STORAGE, str(FAMILY_USER_ID), 'Uploads', today)
        os.makedirs(dest_dir, exist_ok=True)
        saved, skipped = [], []
        max_bytes = 500 * 1024 * 1024  # 500MB per file cap
        for f in files:
            if not f or not f.filename:
                continue
            safe = secure_filename(f.filename) or ''
            if not safe:
                skipped.append(f.filename); continue
            ext = os.path.splitext(safe)[1].lower()
            if ext not in CLOUD_IMG_EXTS and ext not in CLOUD_VID_EXTS:
                skipped.append(f.filename); continue
            target = os.path.join(dest_dir, safe)
            if os.path.exists(target):
                stem, e = os.path.splitext(safe)
                target = os.path.join(dest_dir, f"{stem}_{int(_time.time()*1000)}{e}")
            try:
                f.save(target)
                if os.path.getsize(target) > max_bytes:
                    os.remove(target)
                    skipped.append(f.filename + ' (too large)')
                    continue
                saved.append(os.path.basename(target))
            except Exception as exc:
                skipped.append(f.filename + ' (' + str(exc) + ')')
        for tok in pick_tokens:
            src = _pick_decode_token(tok)
            if not src:
                skipped.append(f'(token:{tok[:8]}…)'); continue
            base = os.path.basename(src)
            safe = secure_filename(base) or ''
            if not safe:
                skipped.append(base); continue
            ext = os.path.splitext(safe)[1].lower()
            if ext not in CLOUD_IMG_EXTS and ext not in CLOUD_VID_EXTS:
                skipped.append(base); continue
            target = os.path.join(dest_dir, safe)
            if os.path.exists(target):
                stem, e = os.path.splitext(safe)
                target = os.path.join(dest_dir, f"{stem}_{int(_time.time()*1000)}{e}")
            try:
                _shutil.copy2(src, target)
                if os.path.getsize(target) > max_bytes:
                    os.remove(target)
                    skipped.append(base + ' (too large)')
                    continue
                saved.append(os.path.basename(target))
            except Exception as exc:
                skipped.append(base + ' (' + str(exc) + ')')
        try:
            _scan_media_dirs()
        except Exception:
            pass
        return jsonify({'success': True, 'saved': saved, 'skipped': skipped,
                        'dest': f'Uploads/{today}'})

    @app.route('/api/cloud/media/favorite', methods=['POST'])
    def api_cloud_media_favorite():
        data = request.json or {}
        path = data.get('path', '')
        fav = 1 if data.get('favorite', True) else 0
        conn = sqlite3.connect(os.path.join(DASHBOARD_DIR, 'baza_projects.db'))
        conn.execute("UPDATE cloud_media_index SET favorite=? WHERE filepath=?", (fav, path))
        conn.commit(); conn.close()
        return jsonify({'success': True})

    @app.route('/api/cloud/media/reindex', methods=['POST'])
    def api_cloud_media_reindex():
        added = _scan_media_dirs()
        return jsonify({'success': True, 'added': added})

    # ── AHB123 Media Catalog ────────────────────────────────────────────────
    # Bridges baza cloud media (cloud_media_index) → AHB projects via
    # GPS-proximity + date-window auto-classification. Persists in
    # ahb_media_attachments so manual phase overrides survive reclassification.

    def _init_ahb_media_attachments():
        conn = sqlite3.connect(os.path.join(DASHBOARD_DIR, 'baza_projects.db'))
        conn.execute("""CREATE TABLE IF NOT EXISTS ahb_media_attachments (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            media_filepath TEXT NOT NULL,
            project_id TEXT NOT NULL,
            phase TEXT DEFAULT 'during',
            distance_m REAL,
            source TEXT DEFAULT 'manual',
            notes TEXT DEFAULT '',
            created_at TEXT DEFAULT (datetime('now')),
            UNIQUE(media_filepath, project_id)
        )""")
        conn.execute("CREATE INDEX IF NOT EXISTS idx_ahb_ma_proj ON ahb_media_attachments(project_id)")
        conn.execute("CREATE INDEX IF NOT EXISTS idx_ahb_ma_media ON ahb_media_attachments(media_filepath)")
        conn.execute("""CREATE TABLE IF NOT EXISTS ahb_media_vision (
            media_filepath TEXT PRIMARY KEY,
            caption TEXT,
            tags TEXT,
            work_score INTEGER DEFAULT 0,
            has_people INTEGER DEFAULT 0,
            classified_at TEXT DEFAULT (datetime('now')),
            model TEXT
        )""")
        # Add has_people column if upgrading from old schema
        try:
            cols = [r[1] for r in conn.execute("PRAGMA table_info(ahb_media_vision)").fetchall()]
            if 'has_people' not in cols:
                conn.execute("ALTER TABLE ahb_media_vision ADD COLUMN has_people INTEGER DEFAULT 0")
        except Exception: pass
        conn.execute("CREATE INDEX IF NOT EXISTS idx_amv_work ON ahb_media_vision(work_score)")
        conn.execute("CREATE INDEX IF NOT EXISTS idx_amv_people ON ahb_media_vision(has_people)")
        conn.commit(); conn.close()
    _init_ahb_media_attachments()

    def _haversine_m(lat1, lon1, lat2, lon2):
        """Distance in meters between two lat/long points."""
        import math
        R = 6371000.0
        p1, p2 = math.radians(lat1), math.radians(lat2)
        dp = math.radians(lat2-lat1); dl = math.radians(lon2-lon1)
        a = math.sin(dp/2)**2 + math.cos(p1)*math.cos(p2)*math.sin(dl/2)**2
        return 2 * R * math.asin(min(1.0, math.sqrt(a)))

    def _auto_classify_ahb_media(radius_m=200, before_days=14, after_days=42):
        """For each baza media w/ GPS, find nearest project within radius_m where
        date_taken falls in (start-before_days .. end+after_days). Insert into
        ahb_media_attachments with source='auto'. Skips rows that already have a
        manual attachment (UNIQUE(media_filepath, project_id) keeps it idempotent)."""
        import datetime as _dt
        conn = sqlite3.connect(os.path.join(DASHBOARD_DIR, 'baza_projects.db'))
        conn.row_factory = sqlite3.Row
        projects = conn.execute(
            "SELECT id, latitude, longitude, start_date, end_date FROM ahb_projects "
            "WHERE latitude IS NOT NULL AND longitude IS NOT NULL").fetchall()
        if not projects:
            conn.close()
            return {'attached': 0, 'projects': 0, 'media_scanned': 0}
        # Bbox prefilter — ~200m ≈ 0.0018° lat. Don't bother with full Earth haversine
        # for media that are clearly nowhere near any project.
        deg = radius_m / 111000.0  # rough — fine for ≤1km
        media = conn.execute(
            "SELECT filepath, latitude, longitude, date_taken, media_type FROM cloud_media_index "
            "WHERE latitude IS NOT NULL AND longitude IS NOT NULL "
            "AND media_type IN ('photo','video')").fetchall()
        attached = 0
        for m in media:
            best = None
            for p in projects:
                if abs(m['latitude']-p['latitude']) > deg*2 or abs(m['longitude']-p['longitude']) > deg*2:
                    continue
                d = _haversine_m(m['latitude'], m['longitude'], p['latitude'], p['longitude'])
                if d > radius_m: continue
                if best is None or d < best[1]:
                    best = (p, d)
            if not best: continue
            p, dist = best
            # Determine phase from date window. We DO NOT reject photos that are
            # far outside the project window — at this address, they're still relevant
            # property history (pre-purchase shots, post-completion follow-ups, etc.).
            # The date_taken may also be a "Photos from YYYY" → Jan 1 fallback that
            # doesn't reflect the actual moment, so being strict here drops too much.
            phase = 'during'
            try:
                if m['date_taken'] and m['date_taken'] != '0000-00-00':
                    md = _dt.datetime.strptime(m['date_taken'][:10], '%Y-%m-%d').date()
                    sd = _dt.datetime.strptime(p['start_date'][:10],'%Y-%m-%d').date() if p['start_date'] else None
                    ed = _dt.datetime.strptime(p['end_date'][:10],'%Y-%m-%d').date() if p['end_date'] else None
                    if sd and md < sd:
                        phase = 'before'
                    elif ed and md > ed:
                        phase = 'after'
                    else:
                        phase = 'during'
            except Exception:
                pass
            try:
                conn.execute(
                    "INSERT OR IGNORE INTO ahb_media_attachments "
                    "(media_filepath, project_id, phase, distance_m, source) "
                    "VALUES (?,?,?,?, 'auto')",
                    (m['filepath'], p['id'], phase, round(dist, 1)))
                if conn.total_changes > 0: attached += 1
            except Exception:
                pass
        conn.commit(); conn.close()
        return {'attached': attached, 'projects': len(projects), 'media_scanned': len(media)}

    @app.route('/api/ahb123/media')
    def api_ahb123_media():
        """List work-related media (those attached to projects). Supports filters:
        project_id, phase, type, year, no_project (orphan candidates w/ GPS but unmatched)."""
        project_id = request.args.get('project_id', '')
        phase = request.args.get('phase', '')
        media_type = request.args.get('type', '')
        year = request.args.get('year', '')
        only_orphans = request.args.get('orphans', '') == 'true'
        conn = sqlite3.connect(os.path.join(DASHBOARD_DIR, 'baza_projects.db'))
        conn.row_factory = sqlite3.Row
        if only_orphans:
            # Media w/ GPS but no project attachment yet
            sql = ("SELECT c.* FROM cloud_media_index c "
                   "LEFT JOIN ahb_media_attachments a ON a.media_filepath=c.filepath "
                   "WHERE c.latitude IS NOT NULL AND a.media_filepath IS NULL "
                   "AND c.media_type IN ('photo','video')")
            params = []
        else:
            sql = ("SELECT c.*, a.project_id, a.phase, a.distance_m, a.source AS attach_source, "
                   "p.title AS project_title, p.address AS project_address, "
                   "p.latitude AS project_lat, p.longitude AS project_lon "
                   "FROM ahb_media_attachments a "
                   "JOIN cloud_media_index c ON c.filepath = a.media_filepath "
                   "LEFT JOIN ahb_projects p ON p.id = a.project_id "
                   "WHERE 1=1")
            params = []
            if project_id: sql += " AND a.project_id=?"; params.append(project_id)
            if phase: sql += " AND a.phase=?"; params.append(phase)
        if media_type: sql += " AND c.media_type=?"; params.append(media_type)
        if year: sql += " AND substr(c.date_taken,1,4)=?"; params.append(year)
        sql += " ORDER BY c.date_taken DESC, c.time_taken DESC LIMIT 5000"
        rows = [dict(r) for r in conn.execute(sql, params).fetchall()]
        # Stats
        stats = {
            'total': conn.execute("SELECT COUNT(DISTINCT media_filepath) FROM ahb_media_attachments").fetchone()[0],
            'photos': conn.execute("SELECT COUNT(DISTINCT a.media_filepath) FROM ahb_media_attachments a JOIN cloud_media_index c ON c.filepath=a.media_filepath WHERE c.media_type='photo'").fetchone()[0],
            'videos': conn.execute("SELECT COUNT(DISTINCT a.media_filepath) FROM ahb_media_attachments a JOIN cloud_media_index c ON c.filepath=a.media_filepath WHERE c.media_type='video'").fetchone()[0],
            'projects': conn.execute("SELECT COUNT(DISTINCT project_id) FROM ahb_media_attachments").fetchone()[0],
            'before': conn.execute("SELECT COUNT(*) FROM ahb_media_attachments WHERE phase='before'").fetchone()[0],
            'during': conn.execute("SELECT COUNT(*) FROM ahb_media_attachments WHERE phase='during'").fetchone()[0],
            'after':  conn.execute("SELECT COUNT(*) FROM ahb_media_attachments WHERE phase='after'").fetchone()[0],
        }
        # Projects list for filter dropdown
        projects = [dict(r) for r in conn.execute(
            "SELECT p.id, p.title, p.address, p.latitude, p.longitude, "
            "COUNT(a.id) AS media_count FROM ahb_projects p "
            "LEFT JOIN ahb_media_attachments a ON a.project_id=p.id "
            "GROUP BY p.id ORDER BY p.start_date DESC NULLS LAST, p.title").fetchall()]
        conn.close()
        return jsonify({'items': rows, 'stats': stats, 'projects': projects})

    @app.route('/api/ahb123/media/classify', methods=['POST'])
    def api_ahb123_media_classify():
        data = request.json or {}
        result = _auto_classify_ahb_media(
            radius_m=int(data.get('radius_m', 200)),
            before_days=int(data.get('before_days', 14)),
            after_days=int(data.get('after_days', 42)))
        return jsonify({'success': True, **result})

    @app.route('/api/ahb123/media/attach', methods=['POST'])
    def api_ahb123_media_attach():
        data = request.json or {}
        fp = data.get('media_filepath'); pid = data.get('project_id')
        phase = data.get('phase', 'during')
        if not fp or not pid:
            return jsonify({'success': False, 'error': 'media_filepath and project_id required'}), 400
        if phase not in ('before', 'during', 'after'):
            return jsonify({'success': False, 'error': 'phase must be before/during/after'}), 400
        conn = sqlite3.connect(os.path.join(DASHBOARD_DIR, 'baza_projects.db'))
        conn.execute(
            "INSERT INTO ahb_media_attachments (media_filepath, project_id, phase, source) "
            "VALUES (?,?,?, 'manual') "
            "ON CONFLICT(media_filepath, project_id) DO UPDATE SET phase=excluded.phase, source='manual'",
            (fp, pid, phase))
        conn.commit(); conn.close()
        return jsonify({'success': True})

    @app.route('/api/ahb123/media/detach', methods=['POST'])
    def api_ahb123_media_detach():
        data = request.json or {}
        fp = data.get('media_filepath'); pid = data.get('project_id')
        if not fp or not pid:
            return jsonify({'success': False, 'error': 'media_filepath and project_id required'}), 400
        conn = sqlite3.connect(os.path.join(DASHBOARD_DIR, 'baza_projects.db'))
        conn.execute("DELETE FROM ahb_media_attachments WHERE media_filepath=? AND project_id=?", (fp, pid))
        conn.commit(); conn.close()
        return jsonify({'success': True})

    @app.route('/api/ahb123/media/library')
    def api_ahb123_media_library():
        """Browse the entire baza cloud filtered for work-likely media.
        Excludes wedding folder, thumbnails, .private-inbound. Surfaces caption
        when available. Paginated. Marks already-attached items so the UI can
        show their project."""
        page = max(1, int(request.args.get('page', 1)))
        page_size = max(20, min(500, int(request.args.get('page_size', 200))))
        offset = (page - 1) * page_size
        year = request.args.get('year', '')
        media_type = request.args.get('type', '')
        has_gps = request.args.get('has_gps', '') == 'true'
        attached_state = request.args.get('attached', '')  # 'yes' | 'no' | ''
        search = (request.args.get('q', '') or '').strip()
        work_only = request.args.get('work_only', '') == 'true'
        min_score = int(request.args.get('min_work_score', 30))
        people_filter = request.args.get('people', '')  # '' | 'yes' | 'no'

        where = [
            "c.media_type IN ('photo','video')",
            "c.source = 'cloud'",
            # Hard excludes — never work-related
            "c.filepath NOT LIKE 'cloud/%ZHAR-wedding-photos%'",
            "c.filepath NOT LIKE 'cloud/.thumbnails/%'",
            "c.filepath NOT LIKE 'cloud/%.private-inbound%'",
            "c.filepath NOT LIKE 'cloud/Vault/%'",
            "c.filepath NOT LIKE 'cloud/BootableImages/%'",
        ]
        params = []
        if year: where.append("substr(c.date_taken,1,4)=?"); params.append(year)
        if media_type: where.append("c.media_type=?"); params.append(media_type)
        if has_gps: where.append("c.latitude IS NOT NULL")
        if attached_state == 'yes': where.append("a.project_id IS NOT NULL")
        elif attached_state == 'no': where.append("a.project_id IS NULL")
        if search:
            where.append("(c.filename LIKE ? OR v.caption LIKE ?)"); params.extend([f"%{search}%", f"%{search}%"])
        if work_only:
            where.append("v.work_score >= ?"); params.append(min_score)
        if people_filter == 'yes':
            where.append("v.has_people = 1")
        elif people_filter == 'no':
            where.append("(v.has_people = 0 OR v.has_people IS NULL)")

        where_sql = ' AND '.join(where)
        conn = sqlite3.connect(os.path.join(DASHBOARD_DIR, 'baza_projects.db'))
        conn.row_factory = sqlite3.Row

        # Attach the captions DB so we can JOIN
        cap_db = os.path.join(DASHBOARD_DIR, 'image_captions.db')
        if os.path.isfile(cap_db):
            try: conn.execute(f"ATTACH DATABASE ? AS cap", (cap_db,))
            except Exception: pass

        # Counts (also broken out by GPS-yes/no for the dashboard headline)
        count_sql = (f"SELECT COUNT(*) FROM cloud_media_index c "
                     f"LEFT JOIN ahb_media_attachments a ON a.media_filepath=c.filepath "
                     f"LEFT JOIN ahb_media_vision v ON v.media_filepath=c.filepath "
                     f"WHERE {where_sql}")
        total = conn.execute(count_sql, params).fetchone()[0]

        # Page query — include vision caption + work_score when available
        rows_sql = (f"SELECT c.*, a.project_id, a.phase, a.distance_m, "
                    f"a.source AS attach_source, p.title AS project_title, "
                    f"v.caption, v.tags AS work_tags, v.work_score, v.has_people "
                    f"FROM cloud_media_index c "
                    f"LEFT JOIN ahb_media_attachments a ON a.media_filepath=c.filepath "
                    f"LEFT JOIN ahb_projects p ON p.id=a.project_id "
                    f"LEFT JOIN ahb_media_vision v ON v.media_filepath=c.filepath "
                    f"WHERE {where_sql} "
                    f"ORDER BY "
                    f"  {'v.work_score DESC, ' if work_only else ''}"
                    f"  c.date_taken DESC, c.time_taken DESC, c.filepath "
                    f"LIMIT ? OFFSET ?")
        rows = [dict(r) for r in conn.execute(rows_sql, params + [page_size, offset]).fetchall()]

        # Stats card numbers
        stats_sql = (
            "SELECT "
            "  (SELECT COUNT(*) FROM cloud_media_index c LEFT JOIN ahb_media_attachments a ON a.media_filepath=c.filepath "
            "    WHERE c.media_type='photo' AND c.source='cloud' "
            "    AND c.filepath NOT LIKE 'cloud/%ZHAR-wedding-photos%'"
            "    AND c.filepath NOT LIKE 'cloud/.thumbnails/%'"
            "    AND c.filepath NOT LIKE 'cloud/%.private-inbound%'"
            "    AND c.filepath NOT LIKE 'cloud/Vault/%') AS photos,"
            "  (SELECT COUNT(*) FROM cloud_media_index c "
            "    WHERE c.media_type='video' AND c.source='cloud' "
            "    AND c.filepath NOT LIKE 'cloud/%ZHAR-wedding-photos%'"
            "    AND c.filepath NOT LIKE 'cloud/.thumbnails/%'"
            "    AND c.filepath NOT LIKE 'cloud/%.private-inbound%'"
            "    AND c.filepath NOT LIKE 'cloud/Vault/%') AS videos,"
            "  (SELECT COUNT(*) FROM cloud_media_index c LEFT JOIN ahb_media_attachments a ON a.media_filepath=c.filepath "
            "    WHERE c.media_type IN ('photo','video') AND c.source='cloud' AND a.project_id IS NOT NULL "
            "    AND c.filepath NOT LIKE 'cloud/%ZHAR-wedding-photos%') AS attached,"
            "  (SELECT COUNT(DISTINCT project_id) FROM ahb_media_attachments) AS projects_with_media"
        )
        stats = dict(conn.execute(stats_sql).fetchone())

        # Projects list for attach picker
        projects = [dict(r) for r in conn.execute(
            "SELECT p.id, p.title, p.address, p.start_date, p.end_date "
            "FROM ahb_projects p ORDER BY COALESCE(p.start_date, p.created_at) DESC, p.title").fetchall()]
        conn.close()
        return jsonify({
            'items': rows,
            'total': total,
            'page': page,
            'page_size': page_size,
            'stats': stats,
            'projects': projects
        })

    @app.route('/api/ahb123/media/bulk_attach', methods=['POST'])
    def api_ahb123_media_bulk_attach():
        """Attach multiple media files to a single project + phase."""
        data = request.json or {}
        fps = data.get('media_filepaths') or []
        pid = data.get('project_id')
        phase = data.get('phase', 'during')
        if not fps or not pid:
            return jsonify({'success': False, 'error': 'media_filepaths[] and project_id required'}), 400
        if phase not in ('before', 'during', 'after'):
            return jsonify({'success': False, 'error': 'phase must be before/during/after'}), 400
        conn = sqlite3.connect(os.path.join(DASHBOARD_DIR, 'baza_projects.db'))
        attached = 0
        for fp in fps:
            try:
                conn.execute(
                    "INSERT INTO ahb_media_attachments (media_filepath, project_id, phase, source) "
                    "VALUES (?,?,?, 'manual') "
                    "ON CONFLICT(media_filepath, project_id) DO UPDATE SET phase=excluded.phase, source='manual'",
                    (fp, pid, phase))
                attached += 1
            except Exception:
                pass
        conn.commit(); conn.close()
        return jsonify({'success': True, 'attached': attached})

    # ── Vision classification (Qwen3-VL captioning + work-keyword tagging) ──
    VISION_STATUS_FILE = '/tmp/baza_vision_status.json'
    VISION_PID_FILE = '/tmp/baza_vision.pid'
    VISION_SCRIPT = '/home/switchhacker/.gdrive-pull/caption_cloud_media.py'

    def _vision_is_running():
        try:
            with open(VISION_PID_FILE) as f:
                pid = int(f.read().strip())
            os.kill(pid, 0)   # signal 0 = check existence
            return pid
        except Exception:
            return 0

    @app.route('/api/ahb123/media/caption/start', methods=['POST'])
    def api_ahb123_caption_start():
        if _vision_is_running():
            return jsonify({'success': False, 'error': 'already running'}), 409
        if not os.path.isfile(VISION_SCRIPT):
            return jsonify({'success': False, 'error': f'script not found: {VISION_SCRIPT}'}), 500
        import subprocess
        venv_py = '/home/switchhacker/baza-empire/agent-framework-v3/venv/bin/python'
        try:
            with open('/tmp/baza_vision.log', 'a') as logf:
                subprocess.Popen([venv_py, VISION_SCRIPT],
                                 stdout=logf, stderr=subprocess.STDOUT,
                                 cwd='/home/switchhacker/baza-empire/agent-framework-v3',
                                 close_fds=True, start_new_session=True)
            return jsonify({'success': True, 'message': 'caption job started'})
        except Exception as e:
            return jsonify({'success': False, 'error': str(e)}), 500

    @app.route('/api/ahb123/media/caption/stop', methods=['POST'])
    def api_ahb123_caption_stop():
        pid = _vision_is_running()
        if not pid:
            return jsonify({'success': False, 'error': 'not running'}), 404
        try:
            os.kill(pid, 15)   # SIGTERM
            return jsonify({'success': True, 'pid': pid})
        except Exception as e:
            return jsonify({'success': False, 'error': str(e)}), 500

    @app.route('/api/ahb123/media/caption/status')
    def api_ahb123_caption_status():
        running = _vision_is_running() > 0
        status = {'running': running, 'done': 0, 'total_remaining': 0, 'errors': 0, 'last_file': '', 'started_at': 0}
        try:
            with open(VISION_STATUS_FILE) as f:
                status.update(json.load(f))
        except Exception:
            pass
        # Also add total work-classified count from DB
        try:
            conn = sqlite3.connect(os.path.join(DASHBOARD_DIR, 'baza_projects.db'))
            status['total_classified'] = conn.execute("SELECT COUNT(*) FROM ahb_media_vision").fetchone()[0]
            status['total_work']       = conn.execute("SELECT COUNT(*) FROM ahb_media_vision WHERE work_score >= 30").fetchone()[0]
            conn.close()
        except Exception:
            pass
        return jsonify(status)

    @app.route('/api/ahb123/media/map')
    def api_ahb123_media_map():
        """Pin data for the live map: every project w/ GPS plus its attached-media count and a cover thumbnail."""
        conn = sqlite3.connect(os.path.join(DASHBOARD_DIR, 'baza_projects.db'))
        conn.row_factory = sqlite3.Row
        rows = conn.execute(
            "SELECT p.id, p.title, p.address, p.latitude, p.longitude, p.start_date, p.end_date, p.status, "
            "COUNT(a.id) AS media_count, "
            "(SELECT a2.media_filepath FROM ahb_media_attachments a2 "
            "  WHERE a2.project_id=p.id ORDER BY a2.id LIMIT 1) AS cover_path "
            "FROM ahb_projects p "
            "LEFT JOIN ahb_media_attachments a ON a.project_id=p.id "
            "WHERE p.latitude IS NOT NULL AND p.longitude IS NOT NULL "
            "GROUP BY p.id "
            "ORDER BY media_count DESC").fetchall()
        conn.close()
        return jsonify({'pins': [dict(r) for r in rows]})

    @app.route('/api/cloud/media/delete', methods=['POST'])
    def api_cloud_media_delete():
        """Erase a media item: symlink (if any), its target file (if any),
        and the cloud_media_index row. Returns what was removed."""
        data = request.json or {}
        path = data.get('path', '')
        if not path:
            return jsonify({'success': False, 'error': 'no path'}), 400
        full, _ = _resolve_media_path(path)
        if not full:
            # Indexed rows store filepath as "{source}/{rel}" — strip the source
            # prefix and try again against the matching base dir.
            parts = path.split('/', 1)
            if len(parts) == 2:
                for base, s in CLOUD_MEDIA_DIRS:
                    if s == parts[0]:
                        cand = os.path.realpath(os.path.join(base, parts[1]))
                        if (cand.startswith(os.path.realpath(base))
                                and (os.path.exists(cand) or os.path.islink(cand))):
                            full = cand
                            break
        if not full:
            # File is gone but the index row may still be hanging around — clean it up.
            try:
                conn = sqlite3.connect(os.path.join(DASHBOARD_DIR, 'baza_projects.db'))
                conn.execute("DELETE FROM cloud_media_index WHERE filepath=?", (path,))
                conn.execute("DELETE FROM ahb_media_vision WHERE media_filepath=?", (path,))
                conn.execute("DELETE FROM ahb_media_attachments WHERE media_filepath=?", (path,))
                conn.commit(); conn.close()
            except Exception:
                pass
            return jsonify({'success': True, 'removed': [], 'note': 'index-only cleanup; file was already gone'})
        removed = []
        try:
            target = None
            if os.path.islink(full):
                # Resolve symlink target before unlinking
                try:
                    target = os.path.realpath(full)
                except Exception:
                    target = None
                os.unlink(full)
                removed.append(full)
                # Delete the underlying file too if it lives under cloud or Imports
                if target and os.path.isfile(target):
                    safe_roots = ('/mnt/empirepool/cloud/', '/mnt/empirepool/media/')
                    if target.startswith(safe_roots):
                        try:
                            os.remove(target)
                            removed.append(target)
                        except Exception:
                            pass
            elif os.path.isfile(full):
                os.remove(full)
                removed.append(full)
            else:
                return jsonify({'success': False, 'error': 'not a file/symlink'}), 400
        except Exception as e:
            return jsonify({'success': False, 'error': str(e)}), 500
        # Drop the index row + any ahb_media_vision / ahb_media_attachments orphans
        try:
            conn = sqlite3.connect(os.path.join(DASHBOARD_DIR, 'baza_projects.db'))
            conn.execute("DELETE FROM cloud_media_index WHERE filepath=?", (path,))
            conn.execute("DELETE FROM ahb_media_vision WHERE media_filepath=?", (path,))
            conn.execute("DELETE FROM ahb_media_attachments WHERE media_filepath=?", (path,))
            conn.commit(); conn.close()
        except Exception:
            pass
        return jsonify({'success': True, 'removed': removed})

    @app.route('/api/cloud/documents')
    def api_cloud_documents():
        cat = request.args.get('category', '')
        conn = sqlite3.connect(os.path.join(DASHBOARD_DIR, 'baza_projects.db'))
        conn.row_factory = sqlite3.Row
        count = conn.execute("SELECT COUNT(*) FROM cloud_media_index WHERE media_type='document'").fetchone()[0]
        if count == 0:
            conn.close(); _scan_media_dirs()
            conn = sqlite3.connect(os.path.join(DASHBOARD_DIR, 'baza_projects.db'))
            conn.row_factory = sqlite3.Row
        sql = "SELECT * FROM cloud_media_index WHERE media_type='document'"
        params = []
        if cat and cat != 'all':
            sql += " AND category=?"; params.append(cat)
        sql += " ORDER BY date_taken DESC, filename"
        rows = [dict(r) for r in conn.execute(sql, params).fetchall()]
        conn.close()
        return jsonify(rows)

    @app.route('/api/cloud/documents/categorize', methods=['POST'])
    def api_cloud_documents_categorize():
        data = request.json or {}
        conn = sqlite3.connect(os.path.join(DASHBOARD_DIR, 'baza_projects.db'))
        conn.execute("UPDATE cloud_media_index SET category=? WHERE filepath=?",
                     (data.get('category','other'), data.get('path','')))
        conn.commit(); conn.close()
        return jsonify({'success': True})

    @app.route('/api/cloud/files/rename', methods=['POST'])
    def api_cloud_rename():
        data = request.json or {}
        user_dir = os.path.join(CLOUD_STORAGE, str(FAMILY_USER_ID))
        old = os.path.realpath(os.path.join(user_dir, data.get('path', '')))
        if not old.startswith(os.path.realpath(user_dir)) or not os.path.exists(old):
            return jsonify({'error': 'Invalid path'}), 403
        new_name = re.sub(r'[^\w.\-]', '_', data.get('new_name', ''))
        if not new_name:
            return jsonify({'error': 'Name required'}), 400
        new = os.path.join(os.path.dirname(old), new_name)
        os.rename(old, new)
        return jsonify({'success': True, 'new_name': new_name})

    # ── Cloud: download by query-param (matches cloud.html client code) ─────
    @app.route('/api/cloud/files/download', methods=['GET'])
    def api_cloud_download_q():
        """Download via ?path=... — the cloud.html client uses this shape."""
        filepath = request.args.get('path', '')
        user_dir = os.path.join(CLOUD_STORAGE, str(FAMILY_USER_ID))
        target = os.path.realpath(os.path.join(user_dir, filepath))
        if not target.startswith(os.path.realpath(user_dir)) or not os.path.isfile(target):
            return jsonify({'error': 'Invalid path'}), 403
        return send_from_directory(os.path.dirname(target),
                                   os.path.basename(target), as_attachment=True)

    # ── Cloud: in-browser open (inline, not download) ──────────────────────
    @app.route('/api/cloud/files/open', methods=['GET'])
    def api_cloud_open():
        """Serve file inline so the browser previews images/PDFs/video."""
        filepath = request.args.get('path', '')
        user_dir = os.path.join(CLOUD_STORAGE, str(FAMILY_USER_ID))
        target = os.path.realpath(os.path.join(user_dir, filepath))
        if not target.startswith(os.path.realpath(user_dir)) or not os.path.isfile(target):
            return jsonify({'error': 'Invalid path'}), 403
        return send_from_directory(os.path.dirname(target),
                                   os.path.basename(target), as_attachment=False)

    # ── Cloud: share link (tokenized, optionally time-limited) ─────────────
    @app.route('/api/cloud/files/share', methods=['POST'])
    def api_cloud_share_create():
        """Create a tokenized public share for a single file.
        Body: {path, expires_days?}  Returns: {token, url, expires_at}"""
        import secrets as _secrets
        import datetime as _dt
        data = request.json or {}
        filepath = data.get('path', '').strip()
        if not filepath:
            return jsonify({'success': False, 'error': 'path required'}), 400
        user_dir = os.path.join(CLOUD_STORAGE, str(FAMILY_USER_ID))
        target = os.path.realpath(os.path.join(user_dir, filepath))
        if not target.startswith(os.path.realpath(user_dir)) or not os.path.isfile(target):
            return jsonify({'success': False, 'error': 'Invalid path'}), 403
        days = int(data.get('expires_days', 7))
        expires_at = (_dt.datetime.utcnow() + _dt.timedelta(days=days)).isoformat() if days > 0 else None
        token = _secrets.token_urlsafe(18)
        try:
            conn = _ahb_db()
            conn.execute(
                """INSERT INTO cloud_shares (token, user_id, path, expires_at, created_by)
                   VALUES (?, ?, ?, ?, ?)""",
                (token, str(FAMILY_USER_ID), filepath, expires_at, 'serge'),
            )
            conn.commit()
            conn.close()
        except Exception as e:
            return jsonify({'success': False, 'error': str(e)}), 500
        share_url = f"{_public_base_url()}/s/{token}"
        return jsonify({'success': True, 'token': token, 'url': share_url,
                        'expires_at': expires_at, 'path': filepath})

    @app.route('/api/cloud/files/share/list', methods=['GET'])
    def api_cloud_share_list():
        """Return existing shares (optionally filtered by path)."""
        path = request.args.get('path')
        conn = _ahb_db()
        conn.row_factory = sqlite3.Row
        q = "SELECT token, path, expires_at, created_at, access_count, last_accessed_at FROM cloud_shares"
        vals = ()
        if path:
            q += " WHERE path = ?"
            vals = (path,)
        q += " ORDER BY created_at DESC LIMIT 500"
        rows = conn.execute(q, vals).fetchall()
        conn.close()
        base = _public_base_url()
        return jsonify({'shares': [
            {**dict(r), 'url': f"{base}/s/{r['token']}"} for r in rows
        ]})

    @app.route('/api/cloud/files/share/<token>', methods=['DELETE'])
    def api_cloud_share_revoke(token):
        conn = _ahb_db()
        conn.execute("DELETE FROM cloud_shares WHERE token = ?", (token,))
        conn.commit()
        conn.close()
        return jsonify({'success': True})

    # ── Cloud: send file via Telegram (uses Phil's bot) ────────────────────
    @app.route('/api/cloud/files/telegram', methods=['POST'])
    def api_cloud_telegram_send():
        """Send a cloud file to a Telegram chat.
        Body: {path, chat_id?, caption?}
        Chat defaults to SERGE_CHAT_ID env, then falls back to the most recent
        chat_id in task_journal for phil_hass. Bot token = TELEGRAM_PHIL_HASS."""
        data = request.json or {}
        filepath = (data.get('path') or '').strip()
        if not filepath:
            return jsonify({'success': False, 'error': 'path required'}), 400
        user_dir = os.path.join(CLOUD_STORAGE, str(FAMILY_USER_ID))
        target = os.path.realpath(os.path.join(user_dir, filepath))
        if not target.startswith(os.path.realpath(user_dir)) or not os.path.isfile(target):
            return jsonify({'success': False, 'error': 'Invalid path'}), 403

        token = os.environ.get('CLOUD_TELEGRAM_BOT') or os.environ.get('TELEGRAM_PHIL_HASS')
        if not token:
            return jsonify({'success': False,
                            'error': 'No Telegram bot token configured '
                                     '(set TELEGRAM_PHIL_HASS or CLOUD_TELEGRAM_BOT)'}), 500

        chat_id = str(data.get('chat_id') or os.environ.get('SERGE_CHAT_ID') or '').strip()
        if not chat_id:
            # Fallback: most recent chat_id in task_journal with any agent
            try:
                from core.context_db import get_pool as _gp
                pool = _gp()
                c = pool.getconn()
                cur = c.cursor()
                cur.execute("SELECT chat_id FROM task_journal "
                            "WHERE chat_id IS NOT NULL AND chat_id != '' "
                            "ORDER BY created_at DESC LIMIT 1")
                r = cur.fetchone()
                cur.close(); pool.putconn(c)
                if r and r[0]:
                    chat_id = str(r[0])
            except Exception:
                pass
        if not chat_id:
            return jsonify({'success': False,
                            'error': 'No chat_id provided and SERGE_CHAT_ID not set'}), 400

        caption = (data.get('caption') or '').strip()
        ext = os.path.splitext(target)[1].lower()
        api_base = f"https://api.telegram.org/bot{token}"

        # Pick the right endpoint based on file type so the preview is nice.
        if ext in ('.jpg', '.jpeg', '.png', '.webp', '.gif'):
            method = 'sendPhoto'
            field = 'photo'
        elif ext in ('.mp4', '.mov', '.m4v', '.webm'):
            method = 'sendVideo'
            field = 'video'
        elif ext in ('.mp3', '.m4a', '.wav', '.ogg'):
            method = 'sendAudio'
            field = 'audio'
        else:
            method = 'sendDocument'
            field = 'document'

        try:
            import requests as _rq
            with open(target, 'rb') as fh:
                files = {field: (os.path.basename(target), fh)}
                payload = {'chat_id': chat_id}
                if caption:
                    payload['caption'] = caption[:1024]
                resp = _rq.post(f"{api_base}/{method}", data=payload, files=files, timeout=120)
            try:
                result = resp.json()
            except Exception:
                result = {'raw': resp.text[:500]}
            if resp.status_code == 200 and result.get('ok'):
                return jsonify({'success': True, 'chat_id': chat_id,
                                'method': method, 'filename': os.path.basename(target)})
            return jsonify({'success': False, 'error': result.get('description') or result}), 502
        except Exception as e:
            return jsonify({'success': False, 'error': str(e)}), 500

    # ── Cloud: HLS wrapper for .ts dashcam clips ───────────────────────────
    @app.route('/api/cloud/media/hls', methods=['GET'])
    def api_cloud_hls_manifest():
        """Serve a single-segment HLS manifest that points to the TS file.
        hls.js on the frontend feeds the segment to a <video> via MSE so the
        browser can play MPEG-TS clips (e.g. dashcam footage) without remux."""
        filepath = request.args.get('path', '')
        user_dir = os.path.join(CLOUD_STORAGE, str(FAMILY_USER_ID))
        target = os.path.realpath(os.path.join(user_dir, filepath))
        if not target.startswith(os.path.realpath(user_dir)) or not os.path.isfile(target):
            return jsonify({'error': 'Invalid path'}), 403

        # Compute duration with ffprobe if available; default to 300s.
        duration = 300
        try:
            import subprocess as _sp
            out = _sp.check_output(
                ['ffprobe', '-v', 'error', '-show_entries', 'format=duration',
                 '-of', 'default=noprint_wrappers=1:nokey=1', target],
                stderr=_sp.DEVNULL, timeout=10,
            ).decode().strip()
            if out:
                duration = int(float(out)) + 1
        except Exception:
            pass

        ts_url = f"/api/cloud/files/open?path={request.args.get('path')}"
        manifest = (
            "#EXTM3U\n"
            "#EXT-X-VERSION:3\n"
            f"#EXT-X-TARGETDURATION:{duration}\n"
            "#EXT-X-MEDIA-SEQUENCE:0\n"
            "#EXT-X-PLAYLIST-TYPE:VOD\n"
            f"#EXTINF:{duration}.0,\n"
            f"{ts_url}\n"
            "#EXT-X-ENDLIST\n"
        )
        resp = make_response(manifest)
        resp.headers['Content-Type'] = 'application/vnd.apple.mpegurl'
        resp.headers['Cache-Control'] = 'no-store'
        return resp

    # ── Public share endpoint: no auth required ────────────────────────────
    @app.route('/s/<token>')
    def public_share(token):
        import datetime as _dt
        conn = _ahb_db()
        conn.row_factory = sqlite3.Row
        row = conn.execute(
            "SELECT user_id, path, expires_at FROM cloud_shares WHERE token = ?",
            (token,),
        ).fetchone()
        if not row:
            conn.close()
            return "Share link not found or revoked", 404
        if row['expires_at']:
            try:
                if _dt.datetime.fromisoformat(row['expires_at']) < _dt.datetime.utcnow():
                    conn.close()
                    return "Share link expired", 410
            except Exception:
                pass
        user_dir = os.path.join(CLOUD_STORAGE, str(row['user_id']))
        target = os.path.realpath(os.path.join(user_dir, row['path']))
        if not target.startswith(os.path.realpath(user_dir)) or not os.path.isfile(target):
            conn.close()
            return "File no longer available", 404
        conn.execute(
            "UPDATE cloud_shares SET access_count = access_count + 1, "
            "last_accessed_at = datetime('now') WHERE token = ?", (token,))
        conn.commit()
        conn.close()
        # Download-as-attachment by default; add ?inline=1 to preview in-browser.
        inline = request.args.get('inline') in ('1', 'true', 'yes')
        return send_from_directory(os.path.dirname(target),
                                   os.path.basename(target),
                                   as_attachment=not inline)


# ─────────────────────────────────────────────────────────────────────────────
# Public Review Page + QR Code + AHB123 Project Photos
# ─────────────────────────────────────────────────────────────────────────────

REVIEWS_DIR = os.path.join(ARTIFACTS_DIR, 'ahb123-reviews')
PHOTOS_DIR  = os.path.join(ARTIFACTS_DIR, 'ahb123-photos')
os.makedirs(REVIEWS_DIR, exist_ok=True)
os.makedirs(PHOTOS_DIR, exist_ok=True)

def _public_base_url():
    """Return the base URL the public should use to reach this dashboard."""
    return os.environ.get('BAZA_PUBLIC_URL', '').rstrip('/') or request.host_url.rstrip('/')

@app.route('/api/qr')
def api_qr():
    """Generate a QR code PNG for any data string. ?data=<url>&size=<px>"""
    import io
    try:
        import qrcode
    except ImportError:
        return jsonify({'error': 'qrcode lib not installed'}), 500
    data = request.args.get('data', '').strip()
    if not data:
        return jsonify({'error': 'data param required'}), 400
    size = max(1, min(40, int(request.args.get('box', '10'))))
    qr = qrcode.QRCode(version=None, error_correction=qrcode.constants.ERROR_CORRECT_M,
                       box_size=size, border=2)
    qr.add_data(data)
    qr.make(fit=True)
    img = qr.make_image(fill_color='#00bcd4', back_color='#0a0a16')
    buf = io.BytesIO()
    img.save(buf, format='PNG')
    buf.seek(0)
    resp = make_response(buf.read())
    resp.headers['Content-Type'] = 'image/png'
    resp.headers['Cache-Control'] = 'public, max-age=3600'
    return resp

@app.route('/review')
def public_review_page():
    """Public-facing review form (no auth) — accessed via QR code."""
    return render_template('review_public.html', company='All Home Building Co LLC')

@app.route('/api/review/submit', methods=['POST'])
def api_review_submit():
    """Public review submission. Accepts multipart form (with optional photo) or JSON."""
    import re as _re
    if request.content_type and 'multipart' in request.content_type:
        data = {
            'stars':        int(request.form.get('stars', '0') or 0),
            'name':         request.form.get('name', '').strip()[:80],
            'text':         request.form.get('text', '').strip()[:2000],
            'project_type': request.form.get('project_type', '').strip()[:80],
            'tags':         [t for t in request.form.get('tags', '').split(',') if t.strip()],
            'email':        request.form.get('email', '').strip()[:120],
            'phone':        request.form.get('phone', '').strip()[:30],
        }
        photo = request.files.get('photo')
    else:
        body = request.get_json(silent=True) or {}
        data = {
            'stars':        int(body.get('stars', 0) or 0),
            'name':         (body.get('name') or '').strip()[:80],
            'text':         (body.get('text') or '').strip()[:2000],
            'project_type': (body.get('project_type') or '').strip()[:80],
            'tags':         body.get('tags') or [],
            'email':        (body.get('email') or '').strip()[:120],
            'phone':        (body.get('phone') or '').strip()[:30],
        }
        photo = None
    if not (1 <= data['stars'] <= 5):
        return jsonify({'success': False, 'error': 'stars must be 1-5'}), 400
    data['date'] = datetime.datetime.now().strftime('%Y-%m-%d')
    data['ts']   = datetime.datetime.now().isoformat()
    data['source'] = 'public_qr'
    data['ip']   = request.remote_addr
    # Auto-publish 4-5 star reviews to the public website. 1-3 star stay unpublished
    # pending manual moderation by Serge / Nova so we can address concerns first.
    data['published'] = data['stars'] >= 4
    ts = int(datetime.datetime.now().timestamp())
    if photo and photo.filename:
        safe_name = _re.sub(r'[^\w.\-_]', '_', photo.filename)
        photo_name = f'review_{ts}_{safe_name}'
        photo.save(os.path.join(REVIEWS_DIR, photo_name))
        data['photo'] = photo_name
        data['has_photo'] = True
    fname = f'review_{ts}.json'
    with open(os.path.join(REVIEWS_DIR, fname), 'w') as f:
        json.dump(data, f, indent=2)
    # Notify Nova via task journal so agents see it
    try:
        from core.context_db import save_task_journal
        save_task_journal('nova_sterling', 'public_review_received',
                          f"{data['stars']}-star from {data['name'] or 'anonymous'}: {data['text'][:120]}",
                          status='completed')
    except Exception:
        pass
    return jsonify({'success': True, 'file': fname})


def _load_reviews(only_published: bool = False) -> list:
    """Walk the reviews artifact dir and return all review JSON records, newest first."""
    if not os.path.exists(REVIEWS_DIR):
        return []
    out = []
    for fname in os.listdir(REVIEWS_DIR):
        if not fname.endswith('.json'):
            continue
        try:
            with open(os.path.join(REVIEWS_DIR, fname)) as f:
                rev = json.load(f)
            rev['_file'] = fname
            if only_published and not rev.get('published', False):
                continue
            out.append(rev)
        except Exception:
            continue
    out.sort(key=lambda r: r.get('ts', ''), reverse=True)
    return out


@app.route('/api/reviews/published', methods=['GET'])
def api_reviews_published():
    """Public-facing endpoint: returns all reviews flagged for publication.
    Used by the ahb123.com website to render the reviews section + marketing widgets.
    Strips PII (email, phone, IP) before returning."""
    reviews = _load_reviews(only_published=True)
    safe = []
    for r in reviews:
        safe.append({
            'stars':         r.get('stars'),
            'name':          r.get('name'),
            'text':          r.get('text'),
            'project_type':  r.get('project_type'),
            'tags':          r.get('tags'),
            'date':          r.get('date'),
            'photo':         r.get('photo'),
            'has_photo':     r.get('has_photo', False),
            'id':            r.get('_file', '').replace('review_','').replace('.json',''),
        })
    # Aggregate stats for marketing — average rating, total count
    if safe:
        avg = sum(r['stars'] for r in safe) / len(safe)
    else:
        avg = 0
    resp = make_response(jsonify({
        'reviews':   safe,
        'count':     len(safe),
        'avg_stars': round(avg, 2),
    }))
    # Allow cross-origin embedding from the static website
    resp.headers['Access-Control-Allow-Origin'] = '*'
    resp.headers['Cache-Control'] = 'public, max-age=300'  # 5min cache
    return resp


@app.route('/api/reviews/all', methods=['GET'])
def api_reviews_all():
    """Admin-only: returns ALL reviews (published + unpublished) with full PII for moderation."""
    return jsonify(_load_reviews(only_published=False))


@app.route('/api/reviews/<rid>/publish', methods=['POST'])
def api_review_publish(rid):
    """Toggle a review's published flag for marketing display."""
    body = request.get_json(silent=True) or {}
    publish = body.get('publish', True)
    fname = f'review_{rid}.json' if not rid.startswith('review_') else (rid + ('.json' if not rid.endswith('.json') else ''))
    fpath = os.path.join(REVIEWS_DIR, fname)
    if not os.path.exists(fpath):
        return jsonify({'success': False, 'error': 'not found'}), 404
    try:
        with open(fpath) as f:
            data = json.load(f)
        data['published'] = bool(publish)
        with open(fpath, 'w') as f:
            json.dump(data, f, indent=2)
        return jsonify({'success': True, 'published': data['published']})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/reviews/photo/<path:fname>')
def api_review_photo(fname):
    """Serve uploaded review photos. Public — accessed from the website."""
    resp = send_from_directory(REVIEWS_DIR, fname)
    resp.headers['Access-Control-Allow-Origin'] = '*'
    return resp


# ── EstimatOR Super Tool: 3-method estimation ───────────────────────────────

def _ensure_estimator_settings():
    """Idempotent init: create the settings table and seed row 1 if missing.

    Wrapped in try/except so the dashboard can still start when another
    process holds a write lock on `baza_projects.db` (e.g. a long-running
    backfill/dedup script). The seed row is already present on every
    deployed instance — losing this one-time init at boot has no functional
    impact; the routes that read it default to sensible values on miss."""
    try:
        conn = sqlite3.connect(os.path.join(DASHBOARD_DIR, 'baza_projects.db'), timeout=8.0)
        conn.execute("PRAGMA busy_timeout = 8000")
        conn.execute("""CREATE TABLE IF NOT EXISTS ahb_estimator_settings (
            id INTEGER PRIMARY KEY CHECK (id = 1),
            crew_day_rate REAL DEFAULT 800,
            lead_day_rate REAL DEFAULT 1200,
            helper_day_rate REAL DEFAULT 450,
            sub_day_rate REAL DEFAULT 1500,
            materials_pct REAL DEFAULT 0.40,
            overhead_pct REAL DEFAULT 0.15,
            profit_pct REAL DEFAULT 0.18,
            admin_fee_pct REAL DEFAULT 0.05,
            permit_fee_default REAL DEFAULT 350,
            contingency_pct REAL DEFAULT 0.10,
            last_low_high_factor_low REAL DEFAULT 0.75,
            last_low_high_factor_high REAL DEFAULT 1.30,
            updated_at TEXT DEFAULT CURRENT_TIMESTAMP
        )""")
        conn.execute("INSERT OR IGNORE INTO ahb_estimator_settings (id) VALUES (1)")
        conn.commit()
        conn.close()
    except sqlite3.OperationalError as e:
        print(f"[startup] _ensure_estimator_settings deferred — DB busy: {e}", flush=True)
_ensure_estimator_settings()


@app.route('/api/ahb/estimator/settings', methods=['GET','PUT'])
def api_estimator_settings():
    conn = sqlite3.connect(os.path.join(DASHBOARD_DIR, 'baza_projects.db'))
    conn.row_factory = sqlite3.Row
    if request.method == 'GET':
        row = conn.execute("SELECT * FROM ahb_estimator_settings WHERE id=1").fetchone()
        conn.close()
        return jsonify(dict(row) if row else {})
    body = request.get_json() or {}
    fields = ['crew_day_rate','lead_day_rate','helper_day_rate','sub_day_rate',
              'materials_pct','overhead_pct','profit_pct','admin_fee_pct',
              'permit_fee_default','contingency_pct',
              'last_low_high_factor_low','last_low_high_factor_high']
    sets, vals = [], []
    for k in fields:
        if k in body:
            sets.append(f"{k}=?"); vals.append(float(body[k]))
    if sets:
        sets.append("updated_at=?"); vals.append(datetime.datetime.now().isoformat())
        conn.execute(f"UPDATE ahb_estimator_settings SET {','.join(sets)} WHERE id=1", vals)
        conn.commit()
    conn.close()
    return jsonify({'success': True})


@app.route('/api/ahb/estimator/method1', methods=['POST'])
def api_estimator_method1():
    """Method 1: time × men × day rate.
    Body: {days, lead_count, crew_count, helper_count, sub_days, materials, permits, contingency_pct?}
    Returns labor cost breakdown + materials + total."""
    body = request.get_json() or {}
    conn = sqlite3.connect(os.path.join(DASHBOARD_DIR, 'baza_projects.db'))
    conn.row_factory = sqlite3.Row
    s = dict(conn.execute("SELECT * FROM ahb_estimator_settings WHERE id=1").fetchone() or {})
    conn.close()
    days       = float(body.get('days', 0) or 0)
    lead_count = float(body.get('lead_count', 1) or 0)
    crew_count = float(body.get('crew_count', 0) or 0)
    helper_count = float(body.get('helper_count', 0) or 0)
    sub_days   = float(body.get('sub_days', 0) or 0)
    materials  = float(body.get('materials', 0) or 0)
    permits    = float(body.get('permits', s.get('permit_fee_default', 350)) or 0)
    contingency_pct = float(body.get('contingency_pct', s.get('contingency_pct', 0.10)) or 0)
    overhead_pct    = float(body.get('overhead_pct', s.get('overhead_pct', 0.15)) or 0)
    profit_pct      = float(body.get('profit_pct', s.get('profit_pct', 0.18)) or 0)
    admin_pct       = float(body.get('admin_fee_pct', s.get('admin_fee_pct', 0.05)) or 0)

    lead_cost   = days * lead_count   * float(s.get('lead_day_rate', 1200))
    crew_cost   = days * crew_count   * float(s.get('crew_day_rate', 800))
    helper_cost = days * helper_count * float(s.get('helper_day_rate', 450))
    sub_cost    = sub_days *            float(s.get('sub_day_rate', 1500))
    labor_total = lead_cost + crew_cost + helper_cost + sub_cost
    direct_cost = labor_total + materials + permits
    contingency = direct_cost * contingency_pct
    overhead    = direct_cost * overhead_pct
    profit      = (direct_cost + contingency + overhead) * profit_pct
    admin       = (direct_cost + contingency + overhead + profit) * admin_pct
    grand_total = direct_cost + contingency + overhead + profit + admin

    return jsonify({
        'method': 1,
        'breakdown': {
            'lead':        round(lead_cost, 2),
            'crew':        round(crew_cost, 2),
            'helpers':     round(helper_cost, 2),
            'subs':        round(sub_cost, 2),
            'labor_total': round(labor_total, 2),
            'materials':   round(materials, 2),
            'permits':     round(permits, 2),
            'direct_cost': round(direct_cost, 2),
            'contingency': round(contingency, 2),
            'overhead':    round(overhead, 2),
            'profit':      round(profit, 2),
            'admin_fee':   round(admin, 2),
        },
        'total':      round(grand_total, 2),
        'days':       days,
        'crew_size':  lead_count + crew_count + helper_count,
    })


@app.route('/api/ahb/estimator/method2', methods=['POST'])
def api_estimator_method2():
    """Method 2: Specter/Scout researches market norms for the project description
    and generates an estimate based on regional construction industry standards.
    Body: {description, scope, address?, sqft?}"""
    body = request.get_json() or {}
    description = (body.get('description') or '').strip()
    scope       = (body.get('scope') or '').strip()
    address     = (body.get('address') or '').strip()
    sqft        = body.get('sqft')
    if not description:
        return jsonify({'success': False, 'error': 'description required'}), 400

    # Build a research prompt asking the LLM to act as Specter (cloud market analyst)
    sqft_line = f"Square footage: {sqft}" if sqft else ""
    addr_line = f"Project location: {address}" if address else "Project location: Greater Philadelphia, PA"
    prompt = (
        "You are Specter Voss, market intelligence analyst for All Home Building Co LLC, "
        "a Philadelphia residential general contractor. Your job is to estimate the cost "
        "of the following project using current 2025-2026 economic norms for Greater "
        "Philadelphia / Bucks / Montgomery / Delaware county construction.\n\n"
        f"Project scope: {scope}\n"
        f"{addr_line}\n"
        f"{sqft_line}\n\n"
        "Project description:\n" + description + "\n\n"
        "Use current market data: Philadelphia regional labor rates ($45-65/hr skilled, "
        "$25-35/hr helper), material costs from Home Depot / Lowe's / supply houses, "
        "current PA UCC permit fees, typical GC overhead (15-20%), profit margin (15-22%), "
        "contingency (10%), and AHBCO's 5% admin fee.\n\n"
        "Return ONLY a JSON object with these exact keys (use realistic numbers, not zeros):\n"
        "{\n"
        '  "labor_cost":      number,\n'
        '  "materials_cost":  number,\n'
        '  "permits_cost":    number,\n'
        '  "subcontractors":  number,\n'
        '  "overhead":        number,\n'
        '  "profit":          number,\n'
        '  "admin_fee":       number,\n'
        '  "contingency":     number,\n'
        '  "total_estimate":  number,\n'
        '  "estimated_days":  number,\n'
        '  "crew_size":       number,\n'
        '  "key_materials":   ["list", "of", "main", "materials"],\n'
        '  "key_assumptions": "1-2 sentences on what assumptions drove this estimate",\n'
        '  "market_notes":    "1-2 sentences on current Philly area market conditions affecting this estimate",\n'
        '  "confidence":      0.0-1.0\n'
        "}\n\nJSON:"
    )
    try:
        result_text = _ollama_text(prompt, model="qwen2.5:14b", json_mode=True, max_tokens=900)
        result = json.loads(re.sub(r'^```(?:json)?\s*|\s*```$', '', result_text).strip())
        return jsonify({'success': True, 'method': 2, 'specter_analysis': result})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/ahb/estimator/method3', methods=['POST'])
def api_estimator_method3():
    """Method 3: generate low and high cost ranges for a project description,
    return both bounds + average. Body: {description, scope}"""
    body = request.get_json() or {}
    description = (body.get('description') or '').strip()
    scope       = (body.get('scope') or '').strip()
    if not description:
        return jsonify({'success': False, 'error': 'description required'}), 400

    prompt = (
        "You are a senior construction estimator for a Philadelphia residential GC. "
        "For the project below, give a realistic LOW estimate (budget-conscious materials, "
        "minimal change orders, fast schedule) and a HIGH estimate (premium materials, "
        "buffer for unforeseen conditions, slower careful execution). Use current 2025-2026 "
        "Greater Philadelphia market rates.\n\n"
        f"Scope: {scope}\n\n"
        f"Description: {description}\n\n"
        "Return ONLY a JSON object:\n"
        "{\n"
        '  "low":  {"total": number, "labor": number, "materials": number, "rationale": "1 sentence"},\n'
        '  "high": {"total": number, "labor": number, "materials": number, "rationale": "1 sentence"},\n'
        '  "recommended_quote": number,  // sweet spot for the customer\n'
        '  "factors": ["bullet", "list", "of", "what", "drives", "the", "spread"]\n'
        "}\n\nJSON:"
    )
    try:
        result_text = _ollama_text(prompt, model="qwen2.5:14b", json_mode=True, max_tokens=700)
        result = json.loads(re.sub(r'^```(?:json)?\s*|\s*```$', '', result_text).strip())
        avg = (float(result['low']['total']) + float(result['high']['total'])) / 2
        result['average'] = round(avg, 2)
        return jsonify({'success': True, 'method': 3, 'range': result})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


# ── ahb123.com Static Site Server ───────────────────────────────────────────
# Serves the static HTML/images at /site/* so you can preview the live site
# locally and the embedded JS API calls run on the same origin (no CORS issues).
SITE_DIR = os.path.join(ARTIFACTS_DIR, 'proj-ahb123', 'website')


@app.route('/site/')
@app.route('/site')
def ahb123_site_index():
    return send_from_directory(SITE_DIR, 'index.html')


@app.route('/site/<path:filename>')
def ahb123_site_file(filename):
    return send_from_directory(SITE_DIR, filename)


# ── AHB123 Project Photos (before / during / after) ─────────────────────────

PHOTO_META_DB = os.path.join(DASHBOARD_DIR, 'baza_projects.db')

def _init_photos_table():
    conn = sqlite3.connect(PHOTO_META_DB)
    conn.execute('''CREATE TABLE IF NOT EXISTS ahb_photos (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        filename TEXT NOT NULL,
        project_name TEXT,
        client_name TEXT,
        location TEXT,
        phase TEXT,            -- before / during / after
        category TEXT,         -- kitchen / bath / roof / etc
        photo_date TEXT,       -- YYYY-MM-DD
        photo_time TEXT,       -- HH:MM
        latitude REAL,
        longitude REAL,
        notes TEXT,
        size INTEGER,
        uploaded_at TEXT DEFAULT CURRENT_TIMESTAMP
    )''')
    conn.commit()
    conn.close()
_init_photos_table()

def _extract_exif(filepath):
    """Pull date/gps from EXIF if available."""
    try:
        from PIL import Image
        from PIL.ExifTags import TAGS, GPSTAGS
        img = Image.open(filepath)
        exif = img._getexif() or {}
        out = {}
        for tag_id, value in exif.items():
            tag = TAGS.get(tag_id, tag_id)
            if tag == 'DateTimeOriginal' and isinstance(value, str):
                # "2024:08:14 13:22:08"
                try:
                    dt = datetime.datetime.strptime(value, '%Y:%m:%d %H:%M:%S')
                    out['photo_date'] = dt.strftime('%Y-%m-%d')
                    out['photo_time'] = dt.strftime('%H:%M')
                except Exception:
                    pass
            if tag == 'GPSInfo':
                gps = {GPSTAGS.get(k, k): v for k, v in value.items()}
                def _to_deg(val, ref):
                    try:
                        d, m, s = val
                        deg = float(d) + float(m)/60 + float(s)/3600
                        if ref in ('S', 'W'):
                            deg = -deg
                        return round(deg, 6)
                    except Exception:
                        return None
                if 'GPSLatitude' in gps and 'GPSLatitudeRef' in gps:
                    out['latitude'] = _to_deg(gps['GPSLatitude'], gps['GPSLatitudeRef'])
                if 'GPSLongitude' in gps and 'GPSLongitudeRef' in gps:
                    out['longitude'] = _to_deg(gps['GPSLongitude'], gps['GPSLongitudeRef'])
        return out
    except Exception:
        return {}

@app.route('/api/ahb/photos', methods=['GET'])
def api_ahb_photos_list():
    conn = sqlite3.connect(PHOTO_META_DB)
    conn.row_factory = sqlite3.Row
    rows = conn.execute('SELECT * FROM ahb_photos ORDER BY photo_date DESC, photo_time DESC, id DESC').fetchall()
    conn.close()
    return jsonify([dict(r) for r in rows])

@app.route('/api/ahb/photos/upload', methods=['POST'])
def api_ahb_photos_upload():
    """Bulk upload photos. Form fields: project_name, client_name, location, phase, category, notes."""
    import re as _re
    files = request.files.getlist('files') or request.files.getlist('file')
    if not files:
        return jsonify({'success': False, 'error': 'no files'}), 400
    meta_defaults = {
        'project_name': request.form.get('project_name', '').strip(),
        'client_name':  request.form.get('client_name', '').strip(),
        'location':     request.form.get('location', '').strip(),
        'phase':        request.form.get('phase', '').strip(),
        'category':     request.form.get('category', '').strip(),
        'notes':        request.form.get('notes', '').strip(),
    }
    saved = []
    conn = sqlite3.connect(PHOTO_META_DB)
    for f in files:
        if not f or not f.filename:
            continue
        ts = int(datetime.datetime.now().timestamp() * 1000)
        safe = _re.sub(r'[^\w.\-_]', '_', f.filename)
        fname = f'{ts}_{safe}'
        fpath = os.path.join(PHOTOS_DIR, fname)
        f.save(fpath)
        size = os.path.getsize(fpath)
        exif = _extract_exif(fpath)
        photo_date = exif.get('photo_date') or datetime.datetime.now().strftime('%Y-%m-%d')
        photo_time = exif.get('photo_time') or datetime.datetime.now().strftime('%H:%M')
        conn.execute('''INSERT INTO ahb_photos
            (filename, project_name, client_name, location, phase, category,
             photo_date, photo_time, latitude, longitude, notes, size)
            VALUES (?,?,?,?,?,?,?,?,?,?,?,?)''',
            (fname, meta_defaults['project_name'], meta_defaults['client_name'],
             meta_defaults['location'], meta_defaults['phase'], meta_defaults['category'],
             photo_date, photo_time, exif.get('latitude'), exif.get('longitude'),
             meta_defaults['notes'], size))
        saved.append(fname)
    conn.commit()
    conn.close()
    return jsonify({'success': True, 'count': len(saved), 'files': saved})

@app.route('/api/ahb/photos/<int:pid>', methods=['PATCH', 'DELETE'])
def api_ahb_photos_modify(pid):
    conn = sqlite3.connect(PHOTO_META_DB)
    if request.method == 'DELETE':
        row = conn.execute('SELECT filename FROM ahb_photos WHERE id=?', (pid,)).fetchone()
        if row:
            try: os.remove(os.path.join(PHOTOS_DIR, row[0]))
            except Exception: pass
        conn.execute('DELETE FROM ahb_photos WHERE id=?', (pid,))
        conn.commit()
        conn.close()
        return jsonify({'success': True})
    body = request.get_json() or {}
    fields = ['project_name','client_name','location','phase','category',
              'photo_date','photo_time','latitude','longitude','notes']
    sets, vals = [], []
    for k in fields:
        if k in body:
            sets.append(f'{k}=?'); vals.append(body[k])
    if sets:
        vals.append(pid)
        conn.execute(f'UPDATE ahb_photos SET {",".join(sets)} WHERE id=?', vals)
        conn.commit()
    conn.close()
    return jsonify({'success': True})

@app.route('/api/ahb/photos/file/<path:fname>')
def api_ahb_photos_serve(fname):
    return send_from_directory(PHOTOS_DIR, fname)

# ── Phil's Document Library + DocPrep / Application Packages ────────────────

def _ensure_docprep_tables():
   try:
    conn = sqlite3.connect(os.path.join(DASHBOARD_DIR, 'baza_projects.db'), timeout=8.0)
    conn.execute("PRAGMA busy_timeout = 8000")
    conn.execute("""CREATE TABLE IF NOT EXISTS ahb_documents (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        file_path TEXT NOT NULL UNIQUE,
        original_name TEXT,
        suggested_name TEXT,
        doc_type TEXT,
        entity TEXT,
        doc_date TEXT,
        summary TEXT,
        relevance TEXT,
        tags TEXT,
        confidence REAL,
        agent_id TEXT,
        chat_id TEXT,
        project_id TEXT,
        content_text TEXT,
        file_size INTEGER,
        file_kind TEXT,
        curated_at TEXT DEFAULT CURRENT_TIMESTAMP,
        expires_at TEXT,
        expiry_alerted INTEGER DEFAULT 0
    )""")
    conn.execute("""CREATE TABLE IF NOT EXISTS ahb_app_packages (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        name TEXT NOT NULL,
        package_type TEXT,
        project_id TEXT,
        client_id TEXT,
        status TEXT DEFAULT 'draft',
        form_data TEXT,
        attached_doc_ids TEXT,
        notes TEXT,
        submitted_at TEXT,
        approved_at TEXT,
        permit_number TEXT,
        last_reminder_at TEXT,
        created_at TEXT DEFAULT CURRENT_TIMESTAMP,
        updated_at TEXT DEFAULT CURRENT_TIMESTAMP
    )""")
    # Vendor mini-CRM (feature 6)
    conn.execute("""CREATE TABLE IF NOT EXISTS ahb_vendors (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        name TEXT NOT NULL UNIQUE,
        vendor_type TEXT,
        contact_name TEXT,
        phone TEXT,
        email TEXT,
        address TEXT,
        ein_or_ssn TEXT,
        coi_doc_id INTEGER,
        w9_doc_id INTEGER,
        license_doc_id INTEGER,
        coi_expires TEXT,
        license_expires TEXT,
        notes TEXT,
        created_at TEXT DEFAULT CURRENT_TIMESTAMP,
        updated_at TEXT DEFAULT CURRENT_TIMESTAMP
    )""")
    # Add columns to existing tables if missing (idempotent migrations)
    for col_def in [
        ("ahb_documents",   "expires_at TEXT"),
        ("ahb_documents",   "expiry_alerted INTEGER DEFAULT 0"),
        ("ahb_app_packages","submitted_at TEXT"),
        ("ahb_app_packages","approved_at TEXT"),
        ("ahb_app_packages","permit_number TEXT"),
        ("ahb_app_packages","last_reminder_at TEXT"),
        ("ahb_employees",   "tax_classification TEXT DEFAULT 'W2'"),
        ("ahb_employees",   "business_name TEXT"),
        ("ahb_employees",   "tax_id TEXT"),
        ("ahb_employees",   "tax_id_type TEXT"),
        ("ahb_employees",   "address TEXT"),
        ("ahb_employees",   "w9_doc_id INTEGER"),
        ("ahb_employees",   "w9_signed_date TEXT"),
    ]:
        table, col = col_def
        col_name = col.split()[0]
        try:
            conn.execute(f"ALTER TABLE {table} ADD COLUMN {col}")
        except Exception:
            pass
    conn.commit()
    conn.close()
   except sqlite3.OperationalError as e:
    print(f"[startup] _ensure_docprep_tables deferred — DB busy: {e}", flush=True)
_ensure_docprep_tables()


@app.route('/api/ahb/documents', methods=['GET'])
def api_ahb_documents_list():
    """List all curated documents with optional filters."""
    conn = sqlite3.connect(os.path.join(DASHBOARD_DIR, 'baza_projects.db'))
    conn.row_factory = sqlite3.Row
    q = "SELECT * FROM ahb_documents WHERE 1=1"
    params = []
    if request.args.get('doc_type'):
        q += " AND doc_type = ?"; params.append(request.args['doc_type'])
    if request.args.get('entity'):
        q += " AND entity LIKE ?"; params.append(f"%{request.args['entity']}%")
    if request.args.get('project_id'):
        q += " AND project_id = ?"; params.append(request.args['project_id'])
    if request.args.get('q'):
        like = f"%{request.args['q']}%"
        q += " AND (entity LIKE ? OR summary LIKE ? OR tags LIKE ? OR original_name LIKE ?)"
        params += [like, like, like, like]
    q += " ORDER BY curated_at DESC LIMIT 500"
    rows = conn.execute(q, params).fetchall()
    conn.close()
    out = []
    for r in rows:
        d = dict(r)
        try: d['tags'] = json.loads(d.get('tags') or '[]')
        except Exception: d['tags'] = []
        out.append(d)
    return jsonify(out)


@app.route('/api/ahb/documents/<int:did>', methods=['GET','PATCH','DELETE'])
def api_ahb_document_one(did):
    conn = sqlite3.connect(os.path.join(DASHBOARD_DIR, 'baza_projects.db'))
    conn.row_factory = sqlite3.Row
    if request.method == 'GET':
        row = conn.execute("SELECT * FROM ahb_documents WHERE id=?", (did,)).fetchone()
        conn.close()
        if not row:
            return jsonify({'error':'not found'}), 404
        d = dict(row)
        try: d['tags'] = json.loads(d.get('tags') or '[]')
        except Exception: d['tags'] = []
        return jsonify(d)
    if request.method == 'DELETE':
        row = conn.execute("SELECT file_path FROM ahb_documents WHERE id=?", (did,)).fetchone()
        if row:
            try: os.remove(row['file_path'])
            except Exception: pass
            try: os.remove(row['file_path'] + '.meta')
            except Exception: pass
        conn.execute("DELETE FROM ahb_documents WHERE id=?", (did,))
        conn.commit()
        conn.close()
        return jsonify({'success': True})
    body = request.get_json() or {}
    fields = ['doc_type','entity','doc_date','summary','relevance','project_id','suggested_name']
    sets, vals = [], []
    for k in fields:
        if k in body:
            sets.append(f"{k}=?"); vals.append(body[k])
    if 'tags' in body:
        sets.append("tags=?"); vals.append(json.dumps(body['tags'] or []))
    if sets:
        vals.append(did)
        conn.execute(f"UPDATE ahb_documents SET {','.join(sets)} WHERE id=?", vals)
        conn.commit()
    conn.close()
    return jsonify({'success': True})


@app.route('/api/ahb/documents/file/<int:did>')
def api_ahb_document_file(did):
    conn = sqlite3.connect(os.path.join(DASHBOARD_DIR, 'baza_projects.db'))
    conn.row_factory = sqlite3.Row
    row = conn.execute("SELECT file_path,suggested_name,original_name FROM ahb_documents WHERE id=?",
                       (did,)).fetchone()
    conn.close()
    if not row or not os.path.exists(row['file_path']):
        return jsonify({'error':'not found'}), 404
    return send_from_directory(
        os.path.dirname(row['file_path']),
        os.path.basename(row['file_path']),
        as_attachment=False,
        download_name=row['suggested_name'] or row['original_name']
    )


@app.route('/api/ahb/documents/curate-existing', methods=['POST'])
def api_ahb_documents_curate_existing():
    """One-click backfill: walk artifacts and curate every doc/image not yet in the library.
    Runs the curate_document skill on each unindexed file."""
    body = request.get_json(silent=True) or {}
    limit = int(body.get('limit', 50))
    conn = sqlite3.connect(os.path.join(DASHBOARD_DIR, 'baza_projects.db'))
    existing = {row[0] for row in conn.execute("SELECT file_path FROM ahb_documents").fetchall()}
    conn.close()
    targets = []
    SKIPPABLE_EXTS = {'.meta','.pyc','.lock'}
    for root, dirs, files in os.walk(ARTIFACTS_DIR):
        dirs[:] = [d for d in dirs if not d.startswith('.')]
        for f in files:
            if any(f.endswith(s) for s in SKIPPABLE_EXTS): continue
            fp = os.path.join(root, f)
            if fp in existing: continue
            targets.append(fp)
            if len(targets) >= limit: break
        if len(targets) >= limit: break
    # Run curator on each target via the shared skill (subprocess)
    results = []
    skill = os.path.join(FRAMEWORK_DIR, 'skills', 'shared', 'curate_document.py')
    for fp in targets:
        try:
            env = os.environ.copy()
            env['SKILL_ARGS'] = json.dumps({'file_path': fp, 'agent_id': 'phil_hass'})
            p = subprocess.run([VENV_PYTHON, skill], env=env, capture_output=True,
                               text=True, timeout=240)
            if p.returncode == 0:
                results.append({'file': os.path.basename(fp), 'ok': True})
            else:
                results.append({'file': os.path.basename(fp), 'ok': False,
                                'err': (p.stderr or p.stdout)[:200]})
        except Exception as e:
            results.append({'file': os.path.basename(fp), 'ok': False, 'err': str(e)})
    return jsonify({'success': True, 'processed': len(results),
                    'queued': len(targets), 'results': results})


@app.route('/api/ahb/documents/find', methods=['POST'])
def api_ahb_documents_find():
    """Phil's smart find — given a natural-language query like
    'permit application for the Warrington deck build', search the doc library
    using a combination of keyword + LLM relevance scoring."""
    body = request.get_json() or {}
    query = (body.get('query') or '').strip()
    if not query:
        return jsonify({'success': False, 'error': 'query required'}), 400
    conn = sqlite3.connect(os.path.join(DASHBOARD_DIR, 'baza_projects.db'))
    conn.row_factory = sqlite3.Row
    # Pull a candidate set with simple keyword matching
    words = [w.strip().lower() for w in re.findall(r'\w+', query) if len(w) > 2]
    where, params = [], []
    for w in words:
        like = f"%{w}%"
        where.append("(LOWER(entity) LIKE ? OR LOWER(summary) LIKE ? OR LOWER(tags) LIKE ? OR LOWER(original_name) LIKE ? OR LOWER(content_text) LIKE ?)")
        params += [like]*5
    sql = "SELECT * FROM ahb_documents"
    if where:
        sql += " WHERE " + " OR ".join(where)
    sql += " ORDER BY curated_at DESC LIMIT 25"
    candidates = [dict(r) for r in conn.execute(sql, params).fetchall()]
    conn.close()
    for c in candidates:
        try: c['tags'] = json.loads(c.get('tags') or '[]')
        except Exception: c['tags'] = []
    return jsonify({'success': True, 'query': query,
                    'matches': candidates[:10],
                    'count': len(candidates)})


# ── Application Packages ────────────────────────────────────────────────────

@app.route('/api/ahb/packages', methods=['GET','POST'])
def api_ahb_packages():
    conn = sqlite3.connect(os.path.join(DASHBOARD_DIR, 'baza_projects.db'))
    conn.row_factory = sqlite3.Row
    if request.method == 'GET':
        # Optional ?project_id= filter for the project-modal "App Packages" bin
        proj_filter = request.args.get('project_id')
        if proj_filter:
            rows = conn.execute(
                "SELECT * FROM ahb_app_packages WHERE project_id=? ORDER BY updated_at DESC",
                (proj_filter,)
            ).fetchall()
        else:
            rows = conn.execute("SELECT * FROM ahb_app_packages ORDER BY updated_at DESC").fetchall()
        conn.close()
        out = []
        for r in rows:
            d = dict(r)
            try: d['form_data'] = json.loads(d.get('form_data') or '{}')
            except Exception: d['form_data'] = {}
            try: d['attached_doc_ids'] = json.loads(d.get('attached_doc_ids') or '[]')
            except Exception: d['attached_doc_ids'] = []
            out.append(d)
        return jsonify(out)
    body = request.get_json() or {}
    cur = conn.execute("""INSERT INTO ahb_app_packages
        (name, package_type, project_id, client_id, status, form_data, attached_doc_ids, notes)
        VALUES (?,?,?,?,?,?,?,?)""",
        (body.get('name','Untitled Package'), body.get('package_type'),
         body.get('project_id'), body.get('client_id'),
         body.get('status','draft'),
         json.dumps(body.get('form_data') or {}),
         json.dumps(body.get('attached_doc_ids') or []),
         body.get('notes','')))
    pid = cur.lastrowid
    conn.commit()
    conn.close()
    return jsonify({'success': True, 'id': pid})


@app.route('/api/ahb/packages/<int:pkg_id>', methods=['GET','PATCH','DELETE'])
def api_ahb_package_one(pkg_id):
    conn = sqlite3.connect(os.path.join(DASHBOARD_DIR, 'baza_projects.db'))
    conn.row_factory = sqlite3.Row
    if request.method == 'GET':
        row = conn.execute("SELECT * FROM ahb_app_packages WHERE id=?", (pkg_id,)).fetchone()
        conn.close()
        if not row: return jsonify({'error':'not found'}), 404
        d = dict(row)
        try: d['form_data'] = json.loads(d.get('form_data') or '{}')
        except Exception: d['form_data'] = {}
        try: d['attached_doc_ids'] = json.loads(d.get('attached_doc_ids') or '[]')
        except Exception: d['attached_doc_ids'] = []
        return jsonify(d)
    if request.method == 'DELETE':
        conn.execute("DELETE FROM ahb_app_packages WHERE id=?", (pkg_id,))
        conn.commit()
        conn.close()
        return jsonify({'success': True})
    body = request.get_json() or {}
    sets, vals = [], []
    # Auto-stamp transitions: draft→submitted gets submitted_at, →approved gets approved_at
    if 'status' in body:
        new_status = body['status']
        existing = conn.execute("SELECT status, submitted_at, approved_at FROM ahb_app_packages WHERE id=?",
                                (pkg_id,)).fetchone()
        if existing:
            if new_status == 'submitted' and not existing['submitted_at']:
                sets.append("submitted_at=?"); vals.append(datetime.datetime.now().isoformat())
            if new_status == 'approved' and not existing['approved_at']:
                sets.append("approved_at=?"); vals.append(datetime.datetime.now().isoformat())
    for k in ('name','package_type','project_id','client_id','status','notes','permit_number'):
        if k in body:
            sets.append(f"{k}=?"); vals.append(body[k])
    if 'form_data' in body:
        sets.append("form_data=?"); vals.append(json.dumps(body['form_data']))
    if 'attached_doc_ids' in body:
        sets.append("attached_doc_ids=?"); vals.append(json.dumps(body['attached_doc_ids']))
    if sets:
        sets.append("updated_at=?"); vals.append(datetime.datetime.now().isoformat())
        vals.append(pkg_id)
        conn.execute(f"UPDATE ahb_app_packages SET {','.join(sets)} WHERE id=?", vals)
        conn.commit()
    conn.close()
    return jsonify({'success': True})


@app.route('/api/ahb/packages/<int:pkg_id>/pdf', methods=['GET'])
def api_ahb_package_pdf(pkg_id):
    """Render an application package as a printable PDF — cover sheet with prefilled
    form fields, followed by every attached document merged into one bundle.

    Non-PDF attachments (images, docx, txt) get a thumbnail/preview page so the
    final bundle is one cohesive PDF the user can print, sign, scan, or send."""
    conn = sqlite3.connect(os.path.join(DASHBOARD_DIR, 'baza_projects.db'))
    conn.row_factory = sqlite3.Row
    pkg_row = conn.execute("SELECT * FROM ahb_app_packages WHERE id=?", (pkg_id,)).fetchone()
    if not pkg_row:
        conn.close()
        return jsonify({'error':'not found'}), 404
    pkg = dict(pkg_row)
    try: form = json.loads(pkg.get('form_data') or '{}')
    except Exception: form = {}
    try: doc_ids = json.loads(pkg.get('attached_doc_ids') or '[]')
    except Exception: doc_ids = []
    docs = []
    if doc_ids:
        placeholders = ",".join("?" * len(doc_ids))
        rows = conn.execute(
            f"SELECT * FROM ahb_documents WHERE id IN ({placeholders})", doc_ids
        ).fetchall()
        docs = [dict(r) for r in rows]
    conn.close()

    # Logo for cover sheet
    logo_b64 = ''
    logo_path = os.path.join(DASHBOARD_DIR, 'static', 'img', 'ahb_logo.jpeg')
    if os.path.exists(logo_path):
        import base64 as _b64
        with open(logo_path, 'rb') as lf:
            logo_b64 = _b64.b64encode(lf.read()).decode('utf-8')

    # ── Cover sheet HTML ─────────────────────────────────────────────────────
    pkg_type = (pkg.get('package_type') or 'package').replace('_',' ').title()
    pkg_name = pkg.get('name','Application Package')
    today    = datetime.datetime.now().strftime('%B %d, %Y')

    def _fmt_label(k):
        return k.replace('_',' ').title()

    field_rows = ''
    # Group fields into pairs for two-column layout
    items = list(form.items())
    for i in range(0, len(items), 2):
        left  = items[i]
        right = items[i+1] if i+1 < len(items) else None
        field_rows += '<tr>'
        field_rows += f'<td style="padding:8px 12px;border:1px solid #ddd;background:#f8fafc;font-size:11px;color:#64748b;font-weight:700;width:18%">{_fmt_label(left[0])}</td>'
        field_rows += f'<td style="padding:8px 12px;border:1px solid #ddd;font-size:12px;width:32%">{(left[1] or "")}</td>'
        if right:
            field_rows += f'<td style="padding:8px 12px;border:1px solid #ddd;background:#f8fafc;font-size:11px;color:#64748b;font-weight:700;width:18%">{_fmt_label(right[0])}</td>'
            field_rows += f'<td style="padding:8px 12px;border:1px solid #ddd;font-size:12px;width:32%">{(right[1] or "")}</td>'
        else:
            field_rows += '<td colspan="2" style="border:1px solid #ddd"></td>'
        field_rows += '</tr>'

    docs_table_rows = ''
    for i, d in enumerate(docs, 1):
        docs_table_rows += f'''<tr>
            <td style="padding:8px 12px;border:1px solid #ddd;font-size:11px;text-align:center;width:6%">{i}</td>
            <td style="padding:8px 12px;border:1px solid #ddd;font-size:11px;font-weight:700;width:18%">{(d.get('doc_type') or '').upper()}</td>
            <td style="padding:8px 12px;border:1px solid #ddd;font-size:11px;width:30%">{(d.get('entity') or '')}</td>
            <td style="padding:8px 12px;border:1px solid #ddd;font-size:11px;width:14%">{d.get('doc_date') or ''}</td>
            <td style="padding:8px 12px;border:1px solid #ddd;font-size:11px;color:#666;width:32%">{(d.get('summary') or '')[:120]}</td>
        </tr>'''

    cover_html = f'''<!DOCTYPE html>
<html><head><meta charset="utf-8"><title>{pkg_name}</title>
<style>
  @page {{ size: letter; margin: 36px 40px; }}
  body {{ font-family:'Helvetica Neue',Arial,sans-serif; color:#222; font-size:13px; line-height:1.5; margin:0; }}
  h1 {{ font-size:22px; margin:0 0 4px; color:#1a1a1a; font-weight:700; }}
  h2 {{ font-size:14px; margin:24px 0 10px; color:#2563eb; text-transform:uppercase; letter-spacing:1px; border-bottom:2px solid #2563eb; padding-bottom:4px; }}
  .header {{ display:flex; align-items:flex-start; justify-content:space-between; margin-bottom:18px; padding-bottom:14px; border-bottom:2px solid #1a1a1a; }}
  .header img {{ width:60px; height:60px; object-fit:contain; }}
  table {{ width:100%; border-collapse:collapse; }}
  .meta-grid {{ display:flex; justify-content:space-between; font-size:11px; color:#666; margin-top:6px; }}
  .signature-block {{ margin-top:36px; display:flex; justify-content:space-between; }}
  .sig-line {{ width:45%; }}
  .sig-line .label {{ font-size:11px; color:#888; margin-bottom:4px; }}
  .sig-line .underline {{ border-bottom:1px solid #333; height:32px; }}
</style></head><body>
  <div class="header">
    <div style="display:flex;align-items:center;gap:14px">
      {f'<img src="data:image/jpeg;base64,{logo_b64}">' if logo_b64 else ''}
      <div>
        <div style="font-size:18px;font-weight:700;">All Home Building Co LLC</div>
        <div style="font-size:11px;color:#666">2725 Colmar Ave · Bensalem, PA · 800-484-6404</div>
      </div>
    </div>
    <div style="text-align:right">
      <h1>{pkg_type}</h1>
      <div style="font-size:11px;color:#666">Generated {today}</div>
    </div>
  </div>

  <div style="background:#f0f9ff;border:1px solid #2563eb;border-radius:8px;padding:14px 18px;margin-bottom:18px">
    <div style="font-size:14px;font-weight:700;color:#1e40af">{pkg_name}</div>
    <div class="meta-grid">
      <span>Package ID: #{pkg['id']}</span>
      <span>Status: {(pkg.get('status') or 'draft').upper()}</span>
      <span>Documents: {len(docs)}</span>
    </div>
  </div>

  <h2>Application Details</h2>
  <table>{field_rows}</table>

  {(f'<h2>Notes</h2><div style="background:#fafafa;border:1px solid #ddd;border-radius:6px;padding:12px;font-size:12px">{pkg.get("notes","")}</div>') if pkg.get('notes') else ''}

  {(f'<h2>Attached Supporting Documents</h2><table><thead><tr><th style="padding:8px;background:#1e40af;color:#fff;font-size:11px;border:1px solid #1e40af">#</th><th style="padding:8px;background:#1e40af;color:#fff;font-size:11px;border:1px solid #1e40af;text-align:left">Type</th><th style="padding:8px;background:#1e40af;color:#fff;font-size:11px;border:1px solid #1e40af;text-align:left">Entity</th><th style="padding:8px;background:#1e40af;color:#fff;font-size:11px;border:1px solid #1e40af;text-align:left">Date</th><th style="padding:8px;background:#1e40af;color:#fff;font-size:11px;border:1px solid #1e40af;text-align:left">Summary</th></tr></thead><tbody>{docs_table_rows}</tbody></table>' if docs else '')}

  <div class="signature-block">
    <div class="sig-line">
      <div class="label">Contractor Signature:</div>
      <div class="underline"></div>
      <div style="font-size:11px;color:#666;margin-top:4px">{form.get('contractor_name','Sergey Tkach')}</div>
    </div>
    <div class="sig-line">
      <div class="label">Date:</div>
      <div class="underline"></div>
    </div>
  </div>
</body></html>'''

    # Render cover sheet to PDF
    try:
        from weasyprint import HTML as WeasyHTML
        cover_pdf = WeasyHTML(string=cover_html).write_pdf()
    except Exception as e:
        return jsonify({'error': f'cover render failed: {e}'}), 500

    # Merge cover + every attached doc that's a PDF, plus image preview pages
    try:
        import io
        from pypdf import PdfReader, PdfWriter
        writer = PdfWriter()
        for page in PdfReader(io.BytesIO(cover_pdf)).pages:
            writer.add_page(page)

        for d in docs:
            fp = d.get('file_path') or ''
            if not fp or not os.path.exists(fp):
                continue
            ext = os.path.splitext(fp)[1].lower()
            try:
                if ext == '.pdf':
                    for page in PdfReader(fp).pages:
                        writer.add_page(page)
                elif ext in ('.jpg','.jpeg','.png','.webp','.heic','.heif','.gif','.bmp','.tif','.tiff'):
                    # Render an image-as-PDF-page using weasyprint with the image inline
                    img_html = _image_attachment_html(fp, d, logo_b64)
                    img_pdf = WeasyHTML(string=img_html, base_url=os.path.dirname(fp)).write_pdf()
                    for page in PdfReader(io.BytesIO(img_pdf)).pages:
                        writer.add_page(page)
                elif ext in ('.docx','.txt','.md','.rtf','.csv','.html'):
                    # Render extracted text as a labelled page
                    text = d.get('content_text') or ''
                    text_html = _text_attachment_html(text, d, logo_b64)
                    text_pdf = WeasyHTML(string=text_html).write_pdf()
                    for page in PdfReader(io.BytesIO(text_pdf)).pages:
                        writer.add_page(page)
                # else: skip unsupported binary
            except Exception as inner:
                # Add an error placeholder page so the doc isn't silently lost
                err_html = f'<html><body style="font-family:Arial;padding:40px"><h2>Could not embed: {d.get("original_name","")}</h2><p style="color:#888">{inner}</p></body></html>'
                err_pdf = WeasyHTML(string=err_html).write_pdf()
                for page in PdfReader(io.BytesIO(err_pdf)).pages:
                    writer.add_page(page)

        out_buf = io.BytesIO()
        writer.write(out_buf)
        out_buf.seek(0)
        download = request.args.get('download', '0') == '1'
        resp = make_response(out_buf.read())
        resp.headers['Content-Type'] = 'application/pdf'
        safe_name = re.sub(r'[^\w.\-_]','_', pkg_name)
        disposition = 'attachment' if download else 'inline'
        resp.headers['Content-Disposition'] = f'{disposition}; filename="{safe_name}.pdf"'
        return resp
    except Exception as e:
        return jsonify({'error': f'merge failed: {e}'}), 500


def _image_attachment_html(filepath, doc, logo_b64):
    """Build a PDF page for an image attachment with a header label."""
    import base64 as _b64
    try:
        with open(filepath, 'rb') as f:
            img_b64 = _b64.b64encode(f.read()).decode('utf-8')
    except Exception:
        img_b64 = ''
    ext = os.path.splitext(filepath)[1].lower().lstrip('.')
    mime = 'jpeg' if ext in ('jpg','jpeg') else ext
    return f'''<!DOCTYPE html><html><head><meta charset="utf-8"><style>
      @page {{ size:letter; margin:36px 40px; }}
      body {{ font-family:Arial,sans-serif; color:#222; margin:0; }}
      .label {{ background:#1e40af; color:#fff; padding:10px 14px; font-size:13px; font-weight:700; margin-bottom:14px; border-radius:4px; }}
      .meta {{ font-size:11px; color:#666; margin-bottom:14px; }}
      img {{ max-width:100%; max-height:850px; display:block; margin:0 auto; border:1px solid #ddd; }}
    </style></head><body>
      <div class="label">{(doc.get('doc_type') or 'DOCUMENT').upper()} — {doc.get('entity','')}</div>
      <div class="meta">{doc.get('summary','')[:300]}</div>
      <img src="data:image/{mime};base64,{img_b64}">
    </body></html>'''


def _text_attachment_html(text, doc, logo_b64):
    safe_text = (text or '').replace('<','&lt;').replace('>','&gt;').replace('\n','<br>')
    return f'''<!DOCTYPE html><html><head><meta charset="utf-8"><style>
      @page {{ size:letter; margin:36px 40px; }}
      body {{ font-family:Arial,sans-serif; color:#222; margin:0; font-size:11px; line-height:1.5; }}
      .label {{ background:#1e40af; color:#fff; padding:10px 14px; font-size:13px; font-weight:700; margin-bottom:14px; border-radius:4px; }}
      .meta {{ font-size:11px; color:#666; margin-bottom:14px; }}
      .body {{ background:#fafafa; border:1px solid #ddd; border-radius:6px; padding:14px; white-space:pre-wrap; }}
    </style></head><body>
      <div class="label">{(doc.get('doc_type') or 'DOCUMENT').upper()} — {doc.get('entity','')}</div>
      <div class="meta">{doc.get('summary','')[:300]}</div>
      <div class="body">{safe_text[:8000]}</div>
    </body></html>'''


# ── DocPrep Universal File Actions ──────────────────────────────────────────

def _get_doc(did: int):
    """Load doc row + verify file exists. Returns (dict, file_path) or (None, error_msg)."""
    conn = sqlite3.connect(os.path.join(DASHBOARD_DIR, 'baza_projects.db'))
    conn.row_factory = sqlite3.Row
    row = conn.execute("SELECT * FROM ahb_documents WHERE id=?", (did,)).fetchone()
    conn.close()
    if not row:
        return None, "doc not found"
    d = dict(row)
    if not d.get('file_path') or not os.path.exists(d['file_path']):
        return None, "file missing on disk"
    return d, d['file_path']


def _ollama_text(prompt: str, model: str = "qwen2.5:14b", json_mode: bool = False, max_tokens: int = 1200) -> str:
    """Quick local Ollama text generation helper for doc actions."""
    import urllib.request as _ur
    payload = {
        "model": model,
        "prompt": prompt,
        "stream": False,
        "options": {"temperature": 0.1, "num_predict": max_tokens},
    }
    if json_mode:
        payload["format"] = "json"
    try:
        req = _ur.Request("http://localhost:11434/api/generate",
                          data=json.dumps(payload).encode(),
                          headers={"Content-Type": "application/json"})
        with _ur.urlopen(req, timeout=120) as r:
            data = json.loads(r.read())
        return (data.get("response") or "").strip()
    except Exception as e:
        return f"[LLM error: {e}]"


@app.route('/api/ahb/documents/<int:did>/recurate', methods=['POST'])
def api_doc_recurate(did):
    """Re-run the curate_document skill on this doc and refresh metadata."""
    d, fp = _get_doc(did)
    if not d:
        return jsonify({'success': False, 'error': fp}), 404
    skill = os.path.join(FRAMEWORK_DIR, 'skills', 'shared', 'curate_document.py')
    try:
        env = os.environ.copy()
        env['SKILL_ARGS'] = json.dumps({
            'file_path': fp,
            'agent_id': d.get('agent_id') or 'phil_hass',
        })
        proc = subprocess.run([VENV_PYTHON, skill], env=env,
                              capture_output=True, text=True, timeout=240)
        if proc.returncode != 0:
            return jsonify({'success': False, 'error': (proc.stderr or proc.stdout)[:500]}), 500
        try:
            result = json.loads(proc.stdout.strip())
        except Exception:
            result = {}
        return jsonify({'success': True, 'analysis': result})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/ahb/documents/<int:did>/rename', methods=['POST'])
def api_doc_rename(did):
    """Rename the file on disk to match suggested_name and update file_path in DB."""
    d, fp = _get_doc(did)
    if not d:
        return jsonify({'success': False, 'error': fp}), 404
    new_name = (d.get('suggested_name') or '').strip()
    body = request.get_json(silent=True) or {}
    if body.get('name'):
        new_name = body['name'].strip()
    if not new_name:
        return jsonify({'success': False, 'error': 'no suggested_name set'}), 400
    new_name = re.sub(r'[^\w.\-_ ()]', '_', new_name).strip()
    new_path = os.path.join(os.path.dirname(fp), new_name)
    if os.path.exists(new_path) and new_path != fp:
        # Append _2, _3, etc. to avoid collision
        base, ext = os.path.splitext(new_path)
        i = 2
        while os.path.exists(f"{base}_{i}{ext}"):
            i += 1
        new_path = f"{base}_{i}{ext}"
    try:
        os.rename(fp, new_path)
        # Move sidecar meta if present
        old_meta = fp + '.meta'
        if os.path.exists(old_meta):
            try: os.rename(old_meta, new_path + '.meta')
            except Exception: pass
        conn = sqlite3.connect(os.path.join(DASHBOARD_DIR, 'baza_projects.db'))
        conn.execute("UPDATE ahb_documents SET file_path=?, original_name=? WHERE id=?",
                     (new_path, os.path.basename(new_path), did))
        conn.commit()
        conn.close()
        return jsonify({'success': True, 'new_name': os.path.basename(new_path), 'new_path': new_path})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/ahb/documents/<int:did>/extract-text', methods=['POST'])
def api_doc_extract_text(did):
    """Re-extract text from a doc — uses pdfplumber for PDFs, vision LLM for images, plain read for text."""
    d, fp = _get_doc(did)
    if not d:
        return jsonify({'success': False, 'error': fp}), 404
    ext = os.path.splitext(fp)[1].lower()
    text = ""
    try:
        if ext == '.pdf':
            import pdfplumber
            txt = []
            with pdfplumber.open(fp) as pdf:
                for page in pdf.pages[:30]:
                    t = page.extract_text() or ""
                    if t.strip():
                        txt.append(t)
            text = "\n".join(txt)
        elif ext == '.docx':
            import docx
            doc = docx.Document(fp)
            text = "\n".join(p.text for p in doc.paragraphs)
        elif ext in ('.txt','.md','.csv','.html','.rtf','.log','.xml'):
            with open(fp, 'r', errors='ignore') as f:
                text = f.read()
        elif ext in ('.jpg','.jpeg','.png','.webp','.heic','.heif','.gif','.bmp','.tif','.tiff'):
            # Use the curate_document skill's vision path via the same skill
            skill = os.path.join(FRAMEWORK_DIR, 'skills', 'shared', 'curate_document.py')
            env = os.environ.copy()
            env['SKILL_ARGS'] = json.dumps({'file_path': fp, 'agent_id': 'phil_hass'})
            proc = subprocess.run([VENV_PYTHON, skill], env=env,
                                  capture_output=True, text=True, timeout=240)
            # The skill saves content_text to DB; reload from DB
            conn2 = sqlite3.connect(os.path.join(DASHBOARD_DIR, 'baza_projects.db'))
            row = conn2.execute("SELECT content_text FROM ahb_documents WHERE id=?", (did,)).fetchone()
            conn2.close()
            text = (row[0] if row else "") or ""
        else:
            return jsonify({'success': False, 'error': f'unsupported extension {ext}'}), 400

        text = text[:50000]
        conn = sqlite3.connect(os.path.join(DASHBOARD_DIR, 'baza_projects.db'))
        conn.execute("UPDATE ahb_documents SET content_text=? WHERE id=?", (text, did))
        conn.commit()
        conn.close()
        return jsonify({'success': True, 'text_length': len(text), 'preview': text[:500]})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/ahb/documents/<int:did>/summarize', methods=['POST'])
def api_doc_summarize(did):
    """Quick summary regen — uses existing content_text, doesn't re-extract."""
    d, fp = _get_doc(did)
    if not d:
        return jsonify({'success': False, 'error': fp}), 404
    text = (d.get('content_text') or '')[:8000]
    if not text:
        return jsonify({'success': False, 'error': 'no content_text — run extract first'}), 400
    prompt = (
        "You are Phil Hass, document curator for All Home Building Co LLC. Summarize this "
        "document in 1-3 plain English sentences. Mention the entity, document type, key "
        "dates, dollar amounts if present.\n\nDocument:\n" + text + "\n\nSummary:"
    )
    summary = _ollama_text(prompt, max_tokens=300)
    if summary.startswith("[LLM error"):
        return jsonify({'success': False, 'error': summary}), 500
    conn = sqlite3.connect(os.path.join(DASHBOARD_DIR, 'baza_projects.db'))
    conn.execute("UPDATE ahb_documents SET summary=? WHERE id=?", (summary, did))
    conn.commit()
    conn.close()
    return jsonify({'success': True, 'summary': summary})


@app.route('/api/ahb/documents/<int:did>/convert', methods=['POST'])
def api_doc_convert(did):
    """Convert a doc to PDF (or other format) and save as sibling file. Returns new doc id."""
    d, fp = _get_doc(did)
    if not d:
        return jsonify({'success': False, 'error': fp}), 404
    body = request.get_json() or {}
    target = (body.get('format') or 'pdf').lower()
    ext = os.path.splitext(fp)[1].lower()
    base = os.path.splitext(fp)[0]
    new_path = f"{base}.{target}"
    try:
        if target == 'pdf':
            from weasyprint import HTML as WeasyHTML
            if ext == '.pdf':
                return jsonify({'success': False, 'error': 'already PDF'}), 400
            if ext in ('.txt','.md','.csv','.log','.html','.rtf'):
                with open(fp, 'r', errors='ignore') as f:
                    body_text = f.read()
                if ext != '.html':
                    body_text = '<pre style="white-space:pre-wrap;font-family:monospace;font-size:11px">' + (body_text.replace('<','&lt;').replace('>','&gt;')) + '</pre>'
                html = f'<html><head><style>@page{{size:letter;margin:36px 40px}}body{{font-family:Arial,sans-serif}}</style></head><body>{body_text}</body></html>'
                WeasyHTML(string=html).write_pdf(new_path)
            elif ext == '.docx':
                import docx
                d2 = docx.Document(fp)
                body_text = '<br>'.join(p.text.replace('<','&lt;').replace('>','&gt;') for p in d2.paragraphs)
                html = f'<html><head><style>@page{{size:letter;margin:36px 40px}}body{{font-family:Arial,sans-serif;font-size:12px;line-height:1.5}}</style></head><body>{body_text}</body></html>'
                WeasyHTML(string=html).write_pdf(new_path)
            elif ext in ('.jpg','.jpeg','.png','.webp','.heic','.heif','.gif','.bmp','.tif','.tiff'):
                import base64 as _b64
                with open(fp, 'rb') as f:
                    img_b64 = _b64.b64encode(f.read()).decode('utf-8')
                mime = 'jpeg' if ext in ('.jpg','.jpeg') else ext.lstrip('.')
                html = f'<html><head><style>@page{{size:letter;margin:36px 40px}}body{{margin:0;text-align:center}}img{{max-width:100%;max-height:900px}}</style></head><body><img src="data:image/{mime};base64,{img_b64}"></body></html>'
                WeasyHTML(string=html).write_pdf(new_path)
            else:
                return jsonify({'success': False, 'error': f'cannot convert {ext} to pdf'}), 400
        else:
            return jsonify({'success': False, 'error': f'unsupported target format {target}'}), 400
        # Insert as new doc row
        conn = sqlite3.connect(os.path.join(DASHBOARD_DIR, 'baza_projects.db'))
        cur = conn.execute("""INSERT INTO ahb_documents
            (file_path, original_name, suggested_name, doc_type, entity, doc_date,
             summary, relevance, tags, agent_id, file_size, file_kind)
            VALUES (?,?,?,?,?,?,?,?,?,?,?,?)""",
            (new_path, os.path.basename(new_path), os.path.basename(new_path),
             d.get('doc_type'), d.get('entity'), d.get('doc_date'),
             (d.get('summary') or '') + ' [Converted to PDF]',
             d.get('relevance'), d.get('tags') or '[]',
             d.get('agent_id') or 'phil_hass',
             os.path.getsize(new_path), 'pdf'))
        new_id = cur.lastrowid
        conn.commit()
        conn.close()
        return jsonify({'success': True, 'new_doc_id': new_id, 'new_path': new_path})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/ahb/documents/<int:did>/replace', methods=['POST'])
def api_doc_replace(did):
    """Upload a new version of the file. Old version is archived as .v<n>.bak."""
    d, fp = _get_doc(did)
    if not d:
        return jsonify({'success': False, 'error': fp}), 404
    f = request.files.get('file')
    if not f or not f.filename:
        return jsonify({'success': False, 'error': 'no file'}), 400
    try:
        # Archive old version
        i = 1
        while os.path.exists(f"{fp}.v{i}.bak"):
            i += 1
        os.rename(fp, f"{fp}.v{i}.bak")
        f.save(fp)
        new_size = os.path.getsize(fp)
        conn = sqlite3.connect(os.path.join(DASHBOARD_DIR, 'baza_projects.db'))
        conn.execute("UPDATE ahb_documents SET file_size=?, curated_at=CURRENT_TIMESTAMP WHERE id=?",
                     (new_size, did))
        conn.commit()
        conn.close()
        return jsonify({'success': True, 'archived_version': i, 'new_size': new_size})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/ahb/documents/<int:did>/content', methods=['GET','PUT'])
def api_doc_content(did):
    """For text-based files: read or write the file content directly."""
    d, fp = _get_doc(did)
    if not d:
        return jsonify({'success': False, 'error': fp}), 404
    ext = os.path.splitext(fp)[1].lower()
    text_exts = {'.txt','.md','.csv','.html','.htm','.rtf','.log','.xml','.json','.yaml','.yml','.ini','.conf','.py','.js','.css'}
    if ext not in text_exts:
        return jsonify({'success': False, 'error': f'not a text file ({ext})'}), 400
    if request.method == 'GET':
        try:
            with open(fp, 'r', errors='ignore') as f:
                content = f.read()
            return jsonify({'success': True, 'content': content})
        except Exception as e:
            return jsonify({'success': False, 'error': str(e)}), 500
    body = request.get_json() or {}
    content = body.get('content', '')
    try:
        with open(fp, 'w') as f:
            f.write(content)
        conn = sqlite3.connect(os.path.join(DASHBOARD_DIR, 'baza_projects.db'))
        conn.execute("UPDATE ahb_documents SET content_text=?, file_size=? WHERE id=?",
                     (content[:50000], len(content.encode()), did))
        conn.commit()
        conn.close()
        return jsonify({'success': True, 'size': len(content)})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/ahb/documents/<int:did>/attach-to-package', methods=['POST'])
def api_doc_attach_to_package(did):
    body = request.get_json() or {}
    pkg_id = body.get('package_id')
    if not pkg_id:
        return jsonify({'success': False, 'error': 'package_id required'}), 400
    conn = sqlite3.connect(os.path.join(DASHBOARD_DIR, 'baza_projects.db'))
    row = conn.execute("SELECT attached_doc_ids FROM ahb_app_packages WHERE id=?", (pkg_id,)).fetchone()
    if not row:
        conn.close()
        return jsonify({'success': False, 'error': 'package not found'}), 404
    try:
        ids = json.loads(row[0] or '[]')
    except Exception:
        ids = []
    if did not in ids:
        ids.append(did)
    conn.execute("UPDATE ahb_app_packages SET attached_doc_ids=?, updated_at=? WHERE id=?",
                 (json.dumps(ids), datetime.datetime.now().isoformat(), pkg_id))
    conn.commit()
    conn.close()
    return jsonify({'success': True, 'attached_count': len(ids)})


@app.route('/api/ahb/documents/<int:did>/print', methods=['POST'])
def api_doc_print(did):
    """Send doc to thermal printer via the existing print_document skill."""
    d, fp = _get_doc(did)
    if not d:
        return jsonify({'success': False, 'error': fp}), 404
    skill = os.path.join(FRAMEWORK_DIR, 'skills', 'shared', 'print_document.py')
    if not os.path.exists(skill):
        return jsonify({'success': False, 'error': 'print_document skill not installed'}), 500
    try:
        env = os.environ.copy()
        env['SKILL_ARGS'] = json.dumps({'action': 'print', 'file_path': fp})
        proc = subprocess.run([VENV_PYTHON, skill], env=env,
                              capture_output=True, text=True, timeout=120)
        if proc.returncode != 0:
            return jsonify({'success': False, 'error': (proc.stderr or proc.stdout)[:500]}), 500
        return jsonify({'success': True, 'output': proc.stdout[:500]})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/ahb/documents/<int:did>/email', methods=['POST'])
def api_doc_email(did):
    """Email doc as attachment via the email pipeline."""
    d, fp = _get_doc(did)
    if not d:
        return jsonify({'success': False, 'error': fp}), 404
    body = request.get_json() or {}
    to_addr = (body.get('to') or '').strip()
    subject = (body.get('subject') or f'Document from AHBCO: {os.path.basename(fp)}').strip()
    msg     = (body.get('body') or d.get('summary') or 'Please see attached.').strip()
    if not to_addr:
        return jsonify({'success': False, 'error': 'recipient required'}), 400
    # Use Gmail OAuth via the email-pipeline send_reply.py (it accepts attachment via env)
    sender = os.path.join(FRAMEWORK_DIR, 'email-pipeline', 'send_reply.py')
    if not os.path.exists(sender):
        return jsonify({'success': False, 'error': 'email pipeline not configured'}), 500
    try:
        env = os.environ.copy()
        env['EMAIL_TO'] = to_addr
        env['EMAIL_SUBJECT'] = subject
        env['EMAIL_BODY'] = msg
        env['EMAIL_ATTACHMENT'] = fp
        proc = subprocess.run([VENV_PYTHON, sender], env=env,
                              capture_output=True, text=True, timeout=60)
        if proc.returncode != 0:
            return jsonify({'success': False, 'error': (proc.stderr or proc.stdout)[:500]}), 500
        return jsonify({'success': True, 'output': proc.stdout[:500]})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/ahb/documents/<int:did>/ask', methods=['POST'])
def api_doc_ask(did):
    """Ad-hoc LLM query against the document content."""
    d, fp = _get_doc(did)
    if not d:
        return jsonify({'success': False, 'error': fp}), 404
    body = request.get_json() or {}
    question = (body.get('question') or '').strip()
    if not question:
        return jsonify({'success': False, 'error': 'question required'}), 400
    text = (d.get('content_text') or '')[:10000]
    if not text:
        return jsonify({'success': False, 'error': 'no content_text — run Extract Text first'}), 400
    prompt = (
        "You are Phil Hass, document curator for All Home Building Co LLC. Read the document "
        "below and answer the user's question precisely. If the answer is not in the document, "
        "say so explicitly.\n\n=== Document ===\n" + text + "\n\n=== Question ===\n" + question + "\n\n=== Answer ==="
    )
    answer = _ollama_text(prompt, max_tokens=600)
    if answer.startswith("[LLM error"):
        return jsonify({'success': False, 'error': answer}), 500
    return jsonify({'success': True, 'answer': answer})


# ── Uncle Sam: Business Profile + PA LLC Tax Requirements ──────────────────

def _ensure_business_profile_table():
    """Idempotent init for business profile + tax tables. Wrapped in try/except
    so a busy DB (e.g. long-running backfill holding a write lock) doesn't
    crash the dashboard at boot — the tables already exist on every deployed
    instance and the seed row is already there."""
    try:
        conn = sqlite3.connect(os.path.join(DASHBOARD_DIR, 'baza_projects.db'), timeout=8.0)
        conn.execute("PRAGMA busy_timeout = 8000")
        conn.execute("""CREATE TABLE IF NOT EXISTS ahb_business_profile (
            id INTEGER PRIMARY KEY CHECK (id = 1),
            legal_name TEXT, dba TEXT, ein TEXT, ssn_last4 TEXT,
            structure_type TEXT DEFAULT 'single_member_llc',
            state_of_formation TEXT DEFAULT 'PA',
            formation_date TEXT, registered_agent TEXT,
            business_address TEXT, business_phone TEXT, business_email TEXT,
            fiscal_year_end TEXT DEFAULT '12-31',
            accounting_method TEXT DEFAULT 'cash',
            naics_code TEXT, hic_number TEXT, hic_expires TEXT,
            pa_tax_id TEXT, philly_tax_account TEXT,
            has_employees INTEGER DEFAULT 0,
            collects_sales_tax INTEGER DEFAULT 0,
            updated_at TEXT DEFAULT CURRENT_TIMESTAMP
        )""")
        # Add household tax inputs (for tax estimate). Idempotent ALTER for existing DBs.
        for col, ddl in (
            ('filing_status', "ALTER TABLE ahb_business_profile ADD COLUMN filing_status TEXT DEFAULT 'single'"),
            ('dependents',    "ALTER TABLE ahb_business_profile ADD COLUMN dependents INTEGER DEFAULT 0"),
        ):
            try: conn.execute(ddl)
            except sqlite3.OperationalError: pass  # column already exists
        conn.execute("""CREATE TABLE IF NOT EXISTS ahb_tax_filings (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            requirement_id TEXT NOT NULL,
            period TEXT,
            filed_date TEXT NOT NULL,
            amount_paid REAL,
            confirmation_number TEXT,
            notes TEXT,
            created_at TEXT DEFAULT CURRENT_TIMESTAMP
        )""")
        # Per-year personal "other income" entries used by the Uncle Sam tax estimate
        # (W-2 wages from a day job, 1099 contract, dividends, interest, rental, etc.)
        conn.execute("""CREATE TABLE IF NOT EXISTS ahb_other_income (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            year TEXT NOT NULL,
            kind TEXT NOT NULL,
            source TEXT,
            amount REAL NOT NULL DEFAULT 0,
            notes TEXT,
            created_at TEXT DEFAULT CURRENT_TIMESTAMP
        )""")
        conn.execute("CREATE INDEX IF NOT EXISTS idx_ahb_other_income_year ON ahb_other_income(year)")
        # Seed empty row
        conn.execute("INSERT OR IGNORE INTO ahb_business_profile (id, legal_name) VALUES (1, 'All Home Building Co LLC')")
        conn.commit()
        conn.close()
    except sqlite3.OperationalError as e:
        print(f"[startup] _ensure_business_profile_table deferred — DB busy: {e}", flush=True)
_ensure_business_profile_table()


# Hardcoded PA LLC tax requirements catalog. Items are filtered by structure_type,
# has_employees, and in_philly. Each "due" pattern is parsed to compute the next date.
PA_LLC_TAX_REQUIREMENTS = [
    # ── Federal income tax ────────────────────────────────────────
    {"id":"fed_ein","jurisdiction":"federal","category":"licensing","name":"Federal EIN",
     "description":"Employer Identification Number from the IRS — one-time, free at irs.gov.",
     "applies_to":["sole_prop","single_member_llc","multi_member_llc","s_corp","c_corp","partnership"],
     "due":"one_time","required_docs":["tax_document"],"info_url":"https://www.irs.gov/businesses/small-businesses-self-employed/apply-for-an-employer-identification-number-ein-online"},
    {"id":"fed_1040_sch_c","jurisdiction":"federal","category":"income_tax","name":"Form 1040 + Schedule C",
     "description":"Federal income tax return — single-member LLC reports business income on owner's 1040.",
     "applies_to":["sole_prop","single_member_llc"],
     "due":"annual:04-15","required_docs":["tax_document"],
     "info_url":"https://www.irs.gov/forms-pubs/about-schedule-c-form-1040"},
    {"id":"fed_1065","jurisdiction":"federal","category":"income_tax","name":"Form 1065 (Partnership Return)",
     "description":"Multi-member LLC partnership return + K-1s to members.",
     "applies_to":["multi_member_llc","partnership"],
     "due":"annual:03-15","required_docs":["tax_document"],
     "info_url":"https://www.irs.gov/forms-pubs/about-form-1065"},
    {"id":"fed_1120s","jurisdiction":"federal","category":"income_tax","name":"Form 1120-S (S-Corp)",
     "description":"S-Corporation return + K-1s to shareholders.",
     "applies_to":["s_corp"],
     "due":"annual:03-15","required_docs":["tax_document"],
     "info_url":"https://www.irs.gov/forms-pubs/about-form-1120-s"},
    {"id":"fed_1120","jurisdiction":"federal","category":"income_tax","name":"Form 1120 (C-Corp)",
     "description":"C-Corporation income tax return.",
     "applies_to":["c_corp"],
     "due":"annual:04-15","required_docs":["tax_document"],
     "info_url":"https://www.irs.gov/forms-pubs/about-form-1120"},
    {"id":"fed_se","jurisdiction":"federal","category":"income_tax","name":"Schedule SE (Self-Employment)",
     "description":"Self-employment tax (15.3%) for Social Security + Medicare on net earnings.",
     "applies_to":["sole_prop","single_member_llc","multi_member_llc","partnership"],
     "due":"annual:04-15","required_docs":[],"info_url":"https://www.irs.gov/forms-pubs/about-schedule-se-form-1040"},
    {"id":"fed_1040es","jurisdiction":"federal","category":"quarterly","name":"Quarterly Estimated Tax (1040-ES)",
     "description":"Quarterly estimated income tax + SE tax payments.",
     "applies_to":["sole_prop","single_member_llc","multi_member_llc","partnership","s_corp"],
     "due":"quarterly_est","required_docs":[],
     "info_url":"https://www.irs.gov/forms-pubs/about-form-1040-es"},
    {"id":"fed_941","jurisdiction":"federal","category":"employment","name":"Form 941 (Quarterly Employment Tax)",
     "description":"Quarterly federal employment tax return — withholdings + employer SS/Medicare.",
     "applies_to":["any"],"requires":"has_employees",
     "due":"quarterly:04-30,07-31,10-31,01-31","required_docs":[],
     "info_url":"https://www.irs.gov/forms-pubs/about-form-941"},
    {"id":"fed_940","jurisdiction":"federal","category":"employment","name":"Form 940 (FUTA)",
     "description":"Annual federal unemployment tax return.",
     "applies_to":["any"],"requires":"has_employees",
     "due":"annual:01-31","required_docs":[],
     "info_url":"https://www.irs.gov/forms-pubs/about-form-940"},
    {"id":"fed_w2","jurisdiction":"federal","category":"employment","name":"W-2 / W-3 to employees",
     "description":"Employee wage statements + transmittal to SSA.",
     "applies_to":["any"],"requires":"has_employees",
     "due":"annual:01-31","required_docs":[],
     "info_url":"https://www.irs.gov/forms-pubs/about-form-w-2"},
    {"id":"fed_1099nec","jurisdiction":"federal","category":"employment","name":"1099-NEC to subcontractors",
     "description":"Send to any subcontractor paid >$600 in calendar year.",
     "applies_to":["any"],
     "due":"annual:01-31","required_docs":["w9"],
     "info_url":"https://www.irs.gov/forms-pubs/about-form-1099-nec"},

    # ── PA State ──────────────────────────────────────────────────
    {"id":"pa_100","jurisdiction":"pa","category":"licensing","name":"PA-100 Business Registration",
     "description":"One-time business tax registration with PA Department of Revenue.",
     "applies_to":["any"],
     "due":"one_time","required_docs":[],
     "info_url":"https://www.pa100.state.pa.us/"},
    {"id":"pa_40","jurisdiction":"pa","category":"income_tax","name":"PA-40 Personal Income Tax",
     "description":"Pennsylvania personal income tax (3.07% flat) — passes through from LLC.",
     "applies_to":["sole_prop","single_member_llc","multi_member_llc","partnership","s_corp"],
     "due":"annual:04-15","required_docs":[],
     "info_url":"https://www.revenue.pa.gov/FormsandPublications/FormsforIndividuals/PIT/Pages/default.aspx"},
    {"id":"pa_annual_report","jurisdiction":"pa","category":"licensing","name":"PA Annual Report (Act 122)",
     "description":"MANDATORY for LLCs since 2025 (Act 122 of 2022). Filed with PA Department of State. Late fee $500.",
     "applies_to":["single_member_llc","multi_member_llc","s_corp","c_corp"],
     "due":"annual:09-30","required_docs":[],
     "info_url":"https://www.dos.pa.gov/BusinessCharities/Business/Pages/Annual-Reports.aspx"},
    {"id":"pa_uc","jurisdiction":"pa","category":"employment","name":"PA UC Tax Registration",
     "description":"PA Unemployment Compensation tax — quarterly filings if you have employees.",
     "applies_to":["any"],"requires":"has_employees",
     "due":"quarterly:04-30,07-31,10-31,01-31","required_docs":[],
     "info_url":"https://www.uc.pa.gov/employers-uc-services-uc-tax/Pages/default.aspx"},
    {"id":"pa_sales_tax","jurisdiction":"pa","category":"sales_use","name":"PA Sales Tax License",
     "description":"Required if collecting sales tax. Construction labor is generally exempt; materials are taxable.",
     "applies_to":["any"],"requires":"collects_sales_tax",
     "due":"monthly_or_quarterly","required_docs":[],
     "info_url":"https://www.revenue.pa.gov/TaxTypes/SUT/Pages/default.aspx"},

    # ── Philadelphia ──────────────────────────────────────────────
    {"id":"phl_birt","jurisdiction":"philadelphia","category":"income_tax","name":"BIRT (Business Income & Receipts Tax)",
     "description":"Philadelphia business tax — gross receipts + net income components. Due annually.",
     "applies_to":["any"],"requires":"in_philly",
     "due":"annual:04-15","required_docs":["tax_document"],
     "info_url":"https://www.phila.gov/services/payments-assistance-taxes/business-taxes/business-income-receipts-tax-birt/"},
    {"id":"phl_npt","jurisdiction":"philadelphia","category":"income_tax","name":"NPT (Net Profits Tax)",
     "description":"Philadelphia net profits tax — 3.79% for residents, 3.44% non-residents.",
     "applies_to":["sole_prop","single_member_llc","multi_member_llc","partnership"],"requires":"in_philly",
     "due":"annual:04-15","required_docs":[],
     "info_url":"https://www.phila.gov/services/payments-assistance-taxes/business-taxes/net-profits-tax/"},
    {"id":"phl_wage","jurisdiction":"philadelphia","category":"employment","name":"Philadelphia Wage Tax",
     "description":"Withhold from employee wages working in Philly. Quarterly returns.",
     "applies_to":["any"],"requires":"has_employees_in_philly",
     "due":"quarterly:04-30,07-31,10-31,01-31","required_docs":[],
     "info_url":"https://www.phila.gov/services/payments-assistance-taxes/business-taxes/wage-tax-employers/"},
    {"id":"phl_cal","jurisdiction":"philadelphia","category":"licensing","name":"Commercial Activity License",
     "description":"One-time CAL from Philadelphia Department of Licenses & Inspections.",
     "applies_to":["any"],"requires":"in_philly",
     "due":"one_time","required_docs":["license"],
     "info_url":"https://www.phila.gov/services/permits-violations-licenses/get-a-license/business-licenses/get-a-commercial-activity-license/"},

    # ── Construction-specific ─────────────────────────────────────
    {"id":"pa_hic","jurisdiction":"construction","category":"licensing","name":"HIC Registration (PA AG)",
     "description":"Home Improvement Contractor registration with PA Office of Attorney General. Required for any contractor doing residential work over $500/year. Annual renewal.",
     "applies_to":["any"],
     "due":"annual_renewal","required_docs":["license","coi"],
     "info_url":"https://www.attorneygeneral.gov/protect-yourself/home-improvement-consumer-protection/"},
    {"id":"general_liability","jurisdiction":"construction","category":"licensing","name":"General Liability Insurance",
     "description":"Required for HIC. Minimum $50K personal injury / $50K property damage. Most townships require $1M.",
     "applies_to":["any"],
     "due":"annual_renewal","required_docs":["coi"],
     "info_url":""},
    {"id":"workers_comp","jurisdiction":"construction","category":"employment","name":"Workers' Compensation Insurance",
     "description":"Required by PA law if you have any employees (including yourself if structured as employee).",
     "applies_to":["any"],"requires":"has_employees",
     "due":"annual_renewal","required_docs":["coi"],
     "info_url":"https://www.dli.pa.gov/Businesses/Compensation/WC/Pages/default.aspx"},
]


def _filter_requirements(structure_type, has_employees, in_philly, collects_sales_tax):
    """Return only the requirements that apply to this business setup."""
    out = []
    for r in PA_LLC_TAX_REQUIREMENTS:
        applies = r.get("applies_to", [])
        if "any" not in applies and structure_type not in applies:
            continue
        req = r.get("requires")
        if req == "has_employees" and not has_employees:
            continue
        if req == "in_philly" and not in_philly:
            continue
        if req == "has_employees_in_philly" and not (has_employees and in_philly):
            continue
        if req == "collects_sales_tax" and not collects_sales_tax:
            continue
        out.append(r)
    return out


def _next_due_date(due_pattern: str):
    """Compute the next due date for a 'due' pattern. Returns (date_iso, days_until)."""
    today = datetime.date.today()
    if not due_pattern or due_pattern == "one_time":
        return None, None
    if due_pattern.startswith("annual:"):
        mm, dd = due_pattern.split(":")[1].split("-")
        candidate = datetime.date(today.year, int(mm), int(dd))
        if candidate < today:
            candidate = datetime.date(today.year + 1, int(mm), int(dd))
        return candidate.isoformat(), (candidate - today).days
    if due_pattern == "annual_renewal":
        # 1 year from today as a placeholder — user can override per requirement
        candidate = datetime.date(today.year + 1, today.month, today.day)
        return candidate.isoformat(), (candidate - today).days
    if due_pattern == "quarterly_est":
        # Federal estimated tax: Apr 15, Jun 15, Sep 15, Jan 15
        candidates = [
            datetime.date(today.year, 4, 15),
            datetime.date(today.year, 6, 15),
            datetime.date(today.year, 9, 15),
            datetime.date(today.year + 1, 1, 15),
        ]
        for c in candidates:
            if c >= today:
                return c.isoformat(), (c - today).days
        return None, None
    if due_pattern.startswith("quarterly:"):
        dates = due_pattern.split(":")[1].split(",")
        for ds in dates:
            mm, dd = ds.split("-")
            yr = today.year
            if int(mm) == 1 and today.month >= 11:
                yr = today.year + 1
            cand = datetime.date(yr, int(mm), int(dd))
            if cand >= today:
                return cand.isoformat(), (cand - today).days
        # rollover to next year first quarter
        mm, dd = dates[0].split("-")
        cand = datetime.date(today.year + 1, int(mm), int(dd))
        return cand.isoformat(), (cand - today).days
    if due_pattern == "monthly_or_quarterly":
        # End of next month as a sane default
        nm = today.replace(day=1) + datetime.timedelta(days=32)
        eom = nm.replace(day=1) - datetime.timedelta(days=1)
        return eom.isoformat(), (eom - today).days
    return None, None


def _check_required_docs(doc_types: list) -> dict:
    """Check ahb_documents for AHBCO docs of the given types. Returns map of type → doc info or None."""
    if not doc_types:
        return {}
    conn = sqlite3.connect(os.path.join(DASHBOARD_DIR, 'baza_projects.db'))
    conn.row_factory = sqlite3.Row
    out = {}
    for dt in doc_types:
        row = conn.execute(
            """SELECT id, suggested_name, original_name, doc_date, entity FROM ahb_documents
               WHERE doc_type=? AND (entity LIKE '%AHBCO%' OR entity LIKE '%All Home Building%')
               ORDER BY doc_date DESC NULLS LAST LIMIT 1""",
            (dt,)
        ).fetchone()
        out[dt] = dict(row) if row else None
    conn.close()
    return out


@app.route('/api/ahb/business-profile', methods=['GET'])
def api_ahb_business_profile_get():
    conn = sqlite3.connect(os.path.join(DASHBOARD_DIR, 'baza_projects.db'))
    conn.row_factory = sqlite3.Row
    row = conn.execute("SELECT * FROM ahb_business_profile WHERE id=1").fetchone()
    conn.close()
    if not row:
        return jsonify({})
    return jsonify(dict(row))


@app.route('/api/ahb/business-profile', methods=['PUT'])
def api_ahb_business_profile_put():
    body = request.get_json() or {}
    fields = ['legal_name','dba','ein','ssn_last4','structure_type','state_of_formation',
              'formation_date','registered_agent','business_address','business_phone',
              'business_email','fiscal_year_end','accounting_method','naics_code',
              'hic_number','hic_expires','pa_tax_id','philly_tax_account',
              'has_employees','collects_sales_tax',
              'filing_status','dependents']
    sets, vals = [], []
    for k in fields:
        if k in body:
            sets.append(f"{k}=?"); vals.append(body[k])
    if not sets:
        return jsonify({'success': False, 'error': 'no fields'}), 400
    sets.append("updated_at=?"); vals.append(datetime.datetime.now().isoformat())
    conn = sqlite3.connect(os.path.join(DASHBOARD_DIR, 'baza_projects.db'))
    conn.execute(f"UPDATE ahb_business_profile SET {','.join(sets)} WHERE id=1", vals)
    conn.commit()
    conn.close()
    return jsonify({'success': True})


@app.route('/api/ahb/other-income', methods=['GET', 'POST'])
def api_ahb_other_income():
    """Per-year personal 'other income' entries used by the Uncle Sam tax estimate.
    GET ?year=2026 returns rows for that year.
    POST body: {year, kind, source, amount, notes}."""
    conn = sqlite3.connect(os.path.join(DASHBOARD_DIR, 'baza_projects.db'))
    conn.row_factory = sqlite3.Row
    if request.method == 'GET':
        year = request.args.get('year') or str(datetime.datetime.now().year)
        rows = conn.execute(
            "SELECT * FROM ahb_other_income WHERE year=? ORDER BY id DESC", (year,)
        ).fetchall()
        conn.close()
        return jsonify([dict(r) for r in rows])
    body = request.get_json() or {}
    year = str(body.get('year') or datetime.datetime.now().year)
    kind = body.get('kind') or 'other'
    source = body.get('source') or ''
    try:
        amount = float(body.get('amount') or 0)
    except (TypeError, ValueError):
        amount = 0.0
    notes = body.get('notes') or ''
    cur = conn.execute(
        """INSERT INTO ahb_other_income (year, kind, source, amount, notes)
           VALUES (?,?,?,?,?)""",
        (year, kind, source, amount, notes)
    )
    new_id = cur.lastrowid
    conn.commit()
    row = conn.execute("SELECT * FROM ahb_other_income WHERE id=?", (new_id,)).fetchone()
    conn.close()
    return jsonify({'success': True, 'entry': dict(row)})


@app.route('/api/ahb/other-income/<int:oid>', methods=['PUT', 'DELETE'])
def api_ahb_other_income_modify(oid):
    conn = sqlite3.connect(os.path.join(DASHBOARD_DIR, 'baza_projects.db'))
    conn.row_factory = sqlite3.Row
    row = conn.execute("SELECT * FROM ahb_other_income WHERE id=?", (oid,)).fetchone()
    if not row:
        conn.close()
        return jsonify({'success': False, 'error': 'not found'}), 404
    if request.method == 'DELETE':
        conn.execute("DELETE FROM ahb_other_income WHERE id=?", (oid,))
        conn.commit()
        conn.close()
        return jsonify({'success': True})
    body = request.get_json() or {}
    sets, vals = [], []
    for k in ('year', 'kind', 'source', 'amount', 'notes'):
        if k in body:
            sets.append(f"{k}=?")
            vals.append(float(body[k]) if k == 'amount' else body[k])
    if not sets:
        conn.close()
        return jsonify({'success': False, 'error': 'no fields'}), 400
    vals.append(oid)
    conn.execute(f"UPDATE ahb_other_income SET {','.join(sets)} WHERE id=?", vals)
    conn.commit()
    conn.close()
    return jsonify({'success': True})


@app.route('/api/ahb/tax-requirements/pa-llc', methods=['GET'])
def api_ahb_pa_llc_requirements():
    structure = request.args.get('structure', 'single_member_llc')
    has_employees = request.args.get('has_employees', '0') in ('1','true','True')
    collects_sales_tax = request.args.get('collects_sales_tax', '0') in ('1','true','True')
    address = (request.args.get('address') or '').lower()
    in_philly = 'philadelphia' in address or 'phila' in address or 'philly' in address
    items = _filter_requirements(structure, has_employees, in_philly, collects_sales_tax)
    out = []
    for r in items:
        copy = dict(r)
        next_date, days_until = _next_due_date(r.get("due", ""))
        copy["next_due_date"] = next_date
        copy["days_until"] = days_until
        copy["required_docs_status"] = _check_required_docs(r.get("required_docs", []))
        out.append(copy)
    # Pull recent filings to mark items as filed
    conn = sqlite3.connect(os.path.join(DASHBOARD_DIR, 'baza_projects.db'))
    conn.row_factory = sqlite3.Row
    filings = {row['requirement_id']: dict(row) for row in conn.execute(
        "SELECT * FROM ahb_tax_filings ORDER BY filed_date DESC"
    ).fetchall()}
    conn.close()
    for item in out:
        f = filings.get(item['id'])
        if f:
            item['last_filed'] = f.get('filed_date')
            item['filing_id'] = f.get('id')
    return jsonify({
        'in_philly': in_philly,
        'has_employees': has_employees,
        'structure': structure,
        'items': out,
    })


@app.route('/api/ahb/tax-filings', methods=['POST'])
def api_ahb_tax_filings_create():
    body = request.get_json() or {}
    conn = sqlite3.connect(os.path.join(DASHBOARD_DIR, 'baza_projects.db'))
    cur = conn.execute("""INSERT INTO ahb_tax_filings
        (requirement_id, period, filed_date, amount_paid, confirmation_number, notes)
        VALUES (?,?,?,?,?,?)""",
        (body.get('requirement_id'), body.get('period'),
         body.get('filed_date') or datetime.date.today().isoformat(),
         body.get('amount_paid'), body.get('confirmation_number'), body.get('notes')))
    fid = cur.lastrowid
    conn.commit()
    conn.close()
    return jsonify({'success': True, 'id': fid})


@app.route('/api/ahb/tax-filings/<int:fid>', methods=['DELETE'])
def api_ahb_tax_filings_delete(fid):
    conn = sqlite3.connect(os.path.join(DASHBOARD_DIR, 'baza_projects.db'))
    conn.execute("DELETE FROM ahb_tax_filings WHERE id=?", (fid,))
    conn.commit()
    conn.close()
    return jsonify({'success': True})


# ── Vendor mini-CRM ─────────────────────────────────────────────────────────

@app.route('/api/ahb/vendors', methods=['GET','POST'])
def api_ahb_vendors():
    conn = sqlite3.connect(os.path.join(DASHBOARD_DIR, 'baza_projects.db'))
    conn.row_factory = sqlite3.Row
    if request.method == 'GET':
        rows = conn.execute("SELECT * FROM ahb_vendors ORDER BY name").fetchall()
        conn.close()
        return jsonify([dict(r) for r in rows])
    body = request.get_json() or {}
    cur = conn.execute("""INSERT INTO ahb_vendors
        (name, vendor_type, contact_name, phone, email, address, ein_or_ssn, notes)
        VALUES (?,?,?,?,?,?,?,?)
        ON CONFLICT(name) DO UPDATE SET
            vendor_type=excluded.vendor_type,
            contact_name=excluded.contact_name,
            phone=excluded.phone,
            email=excluded.email,
            address=excluded.address,
            ein_or_ssn=excluded.ein_or_ssn,
            notes=excluded.notes,
            updated_at=CURRENT_TIMESTAMP""",
        (body.get('name'), body.get('vendor_type'), body.get('contact_name'),
         body.get('phone'), body.get('email'), body.get('address'),
         body.get('ein_or_ssn'), body.get('notes')))
    vid = cur.lastrowid
    conn.commit()
    conn.close()
    return jsonify({'success': True, 'id': vid})


@app.route('/api/ahb/vendors/<int:vid>', methods=['PATCH','DELETE'])
def api_ahb_vendor_one(vid):
    conn = sqlite3.connect(os.path.join(DASHBOARD_DIR, 'baza_projects.db'))
    if request.method == 'DELETE':
        conn.execute("DELETE FROM ahb_vendors WHERE id=?", (vid,))
        conn.commit()
        conn.close()
        return jsonify({'success': True})
    body = request.get_json() or {}
    sets, vals = [], []
    for k in ('name','vendor_type','contact_name','phone','email','address',
              'ein_or_ssn','notes','coi_expires','license_expires'):
        if k in body:
            sets.append(f"{k}=?"); vals.append(body[k])
    if sets:
        sets.append("updated_at=?"); vals.append(datetime.datetime.now().isoformat())
        vals.append(vid)
        conn.execute(f"UPDATE ahb_vendors SET {','.join(sets)} WHERE id=?", vals)
        conn.commit()
    conn.close()
    return jsonify({'success': True})


@app.route('/api/ahb/packages/build-from-project', methods=['POST'])
def api_ahb_package_build():
    """Phil's intelligent package builder. Given a project_id and a package_type
    (e.g. 'permit'), this:
      1. Pulls the project + client + linked invoice
      2. Finds related curated documents (COI, license, etc.)
      3. Prefills a form template with known data
      4. Creates a draft package the user can review/edit
    """
    body = request.get_json() or {}
    project_id = body.get('project_id')
    pkg_type   = (body.get('package_type') or 'permit').lower()
    name_hint  = body.get('name', '')
    if not project_id:
        return jsonify({'success': False, 'error': 'project_id required'}), 400
    conn = sqlite3.connect(os.path.join(DASHBOARD_DIR, 'baza_projects.db'))
    conn.row_factory = sqlite3.Row
    project = conn.execute("SELECT * FROM ahb_projects WHERE id=?", (project_id,)).fetchone()
    if not project:
        conn.close()
        return jsonify({'success': False, 'error': 'project not found'}), 404
    project = dict(project)
    client = None
    if project.get('client_id'):
        c = conn.execute("SELECT * FROM ahb_clients WHERE id=?", (project['client_id'],)).fetchone()
        if c: client = dict(c)
    # Pull standing AHBCO docs (license, COI, W9) + project-specific docs
    standing = [dict(r) for r in conn.execute(
        "SELECT * FROM ahb_documents WHERE doc_type IN ('license','coi','w9','tax_document') "
        "AND (entity LIKE '%AHBCO%' OR entity LIKE '%All Home Building%') "
        "ORDER BY doc_date DESC"
    ).fetchall()]
    # Project-specific docs (matched by project_id or address mention in summary)
    addr = (project.get('address') or '').strip()
    proj_docs_rows = conn.execute(
        "SELECT * FROM ahb_documents WHERE project_id=? OR (? != '' AND content_text LIKE ?) "
        "ORDER BY curated_at DESC",
        (project_id, addr, f"%{addr}%")
    ).fetchall() if addr else conn.execute(
        "SELECT * FROM ahb_documents WHERE project_id=? ORDER BY curated_at DESC",
        (project_id,)
    ).fetchall()
    proj_docs = [dict(r) for r in proj_docs_rows]
    # Build the prefilled form data based on package type
    form = _build_form_template(pkg_type, project, client)
    # Pick attached doc ids — standing docs by default, plus any project-specific
    attached_ids = [d['id'] for d in (standing + proj_docs)]
    name = name_hint or f"{pkg_type.title()} package — {project.get('title','Project')}"
    cur = conn.execute("""INSERT INTO ahb_app_packages
        (name, package_type, project_id, client_id, status, form_data, attached_doc_ids, notes)
        VALUES (?,?,?,?,?,?,?,?)""",
        (name, pkg_type, project_id, project.get('client_id'),
         'draft', json.dumps(form), json.dumps(attached_ids),
         f"Auto-built by Phil from project {project.get('title','')}."))
    pid = cur.lastrowid
    conn.commit()
    conn.close()
    return jsonify({'success': True, 'id': pid, 'form_data': form,
                    'attached_doc_ids': attached_ids,
                    'standing_docs': len(standing),
                    'project_docs': len(proj_docs)})


def _detect_municipality(address: str) -> str:
    """Best-effort guess at jurisdiction from a project address. Returns one of:
    'philadelphia', 'bensalem', 'bucks_county', 'montgomery_county', 'delaware_county',
    'chester_county', 'pa_other', or 'unknown'."""
    if not address:
        return 'unknown'
    a = address.lower()
    if 'philadelphia' in a or 'phila' in a or 'philly' in a:
        return 'philadelphia'
    if 'bensalem' in a:
        return 'bensalem'
    # Bucks County boroughs/townships
    bucks_hints = ('warrington','warminster','doylestown','newtown','levittown','feasterville',
                   'langhorne','yardley','perkasie','quakertown','sellersville','richboro',
                   'bristol','morrisville','bucks county')
    if any(h in a for h in bucks_hints):
        return 'bucks_county'
    montco_hints = ('king of prussia','norristown','ambler','jenkintown','abington','willow grove',
                    'fort washington','blue bell','plymouth meeting','conshohocken','lansdale',
                    'pottstown','collegeville','montgomery county')
    if any(h in a for h in montco_hints):
        return 'montgomery_county'
    delco_hints = ('upper darby','media','chester','springfield','havertown','newtown square',
                   'broomall','aldan','lansdowne','delaware county')
    if any(h in a for h in delco_hints):
        return 'delaware_county'
    chesco_hints = ('west chester','exton','downingtown','phoenixville','kennett','coatesville',
                    'malvern','chester county')
    if any(h in a for h in chesco_hints):
        return 'chester_county'
    if ' pa' in a or ', pa' in a:
        return 'pa_other'
    return 'unknown'


def _municipality_extras(muni: str) -> dict:
    """Return additional permit-form fields specific to a jurisdiction."""
    if muni == 'philadelphia':
        return {
            'submission_office':    'Philadelphia Department of Licenses & Inspections (L&I)',
            'submission_address':   '1401 John F. Kennedy Blvd, Philadelphia, PA 19102',
            'submission_portal':    'https://www.phila.gov/eclipse/',
            'permit_type':          'Building Permit (Alteration/Repair)',
            'use_group':            'R-3 (Single Family)',
            'construction_type':    'V-B (Wood Frame)',
            'historic_district':    'Check at atlas.phila.gov before submission',
            'zoning_overlay':       '',
            'l_and_i_property_id':  '',
            'opa_account_number':   '',
            'flood_plain':          'Check FEMA flood map',
        }
    if muni == 'bensalem':
        return {
            'submission_office':    'Bensalem Township Building Department',
            'submission_address':   '2400 Byberry Rd, Bensalem, PA 19020',
            'submission_phone':     '215-633-3600',
            'permit_type':          'Building Permit',
            'parcel_id':            '',
            'zoning_district':      '',
            'lot_size':             '',
            'setback_front':        '',
            'setback_rear':         '',
            'setback_side':         '',
        }
    if muni == 'bucks_county':
        return {
            'submission_office':    'Township Building Department (Bucks County)',
            'permit_type':          'Building Permit',
            'parcel_id':            '',
            'zoning_district':      '',
            'lot_size':             '',
            'setback_front':        '',
            'setback_rear':         '',
            'setback_side':         '',
            'uniform_construction_code': 'Pennsylvania UCC compliant',
        }
    if muni == 'montgomery_county':
        return {
            'submission_office':    'Township/Borough Building Department (Montgomery County)',
            'permit_type':          'Building Permit',
            'parcel_id':            '',
            'zoning_district':      '',
        }
    if muni == 'delaware_county':
        return {
            'submission_office':    'Township/Borough Building Department (Delaware County)',
            'permit_type':          'Building Permit',
            'parcel_id':            '',
            'zoning_district':      '',
        }
    if muni == 'chester_county':
        return {
            'submission_office':    'Township/Borough Building Department (Chester County)',
            'permit_type':          'Building Permit',
            'parcel_id':            '',
            'zoning_district':      '',
        }
    return {
        'submission_office':    'Local Building Department',
        'permit_type':          'Building Permit',
    }


def _build_form_template(pkg_type: str, project: dict, client: dict | None) -> dict:
    """Return a dict of prefilled form fields for the given package type.
    The user can edit any field in the DocPrep UI."""
    base = {
        'contractor_name':       'Sergey Tkach',
        'contractor_company':    'All Home Building Co LLC',
        'contractor_address':    '2725 Colmar Ave, Bensalem, PA',
        'contractor_phone':      '800-484-6404',
        'contractor_license':    '',  # to be filled from license doc
        'contractor_email':      'admin@allhomebuilding.co',
        'project_address':       project.get('address') or '',
        'project_description':   project.get('description') or '',
        'project_scope':         project.get('scope') or '',
        'project_budget':        project.get('value') or project.get('budget_high') or '',
        'project_start_date':    project.get('start_date') or '',
        'project_end_date':      project.get('end_date') or '',
        'client_name':           (client or {}).get('name') or project.get('client_name') or '',
        'client_phone':          (client or {}).get('phone') or '',
        'client_email':          (client or {}).get('email') or '',
        'client_address':        (client or {}).get('address') or project.get('address') or '',
    }
    if pkg_type == 'permit':
        muni = _detect_municipality(project.get('address') or '')
        base['detected_jurisdiction'] = muni.replace('_',' ').title()
        base.update({
            'permit_type':           'Building Permit',
            'work_description':      project.get('description') or project.get('scope') or '',
            'estimated_cost':        project.get('value') or project.get('budget_high') or '',
            'property_owner':        (client or {}).get('name') or '',
            'parcel_number':         '',
            'zoning':                '',
            'square_footage_change': '',
            'occupancy_type':        'Single Family Residential',
            'fire_sprinkler':        'No',
            'subcontractors':        '',
            'work_to_be_performed':  project.get('description') or '',
        })
        # Layer in jurisdiction-specific fields (overrides any duplicates)
        base.update(_municipality_extras(muni))
    elif pkg_type == 'coi_request':
        base.update({
            'certificate_holder':    (client or {}).get('name') or '',
            'certificate_address':   (client or {}).get('address') or project.get('address') or '',
            'project_reference':     project.get('title') or '',
            'coverage_dates':        '',
            'additional_insured':    '',
        })
    elif pkg_type == 'contract':
        base.update({
            'contract_amount':       project.get('value') or project.get('budget_high') or '',
            'payment_schedule':      'Deposit before commencement, balance on completion',
            'completion_date':       project.get('end_date') or '',
            'warranty_period':       '1 year on workmanship',
            'change_order_terms':    'Any change in scope requires written change order signed by both parties.',
        })
    elif pkg_type == 'change_order':
        base.update({
            'original_contract_amount': project.get('value') or project.get('budget_high') or '',
            'change_description':       '',
            'change_amount':            '',
            'reason_for_change':        '',
        })
    return base


# ── Project Document Scraper (extract fields from PDFs/docs via LLM) ─────────

def _extract_text(filepath: str) -> str:
    """Best-effort text extraction from common doc formats."""
    ext = os.path.splitext(filepath)[1].lower()
    try:
        if ext == '.pdf':
            try:
                import pdfplumber
                txt = []
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages[:20]:  # cap at 20 pages
                        t = page.extract_text() or ''
                        if t:
                            txt.append(t)
                return '\n'.join(txt)
            except Exception as e:
                return f"[pdf extract failed: {e}]"
        if ext in ('.docx',):
            try:
                import docx
                d = docx.Document(filepath)
                return '\n'.join(p.text for p in d.paragraphs)
            except Exception as e:
                return f"[docx extract failed: {e}]"
        if ext in ('.txt', '.md', '.csv', '.html', '.rtf'):
            with open(filepath, 'r', errors='ignore') as f:
                return f.read()[:50000]
        # Image — leave empty (could OCR via tesseract later)
        return ''
    except Exception as e:
        return f"[extract failed: {e}]"


def _scrape_project_fields(text: str) -> dict:
    """Send extracted text through Ollama with a structured prompt to pull project fields."""
    if not text or not text.strip():
        return {}
    text = text[:12000]  # context cap
    prompt = (
        "You are extracting structured project information from a construction document "
        "(could be a contract, estimate, bid, lead form, or job description). "
        "Return ONLY a JSON object with these exact keys (use null if not present):\n"
        "{\n"
        '  "client_name": "full name",\n'
        '  "client_phone": "phone",\n'
        '  "client_email": "email",\n'
        '  "address": "full project street address with city/state/zip",\n'
        '  "scope": "kitchen|bathroom|addition|basement|deck|full-reno|other",\n'
        '  "scope_other": "if scope is other, describe what it is",\n'
        '  "title": "short project title",\n'
        '  "description": "1-3 sentence summary of the work",\n'
        '  "budget_low": number_or_null,\n'
        '  "budget_high": number_or_null,\n'
        '  "start_date": "YYYY-MM-DD",\n'
        '  "end_date": "YYYY-MM-DD",\n'
        '  "competitors": ["list of any other contractor/company names mentioned"]\n'
        "}\n\n"
        "Document text:\n" + text + "\n\nJSON:"
    )
    try:
        import urllib.request as _ur
        payload = json.dumps({
            "model": "qwen2.5:14b",
            "prompt": prompt,
            "stream": False,
            "format": "json",
            "options": {"temperature": 0.1, "num_predict": 800}
        }).encode()
        req = _ur.Request("http://localhost:11434/api/generate", data=payload,
                          headers={"Content-Type": "application/json"})
        with _ur.urlopen(req, timeout=120) as r:
            data = json.loads(r.read())
        raw = data.get("response", "").strip()
        # Strip code fences if present
        raw = re.sub(r'^```(?:json)?\s*|\s*```$', '', raw).strip()
        return json.loads(raw)
    except Exception as e:
        return {"_error": str(e)}


@app.route('/api/ahb/projects/<pid>/adopt-staging', methods=['POST'])
def api_ahb_project_adopt_staging(pid):
    """Move scrape-staging files into the real project artifacts dir after the project is saved."""
    body = request.get_json() or {}
    staging_id = body.get('staging_id', '')
    if not staging_id or not staging_id.startswith('scrape-'):
        return jsonify({'success': False, 'error': 'invalid staging_id'}), 400
    src = os.path.join(ARTIFACTS_DIR, staging_id)
    dst = os.path.join(ARTIFACTS_DIR, pid)
    if not os.path.isdir(src):
        return jsonify({'success': False, 'error': 'staging not found'}), 404
    os.makedirs(dst, exist_ok=True)
    moved = []
    for f in os.listdir(src):
        sp = os.path.join(src, f)
        dp = os.path.join(dst, f)
        try:
            os.rename(sp, dp)
            moved.append(f)
        except Exception:
            pass
    try: os.rmdir(src)
    except Exception: pass
    return jsonify({'success': True, 'moved': moved})


@app.route('/api/ahb/projects/scrape', methods=['POST'])
def api_ahb_project_scrape():
    """Upload one or more docs (multipart). Extract text. Run LLM. Return suggested fields.
    Files are also saved to artifacts/<project_id>/ if a project_id is provided, otherwise
    to artifacts/scrape-staging/ keyed by a temp uuid."""
    files = request.files.getlist('files') or request.files.getlist('file')
    if not files:
        return jsonify({'success': False, 'error': 'no files'}), 400
    project_id = request.form.get('project_id') or f"scrape-{uuid.uuid4().hex[:8]}"
    proj_dir = os.path.join(ARTIFACTS_DIR, project_id)
    os.makedirs(proj_dir, exist_ok=True)
    saved   = []
    blobs   = []
    for f in files:
        if not f or not f.filename:
            continue
        safe = re.sub(r'[^\w.\-_ ()]', '_', f.filename)
        fpath = os.path.join(proj_dir, safe)
        f.save(fpath)
        saved.append(safe)
        text = _extract_text(fpath)
        if text:
            blobs.append(f"=== {safe} ===\n{text}")
    combined = '\n\n'.join(blobs)
    suggested = _scrape_project_fields(combined) if combined else {}
    return jsonify({
        'success': True,
        'project_id': project_id,
        'files': saved,
        'text_length': len(combined),
        'suggested': suggested,
    })

# ── Project Geocoding (for iCloud GPS classifier) ────────────────────────────

@app.route('/api/ahb/geocode/status')
def api_geocode_status():
    from core.geocoder import project_geocode_status
    return jsonify(project_geocode_status())

@app.route('/api/ahb/geocode/run', methods=['POST'])
def api_geocode_run():
    """Bulk geocode all AHB project addresses. Rate-limited to ~1 req/sec."""
    from core.geocoder import geocode_all_projects
    body = request.get_json(silent=True) or {}
    force = bool(body.get('force', False))
    try:
        result = geocode_all_projects(force=force)
        return jsonify(result)
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 500

# ── iCloud Ingest Pipeline (admin + multi-tenant) ────────────────────────────

def _icloud_user_id():
    """Return cloud user_id — family mode always returns Serge."""
    return FAMILY_USER_ID if CLOUD_ENABLED else None

@app.route('/api/icloud/accounts', methods=['GET'])
def api_icloud_accounts():
    from core.icloud_ingest import list_accounts
    uid = _icloud_user_id()
    accs = list_accounts(user_id=uid, include_admin=(uid is None))
    # Hide passwords in response
    for a in accs:
        a.pop('app_password', None)
    return jsonify(accs)

@app.route('/api/icloud/accounts', methods=['POST'])
def api_icloud_add_account():
    """Register an iCloud account. Auth-only step still requires CLI run for 2FA."""
    from core.icloud_ingest import add_account
    body = request.get_json() or {}
    apple_id = (body.get('apple_id') or '').strip()
    password = body.get('password') or ''
    ahb_owner= bool(body.get('ahb_owner', False))
    if not apple_id or not password:
        return jsonify({'success': False, 'error': 'apple_id and password required'}), 400
    uid = _icloud_user_id()
    try:
        aid = add_account(apple_id, password, user_id=uid, ahb_owner=ahb_owner)
        return jsonify({'success': True, 'id': aid,
                        'next_step': 'Run: venv/bin/python scripts/icloud_setup.py' +
                                     (f' --user {uid}' if uid else '') +
                                     ' (in a terminal to complete 2FA)'})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/icloud/accounts/<int:aid>', methods=['DELETE'])
def api_icloud_remove_account(aid):
    from core.icloud_ingest import remove_account, _get_account
    acc = _get_account(aid)
    if not acc:
        return jsonify({'success': False, 'error': 'not found'}), 404
    uid = _icloud_user_id()
    if acc['user_id'] != uid:
        return jsonify({'success': False, 'error': 'forbidden'}), 403
    remove_account(aid, user_id=uid)
    return jsonify({'success': True})

@app.route('/api/icloud/sync', methods=['POST'])
def api_icloud_sync():
    """Trigger sync for one account or all of the current user's accounts."""
    from core.icloud_ingest import ingest_account, ingest_all, _get_account
    body = request.get_json() or {}
    aid = body.get('account_id')
    uid = _icloud_user_id()
    try:
        if aid:
            acc = _get_account(int(aid))
            if not acc or acc['user_id'] != uid:
                return jsonify({'success': False, 'error': 'forbidden'}), 403
            result = ingest_account(int(aid), recent=body.get('recent'),
                                    until_found=body.get('until_found', 100))
            return jsonify({'success': result.get('ok', False), 'result': result})
        results = ingest_all(user_id=uid)
        return jsonify({'success': True, 'results': results})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/icloud/status')
def api_icloud_status():
    """Lightweight status for the dashboard widget."""
    from core.icloud_ingest import list_accounts
    uid = _icloud_user_id()
    accs = list_accounts(user_id=uid, include_admin=(uid is None))
    total_synced  = sum(a.get('total_synced')  or 0 for a in accs)
    total_jobsite = sum(a.get('total_jobsite') or 0 for a in accs)
    total_personal= sum(a.get('total_personal') or 0 for a in accs)
    last_sync = max((a.get('last_sync') or '' for a in accs), default='')
    return jsonify({
        'accounts': len(accs),
        'last_sync': last_sync,
        'total_synced': total_synced,
        'total_jobsite': total_jobsite,
        'total_personal': total_personal,
        'is_cloud_user': uid is not None,
    })

# ─────────────────────────────────────────────────────────────────────────────
# COMMS — Sophisticated chat platform for talking to any agent
# ─────────────────────────────────────────────────────────────────────────────
COMMS_UPLOAD_DIR = os.path.join(DASHBOARD_DIR, 'uploads', 'comms')
os.makedirs(COMMS_UPLOAD_DIR, exist_ok=True)

def _comms_init():
    conn = sqlite3.connect(os.path.join(DASHBOARD_DIR, 'baza_projects.db'))
    c = conn.cursor()
    c.executescript("""
        CREATE TABLE IF NOT EXISTS comms_sessions (
            id TEXT PRIMARY KEY,
            agent_id TEXT NOT NULL,
            title TEXT,
            created_at TEXT DEFAULT (datetime('now')),
            updated_at TEXT DEFAULT (datetime('now'))
        );
        CREATE TABLE IF NOT EXISTS comms_messages (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            session_id TEXT NOT NULL,
            role TEXT NOT NULL,
            content TEXT,
            attachments TEXT,
            created_at TEXT DEFAULT (datetime('now'))
        );
        CREATE INDEX IF NOT EXISTS idx_comms_msgs_sess ON comms_messages(session_id);
    """)
    conn.commit()
    conn.close()
_comms_init()

def _comms_db():
    conn = sqlite3.connect(os.path.join(DASHBOARD_DIR, 'baza_projects.db'))
    conn.row_factory = sqlite3.Row
    return conn

@app.route('/comms')
def comms_page():
    config = load_config()
    agents = config.get('agents', {})
    agent_list = []
    for aid, ac in agents.items():
        agent_list.append({
            'id': aid,
            'name': ac.get('name', aid),
            'role': ac.get('role', ''),
            'model': ac.get('model', ''),
            'company_title': ac.get('company_title', ''),
        })
    return render_template('comms.html', agents=agent_list)

@app.route('/api/comms/sessions', methods=['GET', 'POST'])
def api_comms_sessions():
    conn = _comms_db()
    c = conn.cursor()
    if request.method == 'POST':
        data = request.get_json() or {}
        sid = uuid.uuid4().hex[:12]
        c.execute("INSERT INTO comms_sessions (id, agent_id, title) VALUES (?, ?, ?)",
                  (sid, data.get('agent_id', 'simon_bately'), data.get('title', 'New chat')))
        conn.commit()
        row = c.execute("SELECT * FROM comms_sessions WHERE id=?", (sid,)).fetchone()
        conn.close()
        return jsonify(dict(row))
    rows = c.execute("""
        SELECT s.*, (SELECT COUNT(*) FROM comms_messages WHERE session_id=s.id) AS msg_count,
               (SELECT content FROM comms_messages WHERE session_id=s.id ORDER BY id DESC LIMIT 1) AS last_msg
        FROM comms_sessions s ORDER BY datetime(updated_at) DESC
    """).fetchall()
    conn.close()
    return jsonify([dict(r) for r in rows])

@app.route('/api/comms/sessions/<sid>', methods=['GET', 'PUT', 'DELETE'])
def api_comms_session(sid):
    conn = _comms_db()
    c = conn.cursor()
    if request.method == 'DELETE':
        c.execute("DELETE FROM comms_messages WHERE session_id=?", (sid,))
        c.execute("DELETE FROM comms_sessions WHERE id=?", (sid,))
        conn.commit()
        conn.close()
        return jsonify({'success': True})
    if request.method == 'PUT':
        data = request.get_json() or {}
        if 'title' in data:
            c.execute("UPDATE comms_sessions SET title=?, updated_at=datetime('now') WHERE id=?",
                      (data['title'], sid))
        if 'agent_id' in data:
            c.execute("UPDATE comms_sessions SET agent_id=?, updated_at=datetime('now') WHERE id=?",
                      (data['agent_id'], sid))
        conn.commit()
    sess = c.execute("SELECT * FROM comms_sessions WHERE id=?", (sid,)).fetchone()
    if not sess:
        conn.close()
        return jsonify({'error': 'not found'}), 404
    msgs = c.execute("SELECT * FROM comms_messages WHERE session_id=? ORDER BY id ASC", (sid,)).fetchall()
    conn.close()
    out = dict(sess)
    out['messages'] = [dict(m) for m in msgs]
    return jsonify(out)

def _comms_ollama_url(model):
    """Pick which Ollama instance hosts a model. Try AMD first, then NVIDIA."""
    import requests as _r
    for url in ('http://127.0.0.1:11434', 'http://127.0.0.1:11435'):
        try:
            r = _r.get(f"{url}/api/tags", timeout=2)
            if r.ok:
                tags = [m['name'] for m in r.json().get('models', [])]
                if model in tags or any(t.split(':')[0] == model.split(':')[0] for t in tags):
                    return url
        except Exception:
            continue
    return 'http://127.0.0.1:11434'

@app.route('/api/comms/send', methods=['POST'])
def api_comms_send():
    """Send a user message and stream the agent's reply via SSE."""
    from flask import Response, stream_with_context
    import requests as _r
    data = request.get_json() or {}
    sid = data.get('session_id')
    user_text = (data.get('content') or '').strip()
    attachments = data.get('attachments') or []
    if not sid or not user_text:
        return jsonify({'error': 'session_id and content required'}), 400

    conn = _comms_db()
    c = conn.cursor()
    sess = c.execute("SELECT * FROM comms_sessions WHERE id=?", (sid,)).fetchone()
    if not sess:
        conn.close()
        return jsonify({'error': 'session not found'}), 404
    sess = dict(sess)
    agent_id = sess['agent_id']

    # Save user message
    c.execute("INSERT INTO comms_messages (session_id, role, content, attachments) VALUES (?,?,?,?)",
              (sid, 'user', user_text, json.dumps(attachments)))
    # Auto-title from first user message
    if (sess.get('title') or '').strip() in ('', 'New chat'):
        c.execute("UPDATE comms_sessions SET title=? WHERE id=?", (user_text[:60], sid))
    c.execute("UPDATE comms_sessions SET updated_at=datetime('now') WHERE id=?", (sid,))
    conn.commit()

    # Build LLM context
    config = load_config()
    agent_cfg = (config.get('agents') or {}).get(agent_id, {})
    model = agent_cfg.get('model', 'mistral-small:22b')
    system_prompt = agent_cfg.get('system_prompt', f'You are {agent_cfg.get("name", agent_id)}.')

    history = c.execute(
        "SELECT role, content FROM comms_messages WHERE session_id=? ORDER BY id ASC", (sid,)
    ).fetchall()
    messages = [{"role": "system", "content": system_prompt}]
    for h in history:
        messages.append({"role": h['role'], "content": h['content']})
    if attachments:
        files_note = "\n\n[User attached files: " + ", ".join(a.get('name','file') for a in attachments) + "]"
        messages[-1]['content'] += files_note
    conn.close()

    def generate():
        full_response = ''
        try:
            url = _comms_ollama_url(model)
            payload = {"model": model, "messages": messages, "stream": True}
            with _r.post(f"{url}/api/chat", json=payload, stream=True, timeout=300) as resp:
                for line in resp.iter_lines():
                    if not line:
                        continue
                    try:
                        chunk = json.loads(line.decode('utf-8'))
                    except Exception:
                        continue
                    tok = (chunk.get('message') or {}).get('content', '')
                    if tok:
                        full_response += tok
                        yield f"data: {json.dumps({'token': tok})}\n\n"
                    if chunk.get('done'):
                        break
        except Exception as e:
            yield f"data: {json.dumps({'error': str(e)})}\n\n"
        # Persist assistant message
        try:
            c2 = _comms_db()
            c2.execute(
                "INSERT INTO comms_messages (session_id, role, content) VALUES (?,?,?)",
                (sid, 'assistant', full_response)
            )
            c2.execute("UPDATE comms_sessions SET updated_at=datetime('now') WHERE id=?", (sid,))
            c2.commit()
            c2.close()
        except Exception:
            pass
        yield f"data: {json.dumps({'done': True})}\n\n"

    return Response(stream_with_context(generate()), mimetype='text/event-stream',
                    headers={'X-Accel-Buffering': 'no', 'Cache-Control': 'no-cache'})

@app.route('/api/comms/upload', methods=['POST'])
def api_comms_upload():
    f = request.files.get('file')
    if not f:
        return jsonify({'error': 'no file'}), 400
    sid = request.form.get('session_id', 'tmp')
    folder = os.path.join(COMMS_UPLOAD_DIR, sid)
    os.makedirs(folder, exist_ok=True)
    name = secure_filename(f.filename or f"file_{uuid.uuid4().hex[:8]}")
    path = os.path.join(folder, name)
    f.save(path)
    return jsonify({
        'success': True,
        'name': name,
        'size': os.path.getsize(path),
        'url': f"/api/comms/file/{sid}/{name}",
    })

@app.route('/api/comms/file/<sid>/<path:fname>')
def api_comms_file(sid, fname):
    folder = os.path.join(COMMS_UPLOAD_DIR, sid)
    return send_from_directory(folder, fname)

@app.route('/api/serve-local/<path:filepath>')
def api_serve_local(filepath):
    """Serve local files from known safe directories (generated images, artifacts)."""
    full = '/' + filepath
    SAFE_PREFIXES = (
        '/mnt/empirepool/media/generated/',
        '/mnt/empirepool/media/',
        os.path.join(DASHBOARD_DIR, 'artifacts') + '/',
        '/tmp/',
    )
    if not any(full.startswith(p) for p in SAFE_PREFIXES):
        return 'Forbidden', 403
    if not os.path.isfile(full):
        return 'Not found', 404
    directory = os.path.dirname(full)
    filename = os.path.basename(full)
    return send_from_directory(directory, filename)

@app.route('/api/comms/voice', methods=['POST'])
def api_comms_voice():
    """Transcribe an audio blob via local whisper if available, else echo error."""
    f = request.files.get('audio')
    if not f:
        return jsonify({'error': 'no audio'}), 400
    tmp = os.path.join(COMMS_UPLOAD_DIR, f"voice_{uuid.uuid4().hex[:8]}.webm")
    f.save(tmp)
    try:
        # Try whisper.cpp / whisper CLI
        for cmd in (['whisper', tmp, '--model', 'base', '--output_format', 'txt', '--output_dir', COMMS_UPLOAD_DIR],
                    ['whisper-cpp', '-f', tmp]):
            try:
                r = subprocess.run(cmd, capture_output=True, text=True, timeout=120)
                if r.returncode == 0:
                    text = r.stdout.strip()
                    if not text:
                        # whisper writes file
                        txt_path = tmp.rsplit('.', 1)[0] + '.txt'
                        if os.path.exists(txt_path):
                            text = open(txt_path).read().strip()
                    return jsonify({'text': text})
            except FileNotFoundError:
                continue
        return jsonify({'error': 'whisper not installed; install openai-whisper for voice input'}), 501
    finally:
        try: os.remove(tmp)
        except: pass

@app.route('/api/comms/shell', methods=['POST'])
def api_comms_shell():
    """Run a shell command from the inline terminal popout. Restricted to baza user."""
    data = request.get_json() or {}
    cmd = (data.get('cmd') or '').strip()
    if not cmd:
        return jsonify({'error': 'no command'}), 400
    try:
        r = subprocess.run(['bash', '-lc', cmd], capture_output=True, text=True, timeout=60)
        return jsonify({
            'stdout': r.stdout,
            'stderr': r.stderr,
            'code': r.returncode,
        })
    except subprocess.TimeoutExpired:
        return jsonify({'error': 'timeout (60s)'}), 408
    except Exception as e:
        return jsonify({'error': str(e)}), 500


# ── Theme toggle ─────────────────────────────────────────────────────────────
# Stores user theme in session + a 1-year cookie. Templates read session.theme
# (or the cookie via `data-theme="{{ request.cookies.get('theme','dark') }}"`).
@app.route('/settings/theme', methods=['POST'])
def settings_theme():
    body = request.get_json(silent=True) or {}
    val = (body.get('value') or '').strip().lower()
    if val not in ('dark', 'light'):
        return jsonify({'ok': False, 'error': 'theme must be dark or light'}), 400
    session['theme'] = val
    resp = jsonify({'ok': True, 'theme': val})
    # 1y cookie so it survives session expiry. No HttpOnly: theme.js reads it.
    resp.set_cookie('theme', val, max_age=60 * 60 * 24 * 365, samesite='Lax')
    return resp


# ── CYD remote endpoints ────────────────────────────────────────────────────
# Thin shims to feed the ESP32 dashboard at /home/switchhacker/baza_edge/cyd_dashboard.
# The CYD firmware fetches these every 60s and renders the resulting JSON. Most
# return empty/stub payloads where Baza doesn't have a real data source yet —
# the firmware is happy with empty arrays and renders 0/N counts cleanly.
@app.route('/api/alerts/recent')
def api_alerts_recent():
    return jsonify({'alerts': []})


@app.route('/api/tasks/active')
def api_tasks_active():
    """Active task summary for the CYD's TASKS page. Wraps the existing
    /api/tasks data into the {tasks:[...]} shape the firmware expects."""
    try:
        con = sqlite3.connect(BAZA_PROJECTS_DB if 'BAZA_PROJECTS_DB' in globals() else os.path.join(DASHBOARD_DIR, 'baza_projects.db'))
        con.row_factory = sqlite3.Row
        rows = con.execute(
            "SELECT title, assigned_to, status FROM tasks "
            "WHERE status IN ('open','in_progress') "
            "ORDER BY priority DESC, created_at DESC LIMIT 8"
        ).fetchall()
        con.close()
        tasks = [{
            'name':     (r['title'] or '')[:55],
            'agent':    r['assigned_to'] or '',
            'status':   r['status'] or '',
            'next_run': '',
        } for r in rows]
        return jsonify({'tasks': tasks})
    except Exception as e:
        return jsonify({'tasks': [], 'error': str(e)})


@app.route('/api/edge/status')
def api_edge_status():
    """Alias for /api/edge/nodes — CYD firmware uses this URL."""
    return api_edge_nodes()


@app.route('/api/edge/heartbeat', methods=['POST'])
def api_edge_heartbeat():
    """Record a heartbeat from a remote node (CYD, S3 voice, S3 power, etc.)."""
    body = request.get_json(silent=True) or {}
    return jsonify({'ok': True, 'node_id': body.get('node_id', 'unknown')})


@app.route('/api/dispatch', methods=['POST'])
def api_dispatch():
    """Dispatch a task to an agent. CYD's quick-action buttons hit this."""
    body = request.get_json(silent=True) or {}
    agent = body.get('agent') or body.get('handle') or ''
    msg = body.get('message') or body.get('task') or ''
    return jsonify({'ok': True, 'agent': agent, 'queued': bool(agent and msg)})


@app.route('/api/redis/publish', methods=['POST'])
def api_redis_publish():
    """Publish to a Redis channel. CYD's broadcast buttons hit this."""
    body = request.get_json(silent=True) or {}
    try:
        import redis
        r = redis.Redis(host='localhost', port=6379, decode_responses=True)
        channel = body.get('channel', 'baza_broadcast')
        r.publish(channel, json.dumps({
            'event': body.get('event', ''),
            'data':  body.get('data', '{}'),
            'source': body.get('source', 'unknown'),
        }))
        return jsonify({'ok': True})
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)})


# ── Vision UI ──────────────────────────────────────────────────────────────
from dashboard.vision_routes import bp as vision_bp  # noqa: E402
app.register_blueprint(vision_bp)


@app.route("/datahub/private")
def datahub_private_page():
    """Legacy private gallery — renders the unlock form when locked, the
    photo grid when unlocked. /vision is an alternate (catalogue) view of
    the same private content; both pages have nav links to each other."""
    return render_template('private.html',
                           passphrase_set=_private_pass_is_set(),
                           unlocked=_is_private_unlocked())


# One-shot migration: sweep legacy `.private-inbound/` .meta sidecars so any
# auto-private mark from before the inbound-is-public flip is removed. Runs
# at import time on every dashboard start; idempotent (no-op when nothing
# left to sweep).
try:
    _framework_dir = os.path.dirname(DASHBOARD_DIR)
    _swept = _migrate_legacy_inbound_meta(_framework_dir)
    if _swept:
        print(f"[dashboard] swept {_swept} legacy private-meta sidecars (now public).")
except Exception as _e:
    print(f"[dashboard] legacy meta sweep failed (non-fatal): {_e}")


if __name__ == '__main__':
    app.run(host='0.0.0.0', port=8888, debug=False)

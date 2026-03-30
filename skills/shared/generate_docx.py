#!/usr/bin/env python3
"""
Baza Empire Skill — generate_docx
Generate a Word (.docx) document from structured content.
Phil uses this for contracts, proposals, agreements, checklists, letters.

SKILL_ARGS:
  title       (str)  — Document title (required)
  sections    (list) — [{heading: str, body: str}, ...] — document sections
  filename    (str)  — output filename (default: title_snake.docx)
  project_id  (str)  — dashboard project (default: shared)
  author      (str)  — document author line (default: "Phil Hass — AHBCO LLC")
  footer_text (str)  — optional footer string
  table       (dict) — optional table: {headers: [...], rows: [[...], ...]}
"""
import os
import sys
import json
import re
import time
import urllib.request

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print(json.dumps({"error": "python-docx not installed. Run: venv/bin/pip install python-docx"}))
    sys.exit(1)

DASHBOARD_URL = os.environ.get("BAZA_DASHBOARD_URL", "http://localhost:8888")
ARTIFACTS_BASE = os.path.join(
    os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))),
    "dashboard", "artifacts"
)

args = json.loads(os.environ.get("SKILL_ARGS", "{}"))

title      = args.get("title", "Document")
sections   = args.get("sections", [])
author     = args.get("author", "Phil Hass — All Home Building Co LLC")
footer_text= args.get("footer_text", "")
table_data = args.get("table", None)
project_id = args.get("project_id", "shared")

# Build filename
raw_name = args.get("filename", "")
if not raw_name:
    safe = re.sub(r'[^\w]', '_', title.lower())[:40]
    raw_name = f"{safe}_{int(time.time())}.docx"
if not raw_name.endswith(".docx"):
    raw_name += ".docx"

doc = Document()

# ── Style ────────────────────────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# ── Title ─────────────────────────────────────────────────────────────────────
title_para = doc.add_heading(title, 0)
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ── Author / Date ─────────────────────────────────────────────────────────────
from datetime import date
meta = doc.add_paragraph(f"{author}  |  {date.today().strftime('%B %d, %Y')}")
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.runs[0].font.size = Pt(10)
meta.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()  # spacer

# ── Sections ──────────────────────────────────────────────────────────────────
for sec in sections:
    heading = sec.get("heading", "")
    body    = sec.get("body", "")
    if heading:
        doc.add_heading(heading, level=1)
    if body:
        for line in body.strip().split("\n"):
            line = line.strip()
            if not line:
                doc.add_paragraph()
                continue
            # Detect bullet lines starting with - or •
            if line.startswith(("-", "•", "*")):
                p = doc.add_paragraph(line.lstrip("-•* "), style='List Bullet')
            else:
                doc.add_paragraph(line)

# ── Optional table ────────────────────────────────────────────────────────────
if table_data:
    headers = table_data.get("headers", [])
    rows    = table_data.get("rows", [])
    if headers and rows:
        doc.add_paragraph()
        t = doc.add_table(rows=1, cols=len(headers))
        t.style = 'Light Grid Accent 1'
        hdr_cells = t.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = str(h)
            run = hdr_cells[i].paragraphs[0].runs[0]
            run.font.bold = True
        for row_data in rows:
            row_cells = t.add_row().cells
            for i, val in enumerate(row_data[:len(headers)]):
                row_cells[i].text = str(val)
        doc.add_paragraph()

# ── Footer text ───────────────────────────────────────────────────────────────
if footer_text:
    doc.add_page_break()
    footer_para = doc.add_paragraph(footer_text)
    footer_para.runs[0].font.size = Pt(9)
    footer_para.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

# ── Save ──────────────────────────────────────────────────────────────────────
out_dir = os.path.join(ARTIFACTS_BASE, project_id)
os.makedirs(out_dir, exist_ok=True)
out_path = os.path.join(out_dir, raw_name)
doc.save(out_path)

# Register with dashboard upload API
try:
    boundary = "bazadocxboundary"
    with open(out_path, "rb") as f:
        file_data = f.read()
    body = (
        f"--{boundary}\r\nContent-Disposition: form-data; name=\"project_id\"\r\n\r\n{project_id}\r\n"
        f"--{boundary}\r\nContent-Disposition: form-data; name=\"file\"; filename=\"{raw_name}\"\r\n"
        f"Content-Type: application/vnd.openxmlformats-officedocument.wordprocessingml.document\r\n\r\n"
    ).encode() + file_data + f"\r\n--{boundary}--\r\n".encode()
    req = urllib.request.Request(
        f"{DASHBOARD_URL}/api/artifacts/upload",
        data=body,
        headers={"Content-Type": f"multipart/form-data; boundary={boundary}"},
        method="POST"
    )
    with urllib.request.urlopen(req, timeout=15):
        pass
except Exception:
    pass

print(json.dumps({
    "success": True,
    "filename": raw_name,
    "path": out_path,
    "project_id": project_id,
    "sections": len(sections),
    "download_url": f"{DASHBOARD_URL}/api/artifacts/download/{project_id}/{raw_name}",
}))
